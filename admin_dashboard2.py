import pandas as pd
import streamlit as st
import psycopg2
import ast
import plotly.graph_objects as go
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import unicodedata
import requests
from PIL import Image, ImageDraw
import io
import re
from datetime import datetime
import pytz
import tempfile
import os

st.set_page_config(page_title="Archetypy Krzysztofa Hetmana – panel administratora", layout="wide")

COLOR_NAME_MAP = {
    "#000000": "Czerń",
    "#FFD700": "Złoto",
    "#282C34": "Granat (antracyt)",
    "#800020": "Burgund",
    "#E10600": "Czerwień",
    "#2E3141": "Grafitowy granat",
    "#FFFFFF": "Biel",
    "#4682B4": "Stalowy błękit",
    "#B0C4DE": "Jasny niebieskoszary",
    "#6C7A89": "Popielaty szary",
    "#B4D6B4": "Miętowa zieleń",
    "#A7C7E7": "Pastelowy błękit",
    "#FFD580": "Pastelowy żółty / beżowy",
    "#FA709A": "Róż malinowy",
    "#FEE140": "Jasny żółty",
    "#FFD6E0": "Bardzo jasny róż",
    "#FFB300": "Mocna żółć",
    "#FF8300": "Pomarańcz",
    "#FFD93D": "Pastelowa żółć",
    "#7C53C3": "Fiolet",
    "#3BE8B0": "Miętowy cyjan",
    "#87CEEB": "Błękit (Sky Blue)",
    "#43C6DB": "Turkusowy błękit",
    "#A0E8AF": "Seledyn",
    "#F9D371": "Złocisty żółty",
    "#8F00FF": "Fiolet (intensywny)",
    "#181C3A": "Granat bardzo ciemny",
    "#E0BBE4": "Pastelowy fiolet",
    "#F9F9F9": "Biel bardzo jasna",
    "#6CA0DC": "Pastelowy błękit",
    "#A3C1AD": "Pastelowa zieleń",
    "#FFF6C3": "Jasny kremowy",
    "#AAC9CE": "Pastelowy niebieskoszary",
    "#FFF200": "Żółty (cytrynowy)",
    "#FF0000": "Czerwień intensywna",
    "#FF6F61": "Łososiowy róż",
    "#8C564B": "Ciemy brąz",
    "#D62728": "Czerwień karmazynowa",
    "#1F77B4": "Chabrowy",
    "#9467BD": "Fiolet śliwkowy",
    "#F2A93B": "Miodowy żółty",
    "#17BECF": "Niebieski morski",
    "#E377C2": "Pastelowy róż fioletowy",
    "#7C46C5": "Fiolet szafirowy",
    "#2CA02C": "Zieleń trawiasta",
    "#9BD6F4": "Pastelowy błękit jasny",
    "#FF7F0E": "Jaskrawy pomarańcz",
}

ARCHE_NAMES_ORDER = [
    "Niewinny", "Mędrzec", "Odkrywca", "Buntownik", "Czarodziej", "Bohater",
    "Kochanek", "Błazen", "Towarzysz", "Opiekun", "Władca", "Twórca"
]

archetypes = {
    "Władca":   [1, 2, 3, 4],
    "Bohater":  [5, 6, 7, 8],
    "Mędrzec":  [9, 10, 11, 12],
    "Opiekun":  [13, 14, 15, 16],
    "Kochanek": [17, 18, 19, 20],
    "Błazen":   [21, 22, 23, 24],
    "Twórca":   [25, 26, 27, 28],
    "Odkrywca": [29, 30, 31, 32],
    "Czarodziej": [33, 34, 35, 36],
    "Towarzysz": [37, 38, 39, 40],
    "Niewinny": [41, 42, 43, 44],
    "Buntownik": [45, 46, 47, 48],
}

archetype_features = {
    "Władca": "Potrzeba kontroli, organizacji, zarządzanie, wprowadzanie ładu.",
    "Bohater": "Odwaga, walka z przeciwnościami, mobilizacja do działania.",
    "Mędrzec": "Wiedza, analityczność, logiczne argumenty, racjonalność.",
    "Opiekun": "Empatia, dbanie o innych, ochrona, troska.",
    "Kochanek": "Relacje, emocje, bliskość, autentyczność uczuć.",
    "Błazen": "Poczucie humoru, dystans, lekkość, rozładowywanie napięć.",
    "Twórca": "Kreatywność, innowacja, wyrażanie siebie, estetyka.",
    "Odkrywca": "Niezależność, zmiany, nowe doświadczenia, ekspresja.",
    "Czarodziej": "Transformacja, inspiracja, zmiana świata, przekuwanie idei w czyn.",
    "Towarzysz": "Autentyczność, wspólnota, prostota, bycie częścią grupy.",
    "Niewinny": "Optymizm, ufność, unikanie konfliktów, pozytywne nastawienie.",
    "Buntownik": "Kwestionowanie norm, odwaga w burzeniu zasad, radykalna zmiana."
}

archetype_extended = {
    "Władca": {
        "name": "Władca",
        "tagline": "Autorytet. Kontrola. Doskonałość.",
        "description": (
            "Archetyp Władcy w polityce uosabia siłę przywództwa, stabilność, pewność działania,kontrolę i odpowiedzialność za porządek społeczny. "
            "Władcy dążą do stabilności, bezpieczeństwa i efektywnego zarządzania. Politycy o tym archetypie często podkreślają swoją zdolność do podejmowania trudnych decyzji i utrzymywania porządku, nawet w trudnych czasach. "

            "Jako kandydat na prezydenta Lublina Władca stawia na porządek, wyznaczanie standardów rozwoju i podejmowanie stanowczych decyzji dla dobra wspólnego. "
            "Jest symbolem autentycznego autorytetu, przewodzenia i skutecznego zarządzania miastem. "
            "Buduje zaufanie, komunikując skuteczność, odpowiedzialność i gwarantując bezpieczeństwo mieszkańcom."
        ),
        "storyline": (
            "Narracja kampanii oparta na Władcy podkreśla spójność działań, panowanie nad trudnymi sytuacjami i sprawność w zarządzaniu miastem. "
            "Władca nie podąża za modą – wyznacza nowe standardy w samorządzie. "
            "Akcentuje dokonania, referencje i doświadczenie. Buduje obraz lidera odpowiadającego za przyszłość i prestiż Lublina."
        ),
        "recommendations": [
            "Używaj kolorystyki kojarzącej się z autorytetem – czerń, złoto, ciemny granat, burgund.",
            "Projektuj symbole: sygnety, herby miasta Lublin, podkreślając prestiż i zarządzanie.",
            "Komunikuj się językiem odpowiedzialności i troski o przyszłość miasta.",
            "Przekazuj komunikaty stanowczo, jednoznacznie, jako gospodarz miasta.",
            "Pokazuj osiągnięcia, inwestycje, referencje mieszkańców.",
            "Zadbaj o trwałość i jakość działań – nie obniżaj standardów.",
            "Twórz aurę elitarności: zamknięte konsultacje, spotkania liderów opinii.",
            "Przyciągaj wyborców ceniących bezpieczeństwo, stabilizację i prestiż miasta.",
            "Unikaj luźnego, żartobliwego tonu – postaw na klasę i profesjonalizm."
        ],
        "core_traits": [
            "Przywództwo", "Autorytet", "Stabilność", "Prestiż", "Kontrola", "Inspiracja", "Mistrzostwo"
        ],
        "strengths": [
            "Przywództwo", "zdecydowanie", "umiejętności organizacyjne"
        ],
        "weaknesses": [
            "Autorytaryzm", "kontrola", "oderwanie od rzeczywistości"
        ],
        "examples_person": [
            "Vladimir Putin", "Margaret Thatcher", "Xi Jinping", "Ludwik XIV", "Napoleon Bonaparte",
            "Jarosław Kaczyński"
        ],
        "example_brands": [
            "Rolex", "Mercedes-Benz", "IBM", "Microsoft", "Hugo Boss", "Silny samorząd"
        ],
        "color_palette": [
            "#000000", "#FFD700", "#282C34", "#800020", "#8C564B"
        ],
        "visual_elements": [
            "Korona", "Herb Lublina", "Sygnet", "Monogram", "Geometryczna, masywna typografia", "Symetria"
        ],
        "keyword_messaging": [
            "Lider Lublina", "Siła samorządu", "Stabilność", "Doskonałość działań", "Elita miasta", "Bezpieczeństwo"
        ],
        "watchword": [
            "Silne przywództwo i stabilność w niepewnych czasach."
        ],
        "questions": [
            "Jak komunikujesz mieszkańcom swoją pozycję lidera w Lublinie?",
            "W jaki sposób Twoje działania budują autorytet i zaufanie mieszkańców?",
            "Co robisz, by decyzje były stanowcze i jednoznaczne?",
            "Jak Twoje dokonania i inwestycje wzmacniają prestiż oraz bezpieczeństwo miasta?",
            "Jak zachęcasz wyborców do świadomego, silnego przywództwa?"
        ]
    },
    "Bohater": {
        "name": "Bohater",
        "tagline": "Determinacja. Odwaga. Sukces.",
        "description": (
            "Bohater w polityce to archetyp waleczności, determinacji i odwagi w podejmowaniu trudnych decyzji dla społeczności. "
            "Bohaterowie są gotowi stawić czoła wyzwaniom, pokonywać przeszkody i walczyć o lepszą przyszłość dla wszystkich. Ich celem jest udowodnienie swojej wartości poprzez odważne działania i inspirowanie innych do przekraczania własnych granic. Politycy o tym archetypie często podkreślają swoją gotowość do podejmowania trudnych decyzji i stawiania czoła przeciwnościom w imię dobra wspólnego. "

            "Kandydat Bohater mobilizuje mieszkańców do działania, bierze odpowiedzialność w najtrudniejszych momentach i broni interesów Lublina nawet pod presją."
        ),
        "storyline": (
            "Opowieść Bohatera to historia przezwyciężania kryzysów i stawania po stronie obywateli. "
            "Bohater nie rezygnuje nigdy, nawet w obliczu przeciwności. Jego postawa inspiruje i daje przykład innym samorządowcom."
        ),
        "recommendations": [
            "Komunikuj gotowość do działania, podkreślaj determinację w rozwiązywaniu problemów.",
            "Pokaż sukcesy i przykłady walki o interes mieszkańców.",
            "Stosuj dynamiczny język: zaznaczaj odwagę, mobilizację, sukces.",
            "Kolorystyka: czerwień, granat, biel.",
            "Pokazuj się w trudnych sytuacjach – reaguj natychmiast.",
            "Inspiruj współpracowników i mieszkańców do aktywności.",
            "Unikaj bierności, podkreślaj proaktywność."
        ],
        "core_traits": [
            "Odwaga", "Siła", "Determinacja", "Poświęcenie", "Sukces", "Inspiracja"
        ],
        "strengths": [
            "Odwaga", "determinacja", "kompetencja", "inspirowanie innych"
        ],
        "weaknesses": [
            "Arogancja", "obsesja na punkcie zwycięstwa", "skłonność do przechwalania się",
        ],
        "examples_person": [
            "Winston Churchill", "Wołodymyr Zełenski", "George Washington", "Józef Piłsudski"
        ],
        "example_brands": [
            "Nike", "Duracell", "FedEx", "Polska Husaria", "Patriotyczny samorząd"
        ],
        "color_palette": [
            "#E10600", "#2E3141", "#FFFFFF", "#D62728"
        ],
        "visual_elements": [
            "Peleryna", "Tarcza", "Aura odwagi", "Podniesiona dłoń", "Gwiazda"
        ],
        "keyword_messaging": [
            "Siła", "Zwycięstwo", "Poświęcenie", "Mobilizacja"
        ],
        "watchword": [
            "Odważne przywództwo dla lepszej przyszłości."
        ],
        "questions": [
            "Jak komunikujesz skuteczność w przezwyciężaniu kryzysów?",
            "Jak budujesz wizerunek walczącego o dobro mieszkańców?",
            "Jak pokazać determinację i niezłomność w działaniu publicznym?",
            "Które sukcesy świadczą o Twoim zaangażowaniu w trudnych sprawach?"
        ]
    },
    "Mędrzec": {
        "name": "Mędrzec",
        "tagline": "Wiedza. Racjonalność. Strategia.",
        "description": (
            "Mędrzec w polityce opiera komunikację na wiedzy, argumentacji i logicznym rozumowaniu oraz analitycznym podejściu. "
            "Mędrcy poszukują prawdy i wiedzy, wierząc, że informacja i zrozumienie są kluczem do rozwiązywania problemów. Politycy o tym archetypie często prezentują się jako eksperci, którzy podejmują decyzje w oparciu o fakty i analizy, a nie emocje czy ideologię. "

            "Kandydat Mędrzec wykorzystuje rozsądne analizy, doświadczenie oraz ekspercką wiedzę, by podejmować najlepsze decyzje dla całej społeczności."
        ),
        "storyline": (
            "Opowieść Mędrca to budowanie zaufania kompetencjami, przejrzystym uzasadnieniem propozycji i edukacją mieszkańców. "
            "Mędrzec nie działa pod wpływem impulsu; każda decyzja jest przemyślana i poparta faktami oraz wsłuchaniem się w potrzeby miasta."
        ),
        "recommendations": [
            "Wskazuj kompetencje, doświadczenie i eksperckość w zarządzaniu Lublinem.",
            "Komunikuj zrozumiale zawiłości miejskich inwestycji i decyzji.",
            "Stosuj wykresy, dane, analizy i argumenty – przemawiaj do rozumu obywateli.",
            "Zachowaj spokojny, opanowany ton.",
            "Używaj kolorystyki: błękit, szarość, granat.",
            "Podkreślaj racjonalność decyzji i transparentność działań.",
            "Unikaj populizmu – opieraj komunikację na faktach."
        ],
        "core_traits": [
            "Wiedza", "Rozwój", "Analiza", "Strategia", "Refleksja"
        ],
        "strengths": [
            "Inteligencja", "obiektywizm", "umiejętność analizy złożonych problemów"
        ],
        "weaknesses": [
            "Nadmierna rozwaga", "brak zdecydowania", "oderwanie od codziennych problemów"
        ],
        "examples_person": [
            "Angela Merkel", "Thomas Jefferson", "Lee Kuan Yew", "Bronisław Geremek"
        ],
        "example_brands": [
            "BBC", "Google", "MIT", "CNN", "Audi", "think tanki"
        ],
        "color_palette": [
            "#4682B4", "#B0C4DE", "#6C7A89", "#1F77B4"
        ],
        "visual_elements": [
            "Okulary", "Księga", "Wykres", "Lupa", "Symbole nauki"
        ],
        "keyword_messaging": [
            "Wiedza", "Argument", "Racjonalność", "Rozwój miasta"
        ],
        "watchword": [
            "Mądrość i wiedza w służbie społeczeństwa."
        ],
        "questions": [
            "Jak podkreślasz swoje doświadczenie i kompetencje?",
            "Jak przekonujesz mieszkańców argumentami i faktami?",
            "Jak edukujesz oraz tłumaczysz skomplikowane zmiany w mieście?",
            "W czym wyrażasz przewagę eksperckiej wiedzy nad populizmem?"
        ]
    },
    "Opiekun": {
        "name": "Opiekun",
        "tagline": "Empatia. Troska. Bezpieczeństwo.",
        "description": (
            "Opiekun w polityce to archetyp zaangażowania, wspierania i budowania poczucia wspólnoty. "
            "Archetyp Opiekuna reprezentuje troskę, empatię i chęć pomocy innym. "
            "Opiekunowie pragną chronić obywateli i zapewniać im bezpieczeństwo oraz wsparcie. Politycy o tym archetypie często skupiają się na polityce społecznej, ochronie zdrowia, edukacji i innych usługach publicznych, które poprawiają jakość życia obywateli. "
            "Kandydat Opiekun dba o najsłabszych, promuje działania prospołeczne, wdraża programy pomocowe i społecznie odpowiedzialne."
        ),
        "storyline": (
            "Narracja Opiekuna podkreśla działania integrujące, troskę o seniorów, rodziny, niepełnosprawnych i osoby wykluczone. "
            "Buduje poczucie bezpieczeństwa oraz odpowiedzialności urzędu miasta za wszystkich obywateli."
        ),
        "recommendations": [
            "Akcentuj działania na rzecz integracji i wsparcia mieszkańców.",
            "Pokaż realne efekty programów prospołecznych i pomocowych.",
            "Stosuj ciepłą kolorystykę: zieleń, błękit, żółcie.",
            "Używaj symboliki: dłonie, serca, uścisk.",
            "Komunikuj empatię i autentyczną troskę o każdą grupę mieszkańców.",
            "Prowadź otwarte konsultacje społeczne.",
            "Unikaj twardego, technokratycznego tonu."
        ],
        "core_traits": [
            "Empatia", "Troska", "Wspólnota", "Bezpieczeństwo", "Solidarność"
        ],
        "strengths": [
            "Empatia", "troska o innych", "budowanie zaufania"
        ],
        "weaknesses": [
            "Nadopiekuńczość", "unikanie trudnych decyzji", "podatność na manipulację"
        ],
        "examples_person": [
            "Jacinda Ardern", "Franklin D. Roosevelt", "Clement Attlee", "Władysław Kosiniak-Kamysz", "Jacek Kuroń"
        ],
        "example_brands": [
            "UNICEF", "Nivea", "Caritas", "WOŚP", "Pampers", "hospicja"
        ],
        "color_palette": [
            "#B4D6B4", "#A7C7E7", "#FFD580", "#9467BD"
        ],
        "visual_elements": [
            "Dłonie", "Serce", "Koło wspólnoty", "Symbol opieki"
        ],
        "keyword_messaging": [
            "Bezpieczeństwo mieszkańców", "Troska", "Wspólnota"
        ],
        "watchword": [
            "Troska i wsparcie dla każdego obywatela."
        ],
        "questions": [
            "Jak pokazujesz troskę i empatię wobec wszystkich mieszkańców?",
            "Jakie realne efekty mają wdrożone przez Ciebie programy pomocowe?",
            "W czym przejawia się Twoja polityka integrująca?",
            "Jak oceniasz skuteczność działań społecznych w mieście?"
        ]
    },
    "Kochanek": {
        "name": "Kochanek / Wielbiciel",
        "tagline": "Bliskość. Relacje. Pasja.",
        "description": (
            "Kochanek w polityce buduje pozytywne relacje z mieszkańcami, jest otwarty, komunikatywny i wzbudza zaufanie. "
            "Politycy Kochankowie podkreślają bliskość, autentyczność i partnerski dialog, sprawiając, że wyborcy czują się zauważeni i docenieni. "
            "Kochanek potrafi zbliżyć do siebie wyborców i sprawić, by czuli się zauważeni oraz docenieni."
        ),
        "storyline": (
            "Narracja Kochanka promuje serdeczność, ciepło i partnerskie traktowanie obywateli. "
            "Akcentuje jakość relacji z mieszkańcami, zespołem i innymi samorządami."
        ),
        "recommendations": [
            "Buduj relacje oparte na dialogu i wzajemnym szacunku.",
            "Stosuj ciepły, otwarty ton komunikacji.",
            "Promuj wydarzenia i inicjatywy integrujące społeczność.",
            "Używaj kolorystyki: czerwienie, róże, delikatne fiolety.",
            "Pokazuj, że wyborca jest dla Ciebie ważny.",
            "Doceniaj pozytywne postawy, sukcesy mieszkańców.",
            "Unikaj oficjalnego, zimnego tonu."
        ],
        "core_traits": [
            "Ciepło", "Relacje", "Bliskość", "Pasja", "Akceptacja"
        ],
        "strengths": [
            "Empatia", "bliskość", "autentyczność", "pasja"
        ],
        "weaknesses": [
            "Nadmierna emocjonalność", "faworyzowanie bliskich grup", "podatność na krytykę"
        ],
        "examples_person": [
            "Justin Trudeau", "Sanna Marin", "Eva Perón", "John F. Kennedy", "Benito Juárez", "François Mitterrand",
            "Aleksandra Dulkiewicz"
        ],
        "example_brands": [
            "Playboy", "Magnum", "Victoria's Secrets"
        ],
        "color_palette": [
            "#FA709A", "#FEE140", "#FFD6E0", "#FA709A"
        ],
        "visual_elements": [
            "Serce", "Uśmiech", "Gest bliskości"
        ],
        "keyword_messaging": [
            "Relacje", "Bliskość", "Społeczność"
        ],
        "watchword": [
            "Bliskość i pasja w służbie społeczeństwa."
        ],
        "questions": [
            "Jak komunikujesz otwartość i serdeczność wyborcom?",
            "Jakie działania podejmujesz, aby budować pozytywne relacje w mieście?",
            "Co robisz, by mieszkańcy czuli się ważni i zauważeni?"
        ]
    },
    "Błazen": {
        "name": "Błazen",
        "tagline": "Poczucie humoru. Dystans. Entuzjazm.",
        "description": (
            "Błazen w polityce wnosi lekkość, dystans i rozładowanie napięć. "
            "Używa humoru i autoironii, by rozbrajać napięcia oraz tworzyć wrażenie bliskości z wyborcami."
            "Kandydat-Błazen potrafi rozbawić, rozproszyć atmosferę, ale nigdy nie traci dystansu do siebie i powagi spraw publicznych."
        ),
        "storyline": (
            "Narracja Błazna to umiejętność śmiania się z problemów i codziennych wyzwań miasta, ale też dawania mieszkańcom nadziei oraz pozytywnej energii."
        ),
        "recommendations": [
            "Stosuj humor w komunikacji (ale z umiarem i klasą!).",
            "Rozluźniaj atmosferę podczas spotkań i debat.",
            "Podkreślaj pozytywne aspekty życia w mieście.",
            "Kolorystyka: żółcie, pomarańcze, intensywne kolory.",
            "Nie bój się autoironii.",
            "Promuj wydarzenia integrujące, rozrywkowe.",
            "Unikaj przesadnego formalizmu."
        ],
        "core_traits": [
            "Poczucie humoru", "Entuzjazm", "Dystans", "Optymizm"
        ],
        "strengths": [
            "Buduje rozpoznawalność", "umie odwrócić uwagę od trudnych tematów", "kreuje wizerunek 'swojskiego' lidera"
        ],
        "weaknesses": [
            "Łatwo przekracza granicę powagi", "ryzyko, że wyborcy nie odbiorą go serio"
        ],
        "examples_person": [
            "Boris Johnson", "Silvio Berlusconi", "Janusz Palikot",
        ],
        "example_brands": [
            "Old Spice", "M&M's", "Fanta", "Łomża", "kabarety"
        ],
        "color_palette": [
            "#FFB300", "#FF8300", "#FFD93D", "#F2A93B"
        ],
        "visual_elements": [
            "Uśmiech", "Czapka błazna", "Kolorowe akcenty"
        ],
        "keyword_messaging": [
            "Dystans", "Entuzjazm", "Radość"
        ],
        "watchword": [
            "Rozbraja śmiechem, inspiruje luzem."
        ],
        "questions": [
            "W jaki sposób wykorzystujesz humor w komunikacji publicznej?",
            "Jak rozładowujesz napięcia w sytuacjach kryzysowych?",
            "Co robisz, aby mieszkańcy mogli wspólnie się bawić i śmiać?"
        ]
    },
    "Twórca": {
        "name": "Twórca",
        "tagline": "Kreatywność. Innowacja. Wizja.",
        "description": (
            "Twórca charakteryzuje się innowacyjnością, kreatywnością i wizją. "
            "Twórcy dążą do budowania nowych rozwiązań i struktur, które odpowiadają na wyzwania przyszłości. Politycy o tym archetypie często podkreślają swoje innowacyjne podejście do rządzenia i zdolność do wprowadzania pozytywnych zmian. "

            "Jako prezydent Twórca nie boi się wdrażać oryginalnych, często nieszablonowych strategii."
        ),
        "storyline": (
            "Opowieść Twórcy jest oparta na zmianie, wprowadzaniu kreatywnych rozwiązań oraz inspirowaniu innych do współdziałania dla rozwoju Lublina."
        ),
        "recommendations": [
            "Proponuj i wdrażaj nietypowe rozwiązania w mieście.",
            "Pokazuj przykłady innowacyjnych projektów.",
            "Promuj kreatywność i otwartość na zmiany.",
            "Stosuj kolorystykę: zielenie, lazurowe błękity, fiolety.",
            "Doceniaj artystów, startupy, lokalne inicjatywy.",
            "Buduj wizerunek miasta-innowatora.",
            "Unikaj schematów i powtarzalnych projektów."
        ],
        "core_traits": [
            "Kreatywność", "Odwaga twórcza", "Inspiracja", "Wizja", "Nowatorstwo"
        ],
        "strengths": [
            "Innowacyjność", "wizjonerstwo", "kreatywność"
        ],
        "weaknesses": [
            "Brak realizmu", "ignorowanie praktycznych ograniczeń", "perfekcjonizm"
        ],
        "examples_person": [
            "Emmanuel Macron", "Tony Blair", "Konrad Adenauer", "Deng Xiaoping", "Mustafa Kemal Atatürk"
        ],
        "example_brands": [
            "Apple", "Tesla", "Lego", "Adobe", "startupy"
        ],
        "color_palette": [
            "#7C53C3", "#3BE8B0", "#87CEEB", "#17BECF"
        ],
        "visual_elements": [
            "Kostka Rubika", "Żarówka", "Kolorowe fale"
        ],
        "keyword_messaging": [
            "Innowacja", "Twórczość", "Wizja rozwoju"
        ],
        "watchword": [
            "Innowacyjne rozwiązania dla współczesnych wyzwań."
        ],
        "questions": [
            "Jak promujesz kreatywność i innowacyjność w mieście?",
            "Jakie oryginalne projekty wdrożyłeś lub planujesz wdrożyć?",
            "Jak inspirować mieszkańców do kreatywnego działania?"
        ]
    },
    "Odkrywca": {
        "name": "Odkrywca",
        "tagline": "Odwaga. Ciekawość. Nowe horyzonty.",
        "description": (
            "Archetyp Odkrywcy charakteryzuje się ciekawością, poszukiwaniem nowych możliwości i pragnieniem wolności. "
            "Odkrywcy pragną przełamywać granice i eksplorować nieznane terytoria. Politycy o tym archetypie często prezentują się jako wizjonerzy, którzy mogą poprowadzić społeczeństwo ku nowym horyzontom i możliwościom. "

            "Odkrywca poszukuje nowych rozwiązań, jest otwarty na zmiany i śledzi światowe trendy, które wdraża w Lublinie. "
            "Wybiera nowatorskie, nieoczywiste drogi dla rozwoju miasta i jego mieszkańców."
        ),
        "storyline": (
            "Opowieść Odkrywcy to wędrowanie poza schematami, miasto bez barier, eksperymentowanie z nowościami oraz angażowanie mieszkańców w odkrywcze projekty."
        ),
        "recommendations": [
            "Inicjuj nowe projekty i szukaj innowacji także poza Polską.",
            "Promuj przełamywanie standardów i aktywność obywatelską.",
            "Stosuj kolorystykę: turkusy, błękity, odcienie zieleni.",
            "Publikuj inspiracje z innych miast i krajów.",
            "Wspieraj wymiany młodzieży, startupy, koła naukowe.",
            "Unikaj stagnacji i powielania dawnych schematów."
        ],
        "core_traits": [
            "Odwaga", "Ciekawość", "Niezależność", "Nowatorstwo"
        ],
        "strengths": [
            "Innowacyjność", "adaptacyjność", "odwaga w podejmowaniu ryzyka"
        ],
        "weaknesses": [
            "Brak cierpliwości", "trudności z dokończeniem projektów", "ignorowanie tradycji"
        ],
        "examples_person": [
            "Olof Palme", "Shimon Peres", "Theodore Roosevelt", "Jawaharlal Nehru", "Elon Musk"
        ],
        "example_brands": [
            "NASA", "Jeep", "Red Bull", "National Geographic", "The North Face", "Amazon", "Nomadzi"
        ],
        "color_palette": [
            "#43C6DB", "#A0E8AF", "#F9D371", "#E377C2"
        ],
        "visual_elements": [
            "Mapa", "Kompas", "Droga", "Lupa"
        ],
        "keyword_messaging": [
            "Odkrywanie", "Nowe horyzonty", "Zmiana"
        ],
        "watchword": [
            "Odkrywanie nowych możliwości dla wspólnego rozwoju."
        ],
        "questions": [
            "Jak zachęcasz do odkrywania nowości w mieście?",
            "Jakie projekty wdrażasz, które nie były jeszcze realizowane w innych miastach?",
            "Jak budujesz wizerunek Lublina jako miejsca wolnego od barier?"
        ]
    },
    "Czarodziej": {
        "name": "Czarodziej",
        "tagline": "Transformacja. Inspiracja. Przełom.",
        "description": (
            "Czarodziej w polityce to wizjoner i transformator – wytycza nowy kierunek i inspiruje do zmian niemożliwych na pierwszy rzut oka. "
            "Czarodziej obiecuje głęboką przemianę społeczeństwa i nadaje wydarzeniom niemal magiczny sens. "

            "Dzięki jego inicjatywom Lublin przechodzi metamorfozy, w których niemożliwe staje się możliwe."
        ),
        "storyline": (
            "Opowieść Czarodzieja to zmiana wykraczająca poza rutynę, wyobraźnia, inspiracja, a także odwaga w stawianiu pytań i szukaniu odpowiedzi poza schematami."
        ),
        "recommendations": [
            "Wprowadzaj śmiałe, czasem kontrowersyjne pomysły w życie.",
            "Podkreślaj rolę wizji i inspiracji.",
            "Stosuj symbolikę: gwiazdy, zmiany, światło, 'magiczne' efekty.",
            "Stosuj kolorystykę: fiolety, granaty, akcent perłowy.",
            "Buduj wyobrażenie miasta jako miejsca możliwości.",
            "Unikaj banalnych, powtarzalnych rozwiązań."
        ],
        "core_traits": [
            "Inspiracja", "Przemiana", "Wyobraźnia", "Transcendencja"
        ],
        "strengths": [
            "Porywa wielką ideą", "motywuje do zmian", "potrafi łączyć symbole i narracje w spójny mit założycielski"
        ],
        "weaknesses": [
            "Oczekiwania mogą przerosnąć realne możliwości", "ryzyko oskarżeń o 'czcze zaklęcia'"
        ],
        "examples_person": [
            "Barack Obama", "Václav Klaus", "Nelson Mandela", "Martin Luther King"
        ],
        "example_brands": [
            "Intel", "Disney", "XBox", "Sony", "Polaroid", "Nowoczesny Lublin"
        ],
        "color_palette": [
            "#8F00FF", "#181C3A", "#E0BBE4", "#7C46C5"
        ],
        "visual_elements": [
            "Gwiazda", "Iskra", "Łuk magiczny"
        ],
        "keyword_messaging": [
            "Zmiana", "Inspiracja", "Możliwość"
        ],
        "watchword": [
            "Zmieniam rzeczywistość w to, co dziś wydaje się niemożliwe."
        ],
        "questions": [
            "Jak pokazujesz mieszkańcom, że niemożliwe jest możliwe?",
            "Jakie innowacje budują wizerunek miasta kreatywnego i nowoczesnego?",
            "Jak inspirujesz społeczność do patrzenia dalej?"
        ]
    },
    "Towarzysz": {
        "name": "Towarzysz / Zwykły Człowiek",
        "tagline": "Wspólnota. Prostota. Bliskość.",
        "description": (
            "Towarzysz w polityce stoi blisko ludzi, jest autentyczny, stawia na prostotę, tworzenie bezpiecznej wspólnoty społecznej oraz zrozumienie codziennych problemów obywateli. "
            "Nie udaje, nie buduje dystansu – jest 'swojakiem', na którym można polegać. "
            "Politycy o tym archetypie podkreślają swoje zwyczajne pochodzenie i doświadczenia, pokazując, że rozumieją troski i aspiracje przeciętnych ludzi. "
            "Ich siłą jest umiejętność budowania relacji i tworzenia poczucia wspólnoty."
        ),
        "storyline": (
            "Opowieść Towarzysza koncentruje się wokół wartości rodzinnych, codziennych wyzwań, pracy od podstaw oraz pielęgnowania lokalnej tradycji."
        ),
        "recommendations": [
            "Podkreślaj prostotę i codzienność w komunikacji.",
            "Stosuj jasne, proste słowa i obrazy.",
            "Buduj atmosferę równości (każdy ma głos).",
            "Stosuj kolorystykę: beże, błękity, zielone akcenty.",
            "Doceniaj lokalność i rodzinność.",
            "Promuj wspólnotowe inicjatywy.",
            "Unikaj dystansu i języka eksperckiego."
        ],
        "core_traits": [
            "Autentyczność", "Wspólnota", "Prostota", "Równość"
        ],
        "strengths": [
            "Autentyczność", "empatia", "umiejętność komunikacji z obywatelami"
        ],
        "weaknesses": [
            "Brak wizji", "ograniczona perspektywa", "unikanie trudnych decyzji"
        ],
        "examples_person": [
            "Joe Biden", "Bernie Sanders", "Andrzej Duda", "Pedro Sánchez", "Jeremy Corbyn"
        ],
        "example_brands": [
            "Ikea", "Skoda", "Żabka"
        ],
        "color_palette": [
            "#F9F9F9", "#6CA0DC", "#A3C1AD", "#2CA02C"
        ],
        "visual_elements": [
            "Dom", "Krąg ludzi", "Prosta ikona dłoni"
        ],
        "keyword_messaging": [
            "Bliskość", "Razem", "Prostota"
        ],
        "watchword": [
            "Blisko ludzi i ich codziennych spraw."
        ],
        "questions": [
            "Jak podkreślasz autentyczność i codzienność?",
            "Jak pielęgnujesz lokalność i wspólnotę?",
            "Co robisz, by każdy mieszkaniec czuł się zauważony?"
        ]
    },
    "Niewinny": {
        "name": "Niewinny",
        "tagline": "Optymizm. Nadzieja. Nowy początek.",
        "description": (
            "Niewinny w polityce otwarcie komunikuje pozytywne wartości, niesie nadzieję i podkreśla wiarę w zmiany na lepsze. "
            "Głosi prostą, pozytywną wizję dobra wspólnego i nadziei. "
            "Kandydat–Niewinny buduje zaufanie szczerością i skutecznie apeluje o współpracę dla wspólnego dobra."
        ),
        "storyline": (
            "Opowieść Niewinnego buduje napięcie wokół pozytywnych emocji, odwołuje się do marzeń o lepszym Lublinie i wiary we wspólny sukces."
        ),
        "recommendations": [
            "Komunikuj optymizm, wiarę w ludzi i dobre intencje.",
            "Stosuj jasną kolorystykę: biele, pastele, żółcie.",
            "Dziel się sukcesami społeczności.",
            "Stawiaj na transparentność działań.",
            "Angażuj się w kampanie edukacyjne i społeczne.",
            "Unikaj negatywnego przekazu, straszenia, manipulacji."
        ],
        "core_traits": [
            "Optymizm", "Nadzieja", "Współpraca", "Szlachetność"
        ],
        "strengths": [
            "Łatwo zyskuje zaufanie", "łagodzi polaryzację", "odwołuje się do uniwersalnych wartości."
        ],
        "weaknesses": [
            "Może być postrzegany jako naiwny", "trudniej mu prowadzić twarde negocjacje"
        ],
        "examples_person": [
            "Jimmy Carter", "Václav Havel", "Szymon Hołownia"
        ],
        "example_brands": [
            "Dove", "Milka", "Kinder", "Polska Akcja Humanitarna"
        ],
        "color_palette": [
            "#FFF6C3", "#AAC9CE", "#FFF200", "#9BD6F4"
        ],
        "visual_elements": [
            "Gołąb", "Słońce", "Dziecko"
        ],
        "keyword_messaging": [
            "Nadzieja", "Optymizm", "Wspólnie"
        ],
        "watchword": [
            "Uczciwość i nadzieja prowadzą naprzód."
        ],
        "questions": [
            "Jak budujesz wizerunek pozytywnego samorządowca?",
            "Jak zachęcasz mieszkańców do dzielenia się nadzieją?",
            "Jak komunikujesz szczerość i otwartość?"
        ]
    },
    "Buntownik": {
        "name": "Buntownik",
        "tagline": "Zmiana. Odwaga. Przełom.",
        "description": (
            "Buntownik w polityce odważnie kwestionuje zastane układy, nawołuje do zmiany i walczy o nowe, lepsze reguły gry. "
            "Archetyp Buntownika charakteryzuje się odwagą w kwestionowaniu status quo i dążeniem do fundamentalnych zmian. "
            "Buntownicy sprzeciwiają się istniejącym strukturom władzy i konwencjom, proponując radykalne rozwiązania."
            "Politycy o tym archetypie często prezentują się jako outsiderzy, którzy chcą zburzyć skorumpowany system i wprowadzić nowy porządek."

            "Kandydat Buntownik odważnie kwestionuje zastane układy, nawołuje do zmiany i walczy o nowe, lepsze reguły gry w mieście. "
            "Potrafi ściągnąć uwagę i zjednoczyć mieszkańców wokół śmiałych idei. "
        ),
        "storyline": (
            "Narracja Buntownika podkreśla walkę z niesprawiedliwością i stagnacją, wytykanie błędów władzy i radykalne pomysły na rozwój Lublina."
        ),
        "recommendations": [
            "Akcentuj odwagę do mówienia „nie” starym rozwiązaniom.",
            "Publikuj manifesty i odważne postulaty.",
            "Stosuj wyrazistą kolorystykę: czernie, czerwienie, ostre kolory.",
            "Inspiruj mieszkańców do aktywnego sprzeciwu wobec barier rozwojowych.",
            "Podkreślaj wolność słowa, swobody obywatelskie.",
            "Unikaj koncentrowania się wyłącznie na krytyce – pokazuj pozytywne rozwiązania."
        ],
        "core_traits": [
            "Odwaga", "Bezpardonowość", "Radykalizm", "Niepokorność"
        ],
        "strengths": [
            "Odwaga", "autentyczność", "zdolność inspirowania do zmian"
        ],
        "weaknesses": [
            "Nadmierna konfrontacyjność", "brak kompromisu", "trudności w budowaniu koalicji"
        ],
        "examples_person": [
            "Donald Trump", "Marine Le Pen", "Sławomir Mentzen", "Lech Wałęsa", "Aleksiej Nawalny"
        ],
        "example_brands": [
            "Harley Davidson", "Jack Daniel's", "Greenpeace", "Virgin", "Bitcoin"
        ],
        "color_palette": [
            "#000000", "#FF0000", "#FF6F61", "#FF7F0E"
        ],
        "visual_elements": [
            "Piorun", "Megafon", "Odwrócona korona"
        ],
        "keyword_messaging": [
            "Zmiana", "Rewolucja", "Nowe reguły"
        ],
        "watchword": [
            "Rewolucyjne zmiany dla lepszego jutra."
        ],
        "questions": [
            "Jak komunikujesz odwagę i gotowość do zmiany?",
            "Jak mobilizujesz do zrywania z przeszłością?",
            "Co robisz, by mieszkańcy mieli w Tobie rzecznika zmiany?"
        ]
    }
}

ARCHE_IMG_URL = "https://justynakopec.pl/wp-content/uploads/2024/08/Archetypy-marki-Justyna-Kopec.png"
ARCHE_NAME_TO_IDX = {n.lower(): i for i, n in enumerate(ARCHE_NAMES_ORDER)}

@st.cache_data
def load_base_arche_img():
    resp = requests.get(ARCHE_IMG_URL, stream=True, timeout=30)
    resp.raise_for_status()
    img = Image.open(io.BytesIO(resp.content)).convert("RGBA")
    return img

def mask_for(idx, color):
    base = load_base_arche_img()
    w, h = base.size
    cx, cy = w//2, h//2
    rad = w//2
    mask = Image.new("RGBA", (w, h), (0,0,0,0))
    draw = ImageDraw.Draw(mask)
    start = -90 + idx*30
    end = start + 30
    draw.pieslice([cx-rad, cy-rad, cx+rad, cy+rad], start, end, fill=color)
    return mask

def compose_archetype_highlight(idx_main, idx_aux=None):
    base = load_base_arche_img().copy()
    if idx_aux is not None and idx_aux != idx_main and idx_aux < 12:
        mask_aux = mask_for(idx_aux, (255,210,47,140))
        base.alpha_composite(mask_aux)
    if idx_main is not None:
        mask_main = mask_for(idx_main, (255,0,0,140))
        base.alpha_composite(mask_main)
    return base

def archetype_name_to_img_idx(name):
    try:
        return ARCHE_NAMES_ORDER.index(name)
    except ValueError:
        return None

def clean_pdf_text(text):
    if text is None:
        return ""
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')
    text = text.replace("–", "-").replace("—", "-")
    return text

@st.cache_data(ttl=30)
def load():
    try:
        conn = psycopg2.connect(
            host=st.secrets["db_host"],
            database=st.secrets["db_name"],
            user=st.secrets["db_user"],
            password=st.secrets["db_pass"],
            port=st.secrets.get("db_port", 5432),
            sslmode="require"
        )
        df = pd.read_sql("SELECT * FROM ap48_responses", con=conn)
        conn.close()
        if "created_at" in df.columns:
            df["created_at"] = pd.to_datetime(df["created_at"])
        def parse_answers(x):
            if isinstance(x, list):
                return x
            try:
                import json
                return json.loads(x)
            except:
                try:
                    return ast.literal_eval(x)
                except:
                    return None
        if "answers" in df.columns:
            df["answers"] = df["answers"].apply(parse_answers)
        return df
    except Exception as e:
        st.warning(f"Błąd podczas ładowania danych: {e}")
        return pd.DataFrame()

def archetype_scores(answers):
    if not isinstance(answers, list) or len(answers) < 48:
        return {k: None for k in archetypes}
    out = {}
    for name, idxs in archetypes.items():
        out[name] = sum(answers[i-1] for i in idxs)
    return out

def archetype_percent(scoresum):
    if scoresum is None:
        return None
    return round(scoresum / 20 * 100, 1)

def pick_main_and_aux_archetype(archetype_means, archetype_order):
    vals = list(archetype_means.values())
    max_val = max(vals)
    main_candidates = [k for k, v in archetype_means.items() if v == max_val]
    main_type = next(k for k in archetype_order if k in main_candidates)
    aux_vals = [v for k, v in archetype_means.items() if k != main_type]
    if not aux_vals:
        return main_type, None
    aux_val = max(aux_vals)
    aux_candidates = [k for k, v in archetype_means.items() if v == aux_val and k != main_type]
    second_type = next((k for k in archetype_order if k in aux_candidates), None)
    return main_type, second_type

def add_hyperlink(paragraph, url, text, color="0000EE", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def pil_to_bytesio(img, format='PNG'):
    buf = BytesIO()
    img.save(buf, format=format)
    buf.seek(0)
    return buf

def radar_pil_img(fig):
    buf = BytesIO()
    fig.write_image(buf, format="png")
    buf.seek(0)
    return Image.open(buf)

def archetype_table_img(table_df):
    from matplotlib import pyplot as plt
    import matplotlib
    matplotlib.rcParams['font.size'] = 18
    fig, ax = plt.subplots(figsize=(6,6))
    ax.axis('off')
    t = ax.table(cellText=table_df.values,
                colLabels=table_df.columns,
                cellLoc='center', loc='center')
    t.auto_set_font_size(False)
    t.set_fontsize(17)
    t.auto_set_column_width(col=list(range(len(table_df.columns))))
    fig.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=160)
    plt.close(fig)
    buf.seek(0)
    return Image.open(buf)
def get_image_bytesio_from_url(url):
    response = requests.get(url)
    response.raise_for_status()
    return BytesIO(response.content)

def add_two_column_table_atuty_slabosci(doc, strengths, weaknesses):
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Atuty"
    hdr_cells[1].text = "Słabości"
    # Format nagłówków
    for i, cell in enumerate(hdr_cells):
        for p in cell.paragraphs:
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(11)
            if i == 0:
                run.font.color.rgb = RGBColor(20, 80, 200)  # niebieski
            else:
                run.font.color.rgb = RGBColor(200, 20, 20)  # czerwony
    max_rows = max(len(strengths or []), len(weaknesses or []))
    for i in range(max_rows):
        row_cells = table.add_row().cells
        # Atuty (lewa kolumna)
        if strengths and i < len(strengths):
            p = row_cells[0].paragraphs[0]
            run = p.add_run("● ")
            run.font.color.rgb = RGBColor(15, 90, 180)
            run.bold = True
            p.add_run(str(strengths[i]))
        # Słabości (prawa kolumna)
        if weaknesses and i < len(weaknesses):
            p = row_cells[1].paragraphs[0]
            run = p.add_run("● ")
            run.font.color.rgb = RGBColor(220, 30, 30)
            run.bold = True
            p.add_run(str(weaknesses[i]))
    table.style = "Table Grid"

def add_examples_person_with_images(doc, persons):
    p = doc.add_paragraph("Przykłady polityków: ")
    if not persons:
        p.add_run("brak danych")
        return
    for name in persons:
        p.add_run(name)
        p.add_run(", ")
    # Nie ma automatycznego wstawiania miniaturek do każdego nazwiska bez ręcznej mapy imię->foto
    # Jeśli chcesz, możesz dodać słownik np. {"Angela Merkel": "url_do_foto"}, i w pętli pobrać i wstawić miniaturkę

def export_word(main_type, second_type, features, main, second,
                mean_archetype_scores=None, main_img=None, radar_img=None, table_img=None, num_ankiet=None):
    doc = Document()
    section = doc.sections[0]
    section.page_height = 16840  # A4
    section.page_width = 11900
    section.left_margin = section.right_margin = 900000 // 20

    # Okładka z obrazkiem z URL
    cover_bytes = get_image_bytesio_from_url(ARCHE_IMG_URL)
    doc.add_picture(cover_bytes, width=Inches(3.2))
    title = doc.add_heading("Raport archetypowy Krzysztofa Hetmana", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Wybory na Prezydenta Miasta Lublin").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Data wygenerowania raportu: " + datetime.now().strftime("%Y-%m-%d %H:%M")).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()
    intro = doc.add_paragraph()
    intro.add_run(
        "Archetypy to uniwersalne wzorce osobowości, które od wieków pomagają ludziom rozumieć świat i budować autentyczną tożsamość. "
        "Współczesna psychologia, marketing i polityka potwierdzają, że trafnie zdefiniowany archetyp jest potężnym narzędziem komunikacji, pozwalającym budować rozpoznawalność, zaufanie i emocjonalny kontakt. "
        "\n\nW polityce archetyp pomaga wyeksponować najważniejsze cechy lidera, porządkuje przekaz, wzmacnia spójność strategii oraz wyraźnie różnicuje kandydata na tle konkurencji. "
        "Wyniki badań archetypowych stanowią dziś istotny fundament do tworzenia skutecznej narracji wyborczej, strategii wizerunkowej i komunikacji z wyborcami.\n\n"
        "Badanie AP-48 zrealizowane dla Krzysztofa Hetmana pozwala zidentyfikować archetyp główny oraz archetyp wspierający, a więc dwa najważniejsze wzorce, które mogą wzmocnić jego pozycjonowanie w walce o urząd Prezydenta Miasta Lublin."
    )
    doc.add_paragraph()
    pimg = doc.add_paragraph("Koło archetypów (źródło: justynakopec.pl):")
    add_hyperlink(pimg, ARCHE_IMG_URL, "Zobacz w przeglądarce")
    wheel_bytes = get_image_bytesio_from_url(ARCHE_IMG_URL)
    doc.add_picture(wheel_bytes, width=Inches(5))
    doc.add_page_break()
    doc.add_heading("Podsumowanie wyników badań archetypowych", 1)
    doc.add_paragraph(f"W badaniu udział wzięło {num_ankiet} osób. Na podstawie analiz średnich wyników dla wszystkich archetypów zdefiniowano profil główny i pomocniczy kandydata.")
    doc.add_paragraph("Wizualizacja profilu archetypowego – wykres radarowy:")
    if radar_img:
        doc.add_picture(radar_img, width=Inches(5.5))
    doc.add_paragraph("Podświetlenie głównego (czerwony) i pomocniczego (żółty) archetypu (grafika na bazie koła Justyny Kopeć):")
    if main_img:
        doc.add_picture(main_img, width=Inches(4.2))
    doc.add_paragraph("Liczebność archetypów wśród uczestników badania:")
    if table_img:
        doc.add_picture(table_img, width=Inches(5.6))
    doc.add_page_break()

    # --- GŁÓWNY ARCHETYP ---
    doc.add_heading(f"GŁÓWNY ARCHETYP: {main_type}", 1)
    doc.add_paragraph(f"Cechy kluczowe: {features.get(main_type, '-')}", style="List Bullet")
    doc.add_paragraph(main.get("description", ""), style="Normal")
    doc.add_paragraph("Storyline: " + main.get("storyline", ""), style="Normal")
    # Tabela Atuty/Słabości
    add_two_column_table_atuty_slabosci(doc, main.get("strengths", []), main.get("weaknesses", []))
    # Slogan
    if main.get("watchword"):
        doc.add_paragraph("Slogan: " + "; ".join(main.get("watchword")), style="Intense Quote")
    # Kolory
    if main.get("color_palette"):
        doc.add_paragraph("Kolory: " + ", ".join([COLOR_NAME_MAP.get(c, c) for c in main.get("color_palette")]), style="Normal")
    # Elementy wizualne
    doc.add_paragraph("Elementy wizualne: " + ", ".join(main.get("visual_elements", [])), style="Normal")
    # Przykłady polityków (z miniaturkami, jeśli chcesz rozwinąć – tutaj nazwy)
    add_examples_person_with_images(doc, main.get("examples_person", []))
    # Przykłady marek
    if main.get("example_brands"):
        doc.add_paragraph("Przykłady marek: " + ", ".join(main.get("example_brands")), style="Normal")
    # Rekomendacje
    doc.add_paragraph("Rekomendacje:").bold = True
    for rec in main.get("recommendations", []):
        doc.add_paragraph(rec, style="List Bullet 2")
    # Ciekawostki
    doc.add_paragraph("Ciekawostki:").bold = True
    for kw in main.get("keyword_messaging", []):
        doc.add_paragraph(kw, style="List Bullet 2")
    doc.add_paragraph()  # odstęp

    # --- ARCHETYP POMOCNICZY ---
    if second_type and second_type != main_type:
        doc.add_page_break()
        doc.add_heading(f"ARCHETYP POMOCNICZY: {second_type}", 1)
        doc.add_paragraph(f"Cechy kluczowe: {features.get(second_type, '-')}", style="List Bullet")
        doc.add_paragraph(second.get("description", ""), style="Normal")
        doc.add_paragraph("Storyline: " + second.get("storyline", ""), style="Normal")
        # Tabela Atuty/Słabości
        add_two_column_table_atuty_slabosci(doc, second.get("strengths", []), second.get("weaknesses", []))
        # Slogan
        if second.get("watchword"):
            doc.add_paragraph("Slogan: " + "; ".join(second.get("watchword")), style="Intense Quote")
        # Kolory
        if second.get("color_palette"):
            doc.add_paragraph("Kolory: " + ", ".join([COLOR_NAME_MAP.get(c, c) for c in second.get("color_palette")]), style="Normal")
        # Elementy wizualne
        doc.add_paragraph("Elementy wizualne: " + ", ".join(second.get("visual_elements", [])), style="Normal")
        # Przykłady polityków
        add_examples_person_with_images(doc, second.get("examples_person", []))
        # Przykłady marek
        if second.get("example_brands"):
            doc.add_paragraph("Przykłady marek: " + ", ".join(second.get("example_brands")), style="Normal")
        # Rekomendacje
        doc.add_paragraph("Rekomendacje:").bold = True
        for rec in second.get("recommendations", []):
            doc.add_paragraph(rec, style="List Bullet 2")
        # Ciekawostki
        doc.add_paragraph("Ciekawostki:").bold = True
        for kw in second.get("keyword_messaging", []):
            doc.add_paragraph(kw, style="List Bullet 2")
    doc.add_page_break()
    doc.add_paragraph(
        "Raport został automatycznie wygenerowany na podstawie wyników testu AP-48 oraz aktualnych opracowań z zakresu archetypów liderów i psychologii komunikacji."
    )
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_pdf(main_type, second_type, features, main, second,
               mean_archetype_scores=None, main_img=None, radar_img=None, table_img=None, num_ankiet=None):
    pdf = FPDF()
    # Rejestracja wszystkich wariantów czcionki (ważne!)
    pdf.add_font('DejaVu', '', 'fonts/DejaVuSans.ttf', uni=True)
    pdf.add_font('DejaVu', 'B', 'fonts/DejaVuSans-Bold.ttf', uni=True)
    pdf.add_font('DejaVu', 'I', 'fonts/DejaVuSans-Oblique.ttf', uni=True)
    pdf.add_font('DejaVu', 'BI', 'fonts/DejaVuSans-BoldOblique.ttf', uni=True)
    pdf.set_font("DejaVu", "", 22)
    pdf.add_page()
    pdf.set_text_color(13,85,185)
    pdf.cell(0, 16, "RAPORT ARCHETYPÓW: KRZYSZTOF HETMAN", ln=1, align='C')
    pdf.set_font("DejaVu", "", 14)
    pdf.set_text_color(0)
    pdf.cell(0, 10, "Wybory na Prezydenta Miasta Lublin", ln=1, align='C')
    pdf.set_font("DejaVu", "", 10)
    pdf.cell(0, 9, f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=1, align='C')
    pdf.ln(9)
    pdf.set_font("DejaVu", "B", 16)
    pdf.cell(0, 12, "Wstęp: Czym są archetypy i jak je wykorzystywać?", ln=1)
    pdf.set_font("DejaVu", "", 11)
    pdf.multi_cell(0, 8,
        "Archetypy to uniwersalne wzorce osobowości, które pomagają budować autentyczną tożsamość oraz rozpoznawalność lidera. "
        "W polityce archetypy pomagają zrozumieć, jak kandydat jest postrzegany i jakie emocje budzi jego komunikacja. "
        "Stanowią praktyczne narzędzie do precyzyjnego pozycjonowania, budowy strategii i skutecznego dialogu z wyborcami.\n\n"
        "W badaniu AP-48 przeanalizowano, które archetypy dominują u Krzysztofa Hetmana. Pozwala to tworzyć spójny, inspirujący wizerunek, osadzony w realnych oczekiwaniach społecznych."
    )
    pdf.ln(3)
    pdf.set_font("DejaVu", "I", 10)
    pdf.set_text_color(13, 85, 185)
    pdf.cell(0, 8, "Zobacz wizualizację koła archetypów: https://justynakopec.pl/wp-content/uploads/2024/08/Archetypy-marki-Justyna-Kopec.png", ln=1, align='L')
    pdf.set_text_color(0)
    pdf.ln(4)
    pdf.set_font("DejaVu", "B", 15)
    pdf.cell(0, 11, "Podsumowanie wyników badania:", ln=1)
    pdf.set_font("DejaVu", "", 11)
    pdf.multi_cell(0, 8, f"W badaniu udział wzięło {num_ankiet} osób. Zidentyfikowano następujący profil archetypowy dla kandydata.")
    pdf.ln(1)
    if radar_img:
        radar_img.seek(0)
        img = Image.open(radar_img)
        tmp_radar = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        img.save(tmp_radar, "PNG")
        tmp_radar.flush()
        pdf.image(tmp_radar.name, w=120)
    pdf.ln(2)
    if main_img:
        main_img.seek(0)
        img = Image.open(main_img)
        tmp_main = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        img.save(tmp_main, "PNG")
        tmp_main.flush()
        pdf.image(tmp_main.name, w=72)
    pdf.ln(2)
    if table_img:
        table_img.seek(0)
        img = Image.open(table_img)
        tmp_tbl = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        img.save(tmp_tbl, "PNG")
        tmp_tbl.flush()
        pdf.image(tmp_tbl.name, w=140)
    pdf.ln(3)
    pdf.set_font("DejaVu", "B", 13)
    pdf.set_text_color(210, 0, 0)
    pdf.cell(0, 9, f"Główny archetyp: {main_type}", ln=1)
    pdf.set_font("DejaVu", "", 11)
    pdf.set_text_color(0)
    pdf.multi_cell(0, 7, f"Cechy kluczowe: {features.get(main_type, '-')}")
    pdf.multi_cell(0, 7, main.get('description', ""))
    pdf.multi_cell(0, 7, "Storyline: " + main.get('storyline', ""))

    # Tabela Atuty/Słabości
    strengths = main.get("strengths", [])
    weaknesses = main.get("weaknesses", [])
    pdf.set_font("DejaVu", "B", 11)
    pdf.cell(0, 7, "Atuty         Słabości", ln=1)
    pdf.set_font("DejaVu", "", 10)
    for i in range(max(len(strengths), len(weaknesses))):
        left = f"● {strengths[i]}" if i < len(strengths) else ""
        right = f"● {weaknesses[i]}" if i < len(weaknesses) else ""
        pdf.set_text_color(15, 90, 180)
        pdf.cell(60, 7, left)
        pdf.set_text_color(220, 30, 30)
        pdf.cell(0, 7, right, ln=1)
        pdf.set_text_color(0)
    # Slogan
    if main.get("watchword"):
        pdf.set_font("DejaVu", "I", 11)
        pdf.multi_cell(0, 7, "Slogan: " + "; ".join(main.get("watchword")))
    # Kolory
    if main.get("color_palette"):
        pdf.set_font("DejaVu", "", 10)
        pdf.multi_cell(0, 7, "Kolory: " + ", ".join([COLOR_NAME_MAP.get(c, c) for c in main.get("color_palette")]))
    # Elementy wizualne
    pdf.multi_cell(0, 7, "Elementy wizualne: " + ", ".join(main.get("visual_elements", [])))
    # Przykłady polityków
    if main.get("examples_person"):
        pdf.set_font("DejaVu", "B", 10)
        pdf.cell(0, 7, "Przykłady polityków: " + ", ".join(main.get("examples_person")), ln=1)
    # Przykłady marek
    if main.get("example_brands"):
        pdf.set_font("DejaVu", "", 10)
        pdf.multi_cell(0, 7, "Przykłady marek: " + ", ".join(main.get("example_brands")))
    pdf.set_font("DejaVu", "B", 11)
    pdf.cell(0, 7, "Rekomendacje:", ln=1)
    pdf.set_font("DejaVu", "", 10)
    for rec in main.get('recommendations', []):
        pdf.multi_cell(0, 7, "- " + rec)
    pdf.set_font("DejaVu", "B", 11)
    pdf.cell(0, 7, "Ciekawostki:", ln=1)
    pdf.set_font("DejaVu", "", 10)
    for kw in main.get('keyword_messaging', []):
        pdf.multi_cell(0, 7, "- " + kw)
    pdf.ln(1)

    if second_type and second_type != main_type:
        pdf.set_font("DejaVu", "B", 12)
        pdf.set_text_color(210, 165, 31)
        pdf.cell(0, 9, f"Archetyp pomocniczy: {second_type}", ln=1)
        pdf.set_text_color(0)
        pdf.set_font("DejaVu", "", 11)
        pdf.multi_cell(0, 7, f"Cechy kluczowe: {features.get(second_type, '-')}")
        pdf.multi_cell(0, 7, second.get('description', ""))
        pdf.multi_cell(0, 7, "Storyline: " + second.get('storyline', ""))
        strengths = second.get("strengths", [])
        weaknesses = second.get("weaknesses", [])
        pdf.set_font("DejaVu", "B", 11)
        pdf.cell(0, 7, "Atuty         Słabości", ln=1)
        pdf.set_font("DejaVu", "", 10)
        for i in range(max(len(strengths), len(weaknesses))):
            left = f"● {strengths[i]}" if i < len(strengths) else ""
            right = f"● {weaknesses[i]}" if i < len(weaknesses) else ""
            pdf.set_text_color(15, 90, 180)
            pdf.cell(60, 7, left)
            pdf.set_text_color(220, 30, 30)
            pdf.cell(0, 7, right, ln=1)
            pdf.set_text_color(0)
        if second.get("watchword"):
            pdf.set_font("DejaVu", "I", 11)
            pdf.multi_cell(0, 7, "Slogan: " + "; ".join(second.get("watchword")))
        if second.get("color_palette"):
            pdf.set_font("DejaVu", "", 10)
            pdf.multi_cell(0, 7, "Kolory: " + ", ".join([COLOR_NAME_MAP.get(c, c) for c in second.get("color_palette")]))
        pdf.multi_cell(0, 7, "Elementy wizualne: " + ", ".join(second.get("visual_elements", [])))
        if second.get("examples_person"):
            pdf.set_font("DejaVu", "B", 10)
            pdf.cell(0, 7, "Przykłady polityków: " + ", ".join(second.get("examples_person")), ln=1)
        if second.get("example_brands"):
            pdf.set_font("DejaVu", "", 10)
            pdf.multi_cell(0, 7, "Przykłady marek: " + ", ".join(second.get("example_brands")))
        pdf.set_font("DejaVu", "B", 11)
        pdf.cell(0, 7, "Rekomendacje:", ln=1)
        pdf.set_font("DejaVu", "", 10)
        for rec in second.get('recommendations', []):
            pdf.multi_cell(0, 7, "- " + rec)
        pdf.set_font("DejaVu", "B", 11)
        pdf.cell(0, 7, "Ciekawostki:", ln=1)
        pdf.set_font("DejaVu", "", 10)
        for kw in second.get('keyword_messaging', []):
            pdf.multi_cell(0, 7, "- " + kw)
        pdf.ln(2)
    pdf.set_font("DejaVu", "I", 9)
    pdf.cell(0, 7, "Raport automatycznie wygenerowany na podstawie testu AP-48 oraz aktualnych badań nad archetypami liderów.", ln=1)
    # NIE kodujemy .encode('latin1'), tylko .encode('utf-8')!
    pdf_bytes = bytes(pdf.output(dest='S'), encoding='utf-8')
    buf = BytesIO(pdf_bytes)
    buf.seek(0)
    return buf

# ================= PANEL ADMINA + RAPORTY ===================

with st.expander("Panel administracyjny"):
    df = load()
    st.write("Dane z bazy (ostatnie 500):", df.tail(500))
    if st.button("Eksportuj wszystkie dane do CSV"):
        st.download_button("Pobierz plik CSV", df.to_csv(index=False), file_name="ap48_all_data.csv")

results_df = df
if not results_df.empty and "answers" in results_df.columns:
    results_df["archetype_scores"] = results_df["answers"].apply(archetype_scores)
    for name in archetypes:
        results_df[name] = results_df["archetype_scores"].apply(lambda x: x[name] if x else None)
    archetype_names = list(archetypes.keys())
    mean_archetype_scores = results_df[archetype_names].mean().round(2).to_dict()
    num_ankiet = len(results_df)
    main_type, second_type = pick_main_and_aux_archetype(mean_archetype_scores, ARCHE_NAMES_ORDER)
    main = archetype_extended.get(main_type, {})
    second = archetype_extended.get(second_type, {}) if second_type else {}
    archetype_table = results_df[archetype_names].apply(lambda row: pd.Series(row), axis=1)
    archetype_table["Główny archetyp"] = archetype_table[archetype_names].idxmax(axis=1)
    table_counts = archetype_table["Główny archetyp"].value_counts().reset_index()
    table_counts.columns = ["Archetyp", "Liczba osób"]
    table_counts = table_counts.merge(pd.DataFrame({"Archetyp": ARCHE_NAMES_ORDER}), how="right").fillna(0).astype({"Liczba osób": int})
    table_img = pil_to_bytesio(archetype_table_img(table_counts[["Archetyp","Liczba osób"]]))
    fig_radar = go.Figure(
        data=[
            go.Scatterpolar(
                r=list(mean_archetype_scores.values()) + [list(mean_archetype_scores.values())[0]],
                theta=archetype_names + [archetype_names[0]],
                fill='toself',
                name='Średnia wszystkich',
                line=dict(color="royalblue", width=3),
                marker=dict(size=6)
            ),
        ],
        layout=go.Layout(
            polar=dict(
                radialaxis=dict(visible=True, range=[0, 20]),
                angularaxis=dict(tickfont=dict(size=16), tickvals=archetype_names)
            ),
            width=480, height=480,
            margin=dict(l=10, r=10, t=20, b=20),
            showlegend=False
        )
    )
    radar_img = pil_to_bytesio(radar_pil_img(fig_radar))
    kola_img = pil_to_bytesio(compose_archetype_highlight(
        archetype_name_to_img_idx(main_type),
        archetype_name_to_img_idx(second_type) if second_type != main_type else None
    ))
    docx_buf = export_word(main_type, second_type, archetype_features, main, second,
                           mean_archetype_scores, kola_img, radar_img, table_img, num_ankiet)
    pdf_buf = export_pdf(main_type, second_type, archetype_features, main, second,
                        mean_archetype_scores, kola_img, radar_img, table_img, num_ankiet)
    st.download_button("Pobierz raport DOCX", docx_buf, file_name="Raport_archetypowy_Hetman.docx")
    st.download_button("Pobierz raport PDF", pdf_buf, file_name="Raport_archetypowy_Hetman.pdf")
    st.header("Podsumowanie głównych archetypów:")
    st.plotly_chart(fig_radar, use_container_width=False)
    st.image(kola_img, caption="Główny i pomocniczy archetyp na kole archetypów", width=400)
    st.image(table_img, caption="Liczebność poszczególnych archetypów")
    st.write("Profil Hetmana:")
    st.write(f"Główny archetyp: **{main_type}**")
    if second_type and second_type != main_type:
        st.write(f"Archetyp pomocniczy: **{second_type}**")
    st.write(f"Liczba ankietowanych: {num_ankiet}")
