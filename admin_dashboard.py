# admin_dashboard.py - raporty archetypów

import pandas as pd
import streamlit as st
import psycopg2
import sqlalchemy as sa

def _sa_engine():
    """
    Silnik SQLAlchemy do Postgresa z SSL. Używamy go w pd.read_sql.
    Wymaga pakietu SQLAlchemy (pip install SQLAlchemy).
    """
    url = sa.URL.create(
        "postgresql+psycopg2",
        username=st.secrets["db_user"],
        password=st.secrets["db_pass"],
        host=st.secrets["db_host"],
        port=int(st.secrets.get("db_port", 5432)),
        database=st.secrets["db_name"],
        query={"sslmode": "require"},
    )
    return sa.create_engine(url, pool_pre_ping=True)

import ast
import plotly.graph_objects as go
from fpdf import FPDF
import unicodedata
import colorsys
import requests
from PIL import Image, ImageDraw, ImageFont
import io
import re
import html
import base64
import math
from textwrap import dedent
from datetime import datetime
import pytz
from functools import lru_cache
from urllib.parse import quote
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import os
from pathlib import Path
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from io import BytesIO
import tempfile
import warnings
warnings.filterwarnings(
    "ignore",
    message="pkg_resources is deprecated as an API",
    category=UserWarning,
    module="docxcompose.properties",
)

import streamlit.components.v1 as components
import matplotlib.pyplot as plt
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
import numpy as np
from archetype_docx_loader import load_archetype_extended
from archetype_interpretation import generate_archetype_descriptions
from metryczka_config import normalize_personal_metryczka_config
from public_labels import (
    ARCHETYPE_PUBLIC_VALUES,
    FINAL_VALUES_WHEEL_ARC_LABELS,
    FINAL_VALUES_WHEEL_CENTRAL_FIELDS,
)

import sys
if sys.platform.startswith("linux"):
    import subprocess
else:
    from docx2pdf import convert

_BASE_DIR = Path(__file__).resolve().parent

TEMPLATE_PATH = str(_BASE_DIR / "ap48_raport_template.docx")
TEMPLATE_PATH_NOSUPP = str(_BASE_DIR / "ap48_raport_template_nosupp.docx")  # szablon bez sekcji archetypu pobocznego
TEMPLATE_PATH_SHORT = str(_BASE_DIR / "ap48_raport_template_short.docx")
TEMPLATE_PATH_SHORT_NOSUPP = str(_BASE_DIR / "ap48_raport_template_short_nosupp.docx")

TEMPLATE_PATH_FEMALE = str(_BASE_DIR / "ap48_raport_template_female.docx")
TEMPLATE_PATH_NOSUPP_FEMALE = str(_BASE_DIR / "ap48_raport_template_nosupp_female.docx")
TEMPLATE_PATH_SHORT_FEMALE = str(_BASE_DIR / "ap48_raport_template_short_female.docx")
TEMPLATE_PATH_SHORT_NOSUPP_FEMALE = str(_BASE_DIR / "ap48_raport_template_short_nosupp_female.docx")
logos_dir = str(Path(__file__).with_name("logos_local"))

import plotly.io as pio
import shutil, os

# ====== Kaleido + Chrome/Chromium hardening (serwery/headless) ======
# Szukamy przeglądarki w popularnych lokalizacjach i nazwach (snap, deb, wrappery)
_CANDIDATES = [
    "/usr/bin/chromium-browser",
    "/usr/bin/chromium",
    "/snap/bin/chromium",            # snap (czasem PATH nie ma /snap/bin)
    "/usr/bin/google-chrome",
    "/usr/bin/google-chrome-stable",
    "chromium-browser", "chromium", "google-chrome", "chrome"
]

_chrome = None
for c in _CANDIDATES:
    if os.path.isabs(c):
        if os.path.exists(c):
            _chrome = c; break
    else:
        p = shutil.which(c)
        if p:
            _chrome = p; break

# Ustaw ścieżkę Chromium dla Kaleido (jeśli API jest dostępne w tej wersji plotly)
try:
    if _chrome and hasattr(pio, "kaleido") and hasattr(pio.kaleido, "scope"):
        pio.kaleido.scope.chromium_executable = _chrome
        pio.kaleido.scope.chromium_args = ["--no-sandbox", "--disable-gpu", "--disable-dev-shm-usage"]
except Exception:
    pass

import os, streamlit as st
def _build_sha():
    try: return open(os.path.join(os.path.dirname(__file__), ".deployed_sha")).read().strip()[:7]
    except Exception: return "dev"
st.sidebar.caption(f"Build: {_build_sha()}")


def _is_probably_mobile_client() -> bool:
    """Lekka heurystyka mobile po User-Agent (bez wpływu na desktop)."""
    try:
        ctx = getattr(st, "context", None)
        headers = getattr(ctx, "headers", {}) if ctx else {}
        ua = str(headers.get("user-agent") or headers.get("User-Agent") or "").lower()
        if not ua:
            return False
        mobile_tokens = (
            "mobile",
            "android",
            "iphone",
            "ipod",
            "ipad",
            "miui",
        )
        return any(tok in ua for tok in mobile_tokens)
    except Exception:
        return False


def _logo_norm_key(value: str) -> str:
    if value is None:
        return ""
    normalized = unicodedata.normalize("NFKD", str(value)).encode("ascii", "ignore").decode("ascii")
    normalized = normalized.lower().replace("&", "and")
    normalized = re.sub(r"[^a-z0-9]+", "-", normalized).strip("-")
    return normalized


@lru_cache(maxsize=16)
def _logo_lookup_index(logos_dir: str) -> dict[str, str]:
    index: dict[str, str] = {}
    if not logos_dir or not os.path.isdir(logos_dir):
        return index

    for filename in os.listdir(logos_dir):
        if os.path.splitext(filename)[1].lower() not in {".svg", ".png", ".jpg", ".jpeg", ".webp"}:
            continue
        stem = os.path.splitext(filename)[0]
        abs_path = os.path.join(logos_dir, filename)
        keys = {
            stem.casefold(),
            stem.replace("-", "").casefold(),
            _logo_norm_key(stem),
            re.sub(r"[^a-z0-9]+", "", _logo_norm_key(stem)),
        }
        for key in keys:
            if key:
                index.setdefault(key, abs_path)
    return index


LOGO_BRAND_ALIASES = {
    "Victoria’s Secret": "Victorias Secrets",
    "Victoria's Secret": "Victorias Secrets",
    "PKN ORLEN": "Orlen",
    "PZU SA": "PZU",
}


def get_logo_svg_path(brand_name, logos_dir=None):
    if logos_dir is None:
        logos_dir = str(Path(__file__).with_name("logos_local"))

    if not brand_name:
        return None

    raw = str(brand_name).strip()
    if not raw:
        return None

    raw_variants = [raw]
    alias = LOGO_BRAND_ALIASES.get(raw)
    if alias:
        raw_variants.append(alias)

    # Próby "po staremu" (kompatybilność)
    candidate_filenames = []
    for variant in raw_variants:
        base_kebab = (
            variant.lower()
            .replace(" ", "-")
            .replace("'", "")
            .replace("’", "")
            .replace("ł", "l")
            .replace("ś", "s")
            .replace("ż", "z")
            .replace("ó", "o")
            .replace("ć", "c")
            .replace("ń", "n")
            .replace("ę", "e")
            .replace("ą", "a")
            .replace("ś", "s")
        )
        base_flat = variant.lower().replace(" ", "").replace("'", "").replace("’", "")
        for ext in (".svg", ".png", ".jpg", ".jpeg", ".webp"):
            candidate_filenames.append(base_kebab + ext)
            candidate_filenames.append(base_flat + ext)

    for filename in candidate_filenames:
        path = os.path.join(logos_dir, filename)
        if os.path.exists(path):
            return path

    # Fallback: dopasowanie odporne na case/diakrytyki/znaki specjalne.
    lookup = _logo_lookup_index(logos_dir)
    probe_keys = []
    for variant in raw_variants:
        probe_keys.extend(
            [
                variant.casefold(),
                variant.replace(" ", "").casefold(),
                _logo_norm_key(variant),
                re.sub(r"[^a-z0-9]+", "", _logo_norm_key(variant)),
            ]
        )
    for key in probe_keys:
        if key in lookup:
            return lookup[key]

    return None

from io import BytesIO
from docxtpl import InlineImage
from docx.shared import Mm

import subprocess
from io import BytesIO

person_wikipedia_links = {
    "Aleksandra Dulkiewicz": "https://pl.wikipedia.org/wiki/Aleksandra_Dulkiewicz",
    "Aleksiej Nawalny": "https://pl.wikipedia.org/wiki/Aleksiej_Nawalny",
    "Angela Merkel": "https://pl.wikipedia.org/wiki/Angela_Merkel",
    "Andrzej Duda": "https://pl.wikipedia.org/wiki/Andrzej_Duda",
    "Barack Obama": "https://pl.wikipedia.org/wiki/Barack_Obama",
    "Benito Juárez": "https://pl.wikipedia.org/wiki/Benito_Ju%C3%A1rez",
    "Bernie Sanders": "https://pl.wikipedia.org/wiki/Bernie_Sanders",
    "Boris Johnson": "https://pl.wikipedia.org/wiki/Boris_Johnson",
    "Bronisław Geremek": "https://pl.wikipedia.org/wiki/Bronis%C5%82aw_Geremek",
    "Clement Attlee": "https://pl.wikipedia.org/wiki/Clement_Attlee",
    "Donald Trump": "https://pl.wikipedia.org/wiki/Donald_Trump",
    "Elon Musk": "https://pl.wikipedia.org/wiki/Elon_Musk",
    "Emmanuel Macron": "https://pl.wikipedia.org/wiki/Emmanuel_Macron",
    "Eva Perón": "https://pl.wikipedia.org/wiki/Eva_Per%C3%B3n",
    "François Mitterrand": "https://pl.wikipedia.org/wiki/Fran%C3%A7ois_Mitterrand",
    "Franklin D. Roosevelt": "https://pl.wikipedia.org/wiki/Franklin_D._Roosevelt",
    "George Washington": "https://pl.wikipedia.org/wiki/George_Washington",
    "Jacek Kuroń": "https://pl.wikipedia.org/wiki/Jacek_Kuro%C5%84",
    "Jacinda Ardern": "https://pl.wikipedia.org/wiki/Jacinda_Ardern",
    "Jarosław Kaczyński": "https://pl.wikipedia.org/wiki/Jaros%C5%82aw_Kaczy%C5%84ski",
    "Jawaharlal Nehru": "https://pl.wikipedia.org/wiki/Jawaharlal_Nehru",
    "Janusz Palikot": "https://pl.wikipedia.org/wiki/Janusz_Palikot",
    "Jeremy Corbyn": "https://pl.wikipedia.org/wiki/Jeremy_Corbyn",
    "Jimmy Carter": "https://pl.wikipedia.org/wiki/Jimmy_Carter",
    "Joe Biden": "https://pl.wikipedia.org/wiki/Joe_Biden",
    "John F. Kennedy": "https://pl.wikipedia.org/wiki/John_F._Kennedy",
    "Józef Piłsudski": "https://pl.wikipedia.org/wiki/J%C3%B3zef_Pi%C5%82sudski",
    "Justin Trudeau": "https://pl.wikipedia.org/wiki/Justin_Trudeau",
    "Konrad Adenauer": "https://pl.wikipedia.org/wiki/Konrad_Adenauer",
    "Lee Kuan Yew": "https://pl.wikipedia.org/wiki/Lee_Kuan_Yew",
    "Lech Wałęsa": "https://pl.wikipedia.org/wiki/Lech_Wa%C5%82%C4%99sa",
    "Ludwik XIV": "https://pl.wikipedia.org/wiki/Ludwik_XIV_Burbonski",
    "Margaret Thatcher": "https://pl.wikipedia.org/wiki/Margaret_Thatcher",
    "Marine Le Pen": "https://pl.wikipedia.org/wiki/Marine_Le_Pen",
    "Martin Luther King": "https://pl.wikipedia.org/wiki/Martin_Luther_King",
    "Mustafa Kemal Atatürk": "https://pl.wikipedia.org/wiki/Mustafa_Kemal_Atat%C3%BCrk",
    "Napoleon Bonaparte": "https://pl.wikipedia.org/wiki/Napoleon_Bonaparte",
    "Nelson Mandela": "https://pl.wikipedia.org/wiki/Nelson_Mandela",
    "Olof Palme": "https://pl.wikipedia.org/wiki/Olof_Palme",
    "Pedro Sánchez": "https://pl.wikipedia.org/wiki/Pedro_S%C3%A1nchez",
    "Sanna Marin": "https://pl.wikipedia.org/wiki/Sanna_Marin",
    "Shimon Peres": "https://pl.wikipedia.org/wiki/Shimon_Peres",
    "Silvio Berlusconi": "https://pl.wikipedia.org/wiki/Silvio_Berlusconi",
    "Sławomir Mentzen": "https://pl.wikipedia.org/wiki/S%C5%82awomir_Mentzen",
    "Szymon Hołownia": "https://pl.wikipedia.org/wiki/Szymon_Ho%C5%82ownia",
    "Theodore Roosevelt": "https://pl.wikipedia.org/wiki/Theodore_Roosevelt",
    "Thomas Jefferson": "https://pl.wikipedia.org/wiki/Thomas_Jefferson",
    "Tony Blair": "https://pl.wikipedia.org/wiki/Tony_Blair",
    "Václav Havel": "https://pl.wikipedia.org/wiki/V%C3%A1clav_Havel",
    "Václav Klaus": "https://pl.wikipedia.org/wiki/V%C3%A1clav_Klaus",
    "Vladimir Putin": "https://pl.wikipedia.org/wiki/W%C5%82adimir_Putin",
    "Winston Churchill": "https://pl.wikipedia.org/wiki/Winston_Churchill",
    "Wołodymyr Zełenski": "https://pl.wikipedia.org/wiki/Wo%C5%82odymyr_Ze%C5%82enski",
    "Władysław Kosiniak-Kamysz": "https://pl.wikipedia.org/wiki/W%C5%82adys%C5%82aw_Kosiniak-Kamysz",
    "Xi Jinping": "https://pl.wikipedia.org/wiki/Xi_Jinping",
    "Deng Xiaoping": "https://en.wikipedia.org/wiki/Deng_Xiaoping",
}


def _person_norm_key(name: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(name or "")).encode("ascii", "ignore").decode("ascii")
    normalized = re.sub(r"[^a-z0-9]+", " ", normalized.lower())
    return re.sub(r"\s+", " ", normalized).strip()


PERSON_CANONICAL_BY_NORM = {_person_norm_key(name): name for name in person_wikipedia_links}
PERSON_LINK_BY_NORM = {_person_norm_key(name): url for name, url in person_wikipedia_links.items()}


def canonical_person_name(name: str) -> str:
    clean = str(name or "").strip().strip(".")
    if clean in person_wikipedia_links:
        return clean
    return PERSON_CANONICAL_BY_NORM.get(_person_norm_key(clean), clean)


from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), "0000FF")
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def person_link(name):
    canonical = canonical_person_name(name)
    url = person_wikipedia_links.get(canonical) or PERSON_LINK_BY_NORM.get(_person_norm_key(name))
    if url:
        return f"<a href='{url}' target='_blank'>{canonical}</a>"
    return canonical

def svg_to_png_bytes(svg_path, width_mm=None, height_mm=None):
    import cairosvg

    with open(svg_path, "rb") as svg_file:
        svg_bytes = svg_file.read()

    # Przelicz mm na px (96 dpi ≈ 3.78 px/mm)
    arg_dict = {}
    if width_mm is not None:
        arg_dict['output_width'] = int(width_mm * 3.78 * 4)  # ×4 dla ostrości
    if height_mm is not None:
        arg_dict['output_height'] = int(height_mm * 3.78 * 4)

    png_bytes = cairosvg.svg2png(bytestring=svg_bytes, **arg_dict)
    return png_bytes

# --- Ikony archetypów do Word (InlineImage) ---
def _icon_file_for(archetype_name: str, gender_code: str = "M"):
    """
    Zwraca ścieżkę do pliku ikony w assets/person_icons:
      <slug>_<M/K>.svg|png lub fallback <slug>.svg|png
    """
    import glob
    base_masc = base_masc_from_any(archetype_name)
    slug = ARCHETYPE_BASE_SLUGS.get(base_masc, _slug_pl(base_masc))
    g = (gender_code or "M").upper()
    candidates = [
        ARCHETYPE_ICON_DIR / f"{slug}_{g}.svg",
        ARCHETYPE_ICON_DIR / f"{slug}_{g}.png",
        ARCHETYPE_ICON_DIR / f"{slug}.svg",
        ARCHETYPE_ICON_DIR / f"{slug}.png",
    ]
    for c in candidates:
        if Path(c).exists():
            return Path(c)
        m = glob.glob(str(c))
        if m:
            return Path(m[0])
    return None

def _card_file_for(archetype_name: str):
    """
    Zwraca ścieżkę do karty archetypu w assets/card.
    Dopasowanie jest odporne na wielkość liter, polskie znaki i warianty nazw.
    """
    try:
        base_masc = base_masc_from_any(str(archetype_name or "").strip())
        if not base_masc:
            return None

        slug = ARCHETYPE_BASE_SLUGS.get(base_masc, _slug_pl(base_masc))

        def _norm(v: str) -> str:
            return re.sub(r"[^a-z0-9]+", "", _slug_pl(v or ""))

        target_keys = {_norm(base_masc), _norm(slug)}
        target_keys = {k for k in target_keys if k}
        if not target_keys:
            return None

        if ARCHETYPE_CARD_DIR.exists():
            files = sorted(
                [
                    p
                    for p in ARCHETYPE_CARD_DIR.iterdir()
                    if p.is_file() and p.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}
                ],
                key=lambda p: p.name.lower(),
            )

            # 1) Najpierw tylko dokładne dopasowania nazwy (najbezpieczniejsze)
            for path in files:
                stem_norm = _norm(path.stem)
                if stem_norm and stem_norm in target_keys:
                    return path

            # 2) Dopiero potem fallback prefiksowy (deterministyczny)
            fallback_hits: list[tuple[int, str, Path]] = []
            for path in files:
                stem_norm = _norm(path.stem)
                if not stem_norm:
                    continue
                matched = [k for k in target_keys if stem_norm.startswith(k) or k.startswith(stem_norm)]
                if not matched:
                    continue
                best_delta = min(abs(len(stem_norm) - len(k)) for k in matched)
                fallback_hits.append((best_delta, path.name.lower(), path))
            if fallback_hits:
                fallback_hits.sort(key=lambda x: (x[0], x[1]))
                return fallback_hits[0][2]

        # Fallback: jeśli brak dedykowanej karty w assets/card, pokaż ikonę persony.
        # Działa też gdy katalog assets/card nie istnieje na danym środowisku.
        raw_name = str(archetype_name or "").strip()
        raw_norm = _norm(raw_name)
        fem_norms = {_norm(f): m for f, m in GENDER_MASC_FROM_FEM.items()}
        guessed_gender = "K" if raw_norm and raw_norm in fem_norms else "M"
        icon_fallback = _icon_file_for(base_masc, gender_code=guessed_gender) or _icon_file_for(base_masc, gender_code="M")
        if icon_fallback and Path(icon_fallback).exists():
            return Path(icon_fallback)
    except Exception:
        return None
    return None


def _profile_card_file_for(archetype_name: str, gender_code: str = "M"):
    """
    Zwraca ścieżkę do profilowej grafiki archetypu.
    Dla kobiet najpierw sprawdza assets/archetype_profile_cards_female,
    a potem fallback do assets/archetype_profile_cards_male.
    Dopasowanie odporne na brak polskich znaków i warianty nazw.
    """
    try:
        base_masc = base_masc_from_any(str(archetype_name or "").strip())
        if not base_masc:
            return None

        gender_code = str(gender_code or "M").upper()
        dirs = []
        if gender_code == "K":
            dirs.append(ARCHETYPE_PROFILE_CARD_DIR_FEMALE)
        dirs.append(ARCHETYPE_PROFILE_CARD_DIR_MALE)

        def _norm(v: str) -> str:
            return re.sub(r"[^a-z0-9]+", "", _slug_pl(v or ""))

        target_keys = {_norm(base_masc), _norm(ARCHETYPE_BASE_SLUGS.get(base_masc, base_masc))}
        target_keys = {k for k in target_keys if k}
        if not target_keys:
            return None

        for cards_dir in dirs:
            if not cards_dir.exists():
                continue

            direct = cards_dir / f"{base_masc}.png"
            if direct.exists():
                return direct

            files = sorted(
                [p for p in cards_dir.iterdir() if p.is_file() and p.suffix.lower() == ".png"],
                key=lambda p: p.name.lower(),
            )
            for path in files:
                stem_norm = _norm(path.stem)
                if stem_norm and stem_norm in target_keys:
                    return path
    except Exception:
        return None
    return None


def _profile_card_dark_variant(light_path: Path | None) -> Path | None:
    """
    Dla <nazwa>.png zwraca <nazwa>_dark.png jeśli istnieje.
    """
    try:
        if not light_path:
            return None
        candidate = light_path.with_name(f"{light_path.stem}_dark{light_path.suffix}")
        return candidate if candidate.exists() else None
    except Exception:
        return None


def _img_data_uri_from_path(path: Path) -> str:
    suffix = path.suffix.lower()
    mime = "image/png"
    if suffix in {".jpg", ".jpeg"}:
        mime = "image/jpeg"
    elif suffix == ".webp":
        mime = "image/webp"
    b64 = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{b64}"


def _img_data_uri_from_pil(img: Image.Image) -> str:
    buf = BytesIO()
    img.save(buf, format="PNG")
    b64 = base64.b64encode(buf.getvalue()).decode("ascii")
    return f"data:image/png;base64,{b64}"


def _theme_image_dual_html(
    light_uri: str,
    dark_uri: str,
    *,
    style: str = "width:100%;height:auto;display:block;",
    extra_class: str | None = None,
) -> str:
    extra = f" {extra_class.strip()}" if extra_class else ""
    light_style = style
    dark_style = f"{style};display:none;"
    return (
        "<span class='ap-theme-image-wrap'>"
        f"<img src='{light_uri}' class='ap-theme-image ap-theme-image-light{extra}' style='{light_style}'/>"
        f"<img src='{dark_uri}' class='ap-theme-image ap-theme-image-dark{extra}' style='{dark_style}'/>"
        "</span>"
    )

def arche_icon_inline_for_word(doc, archetype_name: str, gender_code: str = "M", height_mm: float = 18):
    """
    Zwraca InlineImage do docxtpl dla danego archetypu.
    Obsługuje SVG (konwersja do PNG) i PNG.
    """
    path = _icon_file_for(archetype_name, gender_code)
    if not path:
        return ""
    if path.suffix.lower() == ".svg":
        png_bytes = svg_to_png_bytes(str(path), height_mm=height_mm)
        return InlineImage(doc, BytesIO(png_bytes), height=Mm(height_mm))
    else:
        return InlineImage(doc, str(path), height=Mm(height_mm))

def arche_card_inline_for_word(doc, archetype_name: str, width_mm: float = 70):
    """
    Zwraca InlineImage karty archetypu (assets/card) do DOCX/PDF.
    """
    path = _card_file_for(archetype_name)
    if not path:
        return ""
    return InlineImage(doc, str(path), width=Mm(width_mm))

def _load_arche_icon_for_mpl(name: str, size_px: int = 160, tint: str = "#3B82F6", as_array: bool = True):
    """
    Wczytuje ikonę archetypu (PNG/SVG), zachowuje proporcje (contain) bez rozciągania,
    osadza na kwadratowym, przezroczystym płótnie i barwi na `tint`.
    Zwraca numpy array (domyślnie) albo PIL.Image gdy as_array=False.
    """
    from pathlib import Path
    from io import BytesIO
    from PIL import Image, ImageOps, ImageColor

    try:
        import numpy as np  # potrzebne, gdy chcemy array
    except Exception:
        np = None

    # --- Twoje utilsy/ścieżki (zostaw jak masz) ---
    base = base_masc_from_any(name)
    slug = ARCHETYPE_BASE_SLUGS.get(base, _slug_pl(base))
    svg = ARCHE_STACKED_ICON_DIR / f"{slug}.svg"
    png = ARCHE_STACKED_ICON_DIR / f"{slug}.png"
    # ------------------------------------------------

    # Pillow 10+
    RESAMPLE = getattr(Image, "Resampling", Image).LANCZOS

    canvas = Image.new("RGBA", (size_px, size_px), (255, 255, 255, 0))

    try:
        if png.exists():
            im = Image.open(png).convert("RGBA")
        elif svg.exists():
            import cairosvg
            buf = cairosvg.svg2png(url=str(svg), output_width=size_px * 4, output_height=size_px * 4)
            im = Image.open(BytesIO(buf)).convert("RGBA")
        else:
            return None

        # Bez deformacji: "contain" do kwadratu
        im = ImageOps.contain(im, (size_px, size_px), RESAMPLE)

        # Tint: bierzemy tylko alpha z ikony i wlewamy jednolity kolor pod maskę
        color_rgb = ImageColor.getrgb(tint)
        alpha = im.getchannel("A")
        colored = Image.new("RGBA", im.size, color_rgb + (0,))
        colored.putalpha(alpha)

        # Wklejamy na środek kwadratowego, przezroczystego płótna
        x = (size_px - colored.width) // 2
        y = (size_px - colored.height) // 2
        canvas.alpha_composite(colored, (x, y))

        if as_array:
            if np is None:
                raise RuntimeError("numpy is required when as_array=True")
            return np.array(canvas)
        return canvas

    except Exception:
        return None


def make_stacked_bar_png_for_word(
    archetype_names: list[str],
    counts_main: dict[str, int],
    counts_aux: dict[str, int],
    counts_supp: dict[str, int],
    out_path: str = "archetypes_stacked.png",
):
    import os
    import numpy as np
    import matplotlib as mpl
    import matplotlib.pyplot as plt
    from matplotlib import ticker
    from matplotlib.offsetbox import OffsetImage, AnnotationBbox
    from matplotlib import font_manager as fm
    from PIL import Image, ImageOps  # do „kwadratowania” ikon

    # ===== POKRĘTŁA =====
    FIG_SCALE  = 0.90
    FONT_SCALE = 1.25
    DPI = 300

    # Roboto Condensed
    FONT_DIR = "assets/fonts"
    for fname in ("RobotoCondensed-Regular.ttf", "RobotoCondensed-Bold.ttf"):
        p = os.path.join(FONT_DIR, fname)
        if os.path.exists(p):
            fm.fontManager.addfont(p)
    mpl.rcParams["font.family"] = "Roboto Condensed"

    base_fs = 9 * FONT_SCALE
    mpl.rcParams.update({
        "font.size": base_fs,
        "axes.labelsize": 0.85 * base_fs,
        "xtick.labelsize": 0.70 * base_fs,
        "ytick.labelsize": 0.85 * base_fs,
        "legend.fontsize": 0.80 * base_fs,
        "pdf.fonttype": 42, "ps.fonttype": 42,
    })

    # ===== Dane =====
    names = list(archetype_names)
    m = np.array([int(counts_main.get(n, 0) or 0) for n in names], dtype=int)
    a = np.array([int(counts_aux.get(n, 0) or 0)  for n in names], dtype=int)
    s = np.array([int(counts_supp.get(n, 0) or 0) for n in names], dtype=int)
    t = m + a + s

    order = np.argsort(-t)
    names, m, a, s, t = [[arr[i] for i in order] if isinstance(arr, list) else arr[order]
                         for arr in (names, m, a, s, t)]
    y = np.arange(len(names))

    # ===== Figura =====
    fig_w = 8.0 * FIG_SCALE
    HEIGHT_SCALE = 1.03  # wysokość wykresu
    fig_h = HEIGHT_SCALE * ((0.55 * FIG_SCALE) * len(names) + 0.9)
    fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=DPI)
    fig.patch.set_alpha(0)
    ax.set_facecolor((1,1,1,0))
    fig.subplots_adjust(left=0.30, right=0.975, top=0.83, bottom=0.05)

    # ===== Osi / kolory =====
    c_main = COLOR_HEX["Czerwony"]
    c_aux  = COLOR_HEX["Żółty"]
    c_supp = COLOR_HEX["Zielony"]
    axis_col = "#4B5563"
    grid_col = "#E5E7EB"

    ax.xaxis.set_ticks_position("top")
    ax.xaxis.set_label_position("top")
    ax.set_xlabel("Liczba wskazań", fontsize=8.5, color="#8A94A6", labelpad=12)   # większy odstęp
    ax.xaxis.set_label_coords(0.5, 1.06)                             # wyżej nad osią
    ax.xaxis.set_major_locator(ticker.MaxNLocator(integer=True))
    ax.grid(axis="x", color=grid_col, linewidth=0.7)
    ax.set_axisbelow(True)

    ax.tick_params(axis="x", colors=axis_col, labelsize=0.70 * base_fs, length=3, pad=2)
    ax.tick_params(axis="y", colors="#111827", pad=2)

    ax.spines["top"].set_color(axis_col)
    ax.spines["left"].set_color(axis_col)
    ax.spines["top"].set_linewidth(1.0)
    ax.spines["left"].set_linewidth(1.0)
    ax.spines["right"].set_visible(False)
    ax.spines["bottom"].set_visible(False)

    # ===== Słupki =====
    bar_h = 0.70
    ax.barh(y, m, height=bar_h, color=c_main, edgecolor="white", linewidth=0.6, label="Główny")
    ax.barh(y, a, left=m, height=bar_h, color=c_aux,  edgecolor="white", linewidth=0.6, label="Wspierający")
    ax.barh(y, s, left=m+a, height=bar_h, color=c_supp, edgecolor="white", linewidth=0.6, label="Poboczny")

    ax.set_yticks(y)
    ax.set_yticklabels(names, ha="right")
    for lbl in ax.get_yticklabels():
        lbl.set_x(-0.012)  # ciut bliżej osi

    # ===== Ikony – mniejsze, dalej w lewo, zawsze kwadrat =====
    ICON_PX   = int(84 * FIG_SCALE)   # rozdzielczość pola
    ICON_ZOOM = 0.30 * FIG_SCALE      # faktyczna wielkość na wykresie
    ICON_X    = -0.17                 # bardziej ujemne = dalej w lewo

    for i, nme in enumerate(names):
        icon_arr = _load_arche_icon_for_mpl(nme, size_px=ICON_PX)
        if icon_arr is None:
            continue
        im = Image.fromarray(icon_arr)
        im = ImageOps.pad(im, (ICON_PX, ICON_PX), color=(255, 255, 255, 0), centering=(0.5, 0.5))
        icon_arr = np.array(im)

        img = OffsetImage(icon_arr, zoom=ICON_ZOOM)  # isotropowe skalowanie (brak spłaszczania)
        ab = AnnotationBbox(img, (ICON_X, i),
                            xycoords=("axes fraction", "data"),
                            box_alignment=(1.0, 0.5),
                            frameon=False, annotation_clip=False)
        ax.add_artist(ab)

    # ===== Etykiety wewnątrz =====
    def annotate_segment(left_arr, vals):
        for row, val in enumerate(vals):
            if val <= 0:
                continue
            ax.text(left_arr[row] + val/2.0, row, str(int(val)),
                    color="white", ha="center", va="center",
                    fontsize=0.65 * base_fs)

    annotate_segment(np.zeros_like(m, float), m)
    annotate_segment(m, a)
    annotate_segment(m + a, s)

    # ===== Suma na końcu =====
    gap = max(0.02 * (t.max() if t.max() > 0 else 1), 0.10)
    for i in range(len(y)):
        ax.text(float(t[i]) + gap, i, str(int(t[i])),
                color="#374151", ha="left", va="center",
                fontsize=0.80 * base_fs, fontweight="bold")

    # ===== Legenda =====
    leg = ax.legend(
        loc="upper center",
        bbox_to_anchor=(0, -0.03, 1, 0),  # (x, y, width, height) w ułamku osi
        #mode="expand",  # rozkłada etykiety równomiernie na całą szerokość
        ncol=3,
        frameon=False,
        handlelength=1.5,  # dłuższe próbki
        handletextpad=0.6,  # odstęp próbka–tekst
        columnspacing=3.0,  # większy rozstrzał między kolumnami
        borderaxespad=0.0
    )

    fig.subplots_adjust(bottom=0.085)  # ciaśniej na dole

    ax.invert_yaxis()
    fig.savefig(out_path, dpi=DPI, bbox_inches="tight", pad_inches=0.02, transparent=True)
    plt.close(fig)
    return out_path


# --- Generowanie grafiki z paletą kolorów do Word ---
def _luma(hexcode: str) -> float:
    h = hexcode.lstrip('#')
    if len(h) == 3:
        h = ''.join(c*2 for c in h)
    r, g, b = (int(h[i:i+2], 16) for i in (0, 2, 4))
    return 0.2126*r + 0.7152*g + 0.0722*b


def make_palette_png(
    palette: list[str],
    target_width_mm: float = 160,   # szerokość obrazka pod Worda
    cols: int = 5,                  # ile kafelków w rzędzie
    tile_w_mm: float = 36,          # SZEROKOŚĆ kafelka (mm) — ręcznie
    tile_h_mm: float = 28,          # WYSOKOŚĆ kafelka (mm) — ręcznie
    pad_mm: float = 6,
    gap_mm: float = 6,
    corner_mm: float = 3,           # zaokrąglenia rogów
    dpi: int = 300,
    name_font_px: int | None = None,  # ROZMIAR czcionki nazwy (px) — ręcznie
    hex_font_px:  int | None = None,  # ROZMIAR czcionki (#HEX) (px) — ręcznie
):
    if not palette:
        return None

    from PIL import Image, ImageDraw, ImageFont
    import math

    def _mm_to_px(mm): return int(round(mm * dpi / 25.4))

    W      = _mm_to_px(target_width_mm)
    pad    = _mm_to_px(pad_mm)
    gap    = _mm_to_px(gap_mm)
    corner = _mm_to_px(corner_mm)

    cell_w = _mm_to_px(tile_w_mm)
    cell_h = _mm_to_px(tile_h_mm)

    # policz rzędy na podstawie liczby elementów i kolumn
    n    = len(palette)
    rows = math.ceil(n / cols)

    # gdy szerokość 5×kafelek + przerwy przekracza obraz — zmniejsz cell_w proporcjonalnie
    total_w = 2*pad + cols*cell_w + (cols-1)*gap
    if total_w > W:
        scale = (W - 2*pad - (cols-1)*gap) / (cols*cell_w)
        cell_w = int(cell_w * scale)
        cell_h = int(cell_h * scale)

    H = 2*pad + rows*cell_h + (rows-1)*gap

    img = Image.new("RGBA", (W, H), (255, 255, 255, 0))
    drw = ImageDraw.Draw(img)

    # FONTY — ręcznie, albo proporcjonalnie do wysokości kafelka
    def _font(paths, size):
        for p in paths:
            try: return ImageFont.truetype(p, size)
            except: pass
        return ImageFont.load_default()

    # --- STAŁE rozmiary czcionek (bez auto-fit) ---
    NAME_FONT_PX_DEFAULT = 26  # ← rozmiar NAZWY
    HEX_FONT_PX_DEFAULT = 18  # ← rozmiar #HEX

    NAME_FONT_PX = name_font_px if name_font_px is not None else NAME_FONT_PX_DEFAULT
    HEX_FONT_PX = hex_font_px if hex_font_px is not None else HEX_FONT_PX_DEFAULT

    # Jeśli chcesz nazwę pogrubioną, zostaw -Bold; jeśli zwykłą, zamień na Regular.
    name_font = _font(
        [
            "assets/fonts/RobotoCondensed-Bold.ttf",
            "assets/fonts/ArialNovaCond-Bold.ttf",
            "DejaVuSans-Bold.ttf",
        ],
        NAME_FONT_PX,
    )

    hex_font = _font(
        [
            "assets/fonts/RobotoCondensed-Regular.ttf",
            "assets/fonts/ArialNovaCond.ttf",
            "DejaVuSans.ttf",
        ],
        HEX_FONT_PX,
    )

    def _luma(hexcode: str) -> float:
        h = hexcode.lstrip('#')
        if len(h) == 3: h = ''.join(c*2 for c in h)
        r, g, b = (int(h[i:i+2], 16) for i in (0, 2, 4))
        return 0.2126*r + 0.7152*g + 0.0722*b

    def ellipsize_two_lines(text: str, max_w: int) -> list[str]:
        words = text.split()
        if not words: return [text]
        lines, cur = [], ""
        for w in words:
            test = (cur + " " + w).strip()
            if drw.textlength(test, font=name_font) <= max_w or not cur:
                cur = test
            else:
                lines.append(cur); cur = w
            if len(lines) == 1 and drw.textlength(cur, font=name_font) > max_w:
                while drw.textlength(cur + "…", font=name_font) > max_w and len(cur) > 1:
                    cur = cur[:-1]
                cur += "…"
        lines.append(cur)
        if len(lines) > 2:
            first, second = lines[0], " ".join(lines[1:])
            while drw.textlength(second + "…", font=name_font) > max_w and len(second) > 1:
                second = second[:-1]
            lines = [first, second + "…"]
        return lines[:2]

    for i, hexcode in enumerate(palette):
        r, c = divmod(i, cols)
        x0 = pad + c * (cell_w + gap)
        y0 = pad + r * (cell_h + gap)
        x1 = x0 + cell_w
        y1 = y0 + cell_h

        drw.rounded_rectangle([x0, y0, x1, y1], radius=corner, fill=hexcode)

        txt_col = (255, 255, 255) if _luma(hexcode) < 110 else (17, 17, 17)

        from_name = COLOR_NAME_MAP.get(str(hexcode).upper(), str(hexcode).upper())
        max_w = cell_w - _mm_to_px(6)
        lines = ellipsize_two_lines(from_name, max_w)

        # oblicz wysokości
        name_h = sum(drw.textbbox((0,0), ln, font=name_font)[3] - drw.textbbox((0,0), ln, font=name_font)[1] for ln in lines)
        gap_y  = _mm_to_px(1.2)
        hex_txt = f"({str(hexcode).upper()})"
        hb  = drw.textbbox((0,0), hex_txt, font=hex_font)
        hex_h = hb[3]-hb[1]

        total_h = name_h + gap_y + hex_h
        base_y  = y0 + (cell_h - total_h)//2
        cx      = x0 + cell_w//2

        # nazwy (wyśrodkowane)
        cur_y = base_y
        for ln in lines:
            nb = drw.textbbox((0,0), ln, font=name_font)
            nw = nb[2]-nb[0]
            drw.text((cx - nw/2, cur_y), ln, font=name_font, fill=txt_col)
            cur_y += nb[3]-nb[1]

        # hex (wyśrodkowany)
        drw.text((cx - (hb[2]-hb[0])/2, cur_y + gap_y), hex_txt, font=hex_font, fill=txt_col)

    out = BytesIO()
    img.save(out, "PNG"); out.seek(0)
    return out


def palette_inline_for_word(
    doc, palette: list[str],
    width_mm: float = 160,
    cols: int = 4,
    tile_w_mm: float = 36,
    tile_h_mm: float = 20,
    name_font_px: int = 38,
    hex_font_px: int = 34,
    gap_mm: float = 5,  # ← odstęp między kaflami
    pad_mm: float = 3,  # ← zewnętrzny margines obrazka
):
    png = make_palette_png(
        palette,
        target_width_mm=width_mm,
        cols=cols,
        tile_w_mm=tile_w_mm,
        tile_h_mm=tile_h_mm,
        name_font_px=name_font_px,
        hex_font_px=hex_font_px,
        pad_mm=pad_mm,
        gap_mm=gap_mm,
    )
    return InlineImage(doc, png, width=Mm(width_mm)) if png else ""

# === Logo → kafelek PNG o stałym rozmiarze (pod Word) ===
def _mm_to_px_logo(mm: float, dpi: int = 300) -> int:
    return int(round(mm * dpi / 25.4))

def _svg_viewbox_ratio(svg_path: str) -> float | None:
    """Spróbuj odczytać proporcje (szer/wys) ze viewBox lub width/height w pliku SVG."""
    try:
        import re
        with open(svg_path, "r", encoding="utf-8", errors="ignore") as f:
            s = f.read()
        m = re.search(r'viewBox\s*=\s*"([\d.\s-]+)"', s)
        if m:
            nums = [float(x) for x in m.group(1).split()]
            if len(nums) == 4 and nums[3] > 0:
                w, h = nums[2], nums[3]
                if h > 0:
                    return w / h
        mw = re.search(r'width\s*=\s*"([\d.]+)', s)
        mh = re.search(r'height\s*=\s*"([\d.]+)', s)
        if mw and mh:
            w, h = float(mw.group(1)), float(mh.group(1))
            if h > 0:
                return w / h
    except Exception:
        pass
    return None

def make_logo_tile_bytes(
    path: str,
    slot_w_mm: float = 25.0,     # szerokość kafelka (identyczna dla WSZYSTKICH)
    slot_h_mm: float = 21.0,     # wysokość kafelka (identyczna dla WSZYSTKICH)
    pad_mm: float = 10.0,         # wewnętrzny margines
    dpi: int = 300,
    bg_rgba=(255, 255, 255, 0)   # tło przezroczyste (PNG z alfą)
) -> bytes:
    """
    Renderuje logo (SVG/PNG) na płótno o stałym rozmiarze.
    Skalowanie proporcjonalne, centrowanie, padding. Zwraca bytes PNG.
    """
    from PIL import Image, ImageOps
    import os
    from io import BytesIO

    slot_w_px = _mm_to_px_logo(slot_w_mm, dpi)
    slot_h_px = _mm_to_px_logo(slot_h_mm, dpi)
    pad_px    = _mm_to_px_logo(pad_mm, dpi)
    box_w = max(1, slot_w_px - 2 * pad_px)
    box_h = max(1, slot_h_px - 2 * pad_px)

    ext = os.path.splitext(path)[1].lower()
    if ext == ".svg":
        # (opcjonalnie) dopasuj rasteryzację pod aspekt
        ar = _svg_viewbox_ratio(path) or 1.0
        # wyjściowy „cel” w pikselach (zmniejszamy aliasing)
        if ar >= 1.0:
            out_w = box_w
            out_h = max(1, int(round(out_w / ar)))
            if out_h > box_h:
                out_h = box_h
                out_w = max(1, int(round(out_h * ar)))
        else:
            out_h = box_h
            out_w = max(1, int(round(out_h * ar)))
            if out_w > box_w:
                out_w = box_w
                out_h = max(1, int(round(out_w / ar)))

        # rasteryzacja SVG → PNG (bez wymiarów – i tak dociśniemy contain)
        png_bytes = svg_to_png_bytes(path)
        im = Image.open(BytesIO(png_bytes)).convert("RGBA")
        im = ImageOps.contain(im, (out_w, out_h))
    else:
        im = Image.open(path).convert("RGBA")
        im = ImageOps.contain(im, (box_w, box_h))

    tile = Image.new("RGBA", (slot_w_px, slot_h_px), bg_rgba)
    x = (slot_w_px - im.width) // 2
    y = (slot_h_px - im.height) // 2
    tile.alpha_composite(im, (x, y))

    out = BytesIO()
    tile.save(out, "PNG")
    return out.getvalue()

def build_brands_for_word(
    doc,
    brand_list,
    logos_dir,
    slot_w_mm: float = 35.0,   # szerokość kafelka w Wordzie
    slot_h_mm: float = 21.0,   # wysokość kafelka w Wordzie
    pad_mm: float = 10.0,
    dpi: int = 300
):
    """
    Zwraca listę {brand, logo}, gdzie 'logo' to InlineImage kafelka
    o stałym „footprincie” – wszystkie znaki wyglądają spójnie.
    """
    out = []
    import os
    from io import BytesIO
    from docxtpl import InlineImage
    from docx.shared import Mm

    for brand in brand_list:
        logo_path = get_logo_svg_path(brand, logos_dir)
        if logo_path and os.path.exists(logo_path):
            try:
                png_bytes = make_logo_tile_bytes(
                    logo_path,
                    slot_w_mm=slot_w_mm,
                    slot_h_mm=slot_h_mm,
                    pad_mm=pad_mm,
                    dpi=dpi
                )
                img_stream = BytesIO(png_bytes)
                # USTALAMY *SZEROKOŚĆ* kafelka – wszystkie będą identyczne
                img = InlineImage(doc, img_stream, width=Mm(slot_w_mm))
                out.append({"brand": brand, "logo": img})
            except Exception:
                out.append({"brand": brand, "logo": ""})
        else:
            out.append({"brand": brand, "logo": ""})
    return out


import base64

# --- IKONY ARCHETYPÓW (SVG) ---
ARCHETYPE_ICON_DIR = Path(__file__).with_name("assets") / "person_icons"
# Karty archetypów (PNG/JPG) pokazywane m.in. przy "5.1.2. Zestawy praktyczne"
ARCHETYPE_CARD_DIR = Path(__file__).with_name("assets") / "card"
# Karty profilowe archetypów (PNG) używane przy TOP2/TOP3
ARCHETYPE_PROFILE_CARD_DIR_MALE = Path(__file__).with_name("assets") / "archetype_profile_cards_male"
ARCHETYPE_PROFILE_CARD_DIR_FEMALE = Path(__file__).with_name("assets") / "archetype_profile_cards_female"
# Ikony do etykiet wykresu skumulowanego (nie mylić z person_icons)
ARCHE_STACKED_ICON_DIR = Path(__file__).with_name("assets") / "arche_icons"

_PL_MAP = str.maketrans({
    "ą":"a","ć":"c","ę":"e","ł":"l","ń":"n","ó":"o","ś":"s","ź":"z","ż":"z",
    "Ą":"a","Ć":"c","Ę":"e","Ł":"l","Ń":"n","Ó":"o","Ś":"s","Ź":"z","Ż":"z",
    "’":"", "'":""
})

def _slug_pl(s: str) -> str:
    s = (s or "").strip().lower().translate(_PL_MAP)
    # normalizacje nazw używanych w danych
    s = s.replace(" / ", " ").replace("/", " ").replace(",", " ")
    s = " ".join(s.split())
    return s.replace(" ", "-")

def build_report_filenames(study: dict) -> tuple[str, str]:
    """
    Zwraca ('raport_<nazwisko-imie>.docx', 'raport_<nazwisko-imie>.pdf').
    Korzysta z *_nom (mianownik), z fallbackiem na first_name / last_name.
    """
    last_nom  = (study.get("last_name_nom")  or study.get("last_name")  or "").strip()
    first_nom = (study.get("first_name_nom") or study.get("first_name") or "").strip()
    base = (f"{last_nom} {first_nom}").strip() or "raport"
    slug = _slug_pl(base)  # już masz zdefiniowane wyżej
    return (f"raport_{slug}.docx", f"raport_{slug}.pdf")


def build_short_report_filenames(study: dict) -> tuple[str, str]:
    docx_name, pdf_name = build_report_filenames(study)
    return (
        docx_name.replace(".docx", "_skrocony.docx"),
        pdf_name.replace(".pdf", "_skrocony.pdf"),
    )


# Gdyby Twoje pliki miały inne nazwy niż slug (opcjonalnie dopisz mapę wyjątków)
ARCHETYPE_FILENAME_MAP = {
    # "kochanelubwielbiciel": "kochanek.svg"  # przykład, jeśli trzeba
}

def arche_icon_img_html(archetype_name: str, height_px: int = 90, gender_code: str = "M") -> str:
    """
    Szuka w assets/person_icons pliku o nazwie:
      <slug>_<gender>.svg|png   (np. blazen_M.png, wladca_K.svg)
    a gdy go nie znajdzie – próbuje też <slug>.svg|png jako fallback.
    """
    import base64, glob, os
    # archetype_name może już być w formie żeńskiej – cofnij do męskiej bazy:
    base_masc = base_masc_from_any(archetype_name)
    slug = ARCHETYPE_BASE_SLUGS.get(base_masc, _slug_pl(base_masc))
    g = (gender_code or "M").upper()

    patterns = [
        ARCHETYPE_ICON_DIR / f"{slug}_{g}.svg",
        ARCHETYPE_ICON_DIR / f"{slug}_{g}.png",
        ARCHETYPE_ICON_DIR / f"{slug}.svg",   # fallback wspólny
        ARCHETYPE_ICON_DIR / f"{slug}.png",
    ]

    path = None
    for pat in patterns:
        matches = glob.glob(str(pat))
        if matches:
            path = Path(matches[0]); break
        if Path(pat).exists():
            path = Path(pat); break

    if path and path.exists():
        data = path.read_bytes()
        mime = "image/svg+xml" if path.suffix.lower() == ".svg" else "image/png"
        b64 = base64.b64encode(data).decode("ascii")
        return (
            f"<img src='data:{mime};base64,{b64}' alt='{archetype_name}' "
            f"style='height:{height_px}px; width:auto; display:block;'/>"
        )

    # Fallback (gdy nic nie ma)
    return f"<div style='font-size:{int(height_px*0.9)}px;line-height:1'>🔹</div>"


import cairosvg


# (page_config usunięty – ustawiany w app.py)

COLOR_NAME_MAP = {
    "#000000": "czerń",
    "#002366": "granat królewski",
    "#0070B5": "błękit corporate",
    "#0192D3": "turkus morski",
    "#0C223F": "granat atramentowy",
    "#0E0D13": "grafit bardzo ciemny",
    "#17BECF": "niebieski morski",
    "#181C3A": "granat nocny",
    "#1B1715": "brąz prawie czarny",
    "#1E90FF": "błękit dodger",
    "#1F77B4": "chabrowy",
    "#212809": "oliwkowy głęboki",
    "#282C34": "granat antracytowy",
    "#2B2D41": "granat ciemny",
    "#2C7D78": "turkus ciemny",
    "#2CA02C": "zieleń trawiasta",
    "#2D4900": "zieleń boru głęboka",
    "#2E3141": "grafitowy granat",
    "#2F4F4F": "grafit morski",
    "#3BE8B0": "mięta neonowa",
    "#40E0D0": "turkus świetlisty",
    "#43C6DB": "turkusowy błękit",
    "#4682B4": "stalowy błękit",
    "#4B0000": "bordo głębokie",
    "#4B0082": "indygo klasyczne",
    "#556B2F": "oliwka zgaszona",
    "#588A4F": "zieleń średnia",
    "#5B6979": "grafit stalowy",
    "#61681C": "oliwkowy",
    "#663399": "purpura rebeccy",
    "#696812": "oliwkowy ciemny",
    "#6C7A89": "popielaty szary",
    "#6CA0DC": "błękit średni",
    "#7AA571": "zieleń jasna",
    "#7C46C5": "fiolet szafirowy",
    "#7C53C3": "fiolet głęboki",
    "#7D0B0B": "czerwony krwisty",
    "#800020": "burgund",
    "#86725D": "taupe mineralny",
    "#8681E8": "fiolet lawendowy",
    "#87CEEB": "błękit nieba",
    "#8B4513": "brąz siodłowy",
    "#8C564B": "ciemny brąz",
    "#8F00FF": "fiolet intensywny",
    "#906C46": "brąz średni",
    "#9467BD": "fiolet śliwkowy",
    "#9BD6F4": "pastelowy błękit jasny",
    "#A0E8AF": "seledyn",
    "#A1B1C2": "szaroniebieski jasny",
    "#A3C1AD": "pastelowa zieleń",
    "#A7C7E7": "pastelowy błękit",
    "#A9A9A9": "szary ciemny",
    "#AAC9CE": "pastelowy niebieskoszary",
    "#AB3941": "czerwony wiśniowy",
    "#ACE7FF": "błękit lodowy",
    "#ADD8E6": "błękit pastelowy",
    "#B0C4DE": "niebieskoszary jasny",
    "#B22222": "czerwony ceglasty ciemny",
    "#B2F2BB": "mięta jasna",
    "#B4D6B4": "miętowa zieleń",
    "#B6019A": "fuksja",
    "#B8B8B8": "szary satynowy",
    "#BBBDA0": "khaki jasne",
    "#C0C0C0": "srebrny",
    "#C2BCC1": "szary perłowy",
    "#CC3E2F": "czerwony ceglasty",
    "#D3D3D3": "szary jasny",
    "#D4AF37": "złoty klasyczny",
    "#D5C6AF": "beż jasny",
    "#D62728": "czerwień karmazynowa",
    "#DAA520": "złoto stare",
    "#E0BBE4": "lawenda pudrowa",
    "#E10209": "czerwony żywy",
    "#E10600": "czerwień flagowa",
    "#E377C2": "róż fioletowy pastelowy",
    "#EEEEEE": "szary mglisty",
    "#F2A93B": "miodowy żółty",
    "#F4F1ED": "biały ciepły",
    "#F5F5DC": "beż kremowy",
    "#F8BBD0": "róż pudrowy",
    "#F9D371": "złocisty żółty",
    "#F9ED06": "żółty intensywny",
    "#F9F9F9": "biel lodowa",
    "#FA709A": "róż malinowy",
    "#FADADD": "róż blady",
    "#FD4431": "czerwony pomarańczowy żywy",
    "#FE89BE": "róż neonowy jasny",
    "#FEE140": "żółty cytrynowy jasny",
    "#FF0000": "czerwień intensywna",
    "#FF6F61": "łososiowy róż",
    "#FF7F0E": "pomarańcz jaskrawy",
    "#FF8300": "pomarańcz mocny",
    "#FF8C00": "pomarańcz ciemny",
    "#FFB300": "żółty bursztynowy",
    "#FFC0CB": "róż klasyczny",
    "#FFD580": "beż morelowy",
    "#FFD6E0": "róż bardzo jasny",
    "#FFD700": "złoto",
    "#FFD93D": "żółty pastelowy",
    "#FFF200": "żółty cytrynowy",
    "#FFF6C3": "krem waniliowy",
    "#FFF9B0": "żółty bananowy blady",
    "#FFFACD": "cytryna kremowa",
    "#FFFFFF": "biel",
}

ARCHETYPE_REQUIRED_PALETTES = {
    "Władca": ["#800020", "#FFD700", "#282C34", "#800020", "#000000", "#8C564B"],
    "Bohater": ["#E10600", "#2E3141", "#FFFFFF", "#D62728", "#0E0D13", "#2B2D41", "#C2BCC1", "#CC3E2F"],
    "Mędrzec": ["#4682B4", "#B0C4DE", "#6C7A89", "#1F77B4", "#86725D", "#F4F1ED", "#BBBDA0", "#2D4900"],
    "Opiekun": ["#0192D3", "#B4D6B4", "#A7C7E7", "#FFD580", "#9467BD", "#5B6979", "#A1B1C2", "#2C7D78"],
    "Kochanek": ["#FA709A", "#FEE140", "#FFD6E0", "#FA709A"],
    "Błazen": ["#AB3941", "#F2A93B", "#FFB300", "#FFD93D", "#588A4F", "#7AA571", "#61681C", "#FF8300"],
    "Twórca": ["#7C53C3", "#3BE8B0", "#87CEEB", "#17BECF", "#B6019A", "#E10209", "#1B1715", "#F9ED06"],
    "Odkrywca": ["#212809", "#A0E8AF", "#F9D371", "#E377C2", "#D5C6AF", "#906C46", "#43C6DB", "#696812"],
    "Czarodziej": ["#181C3A", "#E0BBE4", "#8F00FF", "#7C46C5", "#0070B5", "#8681E8", "#FE89BE", "#FD4431"],
    "Towarzysz": ["#A3C1AD", "#F9F9F9", "#6CA0DC", "#2CA02C"],
    "Niewinny": ["#9BD6F4", "#FFF6C3", "#AAC9CE", "#FFF200"],
    "Buntownik": ["#FF0000", "#FF6F61", "#000000", "#FF7F0E"],
}

ARCHE_NAMES_ORDER = [
    "Niewinny", "Mędrzec", "Odkrywca", "Buntownik", "Czarodziej", "Bohater",
    "Kochanek", "Błazen", "Towarzysz", "Opiekun", "Władca", "Twórca"
]

def archetype_name_to_img_idx(name):
    try:
        return ARCHE_NAMES_ORDER.index(name)
    except ValueError:
        return None

archetypes = {
    "Władca":   [1, 2, 3, 4],
    "Bohater":  [5, 6, 7, 8],
    "Mędrzec":  [9, 10, 11, 12],
    "Opiekun":  [13, 14, 15, 16],
    "Kochanek": [17, 18, 19, 20],
    "Błazen":   [21, 22, 23, 24],
    "Twórca":   [25, 26, 27, 28],
    "Odkrywca": [29, 30, 31, 32],
    "Czarodziej": [33, 34, 35, 36],
    "Towarzysz": [37, 38, 39, 40],
    "Niewinny": [41, 42, 43, 44],
    "Buntownik": [45, 46, 47, 48],
}

# ---- MAPA EMOJI I FUNKCJE (DAJ JE TU, ZAWSZE PRZED CAŁĄ LOGIKĄ) ----
archetype_emoji = {
    "Władca": "👑", "Bohater": "🦸", "Mędrzec": "📖", "Opiekun": "🤝", "Kochanek": "❤️",
    "Błazen": "🤪", "Twórca": "🧩", "Odkrywca": "🗺️", "Czarodziej": "⭐", "Towarzysz": "🏡",
    "Niewinny": "🕊️", "Buntownik": "🔥"
}
def normalize(name):
    if not isinstance(name, str):
        return name
    return name.split("/")[0].split(",")[0].strip().title()

def get_emoji(name):
    """
    Zwraca emoji dla archetypu, nawet jeśli w nazwie pojawiają się ukośniki lub dodatki.
    """
    return archetype_emoji.get(normalize(name), "🔹")

def zero_to_dash(val):
    return "-" if val == 0 else str(val)

archetype_features = {
    "Władca": "Potrzeba porządku, organizacji, odpowiedzialności i utrzymania ładu.",
    "Bohater": "Odwaga, walka z przeciwnościami, mobilizacja do działania.",
    "Mędrzec": "Rozsądek, analityczność, logiczne argumenty i decyzje oparte na faktach.",
    "Opiekun": "Empatia, dbanie o innych, ochrona, troska.",
    "Kochanek": "Relacje, emocje, bliskość, autentyczność uczuć.",
    "Błazen": "Poczucie humoru, dystans, lekkość, rozładowywanie napięć.",
    "Twórca": "Kreatywność, innowacja, wyrażanie siebie, estetyka.",
    "Odkrywca": "Niezależność, zmiany, nowe doświadczenia, ekspresja.",
    "Czarodziej": "Transformacja, inspiracja, zmiana świata, przekuwanie idei w czyn.",
    "Towarzysz": "Autentyczność, wspólnota, prostota, bycie częścią grupy.",
    "Niewinny": "Optymizm, ufność, unikanie konfliktów, pozytywne nastawienie.",
    "Buntownik": "Kwestionowanie norm, odwaga w burzeniu zasad, radykalna zmiana."
}

CORE_TRIPLET_MAP = {
    "Bohater": "Odwaga. Determinacja. Wyzwanie.",
    "Władca": "Porządek. Odpowiedzialność. Ramy.",
    "Mędrzec": "Rozsądek. Wiedza. Analiza.",
    "Opiekun": "Troska. Empatia. Bezpieczeństwo.",
    "Kochanek": "Relacje. Bliskość. Emocje.",
    "Błazen": "Otwartość. Poczucie humoru. Dystans.",
    "Twórca": "Rozwój. Kreatywność. Innowacja.",
    "Odkrywca": "Wolność. Ciekawość. Nowe horyzonty.",
    "Czarodziej": "Wizja. Transformacja. Inspiracja.",
    "Towarzysz": "Współpraca. Wspólnota. Swojskość.",
    "Niewinny": "Przejrzystość. Optymizm. Uczciwość.",
    "Buntownik": "Odnowa. Zmiana. Sprzeciw.",
}

# === Progi i opisy "🧭 Interpretacja natężenia procentowego archetypu" ===
AR_INTENSITY_SCHEME = [
    #  lo,  hi,  etykieta_krótka,                               etykieta_pełna,                         opis_jakościowy
    (0,  29, "marginalne",       "marginalne natężenie",
     "Archetyp prawie nie występuje. Obszar „cienia” — wzorce zachowań i motywacje tego typu są marginalne lub tłumione. Pojawia się incydentalnie, w stresie lub pod wpływem otoczenia. Brak wpływu na decyzje i styl komunikacji. Nie buduje wizerunku."),
    (30, 49, "słabe",            "słabe natężenie",
     "Archetyp widoczny tylko w tle. Ujawnia się w określonych sytuacjach społecznych — głównie gdy trzeba się dostosować. Ograniczony wpływ na komunikację, zachowania i wizerunek."),
    (50, 59, "umiarkowane",      "umiarkowane natężenie",
     "Archetyp częściowo aktywny – równoważy się z innymi. Poziom neutralny lub „potencjalny rdzeń”, który może się rozwinąć. Zaznacza się w komunikacji, ma wpływ pomocniczy."),
    (60, 69, "znaczące",         "znaczące natężenie",
     "Archetyp wyraźny, widoczny w zachowaniu, decyzjach, wartościach. Zauważalnie kształtuje styl działania i odbiór. Wzmacnia kluczowy archetyp, ale sam nie prowadzi decyzji."),
    (70, 79, "wysokie",          "wysokie natężenie",
     "Archetyp mocno ukształtowany. Centralny motyw sposobu myślenia, postrzegania świata i działania — buduje tożsamość. Nadaje ton komunikacji / przywództwu."),
    (80, 89, "bardzo wysokie",   "bardzo wysokie natężenie (rdzeń)",
     "Archetyp niemal jednoznacznie dominuje nad resztą i w dużym stopniu określa charakter oraz wizerunek. Zwykle prowadzi do jednoznacznego stylu (np. przywódczego, opiekuńczego, buntowniczego, wizjonerskiego). Typowe dla silnych osobowości."),
    (90,100, "ekstremalne",      "ekstremalne natężenie",
     "Archetyp w „czystej” postaci — rzadkie, często idealizowane lub przerysowane ujęcie (np. „czysty Bohater”, „czysty Władca”). Ogromna spójność i autentyczność, ale niska elastyczność i zamknięcie na inne perspektywy."),
]

def interpret_archetype_intensity(pct: float) -> dict:
    """
    Zwraca słownik:
      {'short': 'wysokie', 'full': 'Wysokie natężenie', 'desc': '...'}
    """
    try:
        v = float(pct)
    except Exception:
        v = 0.0
    v = max(0.0, min(100.0, v))
    # Przedziały bez „dziur” dla wartości dziesiętnych:
    # 0-29.999..., 30-49.999..., 50-59.999..., ...
    for lo, hi, short, full, desc in AR_INTENSITY_SCHEME:
        upper_exclusive = hi + 1.0
        if lo <= v < upper_exclusive:
            return {"short": short, "full": full, "desc": desc}
    return {"short": "-", "full": "-", "desc": ""}

def intensity_icon_color(short_label: str) -> str:
    """Kolor kwadracika przy opisie natężenia (większy kontrast między „szarym” i „stalowym”)."""
    m = (short_label or "").strip().lower()
    return {
        "marginalne":     "#9CA3AF",  # szary (jaśniejszy)
        "słabe":          "#6B7280",  # stalowy (wyraźnie ciemniejszy)
        "umiarkowane":    "#FBBF24",  # bursztyn
        "znaczące":       "#F59E0B",  # pomarańcz
        "wysokie":        "#EF4444",  # czerwony
        "bardzo wysokie": "#DC2626",  # ciemny czerwony
        "ekstremalne":    "#7F1D1D",  # bordo
    }.get(m, "#9CA3AF")

def intensity_icon_html(short_label: str) -> str:
    """Zwraca HTML: [kolorowy KWADRAT] + tekst (np. 'umiarkowane')."""
    c = intensity_icon_color(short_label)
    return f"<span class='ap-int-ico' style='background:{c}'></span>{short_label or ''}"

def intensity_help_modal_html() -> str:
    """Modal 'Interpretacja natężenia' – 3 kolumny: Przedział; Interpretacja; Znaczenie i opis jakościowy."""
    rows = "".join(
        (
            "<tr>"
            f"<td>{lo}–{hi}%</td>"
            f"<td><b>{full}</b></td>"
            f"<td><span style='color:#4b5563'>{desc}</span></td>"
            "</tr>"
        )
        for (lo, hi, short, full, desc) in AR_INTENSITY_SCHEME
    )
    return f"""
    <style>
      /* Modal + USTAWIENIA (tu zmieniasz marginesy i szerokości kolumn) */
      #ap-intensity-modal{{
        display:none; position:fixed; inset:0; background:rgba(0,0,0,.35); z-index:9999;
    
        /* 👉 marginesy wierszy tabeli w MODALU */
        --ap-row-pad-top: 16px;     /* góra */
        --ap-row-pad-bottom: 16px;  /* dół */
        --ap-cell-pad-h: 12px;      /* lewo/prawo */
    
        /* 👉 szerokości kolumn (px lub %) */
        --ap-col-w1: 110px;  /* „Przedział” */
        --ap-col-w2: 250px;  /* „Interpretacja” */
        --ap-col-w3: auto;   /* „Znaczenie i opis jakościowy” (reszta miejsca) */
    
        /* 👉 kolory kwadracików */
        --ap-col-1: #9CA3AF;  /* 0–29%   */
        --ap-col-2: #6B7280;  /* 30–49%  */
        --ap-col-3: #FBBF24;  /* 50–59%  */
        --ap-col-4: #F59E0B;  /* 60–69%  */
        --ap-col-5: #EF4444;  /* 70–79%  */
        --ap-col-6: #DC2626;  /* 80–89%  */
        --ap-col-7: #7F1D1D;  /* 90–100% */
      }}
      #ap-intensity-modal:target{{display:block;}}
    
      #ap-intensity-modal .ap-int-modal{{
        position:fixed; left:50%; top:8%; transform:translateX(-50%);
        max-width:980px; background:#fff; border-radius:16px;
        box-shadow:0 20px 60px rgba(0,0,0,.25); padding:22px 24px;
      }}
      #ap-intensity-modal .ap-int-head{{display:flex; align-items:center; gap:10px; margin-bottom:10px;}}
      #ap-intensity-modal .ap-int-title{{font:700 18px/1.2 'Segoe UI',system-ui,Arial;}}
      #ap-intensity-modal .ap-int-close{{margin-left:auto; text-decoration:none; cursor:pointer; font:700 18px/1 monospace; color:#111}}
    
      /* Tabela w MODALU – używa zmiennych powyżej */
      .ap-int-table{{
        width:100%; border-collapse:collapse; font:400 14px/1.45 'Segoe UI',system-ui;
        table-layout: fixed;   /* trzyma szerokości kolumn */
      }}
      .ap-int-table th,
      .ap-int-table td{{
        border-bottom:1px solid #eef2f7;
        padding: var(--ap-row-pad-top) var(--ap-cell-pad-h) var(--ap-row-pad-bottom) var(--ap-cell-pad-h);
        vertical-align:top;
      }}
      .ap-int-table th{{text-align:left; font-weight:700; color:#374151;}}
      .ap-intensity-modal .ap-int-table th:nth-child(1){{ font-weight:800; }}
      .ap-intensity-modal .ap-int-table td:nth-child(1){{ font-weight:700; color:#111; }}

        /* POGRUBIENIE pierwszej kolumny w MODALU */
        #ap-intensity-modal .ap-int-table th:nth-child(1){{ font-weight:800; }}
        #ap-intensity-modal .ap-int-table td:nth-child(1){{ font-weight:700; color:#111; }}
    
      /* Szerokości kolumn (łatwe do zmiany wyżej) */
      .ap-int-table th:nth-child(1), .ap-int-table td:nth-child(1){{ width: var(--ap-col-w1); text-align:center; }}
      .ap-int-table th:nth-child(2), .ap-int-table td:nth-child(2){{ width: var(--ap-col-w2); }}
      .ap-int-table th:nth-child(3), .ap-int-table td:nth-child(3){{ width: var(--ap-col-w3); }}
    
      /* Kwadracik + jego kolory */
      .ap-int-ico{{
        display:inline-block; width:12px; height:12px; border-radius:3px;
        border:1px solid #d1d5db; margin-right:6px; vertical-align:-2px;
      }}
      .ap-i--1{{ background:var(--ap-col-1); }}
      .ap-i--2{{ background:var(--ap-col-2); }}
      .ap-i--3{{ background:var(--ap-col-3); }}
      .ap-i--4{{ background:var(--ap-col-4); }}
      .ap-i--5{{ background:var(--ap-col-5); }}
      .ap-i--6{{ background:var(--ap-col-6); }}
      .ap-i--7{{ background:var(--ap-col-7); border-color:var(--ap-col-7); }}
    </style>

    <div id="ap-intensity-modal">
      <div class="ap-int-modal" role="dialog" aria-modal="true" aria-label="Interpretacja natężenia">
        <div class="ap-int-head">
          <div class="ap-int-title">🧭 Interpretacja natężenia procentowego archetypu</div>
          <a href="#" class="ap-int-close" title="Zamknij">×</a>
        </div>
        <table class="ap-int-table">
          <thead>
            <tr>
              <th>Przedział</th>
              <th>Interpretacja</th>
              <th>Znaczenie i opis jakościowy</th>
            </tr>
          </thead>
          <tbody>
            <tr>
              <td>0–29%</td>
              <td><span class="ap-int-ico ap-i--1"></span><span class="ap-int-label">Marginalne natężenie</span></td>
              <td>Archetyp prawie nie występuje. Obszar „cienia” — wzorce zachowań i motywacje tego typu są marginalne lub tłumione. Pojawia się incydentalnie, w stresie lub pod wpływem otoczenia. Brak wpływu na decyzje i styl komunikacji. Nie buduje wizerunku.</td>
            </tr>
            <tr>
              <td>30–49%</td>
              <td><span class="ap-int-ico ap-i--2"></span><span class="ap-int-label">Słabe natężenie</span></td>
              <td>Archetyp widoczny tylko w tle. Ujawnia się w określonych sytuacjach społecznych — głównie gdy trzeba się dostosować. Ograniczony wpływ na komunikację, zachowania i wizerunek.</td>
            </tr>
            <tr>
              <td>50–59%</td>
              <td><span class="ap-int-ico ap-i--3"></span><span class="ap-int-label">Umiarkowane natężenie</span></td>
              <td>Archetyp częściowo aktywny — równoważy się z innymi. Poziom neutralny lub „potencjalny rdzeń”, który może się rozwinąć. Zaznacza się w komunikacji, ma wpływ pomocniczy.</td>
            </tr>
            <tr>
              <td>60–69%</td>
              <td><span class="ap-int-ico ap-i--4"></span><span class="ap-int-label">Znaczące natężenie</span></td>
              <td>Archetyp wyraźny, widoczny w zachowaniu, decyzjach, wartościach. Zauważalnie kształtuje styl działania i odbiór. Wzmacnia kluczowy archetyp, ale sam nie prowadzi decyzji.</td>
            </tr>
            <tr>
              <td>70–79%</td>
              <td><span class="ap-int-ico ap-i--5"></span><span class="ap-int-label">Wysokie natężenie</span></td>
              <td>Archetyp mocno ukształtowany. Centralny motyw sposobu myślenia, postrzegania świata i działania — buduje tożsamość. Nadaje ton komunikacji / przywództwu.</td>
            </tr>
            <tr>
              <td>80–89%</td>
              <td><span class="ap-int-ico ap-i--6"></span><span class="ap-int-label">Bardzo wysokie natężenie (rdzeń)</span></td>
              <td>Archetyp niemal jednoznacznie dominuje nad resztą i w dużym stopniu określa charakter oraz wizerunek. Zwykle prowadzi do jednoznacznego stylu (np. przywódczego, opiekuńczego, buntowniczego, wizjonerskiego). Typowe dla silnych osobowości.</td>
            </tr>
            <tr>
              <td>90–100%</td>
              <td><span class="ap-int-ico ap-i--7"></span><span class="ap-int-label">Ekstremalne natężenie</span></td>
              <td>Archetyp w „czystej” postaci — rzadkie, często idealizowane lub przerysowane ujęcie (np. „czysty Bohater”, „czysty Władca”). Ogromna spójność i autentyczność, ale bardzo niska (często praktycznie zerowa) elastyczność i zamknięcie na inne perspektywy.</td>
            </tr>
          </tbody>
        </table>
      </div>
    </div>
    """

# --- KOLORY: mapowanie pytań (1..48) do 4 kolorów i kalkulator % ---
COLOR_QUESTION_MAP = {
    "Niebieski": [9,10,11,12,37,38,39,40,41,42,43,44],
    "Zielony":   [13,14,15,16,17,18,19,20,21,22,23,24],
    "Żółty":     [25,26,27,28,33,34,35,36,29,30,31,32],
    "Czerwony":  [1,2,3,4,5,6,7,8,45,46,47,48],
}
COLOR_HEX = {  # kolory pierścieni
    "Czerwony":  "#E53935",
    "Zielony":   "#7ED321",
    "Żółty":     "#FFC107",
    "Niebieski": "#29ABE2",
}

# === 4 bąbelki w linii: średnica ~ udział % w całości, pierścień ~ % względem zwycięzcy ===
def _bubble_svg(value_pct: float, winner_pct: float, color: str,
                diameter_px: int, track="#FFFFFF", text_color="#111") -> str:
    import math
    value_pct  = max(0.0, float(value_pct or 0.0))
    winner_pct = max(0.0001, float(winner_pct or 0.0001))
    ring_pct   = max(0.0, min(100.0, 100.0 * value_pct / winner_pct))  # % względem zwycięzcy

    size   = int(diameter_px)
    stroke = max(10, int(size * 0.14))
    r      = (size - stroke) / 2
    c      = 2 * math.pi * r
    dash   = c * (ring_pct / 100.0)
    gap    = c - dash

    return f"""
    <svg width="{size}" height="{size}" viewBox="0 0 {size} {size}"
         style="display:block; shape-rendering:geometricPrecision;">
      <g transform="rotate(-90 {size/2} {size/2})">
        <circle cx="{size/2}" cy="{size/2}" r="{r}" fill="none"
                stroke="{track}" stroke-width="{stroke}" stroke-linecap="round"/>
        <circle cx="{size/2}" cy="{size/2}" r="{r}" fill="none"
                stroke="{color}" stroke-width="{stroke}" stroke-linecap="round"
                stroke-dasharray="{dash} {gap}"/>
      </g>
      <text x="50%" y="50%" dominant-baseline="middle" text-anchor="middle"
            style="font-family:'Segoe UI',system-ui,-apple-system,Arial;
                   font-size:{int(size*0.30)}px; font-weight:800; fill:{text_color};">
        {int(round(value_pct))}%
      </text>
    </svg>
    """

def color_bubbles_html(pcts: dict[str, float],
                       min_d: int = 110, max_d: int = 240,
                       map_mode: str = "diameter") -> str:
    items = sorted(pcts.items(), key=lambda kv: kv[1], reverse=True)
    if not items:
        return "<div></div>"
    winner_val = max(0.0001, items[0][1])

    def dia(val):
        frac = max(0.0, min(1.0, (val or 0.0) / 100.0))
        if map_mode == "area":
            frac = frac ** 0.5
        return int(min_d + frac * (max_d - min_d))

    blocks = []
    for name, val in items:
        d = dia(val)
        svg = _bubble_svg(val, winner_val, COLOR_HEX[name], d)
        blocks.append(f"""
          <div class="ap-bubble-wrap">
            <div class="ap-bubble" style="width:{d}px;height:{d}px;">{svg}</div>
            <div class="ap-chip">
              <span class="ap-dot" style="background:{COLOR_HEX[name]}"></span>{name}
            </div>
          </div>
        """)

    return f"""
    <style>
      .ap-bubbles-row{{display:flex;justify-content:center;align-items:flex-end;gap:26px;background:transparent;}}
    .ap-bubble-wrap{{display:flex;flex-direction:column;align-items:center;gap:16px;}}
      /* delikatny, neutralny cień tylko CSS – brak kolorowych smug */
      .ap-bubble{{filter: drop-shadow(0 6px 12px rgba(0,0,0,.12)); border-radius:50%;}}
    .ap-chip{{display:inline-flex;align-items:center;gap:8px;padding:6px 10px;
             border:1px solid #eceff3;border-radius:999px;margin-top:4px;
             font:600 13px/1.1 'Segoe UI',system-ui;background:#fff;}}
      .ap-dot{{width:10px;height:10px;border-radius:50%;display:inline-block}}
    </style>
    <div class="ap-bubbles-row">{''.join(blocks)}</div>
    """

def color_progress_bars_html(
    pcts: dict[str, float],
    order: str = "desc",
    label_font: str = "'Roboto','Segoe UI',system-ui,Arial,sans-serif",
    label_size_px: int = 15,                  # mniejsze niż było
    label_color: str = "#31333F",
    row_vmargin_px: int = 17,
    track_height_px: int = 40,
    track_color: str = "#eef2f7",
    value_font: str = "'Roboto','Segoe UI',system-ui,Arial,sans-serif",
    value_size_px: int = 17,                  # większe %
):
    items = sorted(pcts.items(), key=lambda kv: kv[1], reverse=(order == "desc"))

    rows = []
    for name, val in items:
        pct = max(0.0, float(val or 0.0))
        pct_int = int(round(pct))
        width_css = f"{pct:.3f}%"
        color = COLOR_HEX[name]
        inside = pct >= 10.0
        rows.append(f"""
                  <div class="cp-row">
                    <div class="cp-label">
                      <span class="cp-dot" style="background:{color}"></span>
                      <span class="cp-label-text">{name}</span>
                    </div>
                    <div class="cp-track" title="{COLOR_LONG.get(name, {}).get('title', '').replace('"', '&quot;')}">
                      <div class="cp-fill" style="width:{width_css}; background:{color}"
                           title="{COLOR_LONG.get(name, {}).get('title', '').replace('"', '&quot;')}">
                        <div class="cp-badge {'in' if inside else 'out'}">{pct_int}%</div>
                      </div>
                    </div>
                  </div>
                """)
    return f"""
    <style>
      .cp-row{{
        display:grid; grid-template-columns:110px 1fr; gap:8px;
        align-items:center; margin:{row_vmargin_px}px 0;
      }}
      .cp-label{{display:flex; align-items:center; gap:10px;}}
      .cp-label-text{{
        font-family:{label_font};
        font-size:{label_size_px}px;
        color:{label_color};
        font-weight:500;                /* lżej */
        letter-spacing:.0px;
      }}
      .cp-dot{{width:10px; height:10px; border-radius:50%; display:inline-block;}}
      .cp-track{{
        position:relative; height:{track_height_px}px; border-radius:999px;
        background:{track_color}; box-shadow: inset 0 0 0 1px #e1e7f0;
      }}
      .cp-fill{{position:relative; height:100%; border-radius:999px; overflow:visible;}}
      .cp-badge{{
        position:absolute; top:50%; transform:translateY(-50%);
        font-family:{value_font}; font-size:{value_size_px}px;
        font-weight:700;               /* grube % */
        color:#111; white-space:nowrap;
      }}
      .cp-badge.in{{ right:12px; }}    /* do wewnętrznej krawędzi */
      .cp-badge.out{{ left:100%; margin-left:12px; }}
    </style>
    <div class="cp-wrap">{''.join(rows)}</div>
    """


def _sum_color_points_for_answers(answers: list[int]) -> dict[str,int]:
    """Suma punktów na podstawie jednej odpowiedzi (48 pytań)."""
    out = {k: 0 for k in COLOR_QUESTION_MAP}
    if not isinstance(answers, list) or len(answers) < 48:
        return out
    for color, qs in COLOR_QUESTION_MAP.items():
        out[color] += sum(answers[i-1] for i in qs)
    return out

def calc_color_percentages_from_df(df: pd.DataFrame) -> dict[str, float]:
    """Średni % dla każdego koloru przy skali 0–100 per kolor (60 pkt max na kolor i osobę)."""
    totals = {k: 0 for k in COLOR_QUESTION_MAP}
    n = 0
    if "answers" not in df.columns or df.empty:
        return {k: 0.0 for k in COLOR_QUESTION_MAP}

    for _, row in df.iterrows():
        ans = row.get("answers")
        sums = _sum_color_points_for_answers(ans)
        if sums:
            for k, v in sums.items():
                totals[k] += v
            n += 1

    if n == 0:
        return {k: 0.0 for k in COLOR_QUESTION_MAP}

    max_per_color = {c: 5 * len(qs) for c, qs in COLOR_QUESTION_MAP.items()}
    return {
        c: round(100.0 * totals[c] / (max_per_color[c] * n), 1)
        for c in COLOR_QUESTION_MAP.keys()
    }

# --- Pierścień w SVG (transparentne tło + delikatna poświata) ---
def _ring_svg(percent: float, color: str, size: int = 180, stroke: int = 16,
              track="#FFFFFF", text_color="#333") -> str:
    import math, uuid
    pct = max(0.0, min(100.0, float(percent)))
    r = (size - stroke) / 2
    c = 2 * math.pi * r
    dash = c * pct / 100.0
    gap = c - dash

    uid = "g_" + uuid.uuid4().hex[:8]  # 👈 unikalny id filtra

    filter_def = ""
    return f"""
    <svg width="{size}" height="{size}" viewBox="0 0 {size} {size}" style="display:block;">
      {filter_def}
      <g transform="rotate(-90 {size/2} {size/2})">
        <circle cx="{size/2}" cy="{size/2}" r="{r}" fill="none"
                stroke="{track}" stroke-width="{stroke}" stroke-linecap="round"/>
        <circle cx="{size/2}" cy="{size/2}" r="{r}" fill="none"
            stroke="{color}" stroke-width="{stroke}" stroke-linecap="round"
            stroke-dasharray="{dash} {gap}"/>
      </g>
      <g>
        <text x="50%" y="50%" dominant-baseline="middle" text-anchor="middle"
              style="font-family: 'Segoe UI', system-ui, -apple-system, Arial; font-size:{int(size*0.28)}px; font-weight:800; fill:{text_color};">
          {int(round(pct))}%
        </text>
      </g>
    </svg>
    """

from PIL import Image, ImageDraw, ImageFont
from io import BytesIO

def make_color_progress_png_for_word(
    pcts: dict[str, float],
    out_path: str = "color_progress.png",
    # rozmiary/odstępy
    width_px: int = 1600,
    pad: int = 36,
    bar_h: int = 72,            # nieco wyższe pastylki → większe % bez ścisku
    bar_gap: int = 34,
    # etykieta (kropka + tekst) i odstęp do paska
    dot_radius: int = 10,
    label_gap_px: int = 40,
    # typografia (etykieta lżejsza, % duże i grube)
    label_font_size: int = 30,
    pct_font_size: int | None = None,   # auto: 50% wysokości paska
    pct_margin: int = 20                # odsunięcie % od krawędzi
):
    rows = sorted(pcts.items(), key=lambda kv: kv[1], reverse=True)
    n = len(rows)
    H = pad + pad + n * bar_h + (n - 1) * bar_gap
    W = width_px

    img = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    drw = ImageDraw.Draw(img)

    track = (238, 242, 247, 255)
    label_col = (33, 35, 47, 255)
    val_col = (17, 17, 17, 255)

    bold_paths = [
        "assets/fonts/RobotoCondensed-Bold.ttf",
        "assets/fonts/DejaVuSans-Bold.ttf",
        "assets/fonts/Arial Bold.ttf",
    ]
    reg_paths  = [
        "assets/fonts/RobotoCondensed-Regular.ttf",
        "assets/fonts/DejaVuSans.ttf",
        "assets/fonts/Arial.ttf",
    ]

    def _try_font(paths, size):
        for p in paths:
            try:
                return ImageFont.truetype(p, size)
            except Exception:
                continue
        return ImageFont.load_default()

    # etykieta – lżejsza
    f_label = _try_font(reg_paths,  label_font_size)

    def _luma(hexcode: str) -> float:
        h = hexcode.lstrip('#')
        if len(h) == 3:
            h = ''.join(c*2 for c in h)
        r, g, b = (int(h[i:i+2], 16) for i in (0, 2, 4))
        return 0.2126*r + 0.7152*g + 0.0722*b

    def draw_bold_text(drw, xy, txt, font, fill, stroke=1):
        """Proste pogrubienie tekstu przez dorysowanie 4 przesuniętych kopii."""
        x, y = xy
        if stroke > 0:
            for dx, dy in ((-stroke, 0), (stroke, 0), (0, -stroke), (0, stroke)):
                drw.text((x + dx, y + dy), txt, font=font, fill=fill)
        drw.text((x, y), txt, font=font, fill=fill)

    # maks. szerokość etykiety → stały odstęp od paska
    max_label_text_w = 0
    for name, _ in rows:
        bbox = drw.textbbox((0, 0), name, font=f_label)
        max_label_text_w = max(max_label_text_w, bbox[2]-bbox[0])

    dot_w   = 2*dot_radius
    dot_pad = 10
    label_block_w = 6 + dot_w + dot_pad + max_label_text_w
    x0_base = pad + label_block_w + label_gap_px
    bar_w   = W - x0_base - pad
    radius  = bar_h // 2

    y = pad

    # --- KONFIG WARTOŚCI % ---
    PCT_FONT_PX = 33  # rozmiar czcionki %
    PCT_BOLD = True  # True = pogrubiona, False = zwykła
    PCT_RIGHT_PAD = max(10, int(bar_h * 0.35))  # odstęp od prawej krawędzi wypełnienia (px)
    PCT_VSHIFT = 0  # ręczne przesunięcie w pionie (px, + w dół, - w górę)

    # załaduj font dla % (pogrubiony lub regularny)
    def _font(paths, size):
        from PIL import ImageFont
        for p in paths:
            try:
                return ImageFont.truetype(p, size)
            except:
                pass
        return ImageFont.load_default()

    _pct_font_paths_bold = [
        "assets/fonts/RobotoCondensed-Bold.ttf",
        "assets/fonts/ArialNovaCond-Bold.ttf",
        "DejaVuSans-Bold.ttf",
    ]
    _pct_font_paths_reg = [
        "assets/fonts/RobotoCondensed-Regular.ttf",
        "assets/fonts/ArialNovaCond.ttf",
        "DejaVuSans.ttf",
    ]
    f_pct = _font(_pct_font_paths_bold if PCT_BOLD else _pct_font_paths_reg, PCT_FONT_PX)

    for name, val in rows:
        ly = y + bar_h//2
        c_hex = COLOR_HEX[name]

        # kropka
        drw.ellipse([pad+6, ly-dot_radius, pad+6+dot_w, ly+dot_radius], fill=c_hex)

        # etykieta – idealnie na środku (anchor) z fallbackiem
        label_x = pad+6+dot_w+dot_pad
        try:
            drw.text((label_x, ly), name, fill=label_col, font=f_label, anchor="lm")
        except TypeError:
            tb = drw.textbbox((0, 0), name, font=f_label)
            th = tb[3]-tb[1]
            drw.text((label_x, ly - th/2), name, fill=label_col, font=f_label)

        # tor
        x0 = x0_base
        x1 = x0 + bar_w
        y0 = y
        y1 = y + bar_h
        drw.rounded_rectangle([x0, y0, x1, y1], radius=radius, fill=track)

        # wypełnienie
        pct_val = max(0.0, min(100.0, float(val)))
        fill_w = int(bar_w * pct_val / 100.0)
        if fill_w > 0:
            drw.rounded_rectangle([x0, y0, x0+fill_w, y1], radius=radius, fill=c_hex)

        # --- % na prawej wewnętrznej krawędzi wypełnienia, wyśrodkowane pionowo ---
        pct_text = f"{int(round(val))}%"

        # kolor napisu w środku; jasność wypełnienia decyduje o bieli/czerni
        text_fill_inside = (255, 255, 255, 255) if _luma(c_hex) < 110 else (17, 17, 17, 255)

        cy = y0 + bar_h / 2 + PCT_VSHIFT  # środek pionowy paska + ewentualny v-shift

        if fill_w > 0:
            # x przy prawej krawędzi wypełnienia z marginesem do środka (anchor='rm')
            cx = x0 + fill_w - PCT_RIGHT_PAD
            # nie pozwól wyjechać w lewo przy małych słupkach
            cx = max(cx, x0 + PCT_RIGHT_PAD)

            try:
                # right-middle = „przyklej” do prawej krawędzi wypełnienia, wyśrodkuj w pionie
                drw.text((cx, cy), pct_text, font=f_pct, fill=text_fill_inside, anchor="rm")
            except TypeError:
                # fallback dla starszego Pillow bez anchor
                pb = drw.textbbox((0, 0), pct_text, font=f_pct)
                tx = cx - (pb[2] - pb[0])  # prawa krawędź tekstu w cx
                ty = cy - (pb[3] - pb[1]) / 2
                drw.text((tx, ty), pct_text, font=f_pct, fill=text_fill_inside)
        else:
            # 0% – przy początku toru, lekko od lewej, wyśrodkowane pionowo
            try:
                drw.text((x0 + PCT_RIGHT_PAD, cy), pct_text, font=f_pct, fill=(17, 17, 17, 255),
                         anchor="lm")
            except TypeError:
                pb = drw.textbbox((0, 0), pct_text, font=f_pct)
                drw.text((x0 + PCT_RIGHT_PAD, cy - (pb[3] - pb[1]) / 2), pct_text, font=f_pct,
                         fill=(17, 17, 17, 255))

        y += bar_h + bar_gap

    img.save(out_path, "PNG")
    return out_path

from PIL import Image, ImageDraw, ImageFont
from io import BytesIO
import math

# mapowanie archetyp → kolor „rodziny” (jak w całej appce)
ARCHETYPE_TO_COLOR_FAMILY = {
    "Władca":"Czerwony", "Bohater":"Czerwony", "Buntownik":"Czerwony",
    "Twórca":"Żółty", "Czarodziej":"Żółty", "Odkrywca":"Żółty",
    "Opiekun":"Zielony", "Kochanek":"Zielony", "Błazen":"Zielony",
    "Mędrzec":"Niebieski", "Towarzysz":"Niebieski", "Niewinny":"Niebieski",
}

def make_capsule_columns_png_for_word(
    means_pct_by_arche: dict[str, float],
    out_path: str = "arche_capsules.png",
    width_px: int = 1900,
    grid_height_px: int = 620,
    top_title: str | None = None,
):
    """
    Kolumny-kapsuły z osią Y oraz X; % nad kropką; siatka 10%.
    (Poprawki: wyższe % nad kropkami, niższe nazwy, kapsuły do 100%, oś X=Y.)
    """
    if not means_pct_by_arche:
        return None

    from PIL import Image, ImageDraw, ImageFont

    # render w 2× dla ostrości (geometria bez zmian)
    S = 2

    # ------- dane -------
    items = sorted(means_pct_by_arche.items(), key=lambda kv: kv[1], reverse=True)

    # ------- layout (logiczne px) -------
    pad_x = 84
    pad_title = 96 if top_title else 40
    axis_band = 62
    pad_bottom = 118
    grid_left  = pad_x + axis_band
    grid_right = width_px - pad_x
    n = len(items)

    col_w = max(68, int((grid_right - grid_left) / max(n, 1) * 0.72))
    gap   = max(30, int((grid_right - grid_left - n * col_w) / max(n - 1, 1))) if n > 1 else 0

    grid_h    = grid_height_px
    height_px = pad_title + grid_h + pad_bottom

    # przeskalowane płótno
    W, H = width_px * S, height_px * S
    img = Image.new("RGBA", (W, H), (255, 255, 255, 0))
    drw = ImageDraw.Draw(img)

    # ------- fonty -------
    def _try_font(paths, size):
        from PIL import ImageFont
        for p in paths:
            try: return ImageFont.truetype(p, size)
            except Exception: pass
        return ImageFont.load_default()

    f_title = _try_font(["assets/fonts/RobotoCondensed-Bold.ttf","DejaVuSans-Bold.ttf","Arial.ttf"], 42*S)
    f_name  = _try_font(["assets/fonts/RobotoCondensed-Regular.ttf","DejaVuSans.ttf","Arial.ttf"], 28*S)
    f_pct   = _try_font(["assets/fonts/RobotoCondensed-Bold.ttf","DejaVuSans-Bold.ttf","Arial.ttf"], 27*S)
    f_axis  = _try_font(["assets/fonts/RobotoCondensed-Regular.ttf","DejaVuSans.ttf","Arial.ttf"], 21*S)
    f_pct_big = _try_font(
        ["assets/fonts/RobotoCondensed-Bold.ttf", "DejaVuSans-Bold.ttf", "Arial.ttf"], (27 + 5) * S)

    # ------- kolory / linie -------
    label_color = (30, 32, 36, 255)
    axis_col = (138, 144, 153, 255)      # identyczny dla X i Y
    tick_col = (200, 205, 213, 255)
    grid_col = (226, 230, 236, 255)
    cap_bg = (234, 238, 243, 255)
    cap_border = (206, 210, 217, 255)
    AXIS_W = 2 * S

    # ------- tytuł -------
    if top_title:
        tb = drw.textbbox((0,0), top_title, font=f_title)
        drw.text(((W - (tb[2]-tb[0]))/2, 16*S), top_title, font=f_title, fill=label_color)

    # ------- osie -------
    y_top   = pad_title * S
    y_bot   = (pad_title + grid_h) * S
    x_axis_y = y_bot                     # oś X na dnie siatki
    axis_x   = (grid_left - 32) * S

    # Oś Y
    drw.line([(axis_x, y_top), (axis_x, x_axis_y)], fill=axis_col, width=AXIS_W)
    # Oś X (ten sam kolor i grubość co Y)
    drw.line([(grid_left * S, x_axis_y), (grid_right * S, x_axis_y)], fill=axis_col, width=AXIS_W)

    # Linie poziome co 10% + etykiety
    for p in range(0, 101, 10):
        yy = y_bot - int(round(grid_h * S * (p/100.0)))
        drw.line([(grid_left * S, yy), (grid_right * S, yy)], fill=grid_col, width=1 * S)
        drw.line([(axis_x - 6 * S, yy), (axis_x, yy)], fill=tick_col, width=1 * S)
        txt = f"{p}%"
        tb = drw.textbbox((0,0), txt, font=f_axis)
        drw.text((axis_x - 10 * S - (tb[2]-tb[0]), yy - (tb[3]-tb[1])//2), txt, font=f_axis, fill=axis_col)

    # ------- kapsuły (teraz DOCHODZĄ do 100%) -------
    col_w_S = col_w * S
    gap_S   = gap * S
    cap_r   = (col_w // 2) * S
    cap_top = y_top               # góra = dokładnie 100%
    cap_bot = x_axis_y            # dół = oś X (0%)

    def y_for_pct(p):
        p = max(0.0, min(100.0, float(p)))
        return cap_bot - int(round((cap_bot - cap_top) * (p/100.0)))

    # offsety wg Twoich uwag
    PCT_OFFSET_ABOVE_DOT = 42 * S     # % wyżej nad kropką
    NAME_OFFSET_BELOW_X  = 18 * S     # nazwy niżej od osi X

    x = grid_left * S
    for name, val in items:
        # tło kapsuły – top=100%, bottom=0%
        cap_left, cap_right = x, x + col_w_S
        drw.rounded_rectangle([cap_left, cap_top, cap_right, cap_bot],
                              radius=cap_r, fill=cap_bg, outline=cap_border, width=2 * S)

        fam = ARCHETYPE_TO_COLOR_FAMILY.get(name, "Niebieski")
        c_hex = COLOR_HEX[fam]
        c_rgb = tuple(int(c_hex[i:i+2], 16) for i in (1,3,5)) + (255,)

        cx    = (cap_left + cap_right) // 2
        y_val = y_for_pct(val)

        # linia od 0%
        drw.line([(cx, cap_bot), (cx, y_val)], fill=c_rgb, width=6 * S)

        # kropka
        dot_r = max(6 * S, int(col_w * 0.13) * S)
        drw.ellipse([cx-dot_r, y_val-dot_r, cx+dot_r, y_val+dot_r], fill=c_rgb)

        # % NAD kropką
        pct_value = float(val)
        pct_txt = f"{round(pct_value, 1):.1f}%"
        font_pct = f_pct_big if pct_value >= 70.0 else f_pct  # +2 px dla ≥ 70%

        pb = drw.textbbox((0, 0), pct_txt, font=font_pct)
        drw.text((cx - (pb[2] - pb[0]) / 2, y_val - dot_r - PCT_OFFSET_ABOVE_DOT),
                 pct_txt, font=font_pct, fill=c_rgb)

        # ticzki X i nazwa (niżej)
        drw.line([(cx, cap_bot), (cx, cap_bot + 7 * S)], fill=tick_col, width=1 * S)
        nb = drw.textbbox((0,0), name, font=f_name)
        drw.text((cx - (nb[2]-nb[0]) / 2, cap_bot + NAME_OFFSET_BELOW_X),
                 name, font=f_name, fill=label_color)

        x += col_w_S + gap_S

    img.save(out_path, "PNG", dpi=(300, 300))
    return out_path





def mean_pct_by_archetype_from_df(df: pd.DataFrame) -> dict[str, float]:
    """
    Zwraca {archetyp: średni % 0..100} dla całego df.
    Zakłada, że kolumna 'answers' ma listę 48 wartości 0..5.
    """
    if df.empty: return {k: 0.0 for k in archetypes.keys()}
    totals = {k: 0 for k in archetypes.keys()}
    n = 0
    for _, row in df.iterrows():
        ans = row.get("answers")
        if not isinstance(ans, list) or len(ans) < 48:
            continue
        sc = archetype_scores(ans)  # masz wyżej
        for k, v in sc.items():
            totals[k] += v
        n += 1
    if n == 0: return {k: 0.0 for k in archetypes.keys()}
    # 4 pytania na archetyp * 5 pkt = 20 maks → %:
    return {k: round((totals[k] / (20.0 * n)) * 100.0, 2) for k in archetypes.keys()}


def color_gauges_html(pcts: dict[str, float]) -> str:
    """Układ jak na screenie: 1 duży + 3 małe po prawej; full transparent."""
    # sort – pierwszy to największy (duży)
    items = sorted(pcts.items(), key=lambda kv: kv[1], reverse=True)
    big = items[0]
    small = items[1:]

    def chip(name, hex_):
        return (f"<span style='display:inline-flex;align-items:center;"
                f"gap:8px;padding:6px 10px;border:1px solid #e8e8ee;border-radius:999px;"
                f"font:600 13px/1.1 \"Segoe UI\",system-ui; background:rgba(255,255,255,.65)'>"
                f"<span style='width:10px;height:10px;border-radius:50%;background:{hex_};display:inline-block'></span>{name}</span>")

    big_svg = _ring_svg(big[1], COLOR_HEX[big[0]], size=220, stroke=20)
    small_svgs = "".join(
        f"<div style='display:flex;align-items:center;gap:12px'>"
        f"<div style='width:120px'>{_ring_svg(v, COLOR_HEX[k], size=120, stroke=14)}</div>"
        f"{chip(k, COLOR_HEX[k])}"
        f"</div>"
        for k, v in small
    )

    return f"""
    <div style="display:flex; gap:32px; align-items:center; background:transparent;">
      <div style="width:240px">{big_svg}</div>
      <div style="display:flex; flex-direction:column; gap:18px;">{small_svgs}</div>
    </div>
    """

# --- OPISY HEURYSTYCZNE 4 KOLORÓW (treści z Twojej specyfikacji) ---
COLOR_EMOJI = {"Czerwony":"🔴","Zielony":"🟢","Żółty":"🟡","Niebieski":"🔵"}

COLOR_LONG = {
    "Niebieski": {
        "title": "Niebieski – analityczny, proceduralny, precyzyjny",
        "orient": "fakty, dane, logikę i procedury",
        "arche": "Mędrzec, Towarzysz, Niewinny",
        "body": (
            "Ceni fakty, logikę i stabilne procedury. Działa najlepiej, gdy ma jasno określone zasady, "
            "harmonogram i dostęp do danych. Nie lubi chaosu, nagłych zmian i improwizacji – woli działać "
            "według planu. Może sprawiać wrażenie zdystansowanego i nadmiernie ostrożnego, ale wnosi do "
            "zespołu rzetelność, sumienność i dbałość o szczegóły. Niebieski to myślenie."
        ),
        "politics": (
            "W polityce to typ eksperta – skrupulatny analityk, który zamiast haseł pokazuje liczby i tabele. "
            "Budzi zaufanie dzięki przygotowaniu merytorycznemu i pragmatycznym rozwiązaniom. Może być odbierany "
            "jako mało charyzmatyczny, ale daje wyborcom poczucie przewidywalności i bezpieczeństwa instytucjonalnego."
        ),
        "hex": COLOR_HEX["Niebieski"]
    },
    "Zielony": {
        "title": "Zielony – empatyczny, harmonijny, wspierający",
        "orient": "relacje, troskę, zaufanie, wspólnotę",
        "arche": "Opiekun, Kochanek, Błazen",
        "body": (
            "Kieruje się wartościami, relacjami i potrzebą budowania poczucia bezpieczeństwa. Jest empatyczny, "
            "uważny na innych i dąży do zgody. Nie lubi gwałtownych zmian i konfrontacji, czasem brakuje mu "
            "asertywności, ale potrafi tworzyć atmosferę zaufania i współpracy. Wnosi do zespołu stabilność, "
            "lojalność i umiejętność łagodzenia napięć. Zielony to uczucia."
        ),
        "politics": (
            "W polityce to typ mediator-społecznik, który stawia na dialog, kompromis i dobro wspólne. Potrafi "
            "przekonać elektorat stylem „opiekuńczego lidera”, akcentując wartości społeczne, wspólnotowe i "
            "solidarnościowe. Może unikać ostrych sporów, ale umiejętnie buduje mosty i zdobywa poparcie przez "
            "bliskość i troskę o codzienne sprawy ludzi."
        ),
        "hex": COLOR_HEX["Zielony"]
    },
    "Żółty": {
        "title": "Żółty – kreatywny, pełny energii i spontaniczny",
        "orient": "wizję, innowację, możliwości, odkrywanie nowych dróg",
        "arche": "Twórca, Czarodziej, Odkrywca",
        "body": (
            "Osoba wizjonerska i entuzjastyczna – pełna pomysłów, które inspirują innych. Najlepiej czuje się w "
            "środowisku swobodnym, otwartym na eksperymenty i innowacje. Nie przepada za rutyną, schematami i "
            "nadmierną kontrolą. Jego mocną stroną jest umiejętność rozbudzania energii zespołu, improwizacja i "
            "znajdowanie nowych możliwości tam, gdzie inni widzą bariery. Żółty to intuicja."
        ),
        "politics": (
            "W polityce to typ showmana i wizjonera, który potrafi porwać tłumy hasłami zmiany i nowego otwarcia. "
            "Umie przekuć abstrakcyjne idee w obrazowe narracje, które przemawiają do emocji. Bywa odbierany jako "
            "idealista lub ryzykant, ale świetnie nadaje dynamikę kampanii i kreuje „nową nadzieję”."
        ),
        "hex": COLOR_HEX["Żółty"]
    },
    "Czerwony": {
        "title": "Czerwony – decyzyjny, nastawiony na wynik, dominujący",
        "orient": "działanie, sprawczość, szybkie decyzje, forsowanie kierunku",
        "arche": "Władca, Bohater, Buntownik",
        "body": (
            "Ma naturalne zdolności przywódcze i skłonność do szybkiego podejmowania decyzji. Jest niezależny, "
            "ambitny i skoncentrowany na rezultatach. Może być niecierpliwy, zbyt stanowczy i mało elastyczny, "
            "ale dzięki determinacji potrafi przeprowadzić projekt do końca mimo przeszkód. To osoba, która nadaje "
            "kierunek i mobilizuje innych do działania. Czerwony to doświadczenie."
        ),
        "politics": (
            "W polityce to typ lidera-wojownika, który buduje swoją pozycję na sile, determinacji i zdolności "
            "„dowiezienia” obietnic. Sprawdza się w kampaniach, gdzie liczy się mocne przywództwo i szybkie decyzje. "
            "Może odstraszać swoją twardością, ale równocześnie daje poczucie, że „trzyma ster”."
        ),
        "hex": COLOR_HEX["Czerwony"]
    },
}

COLOR_META = {
    name: {
        "emoji": COLOR_EMOJI[name],
        "title": COLOR_LONG[name]["title"],
        "orient": f"Orientacja na: {COLOR_LONG[name]['orient']}",
        "arche": f"Archetypy: {COLOR_LONG[name]['arche']}",
        "desc":   COLOR_LONG[name]["body"] + "\n\n" + " 👉 " + COLOR_LONG[name]["politics"],
    }
    for name in COLOR_LONG.keys()
}

def color_explainer_one_html(name: str, pct: float, dark_mode: bool = False) -> str:
    """Jeden panel z opisem dominującego koloru."""
    meta = COLOR_LONG[name]
    emoji = COLOR_EMOJI[name]
    card_bg = "rgba(12, 23, 41, .72)" if dark_mode else "rgba(255,255,255,.65)"
    card_border = "rgba(148,163,184,.38)" if dark_mode else "#ececf3"
    title_color = "#e8f1ff" if dark_mode else "#1f2937"
    body_color = "#d5e2f3" if dark_mode else "#2a2a2a"
    muted_color = "#d2deef" if dark_mode else "#444"
    pct_color = "#e8f1ff" if dark_mode else "#333"
    return f"""
      <div style="border:1px solid {card_border}; border-left:6px solid {meta['hex']};
                  border-radius:12px; padding:20px 22px; margin:4px 0 6px 0;
                  background:{card_bg};">
        <div style="display:flex;align-items:center;gap:12px;margin-bottom:10px;">
          <span style="font-size:20px">{emoji}</span>
          <div style="font:650 18px/1.2 'Segoe UI',system-ui;color:{title_color};">{meta['title']}</div>
          <div style="margin-left:auto;font:700 14px/1 'Segoe UI',system-ui;color:{pct_color}">{pct:.1f}%</div>
        </div>

        <div style="font:600 16px/1.45 'Segoe UI',system-ui; color:{muted_color};">
          • <b>Orientacja na:</b> {meta['orient']}
        </div>
        
        <div style="font:510 14px/1.40 'Segoe UI',system-ui; color:{muted_color}; margin-top:6px;">
          • <b>Archetypy:</b> {meta['arche']}
        </div>

        <div style="margin-top:12px; font:400 14px/1.6 'Segoe UI',system-ui; color:{body_color};">
          {meta['body']}
        </div>

        <div style="margin-top:12px; font:400 14px/1.6 'Segoe UI',system-ui; color:{body_color};">
          👉 {meta['politics']}
        </div>
      </div>
    """


def color_explainer_html(pcts: dict[str, float]) -> str:
    """Render 4 akapity pod wykresem – kolejność wg udziału (%)."""
    items = sorted(pcts.items(), key=lambda kv: kv[1], reverse=True)
    blocks = []
    for name, val in items:
        meta = COLOR_LONG[name]
        emoji = COLOR_EMOJI[name]
        blocks.append(f"""
        <div style="border:1px solid #ececf3; border-left:6px solid {meta['hex']};
                    border-radius:12px; padding:16px 18px; margin-bottom:14px;
                    background:rgba(255,255,255,.65);">
          <div style="display:flex;align-items:center;gap:12px;margin-bottom:10px;">
             <span style="font-size:20px">{emoji}</span>
             <div style="font:650 18px/1.2 'Segoe UI',system-ui">{meta['title']}</div>
             <div style="margin-left:auto;font:700 14px/1 'Segoe UI',system-ui;color:#333">{val:.1f}%</div>
          </div>
          <div style="font:600 16px/1.45 'Segoe UI',system-ui; color:#444;">
             • <b>Orientacja na:</b> {meta['orient']}<br/>
          </div>
          <div style="font:510 14px/1.40 'Segoe UI',system-ui; color:#444; margin-top:6px;">
             • <b>Archetypy:</b> {meta['arche']}
          </div>
          <div style="margin-top:12px; font:400 14px/1.6 'Segoe UI',system-ui; color:#2a2a2a;">{meta['body']}</div>
          <div style="margin-top:12px; font:400 14px/1.6 'Segoe UI',system-ui; color:#2a2a2a;">👉 {meta['politics']}</div>
        </div>
        """)
    return "<div style='background:transparent'>" + "".join(blocks) + "</div>"


def color_scores_from_answers(answers: list[int]) -> dict[str, int]:
    if not isinstance(answers, list) or len(answers) < 48:
        return {c: 0 for c in COLOR_QUESTION_MAP}
    out = {}
    for color, idxs in COLOR_QUESTION_MAP.items():
        out[color] = sum(answers[i-1] for i in idxs)  # pytania są 1-indexed
    return out

def color_percents_from_scores(scores: dict[str, int]) -> dict[str, float]:
    # 12 pytań × 5 pkt = 60; obliczaj z mapy dla odporności
    max_per_color = {c: 5 * len(qs) for c, qs in COLOR_QUESTION_MAP.items()}
    out = {}
    for c in COLOR_QUESTION_MAP.keys():
        v = max(0, int(scores.get(c, 0)))
        out[c] = round(100.0 * v / max_per_color[c], 1)
    return out


# --- PŁEĆ I ODWZOROWANIA NAZW/PLIKÓW ---

# Feminatywy, zgodnie z Twoją listą
GENDER_FEMININE_MAP = {
    "Władca": "Władczyni",
    "Bohater": "Bohaterka",
    "Mędrzec": "Mędrczyni",
    "Opiekun": "Opiekunka",
    "Kochanek": "Kochanka",
    "Błazen": "Komiczka",
    "Twórca": "Twórczyni",
    "Odkrywca": "Odkrywczyni",
    "Czarodziej": "Czarodziejka",
    "Towarzysz": "Towarzyszka",
    "Niewinny": "Niewinna",
    "Buntownik": "Buntowniczka",
}

# odwrotna mapa (z żeńskich na męskie), przydaje się, gdy wejściowo dostaniemy już żeńską formę
GENDER_MASC_FROM_FEM = {v: k for k, v in GENDER_FEMININE_MAP.items()}

# “bazowe” nazwy plików w assets/person_icons (bez sufiksu _M/_K i rozszerzenia)
# ← podajemy tu męskie formy jako klucz
ARCHETYPE_BASE_SLUGS = {
    "Władca": "wladca",
    "Bohater": "bohater",
    "Mędrzec": "medrzec",
    "Opiekun": "opiekun",
    "Kochanek": "kochanek",
    "Błazen": "blazen",
    "Twórca": "tworca",
    "Odkrywca": "odkrywca",
    "Czarodziej": "czarodziej",
    "Towarzysz": "towarzysz",
    "Niewinny": "niewinny",
    "Buntownik": "buntownik",
}

def normalize_gender(value) -> str:
    """
    Z dowolnej wartości („M”, „K”, „mężczyzna”, „kobieta”, True/False, itp.)
    zwraca kod 'M' albo 'K'. Domyślnie 'M'.
    """
    v = (str(value or "")).strip().lower()
    if v in ("k", "kobieta", "female", "f", "kob"):
        return "K"
    return "M"

def display_name_for_gender(base_masc_name: str, gender_code: str) -> str:
    """Zwraca nazwę do pokazania na ekranie zależnie od płci."""
    if gender_code == "K":
        return GENDER_FEMININE_MAP.get(base_masc_name, base_masc_name)
    return base_masc_name

def base_masc_from_any(name: str) -> str:
    """Normalizuje nazwę archetypu do formy męskiej bazowej (odporne na case/PL znaki/warianty)."""
    raw = str(name or "").strip()
    if not raw:
        return ""
    if raw in GENDER_MASC_FROM_FEM:
        return GENDER_MASC_FROM_FEM[raw]

    def _norm_token(v: str) -> str:
        return re.sub(r"[^a-z0-9]+", "", _slug_pl(v or ""))

    norm_raw = _norm_token(raw)
    if not norm_raw:
        return raw

    masc_by_norm = {_norm_token(k): k for k in ARCHETYPE_BASE_SLUGS.keys()}
    fem_by_norm = {_norm_token(f): m for f, m in GENDER_MASC_FROM_FEM.items()}
    slug_by_norm = {_norm_token(v): k for k, v in ARCHETYPE_BASE_SLUGS.items()}

    if norm_raw in masc_by_norm:
        return masc_by_norm[norm_raw]
    if norm_raw in fem_by_norm:
        return fem_by_norm[norm_raw]
    if norm_raw in slug_by_norm:
        return slug_by_norm[norm_raw]

    for token, masc in masc_by_norm.items():
        if token and (norm_raw.startswith(token) or token.startswith(norm_raw)):
            return masc
    for token, masc in fem_by_norm.items():
        if token and (norm_raw.startswith(token) or token.startswith(norm_raw)):
            return masc
    for token, masc in slug_by_norm.items():
        if token and (norm_raw.startswith(token) or token.startswith(norm_raw)):
            return masc
    return raw


def _template_candidates_for_export(
    show_supplement: bool,
    short_report: bool,
    gender_code: str = "M",
) -> list[str]:
    """Zwraca kandydatów szablonów w kolejności preferencji (najpierw płeć badanej osoby)."""

    def _ordered_set(is_female: bool) -> list[str]:
        full = TEMPLATE_PATH_FEMALE if is_female else TEMPLATE_PATH
        full_nosupp = TEMPLATE_PATH_NOSUPP_FEMALE if is_female else TEMPLATE_PATH_NOSUPP
        short = TEMPLATE_PATH_SHORT_FEMALE if is_female else TEMPLATE_PATH_SHORT
        short_nosupp = TEMPLATE_PATH_SHORT_NOSUPP_FEMALE if is_female else TEMPLATE_PATH_SHORT_NOSUPP

        if short_report:
            return [
                short if show_supplement else short_nosupp,
                short_nosupp if show_supplement else short,
                full if show_supplement else full_nosupp,
                full_nosupp if show_supplement else full,
            ]
        return [
            full if show_supplement else full_nosupp,
            full_nosupp if show_supplement else full,
        ]

    is_female = normalize_gender(gender_code) == "K"
    merged = _ordered_set(is_female) + _ordered_set(not is_female)
    out: list[str] = []
    seen: set[str] = set()
    for item in merged:
        if item and item not in seen:
            seen.add(item)
            out.append(item)
    return out


def _default_template_for_export(show_supplement: bool, gender_code: str = "M") -> str:
    candidates = _template_candidates_for_export(
        show_supplement=show_supplement,
        short_report=False,
        gender_code=gender_code,
    )
    return candidates[0] if candidates else (TEMPLATE_PATH if show_supplement else TEMPLATE_PATH_NOSUPP)


_FEMINIZE_PHRASES_RUNTIME: list[tuple[str, str]] = [
    ("on/ona", "ona"),
    ("On/ona", "Ona"),
    ("Jego", "Jej"),
    ("jego", "jej"),
    ("Niego", "Niej"),
    ("niego", "niej"),
    ("Nim", "Nią"),
    ("nim", "nią"),
    (" on ", " ona "),
    (" On ", " Ona "),
]

_FEMINIZE_WORD_MAP_RUNTIME: dict[str, str] = {
    "Władca": "Władczyni",
    "Władcy": "Władczyni",
    "Bohater": "Bohaterka",
    "Bohatera": "Bohaterki",
    "Mędrzec": "Mędrczyni",
    "Mędrca": "Mędrczyni",
    "Opiekun": "Opiekunka",
    "Opiekuna": "Opiekunki",
    "Kochanek": "Kochanka",
    "Błazen": "Komiczka",
    "Błazna": "Komiczki",
    "Twórca": "Twórczyni",
    "Twórcy": "Twórczyni",
    "Odkrywca": "Odkrywczyni",
    "Odkrywcy": "Odkrywczyni",
    "Czarodziej": "Czarodziejka",
    "Czarodzieja": "Czarodziejki",
    "Towarzysz": "Towarzyszka",
    "Towarzysza": "Towarzyszki",
    "Niewinny": "Niewinna",
    "Niewinnego": "Niewinnej",
    "Buntownik": "Buntowniczka",
    "Buntownika": "Buntowniczki",
}

_FEMINIZE_WORD_PATTERNS_RUNTIME: tuple[tuple[re.Pattern[str], str], ...] = tuple(
    (
        re.compile(rf"(?<!\w){re.escape(old)}(?!\w)"),
        new,
    )
    for old, new in sorted(_FEMINIZE_WORD_MAP_RUNTIME.items(), key=lambda kv: (-len(kv[0]), kv[0]))
)


def _feminize_text_runtime(text: str) -> str:
    out = str(text or "")
    if not out:
        return out
    for old, new in _FEMINIZE_PHRASES_RUNTIME:
        out = out.replace(old, new)
    for pat, repl in _FEMINIZE_WORD_PATTERNS_RUNTIME:
        out = pat.sub(repl, out)
    return out


def _feminize_payload_runtime(value):
    if isinstance(value, str):
        return _feminize_text_runtime(value)
    if isinstance(value, list):
        return [_feminize_payload_runtime(v) for v in value]
    if isinstance(value, tuple):
        return tuple(_feminize_payload_runtime(v) for v in value)
    if isinstance(value, dict):
        return {k: _feminize_payload_runtime(v) for k, v in value.items()}
    return value


def get_archetype_payload_for_gender(archetype_name: str, gender_code: str = "M") -> dict:
    """
    Pobiera dane archetypu preferując wersję zgodną z płcią:
    - dla 'K' najpierw nazwa żeńska, potem fallback do męskiej.
    - dla 'M' odwrotnie.
    """
    base_name = base_masc_from_any(archetype_name)
    gender = normalize_gender(gender_code)
    display_name = display_name_for_gender(base_name, gender)

    if gender == "K":
        # Priorytet: gotowy żeński opis z DOCX (bez dodatkowej, destrukcyjnej feminizacji runtime).
        payload_female = archetype_extended.get(display_name)
        if isinstance(payload_female, dict) and payload_female:
            out = dict(payload_female)
            out["name"] = display_name
            return out

        # Fallback bezpieczeństwa: jeśli żeńskiego pliku brak, feminizuj męski payload.
        payload_male = archetype_extended.get(base_name) or {}
        out = _feminize_payload_runtime(dict(payload_male))
        out["name"] = display_name
        return out

    payload = archetype_extended.get(base_name) or archetype_extended.get(display_name) or {}
    out = dict(payload)
    out["name"] = base_name or out.get("name", "")
    return out


def _dedupe_hex_palette(values) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for raw in (values or []):
        code = str(raw or "").strip().upper()
        if not code:
            continue
        if not code.startswith("#"):
            code = f"#{code}"
        if not re.match(r"^#[0-9A-F]{6}$", code):
            continue
        if code in seen:
            continue
        seen.add(code)
        out.append(code)
    return out


def _enforce_required_palettes(arche_map: dict[str, dict]) -> None:
    for arche_name, required_palette in ARCHETYPE_REQUIRED_PALETTES.items():
        required = _dedupe_hex_palette(required_palette)
        if not required:
            continue

        for code in required:
            COLOR_NAME_MAP.setdefault(code, f"kolor-{code[1:].lower()}")

        payload = arche_map.get(arche_name)
        if not payload:
            continue

        existing = _dedupe_hex_palette(payload.get("color_palette", []))
        merged = list(existing)
        for code in required:
            if code not in merged:
                merged.append(code)

        payload["color_palette"] = merged

        metric_rows = payload.get("metric_rows") or []
        for row in metric_rows:
            if str(row.get("label", "")).strip().casefold() == "paleta kolorów (hex)":
                row["kind"] = "colors"
                row["value"] = merged
                break


# <<<--- TUTAJ WKLEJ własne archetype_extended = {...}
archetype_extended = load_archetype_extended(
    os.path.join(os.path.dirname(__file__), 'opisy_archetypow')
)
_enforce_required_palettes(archetype_extended)
# --- KONIEC archetype_extended ---

from pathlib import Path
from PIL import Image

# (opcjonalnie) osadzenie własnych fontów w widoku Streamlit
import base64, pathlib
def _font_face_css(font_path, family_name, weight="normal", style="normal"):
    try:
        data = pathlib.Path(font_path).read_bytes()
        b64 = base64.b64encode(data).decode("ascii")
        return f"""
        @font-face {{
            font-family: '{family_name}';
            src: url(data:font/ttf;base64,{b64}) format('truetype');
            font-weight: {weight};
            font-style: {style};
            font-display: swap;
        }}
        """
    except Exception:
        return ""
css_ff = ""
css_ff += _font_face_css("fonts/ArialNovaCond.ttf", "Arial Nova Cond", "600")
css_ff += _font_face_css("fonts/ArialNovaCondLight.ttf", "Arial Nova Cond", "300")
css_ff += _font_face_css("fonts/RobotoCondensed-Regular.ttf", "Roboto Condensed", "400")
css_ff += _font_face_css("fonts/RobotoCondensed-Light.ttf", "Roboto Condensed", "300")

# >>> BEGIN CSS INJECTOR (wklej po zbudowaniu css_ff) >>>
import streamlit.components.v1 as components  # masz już import wyżej, ale zostaw – jest idempotentne

# 1) Złóż jeden łańcuch CSS: @font-face (css_ff) + Twoje reguły globalne
GLOBAL_CSS = (css_ff or "") + """
/* delikatna, szara linia */
.soft-hr{ height:1px; border:none; background:#e5e7eb; margin:28px 0 26px 0; }

/* jednolity nagłówek sekcji */
.ap-h2{
  font-family: "Segoe UI", system-ui, -apple-system, Arial, sans-serif;
  font-weight: 600 !important;
  font-size: 1.21rem;
  line-height: 1.3;
  letter-spacing: 0;
  color:#1f2937;
  margin: 10px 0 25px 0;
  white-space: normal;
  word-break: keep-all;
}
.ap-h2.center{ text-align:center; }

/* wymuszone nagłówki sekcji raportu */
.ap-heading-force{
  font-family: "Segoe UI", system-ui, -apple-system, Arial, sans-serif !important;
  font-weight: 640 !important;
  font-size: 1.34rem !important;
  line-height: 1.27 !important;
  letter-spacing: 0 !important;
  color:#1f2937 !important;
  margin: 6px 0 10px 0 !important;
}
@media (max-width: 1280px){
  .ap-heading-force{
    font-size: 1.25rem !important;
  }
}
.ap-heading-force.ap-heading-center{ text-align:center !important; }
.ap-heading-force.ap-heading-left{ text-align:left !important; }

/* tytuły sekcji */
.section-title{
  font-family: "Segoe UI", system-ui, -apple-system, Arial, sans-serif;
  font-weight: 530 !important;
  font-size: 1.22em;
  margin: 15px 0 25px 0;
  line-height: 1.15;
  color:#182433;
}

.section-title--padTop{ margin-top:0px !important; }
.mt-28{ margin-top:0px !important; }
.section-title--blue{ color:#1a93e3; }

/* przyciski skoków */
.jump-btns{ display:flex; flex-wrap:wrap; gap:8px; margin:8px 0 16px 0; }
.jump-btn{
  display:inline-block; padding:8px 14px; border-radius:10px; text-decoration:none;
  border:1px solid #1a93e3; color:#1a93e3; font-weight:600; font-size:0.95em;
  background:#f0f8ff;
}
.jump-btn:hover{ background:#e6f3ff; }
:target{ scroll-margin-top: 90px; }

/* selectbox – ciaśniej pod spodem */
div[data-testid="stSelectbox"]{
  margin-bottom: 4px !important;   /* ← tu ściskasz odstęp (np. 4–10 px) */
}
/* zbicie odstępu pod selectem */
div[data-testid="stSelectbox"] { padding-bottom: 0 !important; margin-bottom: 6px !important; }
div[data-testid="stSelectbox"] label { margin-bottom: 2px !important; }
/* usuń dodatkowy margines pierwszego elementu po selekcie (czasem Streamlit dodaje „poduchę”) */
div[data-testid="stSelectbox"] + div { margin-top: 0 !important; }

div[data-testid="stSelectbox"] > div{
  border:1.5px solid #1a93e3 !important;
  border-radius:10px !important;
  box-shadow: 0 0 0 1px #1a93e333 inset;
}
div[data-testid="stSelectbox"] label{
  color:#1a93e3 !important;
  font-weight:700 !important;
}

/* Streamlitowe nagłówki (w tym st.subheader) */
.stHeading h1, .stHeading h2, .stHeading h3,
h1, h2, h3,
.stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
  font-family: "Segoe UI", system-ui, -apple-system, Arial, sans-serif !important;
  font-weight: 600 !important;
  font-size: 1.23rem !important;
  line-height: 1.3 !important;
  color:var(--ap-heading-color, #1f2937) !important;
  margin: 10px 0 25px 0 !important;
  letter-spacing: 0 !important;
}

/* domyślna rodzina dla body */
body { font-family:'Roboto','Segoe UI','Arial',sans-serif; }
"""

def inject_global_css(css_text: str, style_id: str = "ap-global-css"):
    # Bez iframa z height=0 (potrafi generować artefakty typu "0" w UI).
    st.markdown(f"<style id='{style_id}'>{css_text}</style>", unsafe_allow_html=True)

# 2) Wołaj na każdym rerunie – jest szybkie i bezpieczne
inject_global_css(GLOBAL_CSS)
# <<< END CSS INJECTOR <<<


def ap_section_heading(
    title: str,
    center: bool = False,
    margin_bottom_px: int = 8,
    margin_top_px: int = 6,
    shift_x_px: int = 0,
) -> str:
    align_class = "ap-heading-center" if center else "ap-heading-left"
    align = "center" if center else "left"
    return (
        f"<div class='ap-heading-force {align_class}' "
        f"style='font-family:\"Segoe UI\",system-ui,-apple-system,Arial,sans-serif !important;"
        f"font-weight:640 !important;font-size:1.34rem !important;line-height:1.27 !important;"
        f"color:var(--ap-heading-color, #1f2937) !important;letter-spacing:0 !important;"
        f"text-align:{align} !important;transform:translateX({int(shift_x_px)}px) !important;"
        f"margin:{int(margin_top_px)}px 0 {int(margin_bottom_px)}px 0 !important;'>"
        f"{html.escape(str(title))}</div>"
    )


ARCHE_NAME_TO_IDX = {n.lower(): i for i, n in enumerate(ARCHE_NAMES_ORDER)}


_WHEEL_VALUE_ORDER = [
    "Niewinny",
    "Mędrzec",
    "Odkrywca",
    "Buntownik",
    "Czarodziej",
    "Bohater",
    "Kochanek",
    "Błazen",
    "Towarzysz",
    "Opiekun",
    "Władca",
    "Twórca",
]


@lru_cache(maxsize=64)
def _load_wheel_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    assets_dir = Path(__file__).with_name("assets")
    font_candidates = (
        [assets_dir / "fonts" / "DejaVuSans-Bold.ttf", assets_dir / "DejaVuSans-Bold.ttf"]
        if bold
        else [assets_dir / "fonts" / "DejaVuSans.ttf", assets_dir / "DejaVuSans.ttf"]
    )
    for font_path in font_candidates:
        if font_path.exists():
            try:
                return ImageFont.truetype(str(font_path), size=size)
            except Exception:
                continue
    return ImageFont.load_default()


def _segment_label_text(value: str) -> str:
    manual = {
        "Przejrzystość": "PRZEJRZY-\nSTOŚĆ",
        "Współpraca": "WSPÓŁ-\nPRACA",
    }
    return manual.get(value, value.upper())


def _center_field_texts() -> dict[str, str]:
    return {
        key: str(value).replace(" i ", "\ni ")
        for key, value in FINAL_VALUES_WHEEL_CENTRAL_FIELDS.items()
    }


def _draw_rotated_capsule(
    base: Image.Image,
    text: str,
    center_xy: tuple[float, float],
    angle_deg: float,
    box_size: tuple[int, int],
    font: ImageFont.FreeTypeFont | ImageFont.ImageFont,
) -> None:
    box_w, box_h = box_size
    capsule = Image.new("RGBA", (box_w, box_h), (0, 0, 0, 0))
    drawer = ImageDraw.Draw(capsule, "RGBA")
    radius = max(6, int(box_h * 0.45))
    drawer.rounded_rectangle(
        (0, 0, box_w - 1, box_h - 1),
        radius=radius,
        fill=(235, 238, 241, 248),
        outline=(208, 214, 220, 245),
        width=max(1, box_h // 18),
    )
    drawer.text((box_w / 2, box_h / 2), text, anchor="mm", fill=(44, 51, 58, 255), font=font)
    rotated = capsule.rotate(angle_deg, expand=True, resample=Image.BICUBIC)
    px = int(center_xy[0] - rotated.width / 2)
    py = int(center_xy[1] - rotated.height / 2)
    base.alpha_composite(rotated, (px, py))


def _apply_values_wheel_public_overlay(base: Image.Image) -> Image.Image:
    img = base.copy().convert("RGBA")
    w, h = img.size
    m = float(min(w, h))
    cx = w / 2.0
    cy = h / 2.0

    # 1) Segmenty wartości: czyścimy stare etykiety i rysujemy finalne.
    segment_layer = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    segment_draw = ImageDraw.Draw(segment_layer, "RGBA")
    r_outer = m * 0.36
    r_inner = m * 0.20
    r_text = m * 0.272
    for idx, archetype_name in enumerate(_WHEEL_VALUE_ORDER):
        center_deg = 75 - idx * 30
        theta1 = center_deg - 14
        theta2 = center_deg + 14
        ang = math.radians(center_deg)
        sample_x = int(max(0, min(w - 1, cx + m * 0.312 * math.cos(ang))))
        sample_y = int(max(0, min(h - 1, cy - m * 0.312 * math.sin(ang))))
        fill_rgba = img.getpixel((sample_x, sample_y))

        segment_draw.pieslice(
            [cx - r_outer, cy - r_outer, cx + r_outer, cy + r_outer],
            start=theta1,
            end=theta2,
            fill=fill_rgba,
        )
        segment_draw.pieslice(
            [cx - r_inner, cy - r_inner, cx + r_inner, cy + r_inner],
            start=theta1,
            end=theta2,
            fill=(0, 0, 0, 0),
        )
    img.alpha_composite(segment_layer)

    seg_font = _load_wheel_font(int(m * 0.036), bold=False)
    seg_draw = ImageDraw.Draw(img, "RGBA")
    for idx, archetype_name in enumerate(_WHEEL_VALUE_ORDER):
        value = ARCHETYPE_PUBLIC_VALUES[archetype_name]
        center_deg = 75 - idx * 30
        ang = math.radians(center_deg)
        tx = cx + r_text * math.cos(ang)
        ty = cy - r_text * math.sin(ang)
        seg_draw.multiline_text(
            (tx, ty),
            _segment_label_text(value),
            anchor="mm",
            align="center",
            spacing=max(2, int(m * 0.0025)),
            fill=(22, 24, 28, 248),
            font=seg_font,
        )

    # 2) Pola centralne: finalny zestaw 4 pól.
    center_layer = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    center_draw = ImageDraw.Draw(center_layer, "RGBA")
    center_text = _center_field_texts()
    center_positions = {
        "upper_left": (cx - m * 0.115, cy - m * 0.076),
        "upper_right": (cx + m * 0.115, cy - m * 0.076),
        "lower_left": (cx - m * 0.115, cy + m * 0.076),
        "lower_right": (cx + m * 0.115, cy + m * 0.076),
    }
    half_w = m * 0.108
    half_h = m * 0.064
    center_font = _load_wheel_font(int(m * 0.031), bold=False)
    for key, (px, py) in center_positions.items():
        sample_x = int(max(0, min(w - 1, px)))
        sample_y = int(max(0, min(h - 1, py)))
        fill_rgba = img.getpixel((sample_x, sample_y))
        center_draw.rounded_rectangle(
            (px - half_w, py - half_h, px + half_w, py + half_h),
            radius=max(8, int(m * 0.019)),
            fill=(fill_rgba[0], fill_rgba[1], fill_rgba[2], 244),
        )
        center_draw.multiline_text(
            (px, py),
            center_text[key],
            anchor="mm",
            align="center",
            spacing=max(2, int(m * 0.0022)),
            fill=(22, 24, 28, 248),
            font=center_font,
        )
    img.alpha_composite(center_layer)

    # 3) Szare łuki/grupy: finalne nazwy.
    capsule_font = _load_wheel_font(int(m * 0.028), bold=False)
    cap_w = int(w * 0.29)
    cap_h = int(h * 0.065)
    _draw_rotated_capsule(
        img,
        FINAL_VALUES_WHEEL_ARC_LABELS["upper_left"],
        center_xy=(w * 0.12, h * 0.16),
        angle_deg=-62,
        box_size=(cap_w, cap_h),
        font=capsule_font,
    )
    _draw_rotated_capsule(
        img,
        FINAL_VALUES_WHEEL_ARC_LABELS["upper_right"],
        center_xy=(w * 0.88, h * 0.16),
        angle_deg=62,
        box_size=(cap_w, cap_h),
        font=capsule_font,
    )
    _draw_rotated_capsule(
        img,
        FINAL_VALUES_WHEEL_ARC_LABELS["lower_left"],
        center_xy=(w * 0.12, h * 0.84),
        angle_deg=62,
        box_size=(cap_w, cap_h),
        font=capsule_font,
    )
    _draw_rotated_capsule(
        img,
        FINAL_VALUES_WHEEL_ARC_LABELS["lower_right"],
        center_xy=(w * 0.88, h * 0.84),
        angle_deg=-62,
        box_size=(cap_w, cap_h),
        font=capsule_font,
    )

    return img


@st.cache_data
def load_base_arche_img(gender_code: str = "M", dark_mode: bool = False):
    assets_dir = Path(__file__).with_name("assets")
    g = normalize_gender(gender_code)
    if g == "K":
        candidates = (
            ["archetype_wheel_female_dark.png", "archetype_wheel_female.png", "archetype_wheel.png"]
            if dark_mode
            else ["archetype_wheel_female.png", "archetype_wheel_female_dark.png", "archetype_wheel.png"]
        )
    else:
        candidates = (
            ["archetype_wheel_male_dark.png", "archetype_wheel_male.png", "archetype_wheel.png"]
            if dark_mode
            else ["archetype_wheel_male.png", "archetype_wheel_male_dark.png", "archetype_wheel.png"]
        )
    for name in candidates:
        p = assets_dir.joinpath(name)
        if p.exists():
            img = Image.open(p).convert("RGBA")
            # Assets koła mają już finalne, ręcznie dopracowane etykiety.
            # Nie nakładamy runtime-overlay, żeby nie psuć layoutu.
            return img
    raise FileNotFoundError(f"Brak pliku koła archetypów. Szukano: {candidates}")


def mask_for(idx, color, gender_code: str = "M", dark_mode: bool = False):
    base = load_base_arche_img(gender_code, dark_mode=dark_mode)
    w, h = base.size
    cx, cy = w//2, h//2
    # Klin wychodzi ze środka, ale krócej niż pełna przekątna płótna,
    # żeby nie „zalewał” całego koła i zostawiał czytelne marginesy.
    r_outer = int(min(w, h) * 0.54)
    mask = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    draw = ImageDraw.Draw(mask, "RGBA")
    start = -90 + idx*30
    end = start + 30
    # Podświetlamy pełny klin od środka koła.
    draw.pieslice([cx-r_outer, cy-r_outer, cx+r_outer, cy+r_outer], start, end, fill=color)
    return mask

def compose_archetype_highlight(
    idx_main,
    idx_aux=None,
    idx_supplement=None,
    gender_code: str = "M",
    dark_mode: bool = False,
):
    base = load_base_arche_img(gender_code, dark_mode=dark_mode).copy()

    # Najpierw poboczny (żeby nakryło go potem żółte/czerwone, jeśli overlap)
    if idx_supplement is not None and idx_supplement not in [idx_main, idx_aux] and idx_supplement < 12:
        mask_supplement = mask_for(
            idx_supplement,
            (64, 185, 0, 140),
            gender_code=gender_code,
            dark_mode=dark_mode,
        )  # zielony półtransparentny
        base.alpha_composite(mask_supplement)

    # Potem wspierający
    if idx_aux is not None and idx_aux != idx_main and idx_aux < 12:
        mask_aux = mask_for(
            idx_aux,
            (255, 210, 47, 140),
            gender_code=gender_code,
            dark_mode=dark_mode,
        )  # żółty
        base.alpha_composite(mask_aux)

    # Na końcu główny (przykrywa wszystko)
    if idx_main is not None:
        mask_main = mask_for(
            idx_main,
            (255, 0, 0, 140),
            gender_code=gender_code,
            dark_mode=dark_mode,
        )  # czerwony
        base.alpha_composite(mask_main)

    return base

# ---- KOŁO „osi potrzeb” (inny porządek, start o 12:00) ----
KOLO_NAMES_ORDER = [
    "Buntownik","Błazen","Kochanek","Opiekun","Towarzysz","Niewinny",
    "Władca","Mędrzec","Czarodziej","Bohater","Twórca","Odkrywca"
]
ANGLE_OFFSET_DEG = -15  # przesunięcie o 2.5 min w lewo (2.5 × 6° = 15°)

@st.cache_data
def load_axes_wheel_img(gender_code: str = "M", dark_mode: bool = False):
    base_dir = Path(__file__).with_name("assets")
    g = normalize_gender(gender_code)
    if g == "K":
        if dark_mode:
            candidates = [
                ("png", "archetypy_kolo_female_dark.png"),
                ("png", "archetypy_kolo_female_dark.png_dark.png"),
                ("svg", "archetypy_kolo2_dark.svg"),
                ("png", "archetypy_kolo_dark.png"),
                ("png", "archetypy_kolo_female.png"),
                ("svg", "archetypy_kolo2.svg"),
                ("png", "archetypy_kolo.png"),
            ]
        else:
            candidates = [
                ("png", "archetypy_kolo_female.png"),
                ("png", "archetypy_kolo_female_dark.png"),
                ("png", "archetypy_kolo_female_dark.png_dark.png"),
                ("svg", "archetypy_kolo2.svg"),
                ("png", "archetypy_kolo.png"),
            ]
    else:
        if dark_mode:
            candidates = [
                ("png", "archetypy_kolo_dark.png"),
                ("svg", "archetypy_kolo2_male_dark.svg"),
                ("svg", "archetypy_kolo2_dark.svg"),
                ("png", "archetypy_kolo.png"),
                ("svg", "archetypy_kolo2_male.svg"),
                ("svg", "archetypy_kolo2.svg"),
            ]
        else:
            candidates = [
                ("png", "archetypy_kolo.png"),
                ("svg", "archetypy_kolo2_male.svg"),
                ("svg", "archetypy_kolo2.svg"),
                ("png", "archetypy_kolo_dark.png"),
            ]
    for ftype, name in candidates:
        p = base_dir / name
        if not p.exists():
            continue
        if ftype == "png":
            return Image.open(p).convert("RGBA")
        import cairosvg
        from io import BytesIO
        buf = BytesIO(cairosvg.svg2png(url=str(p)))
        return Image.open(buf).convert("RGBA")
    raise FileNotFoundError(f"Brak pliku koła osi. Szukano: {candidates}")

def _mask_pie_ring(base: Image.Image, idx: int, rgba,
                   r_out_frac=0.39, r_in_frac=0.10):
    """
    Maluje półprzezroczysty 30° sektor TYLKO na pierścieniu (donut), nie na całym płótnie.
    r_out_frac / r_in_frac – promień zewnętrzny/wewnętrzny w ułamku szerokości obrazu
    (w razie potrzeby możesz lekko podregulować, np. 0.44 / 0.18).
    """
    if idx is None:
        return base
    w, h = base.size
    cx, cy = w//2, h//2
    R = int(min(w, h) * r_out_frac)
    r = int(min(w, h) * r_in_frac)

    start = -90 + ANGLE_OFFSET_DEG + idx * 30
    end   = start + 30

    # rysujemy na osobnej warstwie z alfą → potem alpha_composite
    layer = Image.new("RGBA", (w, h), (0,0,0,0))
    d = ImageDraw.Draw(layer, "RGBA")
    d.pieslice([cx-R, cy-R, cx+R, cy+R], start, end, fill=rgba)
    # wycinamy środek (donut)
    d.ellipse([cx-r, cy-r, cx+r, cy+r], fill=(0,0,0,0))
    base.alpha_composite(layer)
    return base

def compose_axes_wheel_highlight(
    main_name,
    aux_name=None,
    supp_name=None,
    gender_code: str = "M",
    dark_mode: bool = False,
) -> Image.Image:
    """Koło z podświetleniem: zielony (poboczny), żółty (wspierający), czerwony (główny)."""
    img = load_axes_wheel_img(gender_code, dark_mode=dark_mode).copy()

    def idx(n):
        try:
            return KOLO_NAMES_ORDER.index(n)
        except:
            return None

    # kolejność: poboczny → wspierający → główny (nakładanie)
    _mask_pie_ring(img, idx(supp_name), (64,185,0,110))
    _mask_pie_ring(img, idx(aux_name),  (255,210,47,110))
    _mask_pie_ring(img, idx(main_name), (255,0,0,110))
    return img


SEGMENT_PROFILE_ITEMS = [
    ("Buntownik", ARCHETYPE_PUBLIC_VALUES["Buntownik"], "buntownik.png", "#c62828"),
    ("Błazen", ARCHETYPE_PUBLIC_VALUES["Błazen"], "blazen.png", "#ef5350"),
    ("Kochanek", ARCHETYPE_PUBLIC_VALUES["Kochanek"], "kochanek.png", "#90caf9"),
    ("Opiekun", ARCHETYPE_PUBLIC_VALUES["Opiekun"], "opiekun.png", "#42a5f5"),
    ("Towarzysz", ARCHETYPE_PUBLIC_VALUES["Towarzysz"], "towarzysz.png", "#1565c0"),
    ("Niewinny", ARCHETYPE_PUBLIC_VALUES["Niewinny"], "niewinny.png", "#81c784"),
    ("Władca", ARCHETYPE_PUBLIC_VALUES["Władca"], "wladca.png", "#43a047"),
    ("Mędrzec", ARCHETYPE_PUBLIC_VALUES["Mędrzec"], "medrzec.png", "#1b5e20"),
    ("Czarodziej", ARCHETYPE_PUBLIC_VALUES["Czarodziej"], "czarodziej.png", "#b39ddb"),
    ("Bohater", ARCHETYPE_PUBLIC_VALUES["Bohater"], "bohater.png", "#7e57c2"),
    ("Twórca", ARCHETYPE_PUBLIC_VALUES["Twórca"], "tworca.png", "#5e35b1"),
    ("Odkrywca", ARCHETYPE_PUBLIC_VALUES["Odkrywca"], "odkrywca.png", "#8e0000"),
]


def _plot_segment_profile_wheel_from_scores(
    outpath: Path,
    mean_scores: dict[str, float],
    label_mode: str = "arche",
    gender_code: str = "M",
    dark_mode: bool = False,
) -> None:
    from matplotlib.patches import Wedge, Circle
    import math

    values = np.asarray([float(mean_scores.get(name, 0.0)) for name, *_ in SEGMENT_PROFILE_ITEMS], dtype=float)
    values = np.where(np.isfinite(values), np.clip(values, 0.0, 100.0), 0.0)

    def _load_icon(path: Path):
        img = Image.open(path).convert("RGBA")
        alpha = img.getchannel("A")
        bbox = alpha.getbbox()
        if bbox:
            img = img.crop(bbox)
        if dark_mode:
            arr = np.asarray(img).copy()
            mask = arr[..., 3] > 0
            arr[..., 0][mask] = 244
            arr[..., 1][mask] = 249
            arr[..., 2][mask] = 255
            img = Image.fromarray(arr, mode="RGBA")
        return np.asarray(img)

    def _draw_text_on_arc(ax, text: str, center_deg: float, radius: float) -> None:
        text = str(text or "").upper()
        if not text:
            return
        n = len(text)
        span = min(24, max(11.5, 1.7 * n))
        center_rad = math.radians(center_deg)
        upper = math.sin(center_rad) >= 0
        if n == 1:
            angles = [center_deg]
        else:
            if upper:
                start = center_deg + span / 2
                end = center_deg - span / 2
            else:
                start = center_deg - span / 2
                end = center_deg + span / 2
            angles = np.linspace(start, end, n)
        for ch, ang in zip(text, angles):
            if ch == " ":
                continue
            ang_rad = math.radians(float(ang))
            rot = (ang - 90) if upper else (ang + 90)
            ax.text(
                radius * math.cos(ang_rad),
                radius * math.sin(ang_rad),
                ch,
                ha="center",
                va="center",
                rotation=rot,
                rotation_mode="anchor",
                fontsize=11.4,
                fontweight="bold",
                color="#EDF4FF" if dark_mode else "#363636",
                zorder=4,
            )

    fig, ax = plt.subplots(figsize=(10.8, 10.8), dpi=220)
    ax.set_aspect("equal")
    ax.axis("off")
    bg = (1, 1, 1, 0)
    fig.patch.set_alpha(0.0)
    ax.set_facecolor("none")

    r_hole = 0.115
    r_data_outer = 0.78
    n_cells = 10
    dr = (r_data_outer - r_hole) / n_cells
    r_label_inner = 0.80
    r_label_outer = 0.92
    r_accent_inner = 0.95
    r_accent_outer = 0.99
    edgecolor = "#7C93AF" if dark_mode else "#d4d4d4"
    ring_grid_color = "#6E84A1" if dark_mode else "#dfdfdf"
    radial_line_color = "#7B92AF" if dark_mode else "#d7d7d7"
    score_color = "#F4F8FF" if dark_mode else "#2f2f2f"
    score_bbox_fc = (0.07, 0.12, 0.20, 0.88) if dark_mode else "white"
    hole_edge = "#8EA5C1" if dark_mode else "#d0d0d0"
    cross_color = "#90A6C1" if dark_mode else "#e1e1e1"

    icon_dir = Path(__file__).with_name("ikony")
    icon_cache: dict[str, np.ndarray] = {}
    for arch, _value, icon_file, _color in SEGMENT_PROFILE_ITEMS:
        icon_path = icon_dir / icon_file
        if icon_path.exists():
            try:
                icon_cache[arch] = _load_icon(icon_path)
            except Exception:
                pass

    for k in range(n_cells + 1):
        r = r_hole + k * dr
        ax.add_patch(Circle((0, 0), r, facecolor="none", edgecolor=ring_grid_color, linewidth=0.7, zorder=0))

    for i in range(12):
        center = 75 - i * 30
        edge_ang = math.radians(center - 15)
        ax.plot(
            [r_hole * math.cos(edge_ang), r_accent_outer * math.cos(edge_ang)],
            [r_hole * math.sin(edge_ang), r_accent_outer * math.sin(edge_ang)],
            color=radial_line_color,
            lw=0.7,
            zorder=0,
        )

    mode_norm = str(label_mode or "arche").strip().lower()

    label_gender = normalize_gender(gender_code)

    for i, (arch, value_label, _icon_file, color) in enumerate(SEGMENT_PROFILE_ITEMS):
        center = 75 - i * 30
        theta1 = center - 15
        theta2 = center + 15
        p = float(values[i])
        fill_outer = r_hole + (p / 100.0) * (r_data_outer - r_hole)

        for k in range(n_cells):
            r_outer = r_hole + (k + 1) * dr
            ax.add_patch(
                Wedge((0, 0), r_outer, theta1, theta2, width=dr, facecolor=bg, edgecolor=edgecolor, linewidth=0.8, zorder=1)
            )
        if fill_outer > r_hole + 1e-9:
            ax.add_patch(
                Wedge(
                    (0, 0),
                    fill_outer,
                    theta1,
                    theta2,
                    width=(fill_outer - r_hole),
                    facecolor=color,
                    edgecolor="none",
                    linewidth=0.0,
                    alpha=0.90,
                    zorder=2,
                )
            )
        for k in range(n_cells):
            r_outer = r_hole + (k + 1) * dr
            ax.add_patch(
                Wedge((0, 0), r_outer, theta1, theta2, width=dr, facecolor="none", edgecolor=edgecolor, linewidth=0.8, zorder=3)
            )
        ax.add_patch(
            Wedge(
                (0, 0),
                r_label_outer,
                theta1,
                theta2,
                width=(r_label_outer - r_label_inner),
                facecolor=bg,
                edgecolor=edgecolor,
                linewidth=0.8,
                zorder=3.2,
            )
        )
        ax.add_patch(
            Wedge(
                (0, 0),
                r_accent_outer,
                theta1,
                theta2,
                width=(r_accent_outer - r_accent_inner),
                facecolor=color,
                edgecolor=bg,
                linewidth=1.1,
                zorder=3.4,
            )
        )

        ring_label = (
            str(value_label)
            if mode_norm.startswith("val")
            else display_name_for_gender(str(arch), label_gender)
        )
        _draw_text_on_arc(ax=ax, text=ring_label, center_deg=center, radius=(r_label_inner + r_label_outer) / 2)

        ang = math.radians(center)
        if arch in icon_cache:
            ax.add_artist(
                AnnotationBbox(
                    OffsetImage(icon_cache[arch], zoom=0.26),
                    (1.14 * math.cos(ang), 1.14 * math.sin(ang)),
                    frameon=False,
                    zorder=5,
                )
            )

        filled_span = max(0.0, fill_outer - r_hole)
        label_r = (r_hole + filled_span * 0.52 + dr * 0.18) if p <= 25.0 else (r_hole + filled_span * 0.52)
        if filled_span <= 1e-9:
            label_r = r_hole + dr * 0.9
        ax.text(
            label_r * math.cos(ang),
            label_r * math.sin(ang),
            f"{p:.0f}",
            ha="center",
            va="center",
            fontsize=12.2,
            color=score_color,
            fontweight="bold",
            zorder=8,
            bbox=dict(
                boxstyle="round,pad=0.24,rounding_size=0.10",
                fc=score_bbox_fc,
                ec=color,
                lw=1.0,
                alpha=0.82 if dark_mode else 0.68,
            ),
        )

    ax.add_patch(Circle((0, 0), r_hole * 0.98, facecolor=bg, edgecolor=hole_edge, linewidth=1.2, zorder=10))
    for angle_deg in [0, 90, 180, 270]:
        ang = math.radians(angle_deg)
        ax.plot([0, 1.02 * math.cos(ang)], [0, 1.02 * math.sin(ang)], color=cross_color, lw=0.9, zorder=0)

    ax.set_xlim(-1.22, 1.22)
    ax.set_ylim(-1.22, 1.22)
    outpath = Path(outpath)
    outpath.parent.mkdir(parents=True, exist_ok=True)
    plt.tight_layout(pad=0.06)
    fig.savefig(outpath, dpi=220, bbox_inches="tight", pad_inches=0.10, transparent=True)
    plt.close(fig)


def make_segment_profile_wheel_png(
    mean_scores: dict[str, float],
    out_path: str = "segment_profile_wheel.png",
    label_mode: str = "arche",
    gender_code: str = "M",
    dark_mode: bool = False,
) -> str:
    _plot_segment_profile_wheel_from_scores(
        Path(out_path),
        mean_scores or {},
        label_mode=label_mode,
        gender_code=gender_code,
        dark_mode=dark_mode,
    )
    return out_path


@st.cache_data(ttl=30)
def load(study_id=None):
    import json
    try:
        engine = _sa_engine()
        base_sql = "SELECT created_at, answers, scores FROM public.responses"
        if study_id:
            query = sa.text(base_sql + " WHERE study_id = :sid ORDER BY created_at")
            with engine.begin() as conn:
                df = pd.read_sql(query, con=conn, params={"sid": study_id})
        else:
            query = sa.text(base_sql + " ORDER BY created_at")
            with engine.begin() as conn:
                df = pd.read_sql(query, con=conn)
        engine.dispose()

        def parse_answers(x):
            if isinstance(x, list):
                return [int(v) for v in x] if x and isinstance(x[0], (int, str)) else x
            if isinstance(x, str):
                for loader in (json.loads, ast.literal_eval):
                    try:
                        val = loader(x)
                        return [int(v) for v in val] if isinstance(val, list) else None
                    except Exception:
                        pass
                return None
            return None

        def parse_scores(x):
            if isinstance(x, dict):
                return x
            if isinstance(x, str):
                for loader in (json.loads, ast.literal_eval):
                    try:
                        val = loader(x)
                        return val if isinstance(val, dict) else {}
                    except Exception:
                        pass
                return {}
            return {}

        def include_in_report(raw_scores):
            if not isinstance(raw_scores, dict):
                return True
            quality = raw_scores.get("survey_quality")
            if not isinstance(quality, dict):
                return True
            flag = quality.get("include_in_report")
            if flag is None:
                return True
            if isinstance(flag, bool):
                return flag
            txt = str(flag).strip().lower()
            if txt in {"0", "false", "f", "no", "n", "off", "nie"}:
                return False
            if txt in {"1", "true", "t", "yes", "y", "on", "tak"}:
                return True
            return True

        if "answers" in df.columns:
            df["answers"] = df["answers"].apply(parse_answers)
        if "scores" in df.columns:
            df["scores"] = df["scores"].apply(parse_scores)
            df = df[df["scores"].apply(include_in_report)].copy()
        return df

    except Exception as e:
        st.warning(f"Błąd podczas ładowania danych: {e}")


def _norm_text_token(value: object) -> str:
    txt = str(value or "").strip().lower()
    repl = (
        ("ą", "a"),
        ("ć", "c"),
        ("ę", "e"),
        ("ł", "l"),
        ("ń", "n"),
        ("ó", "o"),
        ("ś", "s"),
        ("ż", "z"),
        ("ź", "z"),
    )
    for src, dst in repl:
        txt = txt.replace(src, dst)
    return re.sub(r"\s+", " ", txt).strip()


def _extract_personal_metry_payload(raw_scores: object) -> dict[str, str]:
    if not isinstance(raw_scores, dict):
        return {}
    metry = raw_scores.get("metryczka")
    candidate = metry if isinstance(metry, dict) else raw_scores
    out: dict[str, str] = {}
    for key, val in candidate.items():
        k = str(key or "").strip()
        if not k.startswith("M_"):
            continue
        v = str(val or "").strip()
        if v:
            out[k] = v
    return out


def _build_personal_metry_questions(study: dict) -> list[dict[str, object]]:
    cfg = normalize_personal_metryczka_config((study or {}).get("metryczka_config"))
    out: list[dict[str, object]] = []
    for q in list(cfg.get("questions") or []):
        if not isinstance(q, dict):
            continue
        qid = str(q.get("id") or "").strip()
        db_col = str(q.get("db_column") or qid).strip()
        if not db_col.startswith("M_"):
            continue
        options: list[dict[str, str]] = []
        for opt in list(q.get("options") or []):
            if not isinstance(opt, dict):
                continue
            label = str(opt.get("label") or "").strip()
            code = str(opt.get("code") or label).strip()
            if label and code:
                options.append(
                    {
                        "label": label,
                        "code": code,
                        "value_emoji": str(opt.get("value_emoji") or "").strip(),
                    }
                )
        if not options:
            continue
        out.append(
            {
                "id": qid or db_col,
                "db_column": db_col,
                "prompt": str(q.get("prompt") or db_col).strip(),
                "table_label": str(q.get("table_label") or q.get("prompt") or db_col).strip(),
                "variable_emoji": str(q.get("variable_emoji") or "").strip(),
                "options": options,
            }
        )
    return out


def _metry_value_label(question: dict[str, object], code: str) -> str:
    raw = str(code or "").strip()
    if not raw:
        return "brak danych"
    for opt in list(question.get("options") or []):
        if not isinstance(opt, dict):
            continue
        if str(opt.get("code") or "").strip() == raw:
            label = str(opt.get("label") or "").strip()
            return label or raw
    return raw


PERSONAL_METRY_VARIABLE_ICONS = {
    "M_PLEC": "👫",
    "M_WIEK": "⌛",
    "M_WYKSZT": "🎓",
    "M_ZAWOD": "💼",
    "M_MATERIAL": "💰",
    "M_OBSZAR": "📍",
}


PERSONAL_CORE_DEMO_META = {
    "M_PLEC": {
        "order": ["kobieta", "mężczyzna"],
        "value_emoji": {"kobieta": "👩", "mężczyzna": "👨"},
    },
    "M_WIEK": {
        "order": ["15-39", "40-59", "60+"],
        "value_emoji": {"15-39": "🧑", "40-59": "🧑‍💼", "60+": "🧓"},
    },
    "M_WYKSZT": {
        "order": ["podst./gim./zaw.", "średnie", "wyższe"],
        "value_emoji": {"podst./gim./zaw.": "🛠️", "średnie": "📘", "wyższe": "🎓"},
    },
    "M_ZAWOD": {
        "order": ["prac. umysłowy", "prac. fizyczny", "własna firma", "student/uczeń", "bezrobotny", "rencista/emeryt", "inna"],
        "value_emoji": {
            "prac. umysłowy": "🧠",
            "prac. fizyczny": "🛠️",
            "własna firma": "🏢",
            "student/uczeń": "🧑‍🎓",
            "bezrobotny": "🔎",
            "rencista/emeryt": "🌿",
            "inna": "🧩",
        },
    },
    "M_MATERIAL": {
        "order": ["bardzo dobra", "raczej dobra", "przeciętna", "raczej zła", "bardzo zła", "odmowa"],
        "value_emoji": {
            "bardzo dobra": "😄",
            "raczej dobra": "🙂",
            "przeciętna": "😐",
            "raczej zła": "🙁",
            "bardzo zła": "😟",
            "odmowa": "🤐",
        },
    },
}


def _personal_demo_canon_value(field: str, value: object) -> str:
    raw = str(value or "").strip()
    n = _norm_text_token(value)
    if not n:
        return "brak danych"
    if field == "M_PLEC":
        if n in {"1", "k", "kobieta"} or "kobiet" in n:
            return "kobieta"
        if n in {"2", "m", "mezczyzna"} or "mezczyzn" in n:
            return "mężczyzna"
    elif field == "M_WIEK":
        if re.search(r"15\D*39", n):
            return "15-39"
        if re.search(r"40\D*59", n):
            return "40-59"
        if "60" in n:
            return "60+"
    elif field == "M_WYKSZT":
        if "wyzsze" in n:
            return "wyższe"
        if "srednie" in n:
            return "średnie"
        if any(k in n for k in ("podstaw", "gimnaz", "zawod")):
            return "podst./gim./zaw."
    elif field == "M_ZAWOD":
        if "umysl" in n:
            return "prac. umysłowy"
        if "fizycz" in n:
            return "prac. fizyczny"
        if "wlasn" in n and "firm" in n:
            return "własna firma"
        if "student" in n or "uczen" in n:
            return "student/uczeń"
        if "bezrobot" in n:
            return "bezrobotny"
        if "renc" in n or "emery" in n:
            return "rencista/emeryt"
        if "inna" in n or "jaka" in n:
            return "inna"
    elif field == "M_MATERIAL":
        if "odmaw" in n:
            return "odmowa"
        if "bardzo dobrze" in n or "bardzo dobra" in n:
            return "bardzo dobra"
        if "raczej dobrze" in n or "raczej dobra" in n:
            return "raczej dobra"
        if "przeciet" in n or "srednio" in n:
            return "przeciętna"
        if "raczej zle" in n or "raczej zla" in n:
            return "raczej zła"
        if "bardzo zle" in n or "bardzo zla" in n or "ciezk" in n:
            return "bardzo zła"
    return raw or "brak danych"


def _personal_demo_code(db_column: str, value: object, table_label: str = "") -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    field = str(db_column or "").strip().upper()
    if field in PERSONAL_CORE_DEMO_META:
        canon = _personal_demo_canon_value(field, raw)
        return canon if canon and canon != "brak danych" else raw
    return raw


def _personal_metry_var_icon(db_column: str, table_label: str = "", preferred_icon: str = "") -> str:
    pref = str(preferred_icon or "").strip()
    if pref:
        return pref
    key = str(db_column or "").strip().upper()
    if key in PERSONAL_METRY_VARIABLE_ICONS:
        return PERSONAL_METRY_VARIABLE_ICONS[key]
    nk = _norm_text_token(f"{table_label} {key}")
    if "WIEK" in key:
        return "⌛"
    if "PLEC" in key:
        return "👫"
    if "WYKSZ" in key:
        return "🎓"
    if "ZAWOD" in key:
        return "💼"
    if "MATERIAL" in key:
        return "💰"
    if any(k in nk for k in ("obszar", "miejsce", "zamiesz", "lokaliz", "wies", "miasto")):
        return "📍"
    if any(k in nk for k in ("preferencj", "komitet", "wybor", "glos", "parti", "sejm")):
        return "🗳️"
    if "orientac" in nk:
        return "🧭"
    if any(k in nk for k in ("pogl", "politycz", "ideolog")):
        return "⚖️"
    return "📌"


def _personal_metry_cat_icon(
    db_column: str,
    table_label: str,
    code: str,
    preferred_icon: str = "",
    option_icon_map: dict[str, str] | None = None,
) -> str:
    pref = str(preferred_icon or "").strip()
    if pref:
        return pref
    if isinstance(option_icon_map, dict):
        map_key = str(code or "").strip()
        if map_key in option_icon_map:
            mapped_icon = str(option_icon_map.get(map_key) or "").strip()
            if mapped_icon in {"🏙️", "🏙", "🌆"}:
                return "🏬"
            return mapped_icon
    field = str(db_column or "").strip().upper()
    canon = _personal_demo_code(field, code, table_label)
    if str(canon).strip().lower() == "brak danych":
        return "❔"
    core_meta = PERSONAL_CORE_DEMO_META.get(field) or {}
    core_icons = core_meta.get("value_emoji") if isinstance(core_meta.get("value_emoji"), dict) else {}
    if canon in core_icons:
        return str(core_icons.get(canon) or "📌")
    nk_var = _norm_text_token(table_label)
    nk = _norm_text_token(canon)
    nk_sp = nk.replace("-", " ")
    if any(k in nk_var for k in ("obszar", "miejsce", "zamiesz", "lokaliz", "wies", "miasto")):
        if "miasto" in nk:
            return "🏬"
        if "wies" in nk:
            return "🌾"
        return "📍"
    if any(k in nk_var for k in ("orientac", "pogl", "politycz", "ideolog")):
        if "centro prawic" in nk_sp:
            return "↗️"
        if "prawic" in nk:
            return "➡️"
        if "centro lewic" in nk_sp:
            return "↖️"
        if "lewic" in nk:
            return "⬅️"
        if "centr" in nk:
            return "↔️"
        if "odmow" in nk:
            return "🤐"
        if "trudno" in nk:
            return "🤷"
        if "niewaz" in nk:
            return "⭕"
        if "nie wiem" in nk:
            return "❓"
        return "🧭"
    if any(k in nk_var for k in ("preferencj", "komitet", "wybor", "glos", "parti", "sejm")):
        if "odmow" in nk:
            return "🤐"
        if "trudno" in nk:
            return "🤷"
        if "niewaz" in nk:
            return "⭕"
        if "nie wiem" in nk or "niezdecyd" in nk:
            return "❓"
        return "🗳️"
    return "📌"


def _collect_personal_metry_available(
    results_df: pd.DataFrame,
    metry_questions: list[dict[str, object]],
) -> list[dict[str, object]]:
    out: list[dict[str, object]] = []
    for mq in metry_questions:
        db_col = str(mq.get("db_column") or "").strip()
        if not db_col:
            continue
        col_name = f"METRY_{db_col}"
        table_label = str((mq or {}).get("table_label") or (mq or {}).get("prompt") or db_col).strip()
        if col_name in results_df.columns:
            col_series = (
                results_df[col_name]
                .fillna("")
                .astype(str)
                .str.strip()
                .map(lambda v, dbc=db_col, tbl=table_label: _personal_demo_code(dbc, v, tbl))
            )
        else:
            col_series = pd.Series([""] * len(results_df.index), index=results_df.index, dtype=object)
        ordered_codes = [
            _personal_demo_code(db_col, str((opt or {}).get("code") or "").strip(), table_label)
            for opt in list(mq.get("options") or [])
        ]
        option_icons: dict[str, str] = {}
        for opt in list(mq.get("options") or []):
            if not isinstance(opt, dict):
                continue
            canon_code = _personal_demo_code(db_col, str(opt.get("code") or "").strip(), table_label)
            if canon_code and canon_code not in option_icons:
                icon_val = str(opt.get("value_emoji") or "").strip()
                if icon_val in {"🏙️", "🏙", "🌆"}:
                    icon_val = "🏬"
                option_icons[canon_code] = icon_val
        ordered_codes = [c for c in ordered_codes if c]
        seen_codes: set[str] = set()
        ordered_codes = [c for c in ordered_codes if not (c in seen_codes or seen_codes.add(c))]
        extra = sorted(set(col_series[col_series != ""]) - set(ordered_codes))
        codes = ordered_codes + list(extra)
        if not codes:
            continue
        out.append(
            {
                "question": mq,
                "col_name": col_name,
                "codes": codes,
                "variable_emoji": str((mq or {}).get("variable_emoji") or "").strip(),
                "option_icons": option_icons,
            }
        )
    return out


def _render_personal_demography_subpage(
    *,
    study_id: str,
    person_gen: str,
    results_df: pd.DataFrame,
    metry_questions: list[dict[str, object]],
    archetype_names: list[str],
    disp_name_fn,
    is_mobile: bool,
    gender_code: str = "M",
) -> None:
    page_state_key = f"personal_demo_page_{study_id}"
    available = _collect_personal_metry_available(results_df, metry_questions)

    st.markdown(
        """
        <style>
          .pdemo-head{
            display:flex;
            justify-content:space-between;
            align-items:flex-end;
            gap:14px;
            margin:8px 0 10px 0;
            flex-wrap:wrap;
          }
          .pdemo-title{
            font-size:1.75rem;
            font-weight:800;
            color:#0f2748;
            line-height:1.1;
          }
          .pdemo-sub{
            color:#475569;
            font-size:0.93rem;
          }
          .pdemo-context{
            border:1px solid #dbe4ef;
            border-radius:12px;
            background:#f8fbff;
            padding:10px 12px;
            margin:8px 0 16px 0;
            display:flex;
            flex-wrap:wrap;
            gap:8px 10px;
            align-items:center;
          }
          .pdemo-chip{
            display:inline-flex;
            align-items:center;
            gap:6px;
            border:1px solid #d7e0eb;
            background:#ffffff;
            border-radius:999px;
            padding:4px 10px;
            font-size:12.5px;
            font-weight:700;
            color:#334155;
          }
          .pdemo-box{
            border:1px solid #dbe4ef;
            border-radius:12px;
            background:#fff;
            padding:10px 12px;
            margin:0 0 12px 0;
          }
          .pdemo-box-label{
            font-size:15px;
            font-weight:900;
            text-transform:uppercase;
            letter-spacing:.02em;
            color:#334155;
            display:flex;
            align-items:center;
            gap:6px;
          }
          .pdemo-box-note{
            color:#5f6b7a;
            font-size:12px;
            margin:2px 0 6px 0;
          }
          .pdemo-cards{
            display:grid;
            grid-template-columns:repeat(auto-fit,minmax(175px,1fr));
            gap:8px;
            margin:10px 0 2px 0;
          }
          .pdemo-stat{
            border:1px solid #dbe4ef;
            border-radius:10px;
            background:#fff;
            padding:8px 10px;
          }
          .pdemo-stat-label{
            font-size:12px;
            font-weight:800;
            text-transform:uppercase;
            letter-spacing:.03em;
            color:#5f6b7a;
          }
          .pdemo-stat-main{
            margin-top:2px;
            font-size:14px;
            font-weight:900;
            color:#111827;
            line-height:1.2;
          }
          .pdemo-stat-sub{
            margin-top:2px;
            font-size:12.5px;
            color:#3f4954;
          }
          .pdemo-table-wrap{
            overflow-x:auto;
            max-width:940px;
          }
          .pdemo-table{
            margin-top:0;
            width:100%;
            min-width:720px;
            max-width:940px;
            border-collapse:collapse;
            border:3px solid #b8c2cc;
            background:#fff;
            font-size:13.5px;
            color:#334155;
          }
          .pdemo-table th,.pdemo-table td{
            padding:8px 10px;
            border:1px solid #dfe4ea;
            text-align:left;
            vertical-align:middle;
          }
          .pdemo-table th{
            background:#f2f6fb;
            color:#1f2f44;
            font-weight:800;
            font-size:13.5px;
          }
          .pdemo-radar-title{
            margin-top:22px;
            margin-bottom:8px;
            font-size:2.0rem;
            font-weight:900;
            color:#1f2f44;
            line-height:1.05;
          }
          .pdemo-top2-grid{
            display:flex;
            justify-content:center;
            align-items:stretch;
            gap:12px;
            margin:-6px 0 0 0;
            flex-wrap:wrap;
          }
          .pdemo-top2-card{
            border:1px solid #dbe4ef;
            border-radius:12px;
            background:#fff;
            padding:10px 12px;
            text-align:center;
            width:min(380px, 48%);
            min-width:260px;
          }
          .pdemo-top2-title{
            font-size:15px;
            font-weight:800;
            color:#1f2f44;
            margin-bottom:6px;
          }
          .pdemo-top2-line{
            font-size:14px;
            font-weight:800;
            color:#334155;
            line-height:1.28;
            letter-spacing:.01em;
            display:flex;
            justify-content:center;
            align-items:center;
            gap:16px;
            flex-wrap:wrap;
          }
          .pdemo-top2-line .pdemo-role-item{
            display:inline-flex;
            align-items:center;
            gap:8px;
            white-space:nowrap;
          }
          .pdemo-wheel-title{
            margin-top:28px;
            margin-bottom:6px;
            font-size:2.0rem;
            font-weight:900;
            color:#1f2f44;
            line-height:1.05;
          }
          .pdemo-wheel-sep{
            border-top:1px solid #d9e2ef;
            margin:30px 0 16px 0;
            width:100%;
          }
          .pdemo-profile-title{
            font-weight:600;
            font-size:1.4rem;
            color:#1f2f44;
            margin:2px 0 8px 0;
            text-align:center;
          }
          .pdemo-wheel-image{
            display:flex;
            justify-content:center;
            margin-top:6px;
          }
          .pdemo-score-card{border:1px solid #d5dfec;border-radius:12px;background:#ffffff;padding:12px 14px;margin:8px 0 10px 0;}
          .pdemo-score-title{font-size:15px;font-weight:800;color:#334155;margin:0 0 4px 0;}
          .pdemo-score-value{font-size:46px;line-height:1;font-weight:900;color:#0f172a;margin:0 0 8px 0;}
          .pdemo-score-badge{display:inline-block;padding:5px 10px;border-radius:999px;border:1px solid var(--pdemo-score-color,#0ea5e9);background:var(--pdemo-score-bg,#eff6ff);color:var(--pdemo-score-color,#0ea5e9);font-weight:900;font-size:15px;}
          .pdemo-score-desc{margin:8px 0 10px 0;color:#475569;font-size:14px;font-weight:600;}
          .pdemo-score-track{height:14px;border-radius:999px;background:#d5dde8;border:1px solid #aebfd3;overflow:hidden;}
          .pdemo-score-fill{height:100%;border-radius:999px;background:linear-gradient(90deg,#2563eb 0%,#22c55e 100%);}
          .pdemo-score-scale{display:flex;justify-content:space-between;color:#64748b;font-size:11px;margin-top:6px;font-weight:700;}
          .pdemo-wheel-legend-wrap{
            display:flex;
            justify-content:center;
            margin-top:6px;
          }
          .pdemo-wheel-legend{
            border:1px solid #cfd9e8;
            background:#fff;
            border-radius:14px;
            padding:10px 16px;
            display:flex;
            gap:20px;
            align-items:center;
            flex-wrap:wrap;
            justify-content:center;
          }
          .pdemo-wheel-legend span{
            font-size:13px;
            font-weight:700;
            color:#334155;
            display:inline-flex;
            align-items:center;
            gap:8px;
          }
          .pdemo-wheel-legend i{
            width:11px;
            height:11px;
            border-radius:3px;
            display:inline-block;
          }
          @media (max-width:900px){
            .pdemo-title{font-size:1.42rem;}
            .pdemo-table-wrap{max-width:100%;}
            .pdemo-radar-title,.pdemo-wheel-title{font-size:1.35rem;}
            .pdemo-top2-grid{justify-content:stretch;}
            .pdemo-top2-card{width:100%;min-width:0;}
            .pdemo-profile-title{font-size:1.1rem;}
          }
        </style>
        """,
        unsafe_allow_html=True,
    )

    back_col, title_col = st.columns([0.16, 0.84], gap="small")
    with back_col:
        if st.button("← Cofnij", key=f"personal_demo_back_btn_{study_id}", use_container_width=True):
            st.session_state[page_state_key] = False
            st.rerun()
    with title_col:
        st.markdown(
            f"""
            <div class="pdemo-head">
              <div>
                <div class="pdemo-title">Profile demograficzne archetypu</div>
                <div class="pdemo-sub">Raport wielocechowy dla badania archetypu {html.escape(person_gen)}</div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    if not available:
        st.info("Dla tego badania brak skonfigurowanych danych metryczkowych do analizy demograficznej.")
        return

    n_all = int(len(results_df))
    st.markdown(
        f"""
        <div class="pdemo-context">
          <span class="pdemo-chip">👤 Badanie personalne</span>
          <span class="pdemo-chip">🧾 Respondenci: {n_all}</span>
          <span class="pdemo-chip">⚙️ Zmienne metryczkowe: {len(available)}</span>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="pdemo-box">
          <div class="pdemo-box-label">🎛️ Filtr wielocechowy</div>
          <div class="pdemo-box-note">Filtry działają łącznie (AND): np. płeć + wiek + wykształcenie.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    min_demo_n = int(
        st.number_input(
            "Minimalna liczebność podgrupy (N) dla stabilnego wniosku",
            min_value=5,
            max_value=5000,
            value=30,
            step=1,
            key=f"personal_demo_min_n_subpage_{study_id}",
        )
    )

    filt_df = results_df.copy()
    active_filters: list[str] = []
    filter_cols = st.columns(2, gap="medium") if len(available) > 1 else [st]
    all_token = "__ALL__"
    for idx, item in enumerate(available):
        mq = item["question"]
        col_name = str(item["col_name"])
        db_col = str((mq or {}).get("db_column") or "").strip()
        prompt = str((mq or {}).get("prompt") or db_col).strip()
        table_label = str((mq or {}).get("table_label") or prompt or db_col).strip()
        var_icon_pref = str(item.get("variable_emoji") or "").strip()
        option_icon_map = item.get("option_icons") if isinstance(item.get("option_icons"), dict) else {}
        demo_col = f"{col_name}__DEMO"
        if demo_col not in filt_df.columns:
            if col_name in results_df.columns:
                filt_df[demo_col] = (
                    results_df[col_name]
                    .fillna("")
                    .astype(str)
                    .str.strip()
                    .map(lambda v, dbc=db_col, tbl=table_label: _personal_demo_code(dbc, v, tbl))
                )
            else:
                filt_df[demo_col] = pd.Series([""] * len(filt_df.index), index=filt_df.index, dtype=object)
        select_options = [all_token] + [str(code) for code in list(item["codes"])]
        col_ctx = filter_cols[idx % len(filter_cols)]
        with col_ctx:
            selected_code = st.selectbox(
                table_label,
                options=select_options,
                format_func=(
                    lambda code, dbc=db_col, tbl=table_label, om=option_icon_map: (
                        "— brak filtra —"
                        if code == all_token
                        else f"{_personal_metry_cat_icon(dbc, tbl, str(code), option_icon_map=om)} {str(code)}"
                    )
                ),
                key=f"personal_demo_sub_filter_single_{study_id}_{db_col}",
            )
        if selected_code != all_token:
            filt_df = filt_df[filt_df[demo_col].astype(str) == str(selected_code)]
            active_filters.append(
                f"{_personal_metry_var_icon(db_col, table_label, var_icon_pref)} {table_label}: "
                f"{_personal_metry_cat_icon(db_col, table_label, str(selected_code), option_icon_map=option_icon_map)} {str(selected_code)}"
            )

    n_filtered = int(len(filt_df))
    has_active_filters = bool(active_filters)
    if has_active_filters:
        st.markdown("**Aktywne filtry:** " + " | ".join(active_filters))
    else:
        st.caption("Brak aktywnych filtrów: profil odpowiada całej próbie.")
    st.markdown(f"**Liczebność podgrupy:** {n_filtered} / {n_all}")

    if n_filtered <= 0:
        st.warning("Brak odpowiedzi spełniających aktualny zestaw filtrów.")
        return
    if n_filtered < min_demo_n:
        st.warning(f"Wynik ma podwyższoną niepewność (N={n_filtered}, próg stabilności: {min_demo_n}).")

    means_20_all = {
        k: (
            float(pd.to_numeric(results_df[k], errors="coerce").mean())
            if k in results_df.columns and pd.to_numeric(results_df[k], errors="coerce").notna().any()
            else 0.0
        )
        for k in archetype_names
    }
    filtered_means_20 = {
        k: (
            float(pd.to_numeric(filt_df[k], errors="coerce").mean())
            if k in filt_df.columns and pd.to_numeric(filt_df[k], errors="coerce").notna().any()
            else 0.0
        )
        for k in archetype_names
    }

    cards_html_parts: list[str] = []
    table_rows_html: list[str] = []
    has_demography_answers = False
    for item in available:
        mq = item["question"]
        col_name = str(item["col_name"])
        db_col = str((mq or {}).get("db_column") or "").strip()
        prompt = str((mq or {}).get("prompt") or db_col).strip()
        table_label = str((mq or {}).get("table_label") or prompt or db_col).strip()
        var_icon_pref = str(item.get("variable_emoji") or "").strip()
        option_icon_map = item.get("option_icons") if isinstance(item.get("option_icons"), dict) else {}
        icon = _personal_metry_var_icon(db_col, table_label, var_icon_pref)
        demo_col = f"{col_name}__DEMO"
        if col_name in results_df.columns:
            q_all = (
                results_df[col_name]
                .fillna("")
                .astype(str)
                .str.strip()
                .map(lambda v, dbc=db_col, tbl=table_label: _personal_demo_code(dbc, v, tbl))
            )
        else:
            q_all = pd.Series([""] * len(results_df.index), index=results_df.index, dtype=object)
        if demo_col in filt_df.columns:
            q_sub = filt_df[demo_col].fillna("").astype(str).str.strip()
        elif col_name not in filt_df.columns:
            q_sub = pd.Series([""] * len(filt_df.index), index=filt_df.index, dtype=object)
        else:
            q_sub = (
                filt_df[col_name]
                .fillna("")
                .astype(str)
                .str.strip()
                .map(lambda v, dbc=db_col, tbl=table_label: _personal_demo_code(dbc, v, tbl))
            )

        categories: list[dict[str, object]] = []
        for code in list(item["codes"]):
            code_txt = str(code or "").strip()
            if not code_txt:
                continue
            all_nonempty = int(q_all.ne("").sum())
            sub_nonempty = int(q_sub.ne("").sum())
            if all_nonempty > 0 or sub_nonempty > 0:
                has_demography_answers = True
            pct_all = 100.0 * float((q_all == code_txt).sum()) / max(1, all_nonempty)
            pct_sub = 100.0 * float((q_sub == code_txt).sum()) / max(1, sub_nonempty)
            categories.append(
                {
                    "code": code_txt,
                    "label": code_txt,
                    "pct_all": float(pct_all),
                    "pct_sub": float(pct_sub),
                    "diff": float(pct_sub - pct_all),
                }
            )
        if not categories:
            continue

        has_missing_row = any(str(c.get("code") or "").strip().lower() == "brak danych" for c in categories)
        all_zero_known = all(abs(float(c.get("pct_sub") or 0.0)) <= 1e-9 for c in categories)
        if all_zero_known and not has_missing_row:
            pct_all_sum = float(sum(float(c.get("pct_all") or 0.0) for c in categories))
            pct_all_missing = max(0.0, min(100.0, 100.0 - pct_all_sum))
            categories.append(
                {
                    "code": "brak danych",
                    "label": "brak danych",
                    "pct_all": float(pct_all_missing),
                    "pct_sub": 100.0,
                    "diff": float(100.0 - pct_all_missing),
                }
            )

        strongest = max(
            categories,
            key=lambda c: (float(c.get("pct_sub") or 0.0), -int(categories.index(c))),
        )
        cards_html_parts.append(
            f"""
            <div class="pdemo-stat">
              <div class="pdemo-stat-label">{html.escape(icon)} {html.escape(table_label.upper())}</div>
              <div class="pdemo-stat-main">{html.escape(_personal_metry_cat_icon(db_col, table_label, str(strongest.get("label") or ""), option_icon_map=option_icon_map))} {html.escape(str(strongest.get("label") or ""))}</div>
              <div class="pdemo-stat-sub">{float(strongest.get("pct_sub") or 0.0):.1f}% • {float(strongest.get("diff") or 0.0):+,.1f} pp</div>
            </div>
            """
        )

        rowspan = len(categories)
        top_pct_sub = max(float(c.get("pct_sub") or 0.0) for c in categories)
        for idx, cat in enumerate(categories):
            pct_sub = float(cat["pct_sub"])
            pct_all = float(cat["pct_all"])
            diff = float(cat["diff"])
            is_top = (top_pct_sub > 1e-9) and (abs(pct_sub - top_pct_sub) <= 1e-9)
            bar_w = max(0.0, min(100.0, pct_sub))
            fill_color = "#8ecae6" if is_top else "#d8e5f1"
            top_border = "border-top:3px solid #b8c2cc;"
            diff_color = "#0f766e" if diff >= 0 else "#9a3412"
            first_col = (
                "<td "
                f"rowspan='{rowspan}' "
                f"style=\"font-weight:800; text-transform:uppercase; vertical-align:middle; background:#fafafa; border-left:3px solid #b8c2cc; {top_border}\">"
                "<span style='display:inline-flex; align-items:center; gap:6px;'>"
                f"<span>{html.escape(icon)}</span>"
                f"<span>{html.escape(table_label)}</span>"
                "</span>"
                "</td>"
                if idx == 0
                else ""
            )
            table_rows_html.append(
                "<tr>"
                f"{first_col}"
                f"<td style=\"font-size:13.5px; font-weight:{'800' if is_top else '500'}; {top_border if idx == 0 else ''}\">"
                "<span style='display:inline-flex; align-items:center; gap:6px;'>"
                f"<span>{html.escape(_personal_metry_cat_icon(db_col, table_label, str(cat['label']), option_icon_map=option_icon_map))}</span>"
                f"<span>{html.escape(str(cat['label']))}</span>"
                "</span>"
                "</td>"
                f"<td style=\"padding:0; min-width:176px; border:1px solid #dfe4ea; {top_border if idx == 0 else ''}\">"
                "<div style=\"position:relative; height:34px; background:#fff;\">"
                f"<div style=\"position:absolute; left:0; top:0; bottom:0; width:{bar_w:.1f}%; background:{fill_color}; opacity:0.96;\"></div>"
                f"<span style=\"position:absolute; right:6px; top:6px; z-index:2; background:rgba(255,255,255,0.88); padding:1px 5px; border-radius:4px; font-size:13.5px; font-weight:{'900' if is_top else '600'}; color:#111;\">{pct_sub:.1f}%</span>"
                "</div>"
                "</td>"
                f"<td style=\"font-size:13.5px; text-align:right; {top_border if idx == 0 else ''}\">{pct_all:.1f}%</td>"
                f"<td style=\"font-size:13.5px; text-align:right; color:{diff_color}; font-weight:400; border-right:3px solid #b8c2cc; {top_border if idx == 0 else ''}\">{diff:+.1f} pp</td>"
                "</tr>"
            )

    if cards_html_parts:
        st.markdown(
            f"""
            <div class="pdemo-box">
              <div class="pdemo-box-label">📌 STATYSTYCZNY PROFIL DEMOGRAFICZNY</div>
              <div class="pdemo-cards">{''.join(cards_html_parts)}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    if table_rows_html:
        table_html = (
            "<div class='pdemo-box'>"
            "<div class='pdemo-box-label'>👥 PROFIL DEMOGRAFICZNY</div>"
            "<div class='pdemo-box-note'>W tabeli pogrubiona najwyższa kategoria w każdej zmiennej.</div>"
            "<div class='pdemo-table-wrap'><table class='pdemo-table'>"
            "<thead><tr>"
            "<th style='min-width:150px; font-size:13.5px; border-top:3px solid #b8c2cc; border-left:3px solid #b8c2cc;'>Zmienna</th>"
            "<th style='min-width:220px; font-size:13.5px; border-top:3px solid #b8c2cc;'>Kategoria</th>"
            "<th style='min-width:176px; text-align:center; font-size:13.5px; border-top:3px solid #b8c2cc;'>% podgrupa</th>"
            "<th style='min-width:130px; text-align:center; border-top:3px solid #b8c2cc;'>% cała próba</th>"
            "<th style='min-width:120px; text-align:center; border-top:3px solid #b8c2cc; border-right:3px solid #b8c2cc;'>Różnica (w pp.)</th>"
            "</tr></thead><tbody>"
            + "".join(table_rows_html)
            + "</tbody></table></div></div>"
        )
        st.markdown(table_html, unsafe_allow_html=True)

    if not has_demography_answers:
        st.info("Brak odpowiedzi metryczkowych w próbie: pomijam wykres radarowy i profile 0-100 dla podgrupy.")
        return

    def _demo_match_summary(profile_all_20: dict[str, float], profile_sub_20: dict[str, float]) -> dict[str, object]:
        all_100 = {a: max(0.0, min(100.0, float(profile_all_20.get(a, 0.0)) * 5.0)) for a in archetype_names}
        sub_100 = {a: max(0.0, min(100.0, float(profile_sub_20.get(a, 0.0)) * 5.0)) for a in archetype_names}
        diffs = {a: abs(float(all_100.get(a, 0.0)) - float(sub_100.get(a, 0.0))) for a in archetype_names}
        diff_vals = [float(v) for v in diffs.values()]
        mae = float(sum(diff_vals) / max(1, len(diff_vals)))
        rmse = math.sqrt(float(sum(v * v for v in diff_vals) / max(1, len(diff_vals))))
        sorted_gaps_desc = sorted(diff_vals, reverse=True)
        top3_gap_mae = float(sum(sorted_gaps_desc[:3]) / max(1, min(3, len(sorted_gaps_desc))))

        def _key_priority_pool(profile_100: dict[str, float]) -> list[str]:
            ordered = sorted(
                archetype_names,
                key=lambda a: (-float(profile_100.get(a, 0.0)), archetype_names.index(a)),
            )
            top3 = ordered[:3]
            if len(top3) >= 3 and float(profile_100.get(top3[2], 0.0)) < 70.0:
                return top3[:2]
            return top3

        all_top = _key_priority_pool(all_100)
        sub_top = _key_priority_pool(sub_100)
        key_archetypes: list[str] = []
        for arche in all_top + sub_top:
            if arche not in key_archetypes:
                key_archetypes.append(arche)
        key_gap_vals = [float(diffs.get(a, 0.0)) for a in key_archetypes]
        key_gap_mae = float(sum(key_gap_vals) / max(1, len(key_gap_vals)))
        key_gap_max = float(max(key_gap_vals)) if key_gap_vals else 0.0
        shared_priority_count = len(set(all_top).intersection(set(sub_top)))
        main_priority_mismatch_penalty = 2.5 if (all_top and sub_top and all_top[0] != sub_top[0]) else 0.0
        shared_priority_penalty = 5.5 if shared_priority_count == 0 else (2.0 if shared_priority_count == 1 else 0.0)

        score_mae = max(0.0, min(100.0, 100.0 - mae))
        score_rmse = max(0.0, min(100.0, 100.0 - rmse))
        score_top3 = max(0.0, min(100.0, 100.0 - top3_gap_mae))
        score_key = max(0.0, min(100.0, 100.0 - key_gap_mae))
        base_score = 0.40 * score_mae + 0.20 * score_rmse + 0.20 * score_top3 + 0.20 * score_key
        key_penalty = (
            0.56 * key_gap_mae
            + 0.26 * max(0.0, key_gap_max - 10.0)
            + shared_priority_penalty
            + main_priority_mismatch_penalty
        )
        match_score = max(0.0, min(100.0, base_score - key_penalty))
        match_bands: list[tuple[str, str]] = [
            ("Marginalne dopasowanie", "Profile są w dużej mierze rozbieżne; potrzebna gruntowna korekta przekazu i priorytetów."),
            ("Bardzo niskie dopasowanie", "Dopasowanie jest słabe i niestabilne; dominują rozjazdy strategiczne."),
            ("Niskie dopasowanie", "Widać pojedyncze punkty wspólne, ale profil nadal wyraźnie się rozjeżdża."),
            ("Umiarkowane dopasowanie", "Istnieje wspólny rdzeń, ale kluczowe luki nadal wymagają korekty."),
            ("Znaczące dopasowanie", "Dopasowanie jest zauważalne, choć nadal potrzebne są poprawki na kluczowych pozycjach."),
            ("Wysokie dopasowanie", "Profil jest w dużej części zgodny; pozostają pojedyncze luki do domknięcia."),
            ("Bardzo wysokie dopasowanie", "Różnice są niewielkie i dotyczą głównie lokalnych odchyleń."),
            ("Ekstremalnie wysokie dopasowanie", "Profile są niemal zbieżne także na kluczowych archetypach."),
        ]
        if match_score >= 90:
            band_idx = 7
        elif match_score >= 80:
            band_idx = 6
        elif match_score >= 70:
            band_idx = 5
        elif match_score >= 60:
            band_idx = 4
        elif match_score >= 50:
            band_idx = 3
        elif match_score >= 40:
            band_idx = 2
        elif match_score >= 30:
            band_idx = 1
        else:
            band_idx = 0
        band_label, band_desc = match_bands[band_idx]
        return {
            "score": float(match_score),
            "band_label": band_label,
            "band_desc": band_desc,
            "mae": float(mae),
            "rmse": float(rmse),
            "top3_gap_mae": float(top3_gap_mae),
            "key_gap_mae": float(key_gap_mae),
            "key_gap_max": float(key_gap_max),
        }

    if has_active_filters:
        demo_match = _demo_match_summary(means_20_all, filtered_means_20)
        score_pct = max(0.0, min(100.0, float(demo_match.get("score") or 0.0)))
        if score_pct >= 90:
            score_color, score_bg = "#0f766e", "#ecfeff"
        elif score_pct >= 80:
            score_color, score_bg = "#0e7490", "#ecfeff"
        elif score_pct >= 70:
            score_color, score_bg = "#6d28d9", "#f5f3ff"
        elif score_pct >= 60:
            score_color, score_bg = "#1d4ed8", "#eff6ff"
        elif score_pct >= 50:
            score_color, score_bg = "#b45309", "#fffbeb"
        elif score_pct >= 40:
            score_color, score_bg = "#c2410c", "#fff7ed"
        elif score_pct >= 30:
            score_color, score_bg = "#be123c", "#fff1f2"
        else:
            score_color, score_bg = "#7f1d1d", "#fef2f2"
        st.markdown(
            f"""
            <div class="pdemo-score-card" style="--pdemo-score-color:{score_color}; --pdemo-score-bg:{score_bg};">
              <div class="pdemo-score-title">Poziom dopasowania podgrupy do całej próby</div>
              <div class="pdemo-score-value">{score_pct:.1f}%</div>
              <div class="pdemo-score-badge">Ocena: {html.escape(str(demo_match.get('band_label') or ''))}</div>
              <div class="pdemo-score-desc">{html.escape(str(demo_match.get('band_desc') or ''))}</div>
              <div class="pdemo-score-track"><div class="pdemo-score-fill" style="width:{score_pct:.1f}%;"></div></div>
              <div class="pdemo-score-scale"><span>0%</span><span>100%</span></div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    def _priority_top_for_ui_demo(profile_20: dict[str, float], order: list[str]) -> list[str]:
        ordered = sorted(order, key=lambda a: (-float(profile_20.get(a, 0.0)), order.index(a)))
        top3 = [x for x in ordered[:3] if x]
        if len(top3) >= 3 and float(profile_20.get(top3[2], 0.0)) * 5.0 < 70.0:
            return top3[:2]
        return top3

    all_top = _priority_top_for_ui_demo(means_20_all, archetype_names)
    sub_top = _priority_top_for_ui_demo(filtered_means_20, archetype_names) if has_active_filters else []
    theta_labels = [disp_name_fn(n) for n in archetype_names]
    theta_display = [str(lbl) for lbl in theta_labels]
    radar_top_union = set(all_top + sub_top)
    radar_tick_text = [
        (f"<b>{html.escape(str(lbl))}</b>" if arch in radar_top_union else html.escape(str(lbl)))
        for arch, lbl in zip(archetype_names, theta_labels)
    ]
    all_vals = [float(means_20_all.get(n, 0.0)) for n in archetype_names]
    filt_vals = [float(filtered_means_20.get(n, 0.0)) for n in archetype_names]

    def _marker_series(profile: dict[str, float], topn: list[str], palette: dict[str, str]) -> tuple[list[float | None], list[str]]:
        mapping: dict[str, str] = {}
        if len(topn) > 0:
            mapping[topn[0]] = palette["main"]
        if len(topn) > 1:
            mapping[topn[1]] = palette["aux"]
        if len(topn) > 2:
            mapping[topn[2]] = palette["supp"]
        r_vals: list[float | None] = []
        c_vals: list[str] = []
        for arch in archetype_names:
            if arch in mapping:
                r_vals.append(float(profile.get(arch, 0.0)))
                c_vals.append(str(mapping[arch]))
            else:
                r_vals.append(None)
                c_vals.append("rgba(0,0,0,0)")
        return r_vals, c_vals

    all_top_colors = {"main": "#ef4444", "aux": "#f59e0b", "supp": "#22c55e"}
    sub_top_colors = {"main": "#2563eb", "aux": "#8b5cf6", "supp": "#f97316"}
    all_marker_r, all_marker_c = _marker_series(means_20_all, all_top, all_top_colors)
    radar_traces: list[go.Scatterpolar] = [
        go.Scatterpolar(
            r=all_vals + [all_vals[0]],
            theta=theta_display + [theta_display[0]],
            fill="toself",
            fillcolor="rgba(37,99,235,0.18)",
            line=dict(color="#2563eb", width=3),
            marker=dict(size=5, symbol="circle"),
            name=f"profil całej próby (N={n_all})",
            hovertemplate="<b>%{theta}</b><br>Cała próba: %{r:.2f}<extra></extra>",
        ),
        go.Scatterpolar(
            r=all_marker_r,
            theta=theta_display,
            mode="markers",
            marker=dict(size=16, symbol="circle", color=all_marker_c, opacity=0.92, line=dict(color="black", width=2.4)),
            showlegend=False,
            hovertemplate=f"<b>%{{theta}}</b><br>TOP{max(1, len(all_top))} całej próby: %{{r:.2f}}<extra></extra>",
        ),
    ]
    if has_active_filters:
        sub_marker_r, sub_marker_c = _marker_series(filtered_means_20, sub_top, sub_top_colors)
        radar_traces.append(
            go.Scatterpolar(
                r=filt_vals + [filt_vals[0]],
                theta=theta_display + [theta_display[0]],
                fill="toself",
                fillcolor="rgba(15,118,110,0.16)",
                line=dict(color="#0f766e", width=3, dash="dot"),
                marker=dict(size=6, symbol="square"),
                name=f"profil podgrupy filtrowanej (N={n_filtered})",
                hovertemplate="<b>%{theta}</b><br>Podgrupa: %{r:.2f}<extra></extra>",
            )
        )
        radar_traces.append(
            go.Scatterpolar(
                r=sub_marker_r,
                theta=theta_display,
                mode="markers",
                marker=dict(size=14, symbol="square", color=sub_marker_c, opacity=0.94, line=dict(color="#0f172a", width=1.9)),
                showlegend=False,
                hovertemplate=f"<b>%{{theta}}</b><br>TOP{max(1, len(sub_top))} podgrupy: %{{r:.2f}}<extra></extra>",
            )
        )

    demo_fig = go.Figure(data=radar_traces)
    demo_fig.update_layout(
        height=640 if not is_mobile else 450,
        paper_bgcolor="rgba(0,0,0,0)",
        polar=dict(
            bgcolor="rgba(0,0,0,0)",
            radialaxis=dict(visible=True, range=[0, 20]),
            angularaxis=dict(
                type="category",
                rotation=90,
                direction="clockwise",
                tickfont=dict(size=16 if not is_mobile else 10.5),
                tickvals=theta_display,
                ticktext=radar_tick_text,
            ),
        ),
        margin=dict(l=24, r=24, t=66, b=90),
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.15,
            xanchor="center",
            x=0.5,
            font=dict(size=13.5),
            bgcolor="rgba(255,255,255,0.94)",
            bordercolor="#cfd9e8",
            borderwidth=1,
            entrywidthmode="pixels",
            entrywidth=250,
            tracegroupgap=22,
        ),
    )

    radar_title = "Podgląd radarowy podgrupy" if has_active_filters else "Podgląd radarowy całej próby"
    st.markdown(f"<div class='pdemo-radar-title'>{html.escape(radar_title)}</div>", unsafe_allow_html=True)
    st.plotly_chart(
        demo_fig,
        use_container_width=True,
        config={"displaylogo": False, "displayModeBar": False, "responsive": True},
        key=f"personal_demo_profile_radar_subpage_{study_id}",
    )

    def _role_legend_html(palette: dict[str, str], marker: str, count: int) -> str:
        role_defs = [("główny", "main"), ("wspierający", "aux"), ("poboczny", "supp")]
        items: list[str] = []
        for idx, (label, role_key) in enumerate(role_defs):
            if idx >= max(0, int(count)):
                break
            items.append(
                f"<span class=\"pdemo-role-item\"><span style=\"color:{palette[role_key]};\">{marker}</span><span>{label}</span></span>"
            )
        return "".join(items)

    if has_active_filters:
        st.markdown(
            f"""
            <div class="pdemo-top2-grid">
              <div class="pdemo-top2-card">
                <div class="pdemo-top2-title">TOP{max(1, len(all_top))} całej próby</div>
                <div class="pdemo-top2-line">{_role_legend_html(all_top_colors, "●", len(all_top))}</div>
              </div>
              <div class="pdemo-top2-card">
                <div class="pdemo-top2-title">TOP{max(1, len(sub_top))} podgrupy</div>
                <div class="pdemo-top2-line">{_role_legend_html(sub_top_colors, "■", len(sub_top))}</div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            f"""
            <div class="pdemo-top2-grid">
              <div class="pdemo-top2-card">
                <div class="pdemo-top2-title">TOP{max(1, len(all_top))} całej próby</div>
                <div class="pdemo-top2-line">{_role_legend_html(all_top_colors, "●", len(all_top))}</div>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("<div class='pdemo-wheel-sep'></div>", unsafe_allow_html=True)
    st.markdown("<div class='pdemo-wheel-title'>Profile siły archetypów 0-100 (skala: 0-100)</div>", unsafe_allow_html=True)
    try:
        safe_study = re.sub(r"[^A-Za-z0-9_-]+", "_", str(study_id or "study"))
        profile_all_100 = {k: max(0.0, min(100.0, float(means_20_all.get(k, 0.0)) * 5.0)) for k in archetype_names}
        wheel_all = make_segment_profile_wheel_png(
            mean_scores=profile_all_100,
            out_path=f"personal_demo_all_{safe_study}.png",
            label_mode="arche",
            gender_code=gender_code,
            dark_mode=False,
        )
        wheel_all_dark = make_segment_profile_wheel_png(
            mean_scores=profile_all_100,
            out_path=f"personal_demo_all_{safe_study}_dark.png",
            label_mode="arche",
            gender_code=gender_code,
            dark_mode=True,
        )
        wheel_sub = None
        wheel_sub_dark = None
        if has_active_filters:
            profile_sub_100 = {k: max(0.0, min(100.0, float(filtered_means_20.get(k, 0.0)) * 5.0)) for k in archetype_names}
            wheel_sub = make_segment_profile_wheel_png(
                mean_scores=profile_sub_100,
                out_path=f"personal_demo_sub_{safe_study}.png",
                label_mode="arche",
                gender_code=gender_code,
                dark_mode=False,
            )
            wheel_sub_dark = make_segment_profile_wheel_png(
                mean_scores=profile_sub_100,
                out_path=f"personal_demo_sub_{safe_study}_dark.png",
                label_mode="arche",
                gender_code=gender_code,
                dark_mode=True,
            )

        def _show_image_compat(img_path: str, dark_img_path: str | None = None, max_width_px: int = 560) -> None:
            try:
                if dark_img_path and Path(dark_img_path).exists():
                    light_uri = _img_data_uri_from_path(Path(img_path))
                    dark_uri = _img_data_uri_from_path(Path(dark_img_path))
                    st.markdown(
                        (
                            "<div class='pdemo-wheel-image'>"
                            + _theme_image_dual_html(
                                light_uri,
                                dark_uri,
                                style=f"width:min(100%, {int(max_width_px)}px);height:auto;",
                            )
                            + "</div>"
                        ),
                        unsafe_allow_html=True,
                    )
                else:
                    raw = Path(img_path).read_bytes()
                    b64 = base64.b64encode(raw).decode("ascii")
                    st.markdown(
                        (
                            "<div class='pdemo-wheel-image'>"
                            f"<img src='data:image/png;base64,{b64}' "
                            f"style='width:min(100%, {int(max_width_px)}px);height:auto;'/>"
                            "</div>"
                        ),
                        unsafe_allow_html=True,
                    )
            except Exception:
                try:
                    st.image(img_path, width=max_width_px)
                except TypeError:
                    st.image(img_path, use_column_width=True)

        if has_active_filters and wheel_sub:
            wheels_col1, wheels_col2 = st.columns(2, gap="large")
            with wheels_col1:
                st.markdown("<div class='pdemo-profile-title'>Profil siły archetypów całej próby</div>", unsafe_allow_html=True)
                _show_image_compat(wheel_all, dark_img_path=wheel_all_dark, max_width_px=520)
            with wheels_col2:
                st.markdown("<div class='pdemo-profile-title'>Profil siły archetypów podgrupy filtrowanej</div>", unsafe_allow_html=True)
                _show_image_compat(wheel_sub, dark_img_path=wheel_sub_dark, max_width_px=520)
        else:
            st.markdown("<div class='pdemo-profile-title'>Profil siły archetypów całej próby</div>", unsafe_allow_html=True)
            _show_image_compat(wheel_all, dark_img_path=wheel_all_dark, max_width_px=620 if not is_mobile else 520)
        st.markdown(
            """
            <div class="pdemo-wheel-legend-wrap">
              <div class="pdemo-wheel-legend">
                <span><i style="background:#e53935"></i>Zmiana</span>
                <span><i style="background:#1e88e5"></i>Ludzie</span>
                <span><i style="background:#2e7d32"></i>Porządek</span>
                <span><i style="background:#7e57c2"></i>Niezależność</span>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    except Exception as e:
        st.info(f"Nie udało się wygenerować porównania kół 0-100: {e}")



@st.cache_data(ttl=30)
def fetch_studies_list():
    engine = _sa_engine()
    query = sa.text("""
        SELECT
            id,
            first_name_nom, first_name_gen, first_name_dat, first_name_acc, first_name_ins, first_name_loc, first_name_voc,
            last_name_nom,  last_name_gen,  last_name_dat,  last_name_acc,  last_name_ins,  last_name_loc,  last_name_voc,
            city_nom, slug
        FROM public.studies
        WHERE COALESCE(is_active, true)
        ORDER BY created_at DESC
    """)
    with engine.begin() as conn:
        df = pd.read_sql(query, con=conn)
    engine.dispose()
    return df


def archetype_scores(answers):
    if not isinstance(answers, list) or len(answers) < 48:
        return {k: None for k in archetypes}
    a = [int(v) for v in answers]  # 👈 twarde rzutowanie
    return {name: sum(a[i-1] for i in idxs) for name, idxs in archetypes.items()}


def archetype_percent(scoresum):
    if scoresum is None:
        return None
    return round(scoresum / 20 * 100, 1)

def pick_top_3_archetypes(archetype_means, archetype_order):
    """
    Zwraca trzy archetypy o najwyższych wynikach, w kolejności zgodnej z archetype_order.
    """
    # Sortuj po wartości malejąco, a przy remisie wg porządku archetype_order
    sorted_archetypes = sorted(
        archetype_means.items(),
        key=lambda kv: (-kv[1], archetype_order.index(kv[0]))
    )
    main_type = sorted_archetypes[0][0] if len(sorted_archetypes) > 0 else None
    aux_type = sorted_archetypes[1][0] if len(sorted_archetypes) > 1 else None
    supplement_type = sorted_archetypes[2][0] if len(sorted_archetypes) > 2 else None
    return main_type, aux_type, supplement_type

def should_show_supplement(third_name: str | None,
                           means_pct: dict[str, float],
                           threshold: float = 70.0) -> bool:
    """
    Zwraca True, jeśli trzeci archetyp ma % >= threshold.
    'means_pct' to słownik {archetyp: % 0..100}.
    """
    if not third_name:
        return False
    try:
        return float(means_pct.get(third_name, 0.0)) >= (threshold - 1e-9)
    except Exception:
        return False

def add_image(paragraph, img, width):
    """
    Bezpieczne osadzanie obrazka do DOCX.
    - 1) próba bezpośrednia (path / stream),
    - 2) fallback: PIL -> PNG (dla plików błędnie oznaczonych jako .jpg/.png).
    """
    if img is None:
        return

    def _as_png_stream(source):
        try:
            if isinstance(source, (str, os.PathLike)) and os.path.exists(source):
                with Image.open(source) as im:
                    out = BytesIO()
                    if im.mode not in {"RGB", "RGBA"}:
                        im = im.convert("RGBA" if "A" in im.getbands() else "RGB")
                    im.save(out, format="PNG")
                    out.seek(0)
                    return out
            if hasattr(source, "read"):
                try:
                    source.seek(0)
                except Exception:
                    pass
                with Image.open(source) as im:
                    out = BytesIO()
                    if im.mode not in {"RGB", "RGBA"}:
                        im = im.convert("RGBA" if "A" in im.getbands() else "RGB")
                    im.save(out, format="PNG")
                    out.seek(0)
                    return out
        except Exception:
            return None
        return None

    # 1) próba bezpośrednia
    try:
        run = paragraph.add_run()
        if isinstance(img, (str, os.PathLike)) and os.path.exists(img):
            run.add_picture(img, width=width)
            return
        if hasattr(img, "read"):
            try:
                img.seek(0)
            except Exception:
                pass
            run.add_picture(img, width=width)
            return
    except Exception:
        pass

    # 2) fallback: transkodowanie do PNG
    png_stream = _as_png_stream(img)
    if png_stream is None:
        return
    try:
        run = paragraph.add_run()
        run.add_picture(png_stream, width=width)
    except Exception:
        pass

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from io import BytesIO
import os

def build_word_context(
    main_type, second_type, supplement_type, features, main, second, supplement,
    mean_scores=None, radar_image=None, archetype_table=None, num_ankiet=None,
    person: dict | None = None,
    gender_code: str = "M",
):
    """
    person = {
        "NOM": "Imię Nazwisko",
        "GEN": "Imienia Nazwiska",
        "DAT": "Imieniowi Nazwiskowi",
        "ACC": "Imię Nazwisko",
        "INS": "Imieniem Nazwiskiem",
        "LOC": "o Imieniu Nazwisku",
        "VOC": "Imieniu Nazwisku!",
        "CITY_NOM": "Kraków"   # (opcjonalnie)
    }
    """
    person = person or {}
    def p(key, fallback=""):
        return (person.get(key) or fallback).strip()

    def person_links_plain(person_list):
        cleaned = [canonical_person_name(name) for name in (person_list or []) if str(name or "").strip()]
        # zachowaj kolejność i usuń duplikaty
        return list(dict.fromkeys(cleaned))

    def kolor_label_list(palette):
        if not isinstance(palette, list):
            return ""
        out = []
        for code in palette:
            name = COLOR_NAME_MAP.get(str(code).upper(), str(code).upper())
            out.append(f"{name} ({code})")
        return ', '.join(out)

    g = normalize_gender(gender_code)
    leader_term = "liderki" if g == "K" else "lidera"
    candidate_term = "kandydatkę" if g == "K" else "kandydata"
    politician_term = "polityczki" if g == "K" else "polityka"
    possessive_term = "jej" if g == "K" else "jego"

    context = {
        # ——— Meta
        "TYTUL": "Raport Archetypów",
        "IMIE_NAZWISKO": p("GEN") or p("NOM"),   # zgodność wsteczna
        "IMIE_NAZWISKO_NOM": p("NOM"),
        "IMIE_NAZWISKO_GEN": p("GEN"),
        "IMIE_NAZWISKO_DAT": p("DAT"),
        "IMIE_NAZWISKO_ACC": p("ACC"),
        "IMIE_NAZWISKO_INST": p("INS"),
        "IMIE_NAZWISKO_LOC": p("LOC"),
        "IMIE_NAZWISKO_VOC": p("VOC"),
        "CITY_NOM": p("CITY_NOM"),
        "AUTOR": "Piotr Stec",
        "DATA": datetime.now().strftime("%Y-%m-%d"),

        # ——— Wstęp
        "WSTEP": (
            "Archetypy to uniwersalne wzorce osobowości, które od wieków pomagają ludziom rozumieć świat i budować autentyczną tożsamość. "
            "Współczesna psychologia i marketing potwierdzają, że trafnie zdefiniowany archetyp jest potężnym narzędziem komunikacji, pozwalającym budować rozpoznawalność, zaufanie i emocjonalny kontakt. Czas wykorzystać to także w polityce! "
            f"\n\nW polityce archetyp pomaga wyeksponować najważniejsze cechy {leader_term}, porządkuje przekaz, wzmacnia spójność strategii oraz wyraźnie różnicuje {candidate_term} na tle konkurencji. "
            f"Analiza archetypów pozwala lepiej zrozumieć sposób odbioru {politician_term} przez otoczenie, a co się z tym wiąże także motywacje i aspiracje. "
            "Wyniki badań archetypowych stanowią istotny fundament do tworzenia skutecznej narracji wyborczej, strategii wizerunkowej i komunikacji z wyborcami.\n\n"
            "W modelu przez nas opracowanym wykorzystano klasyfikację Mark and Pearson, obejmującą 12 uniwersalnych typów osobowościowych. "
            f"Raport przedstawia wyniki i profil archetypowy dla {p('GEN') or '—'} w oparciu o dane z przeprowadzonego badania. "
            f"Badanie to pozwoliło zidentyfikować archetyp główny i wspierający, a więc dwa najważniejsze wzorce, które mogą wzmocnić {possessive_term} pozycjonowanie. "
            "Zaprezentowano także trzeci w kolejności ważności — archetyp poboczny.\n\n"
            "Dzięki analizie archetypów można precyzyjnie dopasować komunikację do oczekiwań wyborców, podkreślić atuty, a także przewidzieć skuteczność strategii politycznej w dynamicznym środowisku publicznym. "),

        # ——— Tabela + radar + liczebność
        "TABELA_LICZEBNOSCI": archetype_table.to_dict('records') if archetype_table is not None else [],
        "RADAR_IMG": radar_image if radar_image is not None else "",
        "LICZEBNOSC_OSOB": (
            f"W badaniu udział wzięło {num_ankiet} {'osób' if (num_ankiet is None or num_ankiet != 1) else 'osoba'}."
            if num_ankiet is not None else ""
        ),

        # ——— Główny / wspierający / poboczny (bez zmian merytorycznych)
        "ARCHETYPE_MAIN_NAME": main.get("name") or "",
        "ARCHETYPE_MAIN_CORE_TRIPLET": main.get("core_triplet") or CORE_TRIPLET_MAP.get(main_type, ""),
        "ARCHETYPE_MAIN_TAGLINE": main.get("tagline") or "",
        "ARCHETYPE_MAIN_DESC": main.get("description") or "",
        "ARCHETYPE_MAIN_STORYLINE": main.get("storyline") or "",
        "ARCHETYPE_MAIN_TRAITS": main.get("core_traits") or [],
        "ARCHETYPE_MAIN_STRENGTHS": main.get("strengths") or [],
        "ARCHETYPE_MAIN_WEAKNESSES": main.get("weaknesses") or [],
        "ARCHETYPE_MAIN_RECOMMENDATIONS": main.get("recommendations") or [],
        "ARCHETYPE_MAIN_POLITICIANS": person_links_plain(main.get("examples_person", [])),
        "ARCHETYPE_MAIN_BRANDS_IMG": [],
        "ARCHETYPE_MAIN_COLORS": main.get("color_palette") or [],
        "ARCHETYPE_MAIN_COLORS_LABEL": kolor_label_list(main.get("color_palette", [])),
        "ARCHETYPE_MAIN_VISUALS": main.get("visual_elements") or [],
        "ARCHETYPE_MAIN_KEYWORDS": main.get("keyword_messaging") or [],
        "ARCHETYPE_MAIN_SLOGANS": main.get("watchword") or [],
        "ARCHETYPE_MAIN_QUESTIONS": main.get("questions") or [],
        "ARCHETYPE_MAIN_METRIC_TEXT": main.get("report_metric_text") or "",
        "ARCHETYPE_MAIN_EXPANDED_TEXT": main.get("report_expanded_text") or "",

        "ARCHETYPE_AUX_NAME": second.get("name") or "",
        "ARCHETYPE_AUX_CORE_TRIPLET": second.get("core_triplet") or CORE_TRIPLET_MAP.get(second_type, ""),
        "ARCHETYPE_AUX_TAGLINE": second.get("tagline") or "",
        "ARCHETYPE_AUX_DESC": second.get("description") or "",
        "ARCHETYPE_AUX_STORYLINE": second.get("storyline") or "",
        "ARCHETYPE_AUX_TRAITS": second.get("core_traits") or [],
        "ARCHETYPE_AUX_STRENGTHS": second.get("strengths") or [],
        "ARCHETYPE_AUX_WEAKNESSES": second.get("weaknesses") or [],
        "ARCHETYPE_AUX_RECOMMENDATIONS": second.get("recommendations") or [],
        "ARCHETYPE_AUX_POLITICIANS": person_links_plain(second.get("examples_person", [])),
        "ARCHETYPE_AUX_BRANDS_IMG": [],
        "ARCHETYPE_AUX_COLORS": second.get("color_palette") or [],
        "ARCHETYPE_AUX_COLORS_LABEL": kolor_label_list(second.get("color_palette", [])),
        "ARCHETYPE_AUX_VISUALS": second.get("visual_elements") or [],
        "ARCHETYPE_AUX_KEYWORDS": second.get("keyword_messaging") or [],
        "ARCHETYPE_AUX_SLOGANS": second.get("watchword") or [],
        "ARCHETYPE_AUX_QUESTIONS": second.get("questions") or [],
        "ARCHETYPE_AUX_METRIC_TEXT": second.get("report_metric_text") or "",
        "ARCHETYPE_AUX_EXPANDED_TEXT": second.get("report_expanded_text") or "",

        "ARCHETYPE_SUPPLEMENT_NAME": supplement.get("name") or "",
        "ARCHETYPE_SUPPLEMENT_CORE_TRIPLET": supplement.get("core_triplet") or CORE_TRIPLET_MAP.get(supplement_type, ""),
        "ARCHETYPE_SUPPLEMENT_TAGLINE": supplement.get("tagline") or "",
        "ARCHETYPE_SUPPLEMENT_DESC": supplement.get("description") or "",
        "ARCHETYPE_SUPPLEMENT_STORYLINE": supplement.get("storyline") or "",
        "ARCHETYPE_SUPPLEMENT_TRAITS": supplement.get("core_traits") or [],
        "ARCHETYPE_SUPPLEMENT_STRENGTHS": supplement.get("strengths") or [],
        "ARCHETYPE_SUPPLEMENT_WEAKNESSES": supplement.get("weaknesses") or [],
        "ARCHETYPE_SUPPLEMENT_RECOMMENDATIONS": supplement.get("recommendations") or [],
        "ARCHETYPE_SUPPLEMENT_POLITICIANS": person_links_plain(supplement.get("examples_person", [])),
        "ARCHETYPE_SUPPLEMENT_BRANDS_IMG": [],
        "ARCHETYPE_SUPPLEMENT_COLORS": supplement.get("color_palette") or [],
        "ARCHETYPE_SUPPLEMENT_COLORS_LABEL": kolor_label_list(supplement.get("color_palette", [])),
        "ARCHETYPE_SUPPLEMENT_VISUALS": supplement.get("visual_elements") or [],
        "ARCHETYPE_SUPPLEMENT_KEYWORDS": supplement.get("keyword_messaging") or [],
        "ARCHETYPE_SUPPLEMENT_SLOGANS": supplement.get("watchword") or [],
        "ARCHETYPE_SUPPLEMENT_QUESTIONS": supplement.get("questions") or [],
        "ARCHETYPE_SUPPLEMENT_METRIC_TEXT": supplement.get("report_metric_text") or "",
        "ARCHETYPE_SUPPLEMENT_EXPANDED_TEXT": supplement.get("report_expanded_text") or "",
    }
    return context


def _doc_has_style(doc_obj, style_name: str) -> bool:
    try:
        return any(style.name == style_name for style in doc_obj.styles)
    except Exception:
        return False


def _doc_add_paragraph(doc_obj, text: str, style_name: str | None = None):
    if style_name and _doc_has_style(doc_obj, style_name):
        return doc_obj.add_paragraph(text, style=style_name)
    return doc_obj.add_paragraph(text)


def _new_document_from_template(template_candidates: list[str] | tuple[str, ...]) -> Document:
    """
    Tworzy pusty dokument z zachowaniem stylów/sekcji z pierwszego istniejącego szablonu.
    Jeśli żaden szablon nie istnieje, zwraca czysty Document().
    """
    chosen = None
    for candidate in template_candidates:
        if candidate and os.path.exists(candidate):
            chosen = candidate
            break
    doc = Document(chosen) if chosen else Document()
    try:
        body = doc._element.body
        for child in list(body):
            # zachowaj ustawienia sekcji strony
            if child.tag.endswith("sectPr"):
                continue
            body.remove(child)
    except Exception:
        pass
    return doc


def _doc_add_roman_item(doc_obj, text: str):
    """
    Dodaje wiersz z numeracją rzymską w stylu zbliżonym do list Word.
    Treść powinna już zawierać prefiks np. 'I. ...'.
    """
    p = _doc_add_paragraph(doc_obj, text)
    try:
        pf = p.paragraph_format
        pf.left_indent = Pt(20)
        pf.first_line_indent = Pt(-10)
        pf.space_after = Pt(3)
    except Exception:
        pass
    return p


def _doc_split_semicolon_items(src_lines: list[str]) -> list[str]:
    out: list[str] = []
    for ln in src_lines:
        line = str(ln or "").strip()
        if not line:
            continue
        if ";" in line:
            parts = [x.strip(" ;") for x in re.split(r"\s*;\s*", line) if x.strip(" ;")]
            out.extend(parts if parts else [line])
        else:
            out.append(line)
    return [x for x in out if x]


def _doc_add_list_items(doc_obj, items: list[str], numbered: bool = False, italic: bool = False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        txt = str(item or "").strip()
        if not txt:
            continue
        p = _doc_add_paragraph(doc_obj, txt, style_name=style)
        if italic:
            for run in p.runs:
                run.italic = True


def _doc_add_bold_prefix_line(doc_obj, text: str, allowed_labels: set[str], italic_rest: bool = False):
    raw = str(text or "").strip()
    if not raw:
        return
    m_local = re.match(r"^([^:]{1,60}):\s*(.*)$", raw)
    if not m_local:
        _doc_add_paragraph(doc_obj, raw)
        return
    label = m_local.group(1).strip()
    value = m_local.group(2).strip()
    if label not in allowed_labels:
        _doc_add_paragraph(doc_obj, raw)
        return
    p = doc_obj.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_value = p.add_run(value)
    if italic_rest:
        run_value.italic = True


def _append_expanded_subsection_docx(doc_obj, subsection_title: str, content_lines: list[str], archetype_name: str = ""):
    lines: list[str] = []
    for raw in content_lines or []:
        txt = str(raw or "").strip()
        if not txt:
            continue
        txt = re.sub(r"^[•\-\u2013\u2014]\s*", "", txt).strip()
        if txt:
            lines.append(txt)
    if not lines:
        return

    subtitle = str(subsection_title or "").strip()
    m = re.match(r"^(\d+\.\d+(?:\.\d+)?)\.", subtitle)
    sub_prefix = (m.group(1) + ".") if m else ""

    bullet_only_prefixes = {
        "1.6.", "2.1.", "2.3.", "2.4.", "2.5.",
        "3.6.", "3.8.",
        "4.1.", "4.2.", "4.3.",
        "5.2.", "5.3.",
        "6.1.", "6.2.",
        "7.1.", "7.2.", "7.4.",
        "8.2.",
        "9.2.",
    }
    numbered_prefixes = {"3.4.", "3.5.5.", "9.3."}

    if sub_prefix == "1.2.":
        for line in lines:
            _doc_add_paragraph(doc_obj, line)
        return

    if sub_prefix == "1.3.":
        _doc_add_list_items(doc_obj, lines[:2], numbered=False)
        for line in lines[2:]:
            _doc_add_paragraph(doc_obj, line)
        return

    if sub_prefix in {"1.4.", "1.5.", "8.1."}:
        _doc_add_paragraph(doc_obj, lines[0])
        _doc_add_list_items(doc_obj, _doc_split_semicolon_items(lines[1:]), numbered=False)
        return

    if sub_prefix == "2.2.":
        for line in lines:
            _doc_add_bold_prefix_line(doc_obj, line, {"Decyzje", "Tempo", "Priorytety"})
        return

    if sub_prefix == "3.2.":
        for line in lines:
            _doc_add_bold_prefix_line(doc_obj, line, {"Ton", "Emocja po kontakcie"})
        return

    if sub_prefix == "3.3.":
        idx_recommended = None
        idx_limit = None
        for i, ln in enumerate(lines):
            if idx_recommended is None and ln.casefold().startswith("zalecane zwroty"):
                idx_recommended = i
            if ln.casefold().startswith("do ograniczenia"):
                idx_limit = i
                break

        before_limit = lines if idx_limit is None else lines[:idx_limit]
        after_limit = [] if idx_limit is None else lines[idx_limit + 1:]

        if idx_recommended is not None:
            lead = before_limit[:idx_recommended]
            for x in lead:
                _doc_add_paragraph(doc_obj, x)
            p = doc_obj.add_paragraph()
            r = p.add_run("Zalecane zwroty:")
            r.bold = True
            rec_src = before_limit[idx_recommended + 1:]
            rec_inline = re.sub(r"^zalecane zwroty:\s*", "", before_limit[idx_recommended], flags=re.IGNORECASE).strip()
            if rec_inline:
                rec_src = [rec_inline] + rec_src
            _doc_add_list_items(doc_obj, _doc_split_semicolon_items(rec_src), italic=True)
        else:
            for x in before_limit:
                _doc_add_paragraph(doc_obj, x)

        if idx_limit is not None:
            p = doc_obj.add_paragraph()
            r = p.add_run("Do ograniczenia:")
            r.bold = True
            _doc_add_list_items(doc_obj, _doc_split_semicolon_items(after_limit), numbered=False)
        return

    if sub_prefix in numbered_prefixes:
        _doc_add_list_items(doc_obj, lines, numbered=True)
        return

    if sub_prefix == "7.3.":
        numbered_items: list[str] = []
        for line in lines:
            if line.casefold().startswith("zakaz:"):
                _doc_add_bold_prefix_line(doc_obj, line, {"Zakaz"})
            else:
                numbered_items.append(line)
        _doc_add_list_items(doc_obj, numbered_items, numbered=True)
        return

    if sub_prefix == "9.1.":
        do_items: list[str] = []
        dont_items: list[str] = []
        mode: str | None = None
        for ln in lines:
            low = ln.casefold()
            if low.startswith("do:"):
                mode = "do"
                continue
            if ("don't" in low) or ("don’t" in low) or ("dont" in low):
                mode = "dont"
                continue
            if mode == "do":
                do_items.extend(_doc_split_semicolon_items([ln]))
            elif mode == "dont":
                dont_items.extend(_doc_split_semicolon_items([ln]))
        if do_items:
            p = doc_obj.add_paragraph()
            p.add_run("DO:").bold = True
            _doc_add_list_items(doc_obj, do_items, numbered=False)
        if dont_items:
            p = doc_obj.add_paragraph()
            p.add_run("DON'T:").bold = True
            _doc_add_list_items(doc_obj, dont_items, numbered=False)
        if not do_items and not dont_items:
            for line in lines:
                _doc_add_paragraph(doc_obj, line)
        return

    if sub_prefix == "9.4.":
        diag_idx = None
        for i, ln in enumerate(lines):
            if ln.casefold().startswith("pytania diagnostyczne"):
                diag_idx = i
                break
        if diag_idx is None:
            _doc_add_list_items(doc_obj, _doc_split_semicolon_items(lines), numbered=False)
            return
        _doc_add_list_items(doc_obj, _doc_split_semicolon_items(lines[:diag_idx]), numbered=False)
        p = doc_obj.add_paragraph()
        p.add_run("Pytania diagnostyczne:").bold = True
        _doc_add_list_items(doc_obj, _doc_split_semicolon_items(lines[diag_idx + 1:]), numbered=False)
        return

    if sub_prefix == "8.2.":
        for item in _doc_split_semicolon_items(lines):
            p = _doc_add_paragraph(doc_obj, "", style_name="List Bullet")
            m_line = re.match(r"^([A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż]+)\s+[–-]\s+(.+)$", item)
            if m_line:
                p.add_run(m_line.group(1)).bold = True
                p.add_run(f" — {m_line.group(2)}")
            else:
                p.add_run(item)
        return

    if sub_prefix == "5.1.2.":
        groups: list[tuple[str, list[str]]] = []
        current_head = ""
        current_items: list[str] = []
        for ln in lines:
            if re.match(r"^Zestaw\s+\d+", ln, flags=re.IGNORECASE):
                if current_head or current_items:
                    groups.append((current_head, current_items))
                current_head = ln
                current_items = []
            else:
                current_items.extend(_doc_split_semicolon_items([ln]))
        if current_head or current_items:
            groups.append((current_head, current_items))

        for head, items in groups:
            if head:
                _doc_add_paragraph(doc_obj, head)
            if items:
                _doc_add_list_items(doc_obj, items, numbered=False)
        return

    if sub_prefix == "3.7.":
        for ln in lines:
            if re.match(r"^\d+\)\s+", ln):
                p = _doc_add_paragraph(doc_obj, ln)
                try:
                    p.paragraph_format.space_before = Pt(6)
                except Exception:
                    pass
            elif ln.casefold().startswith("technika:"):
                _doc_add_bold_prefix_line(doc_obj, ln, {"Technika"})
            elif ln.casefold().startswith("schemat:"):
                _doc_add_bold_prefix_line(doc_obj, ln, {"Schemat"})
            elif ln.casefold().startswith("po co?"):
                _doc_add_bold_prefix_line(doc_obj, ln, {"Po co?"})
            else:
                _doc_add_paragraph(doc_obj, ln)
        return

    if sub_prefix in bullet_only_prefixes:
        _doc_add_list_items(doc_obj, _doc_split_semicolon_items(lines), numbered=False)
        return

    for line in lines:
        _doc_add_paragraph(doc_obj, line)


def _remove_table_borders(table_obj):
    """Ukrywa obramowanie tabeli (dla siatek zdjęć w DOCX)."""
    try:
        tbl = table_obj._tbl
        tbl_pr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            elem = OxmlElement(f"w:{edge}")
            elem.set(qn("w:val"), "nil")
            borders.append(elem)
        tbl_pr.append(borders)
    except Exception:
        pass


def _add_people_photo_grid_docx(
    doc_obj,
    names: list[str],
    category: str = "person",
    columns: int = 5,
    image_width_mm: float = 20.0,
) -> bool:
    """
    Wstawia do DOCX siatkę zdjęć + podpisów.
    Zwraca True, jeśli udało się wyrenderować przynajmniej jedną pozycję.
    """
    raw = [str(x).strip() for x in (names or []) if str(x).strip()]
    if not raw:
        return False

    unique = list(dict.fromkeys(raw))
    if not unique:
        return False

    cols = max(1, int(columns))
    rows = (len(unique) + cols - 1) // cols
    table = doc_obj.add_table(rows=rows, cols=cols)
    try:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True
    except Exception:
        pass
    _remove_table_borders(table)

    rendered_any = False
    for idx, raw_name in enumerate(unique):
        row_idx = idx // cols
        col_idx = idx % cols
        cell = table.cell(row_idx, col_idx)
        try:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        except Exception:
            pass

        display_name = canonical_person_name(raw_name) if category == "person" else raw_name
        photo_path = _resolve_photo_path(raw_name, category=category)

        # domyślny pierwszy akapit w komórce
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.text = ""

        if photo_path and os.path.exists(photo_path):
            square_stream = _photo_square_top_stream(photo_path)
            if square_stream is not None:
                add_image(p_img, square_stream, width=Mm(image_width_mm))
            else:
                add_image(p_img, photo_path, width=Mm(image_width_mm))
            rendered_any = True

        p_cap = cell.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        link_url = None
        if category == "person":
            link_url = person_wikipedia_links.get(canonical_person_name(raw_name))
        if link_url:
            add_hyperlink(p_cap, display_name, link_url)
        else:
            run = p_cap.add_run(display_name)
            run.font.size = Pt(8.5)

    return rendered_any


def _append_examples_block_docx(doc_obj, label: str, examples_map: dict):
    """Render sekcji 'Przykłady archetypów' do DOCX, razem ze zdjęciami osób."""
    _doc_add_paragraph(doc_obj, f"{label}:", "Heading 4")
    examples_map = examples_map or {}

    groups = (
        ("Politycy", "person"),
        ("Marki/organizacje", None),
        ("Popkultura/postacie", "popculture"),
    )

    for group_name, category in groups:
        values = [str(v).strip() for v in (examples_map.get(group_name) or []) if str(v).strip()]
        if not values:
            continue

        _doc_add_paragraph(doc_obj, f"{group_name}:", "Heading 4")
        if category in {"person", "popculture"}:
            rendered = _add_people_photo_grid_docx(
                doc_obj,
                values,
                category=category,
                columns=5,
                image_width_mm=20.0,
            )
            if not rendered:
                _doc_add_paragraph(doc_obj, ", ".join(values))
        else:
            _doc_add_paragraph(doc_obj, ", ".join(values))


def _append_archetype_appendix(doc_tpl, appendix_items: list[tuple[str, dict]]):
    valid_items = []
    for role_label, arche_data in appendix_items:
        if not arche_data:
            continue
        if not (arche_data.get("metric_rows") or arche_data.get("expanded_sections")):
            continue
        valid_items.append((role_label, arche_data))

    if not valid_items:
        return

    doc_obj = doc_tpl.docx
    doc_obj.add_page_break()
    _doc_add_paragraph(doc_obj, "Załącznik: Rozbudowane opisy archetypów", "Heading 1")

    for role_label, arche_data in valid_items:
        arche_name = arche_data.get("name") or "Archetyp"
        _doc_add_paragraph(doc_obj, f"{role_label}: {arche_name}", "Heading 2")
        card_path = _card_file_for(arche_name)
        if card_path and os.path.exists(card_path):
            p_img = doc_obj.add_paragraph()
            add_image(p_img, str(card_path), width=Mm(84))

        metric_rows = arche_data.get("metric_rows") or []
        if metric_rows:
            _doc_add_paragraph(doc_obj, "Metryka archetypu", "Heading 3")
            for row in metric_rows:
                label = str(row.get("label", "")).strip()
                kind = row.get("kind", "text")
                value = row.get("value")
                if not label:
                    continue
                if label.casefold() == "core triplet":
                    continue

                if kind == "examples":
                    _append_examples_block_docx(doc_obj, label, value or {})
                    continue

                if isinstance(value, list):
                    vals = [str(v).strip() for v in value if str(v).strip()]
                    vals = _romanize_metric_items_if_needed(label, vals)
                    if not vals:
                        continue
                    _doc_add_paragraph(doc_obj, f"{label}:", "Heading 4")
                    for item in vals:
                        if label.casefold() == "4 filary wartości":
                            _doc_add_roman_item(doc_obj, item)
                        else:
                            _doc_add_paragraph(doc_obj, f"- {item}")
                    continue

                text_value = str(value or "").strip()
                if text_value:
                    _doc_add_paragraph(doc_obj, f"{label}: {text_value}")

        expanded_sections = arche_data.get("expanded_sections") or []
        if expanded_sections:
            _doc_add_paragraph(doc_obj, "Rozbudowany opis", "Heading 3")
            for section in expanded_sections:
                title = str(section.get("title", "")).strip()
                if title:
                    _doc_add_paragraph(doc_obj, title, "Heading 3")
                for subsection in section.get("subsections", []):
                    subtitle = str(subsection.get("title", "")).strip()
                    if subtitle:
                        _doc_add_paragraph(doc_obj, subtitle, "Heading 4")
                    _append_expanded_subsection_docx(
                        doc_obj,
                        subsection_title=subtitle,
                        content_lines=subsection.get("content", []) or [],
                        archetype_name=arche_name,
                    )


def _append_segment_profile_page(doc_tpl, segment_profile_img_path: str | None, subject_gen: str = ""):
    if not segment_profile_img_path or not os.path.exists(segment_profile_img_path):
        return
    doc_obj = doc_tpl.docx
    doc_obj.add_page_break()
    title = re.sub(r"\s{2,}", " ", f"Profil siły archetypów {subject_gen} (skala: 0-100)").strip()
    _doc_add_paragraph(doc_obj, title, "Heading 1")
    p = doc_obj.add_paragraph()
    add_image(p, segment_profile_img_path, width=Mm(170))


def _append_action_profiles_page(
    doc_tpl,
    action_profiles: list[dict[str, str]] | None,
    subject_gen: str = "",
):
    items = []
    for entry in list(action_profiles or []):
        if not isinstance(entry, dict):
            continue
        path = str(entry.get("path") or "").strip()
        if not path or not os.path.exists(path):
            continue
        role = str(entry.get("role") or "").strip()
        name = str(entry.get("name") or "").strip()
        items.append({"role": role, "name": name, "path": path})
    if not items:
        return

    doc_obj = doc_tpl.docx
    doc_obj.add_page_break()
    title = re.sub(r"\s{2,}", " ", f"Profile działania archetypów {subject_gen}").strip()
    _doc_add_paragraph(doc_obj, title, "Heading 1")

    for item in items:
        role_txt = str(item.get("role") or "").strip()
        name_txt = str(item.get("name") or "").strip()
        subtitle = f"{role_txt}: {name_txt} - profil działania".strip(": ")
        _doc_add_paragraph(doc_obj, subtitle, "Heading 3")
        p = doc_obj.add_paragraph()
        add_image(p, item["path"], width=Mm(170))


def _append_generated_descriptions_page(
    doc_tpl,
    generated_descriptions: dict[str, str] | None,
    subject_gen: str = "",
):
    if not isinstance(generated_descriptions, dict):
        return
    values_desc = str(generated_descriptions.get("valuesWheelDescription") or "").strip()
    needs_desc = str(generated_descriptions.get("needsWheelDescription") or "").strip()
    action_desc = str(generated_descriptions.get("actionProfileDescription") or "").strip()
    if not any((values_desc, needs_desc, action_desc)):
        return

    doc_obj = doc_tpl.docx
    doc_obj.add_page_break()
    _doc_add_paragraph(doc_obj, "Opisy interpretacyjne (widok panelu)", "Heading 1")

    if values_desc:
        _doc_add_paragraph(doc_obj, "Koło pragnień i wartości", "Heading 2")
        _doc_add_paragraph(doc_obj, values_desc)

    if needs_desc:
        _doc_add_paragraph(doc_obj, "Koło potrzeb", "Heading 2")
        _doc_add_paragraph(doc_obj, needs_desc)

    if action_desc:
        title = re.sub(r"\s{2,}", " ", f"Profile działania archetypów {subject_gen}").strip()
        _doc_add_paragraph(doc_obj, title, "Heading 2")
        _doc_add_paragraph(doc_obj, action_desc)


def export_word_metrics_only(
    main_type,
    second_type,
    supplement_type,
    main,
    second,
    supplement,
    person: dict | None = None,
    show_supplement: bool = True,
    segment_profile_img_path: str | None = None,
    template_path: str | None = None,
    features: dict | None = None,
    mean_scores=None,
    radar_img_path=None,
    archetype_table=None,
    num_ankiet=None,
    panel_img_path: str | None = None,
    gender_code: str = "M",
    axes_wheel_img_path: str | None = None,
    dom_color: dict | None = None,
    color_progress_img_path: str | None = None,
    archetype_stacked_img_path: str | None = None,
    capsule_columns_img_path: str | None = None,
):
    def _resolve_template(candidate: str | None) -> str | None:
        if not candidate:
            return None
        p = Path(candidate)
        if p.is_absolute():
            return str(p)
        return str((_BASE_DIR / p).resolve())

    preferred_templates: list[str] = []
    if template_path:
        resolved = _resolve_template(template_path)
        if resolved:
            preferred_templates.append(resolved)

    # Zawsze dołóż fallbacki (również gdy template_path wskazuje brakujący plik).
    # Najpierw template zgodny z płcią badanej osoby, potem fallback na drugi wariant.
    fallback_templates = _template_candidates_for_export(
        show_supplement=show_supplement,
        short_report=True,
        gender_code=gender_code,
    )
    for cand in fallback_templates:
        resolved = _resolve_template(cand)
        if resolved and resolved not in preferred_templates:
            preferred_templates.append(resolved)

    chosen_short_template = next(
        (cand for cand in preferred_templates if cand and os.path.exists(cand)),
        None,
    )
    if not chosen_short_template:
        raise FileNotFoundError(
            "Brak szablonu raportu skróconego. "
            f"Sprawdź pliki: {', '.join([x for x in preferred_templates if x])}"
        )

    # Raport skrócony renderujemy przez dedykowany szablon short,
    # aby układ/formatowanie wynikały bezpośrednio z pliku .docx.
    return export_word_docxtpl(
        main_type,
        second_type,
        supplement_type,
        features or {},
        main,
        second,
        supplement,
        mean_scores=mean_scores,
        radar_img_path=radar_img_path,
        archetype_table=archetype_table,
        num_ankiet=num_ankiet,
        panel_img_path=panel_img_path,
        person=person,
        gender_code=gender_code,
        axes_wheel_img_path=axes_wheel_img_path,
        dom_color=dom_color,
        color_progress_img_path=color_progress_img_path,
        archetype_stacked_img_path=archetype_stacked_img_path,
        capsule_columns_img_path=capsule_columns_img_path,
        segment_profile_img_path=segment_profile_img_path,
        show_supplement=show_supplement,
        template_path=chosen_short_template,
    )


def export_word_docxtpl(
    main_type,
    second_type,
    supplement_type,
    features,
    main,
    second,
    supplement,
    mean_scores=None,
    radar_img_path=None,
    archetype_table=None,
    num_ankiet=None,
    panel_img_path=None,
    person: dict | None = None,
    gender_code: str = "M",
    axes_wheel_img_path: str | None = None,
    dom_color: dict | None = None,
    color_progress_img_path: str | None = None,
    archetype_stacked_img_path: str | None = None,
    capsule_columns_img_path: str | None = None,
    segment_profile_img_path: str | None = None,
    action_profiles: list[dict[str, str]] | None = None,
    generated_descriptions: dict[str, str] | None = None,
    show_supplement: bool = True,                       # ⬅️ NOWY ARGUMENT
    template_path: str | None = None,                   # ⬅️ (opcjonalny override)
):
    # Wybór szablonu zależnie od widoczności pobocznego i płci badanej osoby
    _template = template_path or _default_template_for_export(
        show_supplement=show_supplement,
        gender_code=gender_code,
    )
    doc = DocxTemplate(_template)

    # [AUTO-FIT] policz pole składu i zaplanuj szerokości obrazków
    try:
        sect = doc.docx.part.document.sections[0]
        page_w_mm = sect.page_width.mm
        left_mm = sect.left_margin.mm
        right_mm = sect.right_margin.mm
        content_mm = page_w_mm - left_mm - right_mm  # efektywna szerokość wiersza

        # Załóż, że lewa kolumna (tabela z liczebnościami) ma ~55 mm.
        # Zostaje miejsce na 2 obrazki (radar + koło/panel) w jednym wierszu.
        left_col_mm = 55.0
        gap_between_imgs_mm = 6.0  # szpara między kolumnami w tabeli
        rest_mm = max(0.0, content_mm - left_col_mm)

        # Maksymalna szerokość jednego obrazka (po równo) – ale nie więcej niż 92 mm.
        one_img_mm = min(92.0, max(70.0, (rest_mm - gap_between_imgs_mm) / 2.0))
    except Exception:
        # Fallback gdyby sekcja nie była dostępna
        one_img_mm = 115.0

    # Radar image (mniejsze, żeby zmieściły się 3 elementy w wierszu Worda)
    if radar_img_path and os.path.exists(radar_img_path):
        radar_image = InlineImage(doc, radar_img_path, width=Mm(one_img_mm))
    else:
        radar_image = ""

    # Panel image (na bazie koła archetypów)
    panel_image = InlineImage(doc, panel_img_path, width=Mm(105)) if panel_img_path and os.path.exists(panel_img_path) else ""

    # Ikony archetypów do Word (główny/wspierający/poboczny)
    ARCHETYPE_MAIN_ICON = arche_icon_inline_for_word(doc, main_type, gender_code, height_mm=26) if main_type else ""
    ARCHETYPE_AUX_ICON = arche_icon_inline_for_word(doc, second_type, gender_code, height_mm=26) if second_type else ""
    ARCHETYPE_SUPP_ICON = arche_icon_inline_for_word(doc, supplement_type, gender_code, height_mm=26) if supplement_type else ""
    ARCHETYPE_MAIN_CARD = arche_card_inline_for_word(doc, main_type, width_mm=84) if main_type else ""
    ARCHETYPE_AUX_CARD = arche_card_inline_for_word(doc, second_type, width_mm=84) if second_type else ""
    ARCHETYPE_SUPP_CARD = arche_card_inline_for_word(doc, supplement_type, width_mm=84) if supplement_type else ""

    # Grafiki palet kolorów
    ARCHETYPE_MAIN_PALETTE_IMG = palette_inline_for_word(doc, main.get("color_palette", []))
    ARCHETYPE_AUX_PALETTE_IMG = palette_inline_for_word(doc, second.get("color_palette", []))
    ARCHETYPE_SUPP_PALETTE_IMG = palette_inline_for_word(doc, supplement.get("color_palette", []))

    # ——— najważniejsze: przekaż person →
    context = build_word_context(
        main_type, second_type, supplement_type, features, main, second, supplement,
        mean_scores, radar_image, archetype_table, num_ankiet,
        person=person,
        gender_code=gender_code,
    )

    # Wstrzyknięcie ikon i palet do szablonu DOCX
    context["ARCHETYPE_MAIN_ICON"] = ARCHETYPE_MAIN_ICON
    context["ARCHETYPE_AUX_ICON"] = ARCHETYPE_AUX_ICON
    context["ARCHETYPE_SUPP_ICON"] = ARCHETYPE_SUPP_ICON
    context["ARCHETYPE_MAIN_CARD"] = ARCHETYPE_MAIN_CARD
    context["ARCHETYPE_AUX_CARD"] = ARCHETYPE_AUX_CARD
    context["ARCHETYPE_SUPP_CARD"] = ARCHETYPE_SUPP_CARD
    context["ARCHETYPE_MAIN_CARD_IMG"] = ARCHETYPE_MAIN_CARD
    context["ARCHETYPE_AUX_CARD_IMG"] = ARCHETYPE_AUX_CARD
    context["ARCHETYPE_SUPP_CARD_IMG"] = ARCHETYPE_SUPP_CARD
    context["ARCHETYPE_SUPPLEMENT_CARD"] = ARCHETYPE_SUPP_CARD
    context["ARCHETYPE_SUPPLEMENT_CARD_IMG"] = ARCHETYPE_SUPP_CARD

    context["ARCHETYPE_MAIN_PALETTE_IMG"] = ARCHETYPE_MAIN_PALETTE_IMG
    context["ARCHETYPE_AUX_PALETTE_IMG"] = ARCHETYPE_AUX_PALETTE_IMG
    context["ARCHETYPE_SUPP_PALETTE_IMG"] = ARCHETYPE_SUPP_PALETTE_IMG

    # Jeśli nie pokazujemy pobocznego – wyczyść jego pola (na wszelki wypadek)
    if not show_supplement:
        for k in list(context.keys()):
            if k.startswith("ARCHETYPE_SUPPLEMENT_"):
                context[k] = "" if not isinstance(context[k], list) else []
        context["ARCHETYPE_SUPP_ICON"] = ""
        context["ARCHETYPE_SUPP_CARD"] = ""
        context["ARCHETYPE_SUPP_CARD_IMG"] = ""
        context["ARCHETYPE_SUPPLEMENT_CARD"] = ""
        context["ARCHETYPE_SUPPLEMENT_CARD_IMG"] = ""
        context["ARCHETYPE_SUPP_PALETTE_IMG"] = ""

    # Logotypy do Worda
    context["ARCHETYPE_MAIN_BRANDS_IMG"] = build_brands_for_word(
        doc, main.get("example_brands", []), logos_dir=logos_dir, slot_w_mm=30.0, slot_h_mm=20.0,
        pad_mm=2.0)
    context["ARCHETYPE_AUX_BRANDS_IMG"] = build_brands_for_word(
        doc, second.get("example_brands", []), logos_dir=logos_dir, slot_w_mm=30.0, slot_h_mm=20.0,
        pad_mm=2.0)
    context["ARCHETYPE_SUPPLEMENT_BRANDS_IMG"] = build_brands_for_word(
        doc, supplement.get("example_brands", []), logos_dir=logos_dir, slot_w_mm=30.0,
        slot_h_mm=20.0, pad_mm=2.0)

    context["PANEL_IMG"] = panel_image

    # Grafikę pierścienia już masz:
    context["COLOR_RING_IMG"] = InlineImage(doc, "color_ring.png", width=Mm(105))

    # Tekstowy opis dominującego koloru:
    if dom_color:
        context["DOM_COLOR_NAME"] = dom_color["name"]
        context["DOM_COLOR_EMOJI"] = dom_color["emoji"]
        context["DOM_COLOR_TITLE"] = dom_color["title"]
        context["DOM_COLOR_PCT"] = f"{dom_color['pct']:.1f}%"
        context["DOM_COLOR_ORIENT"] = dom_color["orient"]
        context["DOM_COLOR_BODY"] = dom_color["body"]
        context["DOM_COLOR_POLITICS"] = dom_color["politics"]
        context["DOM_COLOR_ARCHE"] = dom_color.get("arche") or COLOR_LONG.get(dom_color["name"],
                                                                              {}).get("arche", "")
    else:
        # Bezpieczne puste wartości, gdyby kiedyś nie było danych
        for k in ("DOM_COLOR_NAME", "DOM_COLOR_EMOJI", "DOM_COLOR_TITLE",
                  "DOM_COLOR_PCT", "DOM_COLOR_ORIENT", "DOM_COLOR_BODY", "DOM_COLOR_POLITICS"):
            context[k] = ""

    AXES_WHEEL_IMG = (
        InlineImage(doc, axes_wheel_img_path, width=Mm(105))
        if (axes_wheel_img_path and os.path.exists(axes_wheel_img_path))
        else "")

    context["AXES_WHEEL_IMG"] = AXES_WHEEL_IMG

    # Wykres pigułek (PNG na przezroczystym tle)
    if color_progress_img_path and os.path.exists(color_progress_img_path):
        context["COLOR_PROGRESS_IMG"] = (
            InlineImage(doc, color_progress_img_path, width=Mm(160))
            if color_progress_img_path and os.path.exists(color_progress_img_path) else "")

    # Skumulowany wykres słupkowy archetypów (główny/ wspierający/ poboczny)
    context["ARCHETYPE_STACKED_IMG"] = (
        InlineImage(doc, archetype_stacked_img_path, width=Mm(160))
        if (archetype_stacked_img_path and os.path.exists(archetype_stacked_img_path))
        else "")

    context["ARCHETYPE_CAPSULES_IMG"] = (
        InlineImage(doc, capsule_columns_img_path, width=Mm(170))
        if (capsule_columns_img_path and os.path.exists(capsule_columns_img_path))
        else "")

    # === Natężenie archetypu: etykiety i opisy dla Worda ===
    # Używamy 'mean_scores' przekazanych do funkcji (to słownik {archetyp: %}).
    _ms = mean_scores or {}

    def _lab_desc_for(name: str):
        if not name:
            return {"short": "", "full": "", "desc": ""}
        pct = float(_ms.get(name, 0.0))
        return interpret_archetype_intensity(pct)

    _main = _lab_desc_for(main_type)
    _aux = _lab_desc_for(second_type)
    _supp = _lab_desc_for(supplement_type)

    # Krótkie linie „Natężenie archetypu: …”
    context["ARCHETYPE_MAIN_INTENSITY_LINE"] = f"Natężenie archetypu: {_main['short']}" if _main[
        "short"] else ""
    context["ARCHETYPE_AUX_INTENSITY_LINE"] = f"Natężenie archetypu: {_aux['short']}" if _aux[
        "short"] else ""
    context["ARCHETYPE_SUPP_INTENSITY_LINE"] = f"Natężenie archetypu: {_supp['short']}" if _supp[
        "short"] else ""

    # Same etykiety (np. „Wysokie natężenie”) — gdybyś wolał użyć oddzielnie
    context["ARCHETYPE_MAIN_INTENSITY_LABEL"] = _main["full"]
    context["ARCHETYPE_AUX_INTENSITY_LABEL"] = _aux["full"]
    context["ARCHETYPE_SUPP_INTENSITY_LABEL"] = _supp["full"]

    # Opisy jakościowe (kolumna „Znaczenie i opis jakościowy”)
    context["ARCHETYPE_MAIN_INTENSITY_DESC"] = _main["desc"]
    context["ARCHETYPE_AUX_INTENSITY_DESC"] = _aux["desc"]
    context["ARCHETYPE_SUPP_INTENSITY_DESC"] = _supp["desc"]

    doc.render(context)

    appendix_payload = [
        ("Archetyp główny", main),
        ("Archetyp wspierający", second),
    ]
    if show_supplement:
        appendix_payload.append(("Archetyp poboczny", supplement))
    _append_archetype_appendix(doc, appendix_payload)
    subject_gen = (person or {}).get("GEN") or (person or {}).get("NOM") or ""
    _append_segment_profile_page(doc, segment_profile_img_path, subject_gen=subject_gen)
    _append_action_profiles_page(doc, action_profiles, subject_gen=subject_gen)
    _append_generated_descriptions_page(doc, generated_descriptions, subject_gen=subject_gen)

    # Hiperłącza do osób: tylko jeśli akapit zawiera wyłącznie nazwisko (z opcjonalnym prefiksem "- ").
    for para in doc.paragraphs:
        raw_text = (para.text or "").strip()
        if not raw_text:
            continue

        bullet_prefix = ""
        candidate = raw_text
        if raw_text.startswith("- "):
            bullet_prefix = "- "
            candidate = raw_text[2:].strip()

        candidate = canonical_person_name(candidate)
        if candidate in person_wikipedia_links:
            para.clear()
            if bullet_prefix:
                para.add_run(bullet_prefix)
            add_hyperlink(para, candidate, person_wikipedia_links[candidate])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def _find_soffice() -> str | None:
    """
    Znajdź binarkę 'soffice' dla LibreOffice (Linux). Na Windows zwraca None.
    Respektuje zmienną środowiskową SOFFICE_BIN.
    """
    import sys, os, shutil
    if sys.platform.startswith("win"):
        return None
    # 1) ręcznie wskazana ścieżka
    env = os.environ.get("SOFFICE_BIN")
    if env and os.path.isfile(env):
        return env
    # 2) PATH / najczęstsze lokalizacje
    candidates = [
        shutil.which("soffice"),
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
        "/snap/bin/libreoffice",
    ]
    for c in candidates:
        if c and os.path.exists(c):
            return c
    return None

def word_to_pdf(docx_bytes_io, soffice_bin: str | None = None):
    import sys
    import tempfile, os, shutil
    from io import BytesIO
    from pathlib import Path
    if sys.platform.startswith("win32"):
        # --- WINDOWS: bez zmian, Word + docx2pdf (honoruje osadzone fonty) ---
        import pythoncom
        from docx2pdf import convert
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "raport.docx")
            pdf_path = os.path.join(tmpdir, "raport.pdf")
            with open(docx_path, "wb") as f:
                f.write(docx_bytes_io.getbuffer())
            pythoncom.CoInitialize()
            convert(docx_path, pdf_path)
            with open(pdf_path, "rb") as f:
                return BytesIO(f.read())
    else:
        # --- LINUX: najpierw do-instaluj fonty do ~/.local/share/fonts, odśwież fontconfig ---
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "raport.docx")
            pdf_path = os.path.join(tmpdir, "raport.pdf")
            with open(docx_path, "wb") as f:
                f.write(docx_bytes_io.getbuffer())

            # 1) Skopiuj TTF-y do prywatnego katalogu fontów
            project_font_dir = Path(__file__).with_name("assets") / "fonts"
            user_font_dir = Path.home() / ".local" / "share" / "fonts" / "ap48"
            user_font_dir.mkdir(parents=True, exist_ok=True)
            if project_font_dir.exists():
                for src in project_font_dir.glob("*.ttf"):
                    dst = user_font_dir / src.name
                    try:
                        shutil.copyfile(src, dst)
                    except Exception:
                        pass

            # 2) Odśwież cache fontów
            try:
                import subprocess
                subprocess.run(["fc-cache", "-f", "-v"], check=False, capture_output=True)
            except Exception:
                pass  # jeśli fc-cache niedostępny w środowisku, i tak próbujemy dalej

            # 3) Konwersja LibreOffice (z możliwością podania ścieżki)
            try:
                cmd = [soffice_bin or "soffice", "--headless", "--convert-to", "pdf",
                       "--outdir", tmpdir, docx_path]
                result = subprocess.run(cmd, capture_output=True)
                if result.returncode != 0 or not os.path.isfile(pdf_path):
                    raise RuntimeError(
                        "LibreOffice PDF error: " + result.stderr.decode(errors="ignore"))
            except FileNotFoundError:
                raise RuntimeError("LibreOffice (soffice) nie jest dostępny w systemie.")
            with open(pdf_path, "rb") as f:
                return BytesIO(f.read())



def is_color_dark(color_hex):
    if color_hex is None:
        return False
    if not color_hex.startswith('#') or len(color_hex) not in (7, 4):
        return False
    h = color_hex.lstrip('#')
    if len(h) == 3:
        h = ''.join([c*2 for c in h])
    r, g, b = tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
    lum = 0.2126*r + 0.7152*g + 0.0722*b
    return lum < 110

def _hex_to_rgb01(color_hex: str) -> tuple[float, float, float]:
    h = str(color_hex or "").strip().lstrip("#")
    if len(h) == 3:
        h = "".join([c * 2 for c in h])
    if len(h) != 6:
        return (0.0, 0.0, 0.0)
    try:
        r, g, b = tuple(int(h[i:i+2], 16) / 255.0 for i in (0, 2, 4))
    except Exception:
        return (0.0, 0.0, 0.0)
    return (r, g, b)

def _rgb01_to_hex(r: float, g: float, b: float) -> str:
    rr = max(0, min(255, int(round(r * 255))))
    gg = max(0, min(255, int(round(g * 255))))
    bb = max(0, min(255, int(round(b * 255))))
    return f"#{rr:02X}{gg:02X}{bb:02X}"

def _blend_hex(color_a: str, color_b: str, mix_b: float) -> str:
    mix = max(0.0, min(1.0, float(mix_b)))
    ra, ga, ba = _hex_to_rgb01(color_a)
    rb, gb, bb = _hex_to_rgb01(color_b)
    r = (1.0 - mix) * ra + mix * rb
    g = (1.0 - mix) * ga + mix * gb
    b = (1.0 - mix) * ba + mix * bb
    return _rgb01_to_hex(r, g, b)

def _soften_aggressive_main_bg(primary_hex: str, secondary_hex: str | None = None) -> str:
    """
    Dla bardzo agresywnych czerwieni (jak #E10600) przyciemnia tło,
    żeby karta pozostała czytelna i mniej męcząca wizualnie.
    """
    base = str(primary_hex or "").upper().strip()
    if not base.startswith("#") or len(base) not in (4, 7):
        return base or "#2B2D41"

    r, g, b = _hex_to_rgb01(base)
    hue, light, sat = colorsys.rgb_to_hls(r, g, b)
    hue_deg = (hue * 360.0) % 360.0
    is_red_zone = (hue_deg <= 24.0) or (hue_deg >= 346.0)
    is_aggressive_red = is_red_zone and (r >= 0.78 and g <= 0.30 and b <= 0.30) and sat >= 0.72 and light >= 0.22
    if not is_aggressive_red:
        return base

    target = str(secondary_hex or "").upper().strip()
    if not (target.startswith("#") and len(target) in (4, 7)):
        target = "#1F2937"
    return _blend_hex(base, target, 0.58)

def palette_boxes_html(palette, color_name_map=COLOR_NAME_MAP):
    if not isinstance(palette, list) or not palette:
        return ""

    def _sort_palette_logical(values: list[str]) -> list[str]:
        cleaned = _dedupe_hex_palette(values)
        if not cleaned:
            return []

        def _key(hexcode: str):
            h = hexcode.lstrip("#")
            r = int(h[0:2], 16) / 255.0
            g = int(h[2:4], 16) / 255.0
            b = int(h[4:6], 16) / 255.0
            hue, light, sat = colorsys.rgb_to_hls(r, g, b)
            # Najpierw kolory (po kole barw), na końcu neutrals (szarości/czerń/biel)
            if sat < 0.085:
                return (2, light, hue)
            hue_deg = (hue * 360.0) % 360.0
            hue_bucket = int((hue_deg + 15.0) // 30.0) % 12
            return (1, hue_bucket, -sat, light)

        return sorted(cleaned, key=_key)

    def label_for(code):
        name = color_name_map.get(str(code).upper(), str(code))
        return f"{name} ({code})"

    boxes = []
    logical_palette = _sort_palette_logical([str(c).upper() for c in palette])
    for hexcode in logical_palette:
            txt = "#111111" if not is_color_dark(hexcode) else "#FFFFFF"
            shadow = ""  # ← zero cienia pod tekstem
            boxes.append(
                f"<div class='ap-box' style='background:{hexcode};'>"
                f"<span style='color:{txt};'>{label_for(hexcode)}</span>"
                f"</div>"
            )
    return (
                "<style>"
                ".ap-palette{display:flex;gap:10px;flex-wrap:wrap;margin:8px 0 12px 0;}"
                ".ap-palette .ap-box{width:110px;height:65px;border-radius:8px;"
                "display:flex;align-items:center;justify-content:center;"
                "text-align:center;padding:6px;border:1px solid rgba(0,0,0,.08);"
                "box-shadow:none;}"  # ← usunięty cień
                ".ap-palette .ap-box span{font-weight:700;font-size:12.5px;line-height:1.25}"
                "@media (max-width:700px){ .ap-palette .ap-box{width:120px;height:76px} }"
                "</style>"
                "<div class='ap-palette'>" + "".join(boxes) + "</div>"
        )


import base64

@lru_cache(maxsize=512)
def _file_to_data_uri(path: str, cache_buster: str = "") -> str:
    _ = cache_buster  # parametr celowo używany tylko do klucza cache
    ext = Path(path).suffix.lower()
    mime = {
        ".svg": "image/svg+xml",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
    }.get(ext, "application/octet-stream")
    blob = Path(path).read_bytes()
    b64 = base64.b64encode(blob).decode("ascii")
    return f"data:{mime};base64,{b64}"


@lru_cache(maxsize=512)
def _photo_to_data_uri(path: str, max_px: int = 220) -> str:
    ext = Path(path).suffix.lower()
    if ext not in {".jpg", ".jpeg", ".png", ".webp", ".gif", ".tif", ".tiff"}:
        return _file_to_data_uri(path)
    try:
        with Image.open(path) as im:
            im = im.convert("RGB")
            resample = getattr(Image, "Resampling", Image).LANCZOS
            im.thumbnail((max_px, max_px), resample)
            out = BytesIO()
            im.save(out, format="JPEG", quality=82, optimize=True)
            b64 = base64.b64encode(out.getvalue()).decode("ascii")
            return f"data:image/jpeg;base64,{b64}"
    except Exception:
        return _file_to_data_uri(path)


def build_brand_icons_html(brand_names, logos_dir, label_color: str = "#f3f6ff"):
    import os
    out_html = (
        '<div style="display:flex;flex-wrap:wrap;gap:12px 18px;align-items:flex-start;'
        'margin-top:6px;margin-bottom:8px;">'
    )
    for brand in brand_names:
        path = get_logo_svg_path(brand, logos_dir)
        if path and os.path.exists(path):
            try:
                data_uri = _file_to_data_uri(path)
                img_tag = (
                    f"<img src='{data_uri}' alt='{html.escape(str(brand))}' "
                    "style='height:38px;max-width:96px;object-fit:contain;display:block;margin:0 auto 3px;'/>"
                )
                out_html += (
                    f"<span title='{html.escape(str(brand))}' style='display:flex;flex-direction:column;"
                    "align-items:center;min-width:56px;padding:2px 4px;'>"
                    f"{img_tag}"
                    f"<span style='font-size:0.86em;line-height:1.2;color:{label_color};'>{html.escape(str(brand))}</span>"
                    "</span>"
                )
            except Exception:
                out_html += (
                    f"<span style='font-size:0.93em;color:{label_color};opacity:.95;padding:2px 4px;'>"
                    f"{html.escape(str(brand))}</span>"
                )
        else:
            out_html += (
                f"<span style='font-size:0.93em;color:{label_color};opacity:.80;padding:2px 4px;'>"
                f"{html.escape(str(brand))}</span>"
            )
    out_html += '</div>'
    return out_html

def person_links_html(person_list):
    if not person_list:
        return ""
    return ', '.join(person_link(name) for name in person_list)


PHOTO_DIR_PERSON = Path(__file__).with_name("assets") / "foto_person"
PHOTO_DIR_POPCULTURE = Path(__file__).with_name("assets") / "foto_popculture"
PHOTO_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".tif", ".tiff"}

POPCULTURE_IMAGE_ALIASES = {
    "Jaś Fasola": "Mr. Bean",
    "Morfeusz (z „Matrixa”)": "Morpheus (The Matrix)",
    "Mistrz Yoda": "Yoda",
    "Oprah": "Oprah Winfrey",
    "księżna Diana": "Diana Princess of Wales",
    "Matka Teresa": "Mother Teresa",
    "Forest Gump": "Forrest Gump",
    "Ojciec Chrzestny": "The Godfather",
    "Indiana Jones": "Indiana Jones (character)",
    "Rambo": "John Rambo",
    "Banksy": "Girl with Balloon",
    "Kuba Wojewódzki": "Kuba Wojewódzki",
    "Marilyn Monroe": "Marilyn Monroe",
    "Steven Spielberg": "Steven Spielberg",
    "Samwise Gamgee": "Sean Astin",
    'Samwise Gamgee ("Sam") towarzysz Frodo z Władcy Pierścieni': "Sean Astin",
    "Dr Quinn": "Dr Quinn",
    "Dr. Quinn": "Dr Quinn",
}


def _photo_norm_key(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(value or "")).encode("ascii", "ignore").decode("ascii")
    normalized = normalized.lower()
    normalized = normalized.replace("’", "").replace("'", "").replace('"', "")
    normalized = re.sub(r"\(.*?\)", " ", normalized)
    normalized = re.sub(r"[^a-z0-9]+", " ", normalized)
    return re.sub(r"\s+", " ", normalized).strip()


def _photo_variants(name: str, category: str) -> list[str]:
    base = str(name or "").strip()
    variants = [base]
    if category == "person":
        canonical = canonical_person_name(base)
        if canonical:
            variants.append(canonical)
    stripped = re.sub(r"\(.*?\)", "", base).strip(" -")
    if stripped:
        variants.append(stripped)
    short = re.split(r"\s+-\s+|\s+towarzysz\s+", stripped, maxsplit=1)[0].strip()
    if short:
        variants.append(short)
    if category == "popculture":
        alias = POPCULTURE_IMAGE_ALIASES.get(base) or POPCULTURE_IMAGE_ALIASES.get(stripped)
        if alias:
            variants.append(alias)
    out: list[str] = []
    seen: set[str] = set()
    for item in variants:
        key = _photo_norm_key(item)
        if key and key not in seen:
            seen.add(key)
            out.append(item)
    return out


def _photo_lookup_index(folder: str) -> dict[str, str]:
    idx: dict[str, tuple[int, int, str]] = {}
    folder_path = Path(folder)
    if not folder_path.exists():
        return {}
    for p in folder_path.iterdir():
        if p.suffix.lower() not in PHOTO_EXTS or not p.is_file():
            continue
        stem = p.stem
        stem_base = re.sub(r"-\d+$", "", stem)
        base_key = _photo_norm_key(stem_base.replace("_", " "))
        if not base_key:
            continue
        compact_key = re.sub(r"[^a-z0-9]+", "", base_key)
        rank = (0 if stem == stem_base else 1, len(stem))
        for key in {base_key, compact_key}:
            if not key:
                continue
            current = idx.get(key)
            if current is None or rank < (current[0], current[1]):
                idx[key] = (rank[0], rank[1], str(p))
    return {k: v[2] for k, v in idx.items()}


def _find_local_photo(name: str, category: str) -> str | None:
    folder = PHOTO_DIR_POPCULTURE if category == "popculture" else PHOTO_DIR_PERSON
    index = _photo_lookup_index(str(folder))
    for variant in _photo_variants(name, category):
        key = _photo_norm_key(variant)
        compact_key = re.sub(r"[^a-z0-9]+", "", key)
        path = index.get(key) or index.get(compact_key)
        if path:
            return path
    tokens = [t for t in _photo_norm_key(name).split(" ") if len(t) >= 4]
    if len(tokens) >= 2:
        for key, path in index.items():
            if all(token in key for token in tokens[:2]):
                return path
    return None


@lru_cache(maxsize=256)
def _download_wikipedia_photo(name: str, category: str) -> str | None:
    target_dir = PHOTO_DIR_POPCULTURE if category == "popculture" else PHOTO_DIR_PERSON
    target_dir.mkdir(parents=True, exist_ok=True)

    user_agent = {"User-Agent": "archetypy-admin/1.0 (image-fetch)"}
    for variant in _photo_variants(name, category):
        title = quote(variant.replace(" ", "_"))
        for lang in ("pl", "en"):
            summary_url = f"https://{lang}.wikipedia.org/api/rest_v1/page/summary/{title}"
            try:
                resp = requests.get(summary_url, headers=user_agent, timeout=3)
                if resp.status_code != 200:
                    continue
                payload = resp.json()
                thumb = (payload.get("thumbnail") or {}).get("source") or (payload.get("originalimage") or {}).get(
                    "source"
                )
                if not thumb:
                    continue
                img_resp = requests.get(thumb, headers=user_agent, timeout=6)
                if img_resp.status_code != 200 or not img_resp.content:
                    continue
                ext = Path(thumb.split("?")[0]).suffix.lower()
                if ext not in PHOTO_EXTS:
                    ext = ".jpg"
                filename = f"{_logo_norm_key(variant) or 'photo'}{ext}"
                out_path = target_dir / filename
                out_path.write_bytes(img_resp.content)
                return str(out_path)
            except Exception:
                continue
    return None


def _resolve_photo_path(name: str, category: str) -> str | None:
    local = _find_local_photo(name, category)
    if local:
        return local
    return _download_wikipedia_photo(name, category)


def _photo_square_top_stream(path: str, size_px: int = 360) -> BytesIO | None:
    """
    Kadruje zdjęcie do 1:1 bez rozciągania.
    - pionowe: obcina z dołu (priorytet góry),
    - poziome: centralnie obcina boki.
    """
    try:
        with Image.open(path) as im:
            if im.mode not in {"RGB", "RGBA"}:
                im = im.convert("RGBA" if "A" in im.getbands() else "RGB")

            w, h = im.size
            side = min(w, h)
            if w > h:
                left = (w - side) // 2
                box = (left, 0, left + side, side)
            else:
                box = (0, 0, side, side)

            im = im.crop(box)
            resample = getattr(Image, "Resampling", Image).LANCZOS
            im = im.resize((size_px, size_px), resample)

            out = BytesIO()
            im.save(out, format="PNG")
            out.seek(0)
            return out
    except Exception:
        return None


def _people_photo_grid_html(names: list[str], category: str = "person") -> str:
    unique_names = list(dict.fromkeys([str(x).strip() for x in (names or []) if str(x).strip()]))
    if not unique_names:
        return ""

    cards: list[str] = []
    for raw_name in unique_names:
        caption = canonical_person_name(raw_name) if category == "person" else raw_name
        path = _resolve_photo_path(raw_name, category=category)

        if path and Path(path).exists():
            try:
                data_uri = _photo_to_data_uri(path)
                avatar = f"<img src='{data_uri}' alt='{html.escape(caption)}' />"
            except Exception:
                avatar = "<div class='ap-face-ph'>?</div>"
        else:
            initials = "".join(part[0] for part in caption.split()[:2]).upper() or "?"
            avatar = f"<div class='ap-face-ph'>{html.escape(initials)}</div>"

        name_html = html.escape(caption)
        if category == "person":
            link_url = person_wikipedia_links.get(canonical_person_name(raw_name))
            if link_url:
                name_html = f"<a href='{link_url}' target='_blank'>{name_html}</a>"

        cards.append(
            "<div class='ap-face-card'>"
            f"<div class='ap-face-avatar'>{avatar}</div>"
            f"<div class='ap-face-name'>{name_html}</div>"
            "</div>"
        )

    if not cards:
        return ""
    return f"<div class='ap-face-grid'>{''.join(cards)}</div>"


def _fallback_metric_rows(archetype_data: dict) -> list[dict]:
    return [
        {"label": "Esencja", "kind": "text", "value": archetype_data.get("description", "")},
        {"label": "Rozbudowany opis", "kind": "text", "value": archetype_data.get("storyline", "")},
        {"label": "Kluczowe atrybuty", "kind": "list", "value": archetype_data.get("core_traits", [])},
        {"label": "Atuty", "kind": "list", "value": archetype_data.get("strengths", [])},
        {"label": "Słabości", "kind": "list", "value": archetype_data.get("weaknesses", [])},
        {
            "label": "Słowa-klucze (Talking points)",
            "kind": "list",
            "value": archetype_data.get("keyword_messaging", []),
        },
        {"label": "Sygnatury wizualne", "kind": "list", "value": archetype_data.get("visual_elements", [])},
        {
            "label": "Przykłady archetypów",
            "kind": "examples",
            "value": {
                "Politycy": archetype_data.get("examples_person", []),
                "Marki/organizacje": archetype_data.get("example_brands", []),
                "Popkultura/postacie": archetype_data.get("metric_examples", {}).get("Popkultura/postacie", []),
            },
        },
        {"label": "Paleta kolorów (HEX)", "kind": "colors", "value": archetype_data.get("color_palette", [])},
    ]


def _split_metric_line_items(text: str) -> list[str]:
    pieces: list[str] = []
    for chunk in re.split(r"\n+|[;]", str(text or "")):
        clean = re.sub(r"^[•\-\u2013\u2014]\s*", "", chunk.strip())
        clean = clean.strip(" ,.")
        if clean:
            pieces.append(clean)
    return pieces


def _story_antagonist_html(value: str) -> str:
    storyline: list[str] = []
    antagonist: list[str] = []
    mode = "storyline"
    for line in _split_metric_line_items(value):
        low = line.casefold()
        if low.startswith("storyline"):
            mode = "storyline"
            rest = line.split(":", 1)[1].strip() if ":" in line else ""
            if rest:
                storyline.append(rest)
            continue
        if low.startswith("antagonista"):
            mode = "antagonista"
            rest = line.split(":", 1)[1].strip() if ":" in line else ""
            if rest:
                antagonist.append(rest)
            continue
        if mode == "antagonista":
            antagonist.append(line)
        else:
            storyline.append(line)

    if not storyline and value:
        storyline = [str(value).strip()]

    story_html = (
        "<ul class='ap-simple-list ap-story-list'>"
        + "".join(f"<li>{html.escape(item)}</li>" for item in storyline if item)
        + "</ul>"
        if storyline
        else "<span style='color:#7c8799;'>—</span>"
    )
    ant_html = (
        "<ul class='ap-simple-list ap-story-list'>"
        + "".join(f"<li>{html.escape(item)}</li>" for item in antagonist if item)
        + "</ul>"
        if antagonist
        else "<span style='color:#7c8799;'>—</span>"
    )
    return (
        "<div class='ap-story-antag'>"
        "<div class='ap-story-head'>Storyline:</div>"
        f"{story_html}"
        "<div class='ap-story-head'>Antagonista (z czym walczy?):</div>"
        f"{ant_html}"
        "</div>"
    )


COMPACT_METRIC_LABELS_CF = {
    "grupa",
    "podstawowe pragnienie",
    "cel",
    "największa obawa",
    "strategia",
    "pułapka",
    "dar",
    "cień",
}


def _metric_text_plain(value) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    if not text:
        return ""
    parts = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    return " ".join(parts).strip()


def _to_roman(num: int) -> str:
    """Konwersja liczby całkowitej (>=1) do zapisu rzymskiego."""
    if num <= 0:
        return str(num)
    mapping = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    out = []
    n = int(num)
    for value, sym in mapping:
        while n >= value:
            out.append(sym)
            n -= value
    return "".join(out)


def _romanize_metric_items_if_needed(label_raw: str, items: list[str]) -> list[str]:
    """
    Dla pola '4 filary wartości' wymusza numerację: I., II., III., IV.
    Dla innych pól zwraca bez zmian.
    """
    label_cf = str(label_raw or "").strip().casefold()
    if label_cf != "4 filary wartości":
        return items
    normalized: list[str] = []
    for idx, item in enumerate(items, start=1):
        clean = re.sub(r"^\s*(?:[IVXLCDM]+\.|\d+\.)\s*", "", str(item or "").strip(), flags=re.IGNORECASE)
        if clean:
            normalized.append(f"{_to_roman(idx)}. {clean}")
    return normalized


def _metric_row_html(row: dict, archetype_data: dict, text_color: str) -> str:
    label_raw = str(row.get("label", "")).strip()
    if label_raw.casefold() == "core triplet":
        return ""
    label = html.escape(label_raw)
    kind = row.get("kind", "text")
    value = row.get("value")

    if kind == "examples":
        examples = value or {}
        politicians = [canonical_person_name(x) for x in (examples.get("Politycy") or [])]
        brands = examples.get("Marki/organizacje") or []
        popculture = examples.get("Popkultura/postacie") or []

        pol_html = ", ".join(person_link(name) for name in politicians) if politicians else "—"
        pol_faces_html = _people_photo_grid_html(politicians, category="person")
        brands_html = (
            build_brand_icons_html(brands, logos_dir, label_color=text_color)
            if brands
            else "<span style='color:#7c8799;font-size:.97em;'>—</span>"
        )
        pop_faces_html = _people_photo_grid_html(popculture, category="popculture")
        pop_html = ", ".join(html.escape(x) for x in popculture) if popculture else "—"
        return f"""
        <div class="ap-metric-row">
            <div class="ap-metric-label">{label}</div>
            <div class="ap-metric-value">
                <div class="ap-example-head ap-example-head-first"><b>Politycy:</b></div>
                {pol_faces_html if pol_faces_html else f"<div>{pol_html}</div>"}
                <div class="ap-example-head"><b>Marki/organizacje:</b></div>
                {brands_html}
                <div class="ap-example-head"><b>Popkultura/postacie:</b></div>
                {pop_faces_html if pop_faces_html else f"<div>{pop_html}</div>"}
            </div>
        </div>
        """

    if kind == "colors":
        colors = [str(c).upper() for c in (value or archetype_data.get("color_palette", [])) if str(c).strip()]
        palette_html = palette_boxes_html(colors) if colors else "<span style='color:#7c8799;'>—</span>"
        return f"""
        <div class="ap-metric-row">
            <div class="ap-metric-label">{label}</div>
            <div class="ap-metric-value">
                {palette_html}
            </div>
        </div>
        """

    if isinstance(value, list):
        items = [str(item).strip() for item in value if str(item).strip()]
        items = _romanize_metric_items_if_needed(label_raw, items)
        label_cf = label_raw.casefold()
        row_extra_cls = ""

        if not items:
            value_html = "<span style='color:#7c8799;'>—</span>"
        elif label_cf == "4 filary wartości":
            clean_items = [
                re.sub(r"^\s*(?:[IVXLCDM]+\.\s*|\d+\.\s*)", "", item, flags=re.IGNORECASE).strip()
                for item in items
            ]
            clean_items = [x for x in clean_items if x]
            value_html = (
                "<ol class='ap-roman-list'>"
                + "".join(f"<li>{html.escape(item)}</li>" for item in clean_items)
                + "</ol>"
            )
        elif label_cf in {"kluczowe atrybuty", "słowa-klucze (talking points)"}:
            row_extra_cls = " ap-metric-row-chips"
            chips = "".join(f"<span class='ap-pill'>{html.escape(item)}</span>" for item in items)
            value_html = f"<div class='ap-pill-wrap'>{chips}</div>"
        elif label_cf == "slogany (taglines)":
            value_html = (
                "<ul class='ap-simple-list'>"
                + "".join(f"<li><i>{html.escape(item)}</i></li>" for item in items)
                + "</ul>"
            )
        elif label_cf == "atuty":
            value_html = (
                "<ul class='ap-qual-list ap-qual-pos'>"
                + "".join(
                    f"<li><span class='ap-qual-ico'>✅</span><span>{html.escape(item)}</span></li>" for item in items
                )
                + "</ul>"
            )
        elif label_cf == "słabości":
            value_html = (
                "<ul class='ap-qual-list ap-qual-neg'>"
                + "".join(
                    f"<li><span class='ap-qual-ico'>❌</span><span>{html.escape(item)}</span></li>" for item in items
                )
                + "</ul>"
            )
        elif label_cf == "sygnatury wizualne":
            value_html = "<p>" + html.escape(", ".join(items)) + "</p>"
        else:
            bullets = "".join(f"<li>{html.escape(item)}</li>" for item in items)
            value_html = f"<ul class='ap-simple-list'>{bullets}</ul>"
    else:
        row_extra_cls = ""
        text = str(value or "").strip()
        if not text:
            value_html = "<span style='color:#7c8799;'>—</span>"
        elif label_raw.casefold() == "rozbudowany opis":
            paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
            value_html = "".join(f"<p><i>{html.escape(p)}</i></p>" for p in paragraphs)
        elif label_raw.casefold() == "slogany (taglines)":
            items = _split_metric_line_items(text) or [text]
            value_html = (
                "<ul class='ap-simple-list'>"
                + "".join(f"<li><i>{html.escape(item)}</i></li>" for item in items if str(item).strip())
                + "</ul>"
            )
        elif label_raw.casefold() == "oś narracyjna i antagonista":
            value_html = _story_antagonist_html(text)
        elif label_raw.casefold() == "4 filary wartości":
            items = _split_metric_line_items(text)
            items = _romanize_metric_items_if_needed(label_raw, items)
            if items:
                clean_items = [
                    re.sub(r"^\s*(?:[IVXLCDM]+\.\s*|\d+\.\s*)", "", item, flags=re.IGNORECASE).strip()
                    for item in items
                ]
                clean_items = [x for x in clean_items if x]
                value_html = (
                    "<ol class='ap-roman-list'>"
                    + "".join(f"<li>{html.escape(item)}</li>" for item in clean_items)
                    + "</ol>"
                )
            else:
                value_html = f"<p>{html.escape(text)}</p>"
        else:
            paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
            value_html = "".join(f"<p>{html.escape(p)}</p>" for p in paragraphs)

    return f"""
    <div class="ap-metric-row{row_extra_cls}">
        <div class="ap-metric-label">{label}</div>
        <div class="ap-metric-value">{value_html}</div>
    </div>
    """


SECTION_ICON_MAP = {
    "1": "🧠",
    "2": "🏛️",
    "3": "🗣️",
    "4": "✅",
    "5": "🎨",
    "6": "🧭",
    "7": "🛡️",
    "8": "⚖️",
    "9": "🧰",
}


def _norm_archetype_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", _slug_pl(value or ""))


@lru_cache(maxsize=96)
def _archetype_card_data_uri(archetype_name: str, card_cache_token: str = "") -> str:
    _ = card_cache_token  # parametr celowo używany tylko do klucza cache
    try:
        path = _card_file_for(archetype_name)
        if path and path.exists():
            stat = path.stat()
            file_token = f"{stat.st_mtime_ns}:{stat.st_size}"
            return _file_to_data_uri(str(path), file_token)
    except Exception:
        return ""
    return ""


def _archetype_card_cache_token(archetype_name: str) -> str:
    try:
        path = _card_file_for(archetype_name)
        if path and path.exists():
            stat = path.stat()
            return f"{path}:{stat.st_mtime_ns}:{stat.st_size}"
    except Exception:
        pass
    return f"missing:{_norm_archetype_key(archetype_name)}"


def _expanded_subsection_content_html(content_lines: list[str], subsection_title: str = "", archetype_name: str = "") -> str:
    if not content_lines:
        return ""

    lines = []
    for raw in content_lines:
        txt = str(raw or "").strip()
        if not txt:
            continue
        txt = re.sub(r"^[•\-\u2013\u2014]\s*", "", txt).strip()
        if txt:
            lines.append(txt)
    if not lines:
        return ""

    subtitle = str(subsection_title or "").strip()
    m = re.match(r"^(\d+\.\d+(?:\.\d+)?)\.", subtitle)
    sub_prefix = (m.group(1) + ".") if m else ""

    bullet_only_prefixes = {
        "1.6.", "2.1.", "2.3.", "2.4.", "2.5.",
        "3.6.", "3.8.",
        "4.1.", "4.2.", "4.3.",
        "5.2.", "5.3.",
        "6.1.", "6.2.",
        "7.1.", "7.2.", "7.4.",
        "8.2.",
        "9.2.",
    }
    numbered_prefixes = {"3.4.", "3.5.5.", "9.3."}

    def p(text: str, cls: str = "") -> str:
        cls_attr = f" class='{cls}'" if cls else ""
        return f"<p{cls_attr}>{html.escape(text)}</p>"

    def p_bold_prefix(text: str, allowed_labels: set[str], cls: str = "") -> str:
        m_local = re.match(r"^([^:]{1,40}):\s*(.*)$", text)
        if not m_local:
            return p(text, cls=cls)
        label = m_local.group(1).strip()
        value = m_local.group(2).strip()
        if label not in allowed_labels:
            return p(text, cls=cls)
        cls_attr = f" class='{cls}'" if cls else ""
        return f"<p{cls_attr}><b>{html.escape(label)}:</b> {html.escape(value)}</p>"

    def split_semicolon_items(src_lines: list[str]) -> list[str]:
        out: list[str] = []
        for ln in src_lines:
            if ";" in ln:
                parts = [x.strip(" ;") for x in re.split(r"\s*;\s*", ln) if x.strip(" ;")]
                out.extend(parts if parts else [ln])
            else:
                out.append(ln)
        return [x for x in out if x]

    def ul(items: list[str], cls: str = "ap-ext-list") -> str:
        if not items:
            return ""
        return f"<ul class='{cls}'>" + "".join(f"<li>{html.escape(i)}</li>" for i in items if i.strip()) + "</ul>"

    def ol(items: list[str], cls: str = "ap-ext-ol") -> str:
        if not items:
            return ""
        return f"<ol class='{cls}'>" + "".join(f"<li>{html.escape(i)}</li>" for i in items if i.strip()) + "</ol>"

    if sub_prefix == "1.2.":
        return "".join(p(x) for x in lines)

    if sub_prefix == "1.3.":
        first_two = lines[:2]
        rest = lines[2:]
        return ul(first_two) + "".join(p(x) for x in rest)

    if sub_prefix in {"1.4.", "1.5.", "8.1."}:
        lead = lines[0]
        rest_items = split_semicolon_items(lines[1:])
        return p(lead) + ul(rest_items)

    if sub_prefix == "2.2.":
        allowed = {"Decyzje", "Tempo", "Priorytety"}
        return "".join(p_bold_prefix(x, allowed) for x in lines)

    if sub_prefix == "3.2.":
        allowed = {"Ton", "Emocja po kontakcie"}
        return "".join(p_bold_prefix(x, allowed) for x in lines)

    if sub_prefix == "3.3.":
        blocks: list[str] = []
        idx_recommended = None
        idx_limit = None
        for i, ln in enumerate(lines):
            if idx_recommended is None and ln.casefold().startswith("zalecane zwroty"):
                idx_recommended = i
            if ln.casefold().startswith("do ograniczenia"):
                idx_limit = i
                break

        before_limit = lines if idx_limit is None else lines[:idx_limit]
        after_limit = [] if idx_limit is None else lines[idx_limit + 1:]

        if idx_recommended is not None:
            lead = before_limit[:idx_recommended]
            rec_src = before_limit[idx_recommended + 1:]
            rec_inline = re.sub(r"^zalecane zwroty:\s*", "", before_limit[idx_recommended], flags=re.IGNORECASE).strip()
            if rec_inline:
                rec_src = [rec_inline] + rec_src
            if lead:
                blocks.extend(p(x) for x in lead)
            blocks.append("<div class='ap-ext-topic ap-ext-topic-strong'><b>Zalecane zwroty:</b></div>")
            rec_items = split_semicolon_items(rec_src)
            if rec_items:
                blocks.append(ul(rec_items, cls="ap-ext-list ap-ext-list-italic"))
        else:
            if before_limit:
                blocks.extend(p(x) for x in before_limit)

        if idx_limit is not None:
            blocks.append("<div class='ap-ext-topic'><b>Do ograniczenia:</b></div>")
            blocks.append(ul(split_semicolon_items(after_limit)))
        return "".join(blocks)

    if sub_prefix in numbered_prefixes:
        return ol(lines)

    if sub_prefix == "7.3.":
        num_items: list[str] = []
        tail: list[str] = []
        for ln in lines:
            if ln.casefold().startswith("zakaz:"):
                tail.append(ln)
            else:
                num_items.append(ln)
        out = ol(num_items)
        for t in tail:
            out += p_bold_prefix(t, {"Zakaz"})
        return out

    if sub_prefix == "9.1.":
        do_items: list[str] = []
        dont_items: list[str] = []
        mode: str | None = None
        for ln in lines:
            low = ln.casefold()
            if low.startswith("do:"):
                mode = "do"
                continue
            if ("don't" in low) or ("don’t" in low) or ("dont" in low):
                mode = "dont"
                continue
            if mode == "do":
                do_items.extend(split_semicolon_items([ln]))
            elif mode == "dont":
                dont_items.extend(split_semicolon_items([ln]))
        blocks = []
        if do_items:
            blocks.append("<div class='ap-ext-topic'><b>DO:</b></div>")
            blocks.append(ul(do_items))
        if dont_items:
            blocks.append("<div class='ap-ext-topic'><b>DON'T:</b></div>")
            blocks.append(ul(dont_items))
        return "".join(blocks) if blocks else "".join(p(x) for x in lines)

    if sub_prefix == "9.4.":
        diag_idx = None
        for i, ln in enumerate(lines):
            if ln.casefold().startswith("pytania diagnostyczne"):
                diag_idx = i
                break
        if diag_idx is None:
            return ul(split_semicolon_items(lines))
        first = split_semicolon_items(lines[:diag_idx])
        second = split_semicolon_items(lines[diag_idx + 1:])
        blocks = []
        if first:
            blocks.append(ul(first))
        blocks.append("<div class='ap-ext-topic'><b>Pytania diagnostyczne:</b></div>")
        if second:
            blocks.append(ul(second))
        return "".join(blocks)

    if sub_prefix == "8.2.":
        items = split_semicolon_items(lines)
        rendered = []
        for item in items:
            m_line = re.match(r"^([A-Za-zĄĆĘŁŃÓŚŹŻąćęłńóśźż]+)\s+[–-]\s+(.+)$", item)
            if m_line:
                rendered.append(f"<li><b>{html.escape(m_line.group(1))}</b> — {html.escape(m_line.group(2))}</li>")
            else:
                rendered.append(f"<li>{html.escape(item)}</li>")
        return "<ul class='ap-ext-list'>" + "".join(rendered) + "</ul>"

    if sub_prefix == "5.1.2.":
        groups: list[tuple[str, list[str]]] = []
        current_head = ""
        current_items: list[str] = []
        for ln in lines:
            if re.match(r"^Zestaw\s+\d+", ln, flags=re.IGNORECASE):
                if current_head or current_items:
                    groups.append((current_head, current_items))
                current_head = ln
                current_items = []
            else:
                current_items.extend(split_semicolon_items([ln]))
        if current_head or current_items:
            groups.append((current_head, current_items))

        out = []
        for head, items in groups:
            if head:
                out.append(p(head, cls="ap-ext-step-head"))
            if items:
                out.append(ul(items))
        left_html = "".join(out)
        card_uri = (
            _archetype_card_data_uri(archetype_name, _archetype_card_cache_token(archetype_name))
            if archetype_name
            else ""
        )
        if not card_uri:
            return left_html
        modal_id = f"ap-card-modal-{_norm_archetype_key(archetype_name)}-{abs(hash('||'.join(lines))) % 1000000}"
        alt_txt = html.escape(f"Karta archetypu {archetype_name}".strip())
        right_html = (
            "<div class='ap-ext-zestawy-card-wrap'>"
            f"<a class='ap-ext-zestawy-card-link' href='#{modal_id}' title='Kliknij, aby powiększyć kartę'>"
            f"<img class='ap-ext-zestawy-card' src='{card_uri}' alt='{alt_txt}'/>"
            "</a>"
            f"<div id='{modal_id}' class='ap-ext-card-modal'>"
            f"<a href='#' class='ap-ext-card-modal-backdrop' aria-label='Zamknij podgląd'></a>"
            "<div class='ap-ext-card-modal-content'>"
            "<a href='#' class='ap-ext-card-modal-close' aria-label='Zamknij podgląd'>&times;</a>"
            f"<img class='ap-ext-card-modal-img' src='{card_uri}' alt='{alt_txt}'/>"
            "</div>"
            "</div>"
            "</div>"
        )
        return (
            "<div class='ap-ext-zestawy-wrap'>"
            f"<div class='ap-ext-zestawy-left'>{left_html}</div>"
            f"{right_html}"
            "</div>"
        )

    if sub_prefix == "3.7.":
        blocks: list[str] = []
        prev_technika = False
        for ln in lines:
            if re.match(r"^\d+\)\s+", ln):
                blocks.append(p(ln, cls="ap-ext-step-head ap-ext-tech-head"))
                prev_technika = False
                continue
            if ln.casefold().startswith("technika:"):
                blocks.append(p_bold_prefix(ln, {"Technika"}))
                prev_technika = True
                continue
            if ln.casefold().startswith("schemat:"):
                blocks.append(p_bold_prefix(ln, {"Schemat"}))
                prev_technika = False
                continue
            if ln.casefold().startswith("po co?"):
                blocks.append(p_bold_prefix(ln, {"Po co?"}))
                prev_technika = False
                continue
            if prev_technika and "," in ln and ("„" in ln or "\"" in ln):
                parts = [x.strip() for x in re.split(r"\s*,\s*", ln) if x.strip()]
                blocks.append(ul(parts))
                prev_technika = False
                continue
            blocks.append(p(ln))
            prev_technika = False
        return "".join(blocks)

    if sub_prefix in {"2.3.", "2.4.", "4.3."}:
        lead = lines[0] if lines else ""
        if lead.endswith(":"):
            rest_items = split_semicolon_items(lines[1:])
            return p(lead) + ul(rest_items)
        return ul(split_semicolon_items(lines))

    if sub_prefix in bullet_only_prefixes:
        return ul(split_semicolon_items(lines))

    # Domyślnie: akapity (bez agresywnego dzielenia po przecinkach)
    return "".join(p(x) for x in lines)


def _expanded_sections_html(archetype_data: dict) -> str:
    sections = archetype_data.get("expanded_sections") or []
    if not sections:
        return "<div class='ap-ext-empty'>Brak rozbudowanego opisu dla tego archetypu.</div>"

    section_blocks: list[str] = []
    for section in sections:
        sec_title_raw = str(section.get("title", "")).strip()
        sec_title = html.escape(sec_title_raw)
        icon = ""
        m_sec = re.match(r"^([1-9])\.\s+", sec_title_raw)
        if m_sec:
            icon = SECTION_ICON_MAP.get(m_sec.group(1), "")
        subsection_blocks: list[str] = []
        for subsection in section.get("subsections", []):
            sub_title_raw = str(subsection.get("title", "")).strip()
            sub_title = html.escape(sub_title_raw)
            content_lines = [str(line).strip() for line in subsection.get("content", []) if str(line).strip()]
            if not content_lines and not sub_title:
                continue
            content_html = _expanded_subsection_content_html(
                content_lines,
                sub_title_raw,
                str(archetype_data.get("name", "")).strip(),
            )
            subtitle_cls = "ap-ext-subtitle ap-ext-subtitle-l2"
            if sub_title_raw:
                m_sub = re.match(r"^(\d+(?:\.\d+)*)\.\s+", sub_title_raw)
                if m_sub:
                    depth = len([p for p in m_sub.group(1).split(".") if p.strip()])
                    if depth >= 3:
                        subtitle_cls = "ap-ext-subtitle ap-ext-subtitle-l3"

            subsection_blocks.append(
                f"""
                <div class="ap-ext-subsection">
                    {f"<div class='{subtitle_cls}'>{sub_title}</div>" if sub_title else ""}
                    <div class="ap-ext-content">{content_html}</div>
                </div>
                """
            )

        section_blocks.append(
            f"""
            <section class="ap-ext-section">
                <div class="ap-ext-title">{f"<span class='ap-ext-title-icon'>{icon}</span>" if icon else ""}<span>{sec_title}</span></div>
                <div class="ap-ext-body">{''.join(subsection_blocks)}</div>
            </section>
            """
        )
    return "".join(section_blocks)


def render_archetype_card(archetype_data, main=True, supplement=False, gender_code="M"):
    if not archetype_data:
        st.warning("Brak danych o archetypie.")
        return

    palette = [str(c).upper() for c in (archetype_data.get("color_palette", []) or []) if str(c).strip()]
    if main:
        bg_color_raw = palette[0] if len(palette) >= 1 else "#2B2D41"
        border_color = palette[1] if len(palette) >= 2 else (palette[0] if palette else "#E99836")
        bg_color = _soften_aggressive_main_bg(bg_color_raw, border_color)
        text_color = "#F8FAFC" if is_color_dark(bg_color) else "#1F2937"
        tagline_color = "#FFE082" if is_color_dark(bg_color) else "#7A2037"
    else:
        bg_color = "#F3F4F6"
        border_color = palette[0] if palette else ("#40b900" if supplement else "#FFD22F")
        text_color = "#1F2937"
        tagline_color = border_color

    is_dark_card = is_color_dark(bg_color)
    link_color = "#DDEAFF" if is_color_dark(bg_color) else "#144FA8"
    details_border_color = "rgba(255,255,255,.30)" if is_dark_card else "rgba(17,24,39,.20)"
    details_bg_color = "rgba(255,255,255,.07)" if is_dark_card else "rgba(255,255,255,.74)"
    details_title_bg = (
        "linear-gradient(90deg, rgba(255,210,47,.20), rgba(255,255,255,.03))"
        if is_dark_card
        else "linear-gradient(90deg, rgba(17,24,39,.08), rgba(17,24,39,.02))"
    )
    details_subtitle_color = "#FFE082" if is_dark_card else "#1F2937"
    details_text_color = "#EDF2FF" if is_dark_card else "#1F2937"
    pill_bg = "rgba(255,255,255,.12)" if is_dark_card else "#F8FAFF"
    pill_border = "rgba(255,255,255,.35)" if is_dark_card else "#A4B4D4"
    pill_text = "#F8FAFF" if is_dark_card else "#1F355E"
    box_shadow = f"0 12px 28px 0 {border_color}33"

    base_name = base_masc_from_any(str(archetype_data.get("name", "")).strip())
    core_triplet = (
        archetype_data.get("core_triplet")
        or CORE_TRIPLET_MAP.get(base_name, "")
        or archetype_data.get("tagline")
        or ""
    )

    width_card = "86vw"
    min_h = "860px"

    metric_rows_src = list(archetype_data.get("metric_rows") or _fallback_metric_rows(archetype_data))
    metric_rows = [
        row for row in metric_rows_src
        if str((row or {}).get("label", "")).strip().casefold() != "core triplet"
    ]
    metric_blocks: list[str] = []
    compact_buffer: list[tuple[str, str]] = []

    def _flush_compact_metric_buffer() -> None:
        nonlocal compact_buffer
        if not compact_buffer:
            return
        tiles = "".join(
            f"<div class='ap-metric-tile'><div class='ap-metric-tile-label'>{html.escape(lbl)}</div>"
            f"<div class='ap-metric-tile-value'>{html.escape(val)}</div></div>"
            for lbl, val in compact_buffer
        )
        metric_blocks.append(f"<div class='ap-metric-grid'>{tiles}</div>")
        compact_buffer = []

    for row in metric_rows:
        label_raw = str((row or {}).get("label", "")).strip()
        label_cf = label_raw.casefold()
        kind = (row or {}).get("kind", "text")
        value = (row or {}).get("value")

        is_compact = (
            label_cf in COMPACT_METRIC_LABELS_CF
            and kind not in {"examples", "colors"}
            and not isinstance(value, list)
        )
        if is_compact:
            plain_value = _metric_text_plain(value)
            if plain_value:
                compact_buffer.append((label_raw, plain_value))
            continue

        _flush_compact_metric_buffer()
        row_html = _metric_row_html(row, archetype_data, text_color)
        if row_html.strip():
            metric_blocks.append(row_html)

    _flush_compact_metric_buffer()
    metric_rows_html = "".join(metric_blocks)
    expanded_html = _expanded_sections_html(archetype_data)
    name_slug = re.sub(r"[^a-z0-9]+", "-", _logo_norm_key(archetype_data.get("name", "archetyp"))).strip("-")
    role_slug = "main" if main else ("supp" if supplement else "aux")
    card_dom_id = f"ap-arch-{name_slug}-{role_slug}"

    card_html_raw = dedent(f"""
        <style>
            #{card_dom_id}.ap-card-wrap {{
                max-width:{width_card};
                border: 3px solid {border_color};
                border-radius: 22px;
                background: {bg_color};
                box-shadow: {box_shadow};
                min-height: {min_h};
                padding: 2.8em 2.3em 2.2em 4.7em;
                margin-bottom: 32px;
                color: {text_color};
            }}
            #{card_dom_id} .ap-card-head {{
                display:flex;
                align-items:flex-start;
                gap:24px;
                margin-bottom:18px;
            }}
            #{card_dom_id} .ap-card-name {{
                font-size:2.38em;
                font-weight:700;
                line-height:1.08;
                margin-top:12px;
                margin-bottom:7px;
                color:{text_color};
            }}
            #{card_dom_id} .ap-card-tagline {{
                font-size:1.32em;
                font-style:italic;
                color:{tagline_color};
                margin-top:4px;
                margin-bottom:8px;
                font-weight:600;
            }}
            #{card_dom_id} .ap-metric-title {{
                font-size:1.54em;
                font-weight:700;
                color:{text_color};
                margin:22px 0 16px;
                letter-spacing:.01em;
            }}
            #{card_dom_id} .ap-card-icon-wrap {{
                margin-left:-14px;
            }}
            #{card_dom_id} .ap-metric-row {{
                border:none;
                border-radius:0;
                padding:0;
                margin-top:40px;
                margin-bottom:0;
                background:transparent;
            }}
            #{card_dom_id} .ap-metric-row:first-child {{
                margin-top:0;
            }}
            #{card_dom_id} .ap-metric-label {{
                font-size:1.15em;
                font-weight:700;
                color:{text_color};
                margin-top:0;
                margin-bottom:1px;
                letter-spacing:.02em;
            }}
            #{card_dom_id} .ap-example-head {{
                margin-top:24px;
                margin-bottom:12px;
                font-weight:700;
            }}
            #{card_dom_id} .ap-example-head-first {{
                margin-top:8px;
            }}
            #{card_dom_id} .ap-metric-value {{
                font-size:1em;
                line-height:1.56;
                color:{text_color};
            }}
            #{card_dom_id} .ap-metric-value p {{
                margin:0 0 6px 0;
            }}
            #{card_dom_id} .ap-roman-list {{
                margin:6px 0 10px 0;
                padding-left:21px;
                list-style:none;
                counter-reset: ap-roman-counter;
            }}
            #{card_dom_id} .ap-roman-list li {{
                margin:0 0 8px 0;
                padding-left:34px;
                position:relative;
            }}
            #{card_dom_id} .ap-roman-list li::before {{
                counter-increment: ap-roman-counter;
                content: counter(ap-roman-counter, upper-roman) ".";
                position:absolute;
                left:0;
                width:28px;
                text-align:right;
                font-weight:400;
            }}
            #{card_dom_id} .ap-roman-list li:last-child {{
                margin-bottom:0;
            }}
            #{card_dom_id} .ap-simple-list {{
                margin:2px 0 4px 0;
                padding-left:21px;
            }}
            #{card_dom_id} .ap-story-antag {{
                margin-top:2px;
                padding-left:4px;
            }}
            #{card_dom_id} .ap-story-head {{
                font-weight:700;
                margin-top:10px;
                margin-bottom:2px;
                margin-left:10px;
            }}
            #{card_dom_id} .ap-story-list {{
                margin:2px 0 8px 0;
                padding-left:34px;
            }}
            #{card_dom_id} .ap-simple-list li {{
                margin-bottom:4px;
            }}
            #{card_dom_id} .ap-qual-list {{
                margin:4px 0 6px 0;
                padding:0;
                list-style:none;
            }}
            #{card_dom_id} .ap-qual-list li {{
                display:flex;
                align-items:flex-start;
                gap:8px;
                margin-bottom:5px;
                line-height:1.42;
            }}
            #{card_dom_id} .ap-qual-ico {{
                display:inline-block;
                width:20px;
                text-align:center;
            }}
            #{card_dom_id} .ap-pill-wrap {{
                display:flex;
                flex-wrap:wrap;
                gap:7px;
            }}
            #{card_dom_id} .ap-pill {{
                display:inline-block;
                border-radius:999px;
                padding:6px 12px;
                border:1px solid {pill_border};
                background:{pill_bg};
                color:{pill_text};
                box-shadow:0 1px 0 rgba(0,0,0,.08);
                font-size:.95em;
                line-height:1.2;
            }}
            #{card_dom_id} .ap-metric-row-chips .ap-metric-label {{
                margin-bottom:16px;
            }}
            #{card_dom_id} .ap-metric-grid {{
                display:grid;
                grid-template-columns:repeat(2,minmax(0,1fr));
                gap:12px 12px;
                margin-top:34px;
            }}
            #{card_dom_id} .ap-metric-tile {{
                border:1px solid {details_border_color};
                border-radius:12px;
                background:{details_bg_color};
                padding:12px 13px;
            }}
            #{card_dom_id} .ap-metric-tile-label {{
                font-size:.99em;
                font-weight:700;
                margin-bottom:4px;
                color:{text_color};
            }}
            #{card_dom_id} .ap-metric-tile-value {{
                font-size:.96em;
                line-height:1.42;
                color:{text_color};
            }}
            #{card_dom_id} .ap-face-grid {{
                display:flex;
                flex-wrap:wrap;
                gap:12px;
                margin-top:8px;
                align-items:flex-start;
            }}
            #{card_dom_id} .ap-face-card {{
                width:114px;
                text-align:center;
                display:flex;
                flex-direction:column;
                align-items:center;
                justify-content:flex-start;
            }}
            #{card_dom_id} .ap-face-avatar img,
            #{card_dom_id} .ap-face-avatar .ap-face-ph {{
                width:94px;
                height:94px;
                border-radius:12px;
                object-fit:cover;
                object-position:center 0%;
                border:1px solid rgba(0,0,0,.18);
                background:rgba(255,255,255,.85);
                margin:0 auto;
                display:flex;
                align-items:center;
                justify-content:center;
                font-weight:700;
            }}
            #{card_dom_id} .ap-face-name {{
                margin-top:4px;
                font-size:.80em;
                line-height:1.25;
            }}
            #{card_dom_id} .ap-face-name a {{
                color:{link_color};
                text-decoration:underline;
            }}
            #{card_dom_id} .ap-details-row {{
                margin-top:18px;
                position:relative;
            }}
            #{card_dom_id} .ap-expand-cue {{
                position:absolute;
                left:-18px;
                top:50%;
                transform:translateY(-50%);
                width:0;
                height:0;
                display:inline-flex;
                align-items:center;
                justify-content:center;
                pointer-events:none;
            }}
            #{card_dom_id} .ap-expand-cue-arrow {{
                color:{tagline_color};
                font-size:1.09em;
                font-weight:800;
                line-height:1;
                opacity:.9;
                animation:ap-expand-nudge 1.2s ease-in-out infinite;
                user-select:none;
            }}
            @keyframes ap-expand-nudge {{
                0% {{ transform:translateX(0); opacity:.82; }}
                45% {{ transform:translateX(4px); opacity:1; }}
                100% {{ transform:translateX(0); opacity:.82; }}
            }}
            #{card_dom_id} .ap-details {{
                margin-top:0;
                border:1px solid {details_border_color};
                border-radius:14px;
                background:{details_bg_color};
                overflow:hidden;
            }}
            #{card_dom_id} .ap-details > summary {{
                cursor:pointer;
                list-style:none;
                padding:13px 14px;
                font-size:1em;
                font-weight:700;
                color:{text_color};
                border-bottom:1px solid {details_border_color};
                display:flex;
                align-items:center;
                justify-content:space-between;
                gap:10px;
            }}
            #{card_dom_id} .ap-summary-left {{
                display:flex;
                align-items:center;
                gap:9px;
                min-width:0;
            }}
            #{card_dom_id} .ap-summary-badge {{
                display:inline-flex;
                align-items:center;
                gap:6px;
                padding:3px 8px;
                border-radius:999px;
                font-size:.78em;
                font-weight:760;
                letter-spacing:.01em;
                color:{text_color};
                background:rgba(255,255,255,.16);
                border:1px solid {details_border_color};
                white-space:nowrap;
            }}
            #{card_dom_id} .ap-summary-text {{
                display:inline-flex;
                align-items:center;
                min-width:0;
            }}
            #{card_dom_id} .ap-summary-arrow {{
                font-size:1.06em;
                opacity:.9;
                transition:transform .17s ease;
            }}
            #{card_dom_id} .ap-details > summary::-webkit-details-marker {{ display:none; }}
            #{card_dom_id} .ap-details .ap-summary-close {{ display:none; }}
            #{card_dom_id} .ap-details[open] .ap-summary-open {{ display:none; }}
            #{card_dom_id} .ap-details[open] .ap-summary-close {{ display:inline; }}
            #{card_dom_id} .ap-details[open] .ap-summary-arrow {{
                transform:rotate(180deg);
            }}
            #{card_dom_id} .ap-expanded-wrap {{
                padding:12px 12px 8px;
            }}
            #{card_dom_id} .ap-ext-section {{
                border:1px solid {details_border_color};
                border-radius:12px;
                margin-bottom:10px;
                background:{details_bg_color};
                overflow:hidden;
            }}
            #{card_dom_id} .ap-ext-title {{
                background:{details_title_bg};
                border-left:4px solid {tagline_color};
                padding:10px 12px;
                font-size:1.03em;
                font-weight:700;
                color:{text_color};
                display:flex;
                align-items:center;
                gap:8px;
            }}
            #{card_dom_id} .ap-ext-title-icon {{
                font-size:1.08em;
                line-height:1;
            }}
            #{card_dom_id} .ap-ext-body {{
                padding:10px 12px 6px;
            }}
            #{card_dom_id} .ap-ext-subsection {{
                margin-bottom:12px;
                padding-bottom:11px;
                border-bottom:1px dashed {details_border_color};
            }}
            #{card_dom_id} .ap-ext-subsection:last-child {{
                border-bottom:none;
                margin-bottom:0;
                padding-bottom:4px;
            }}
            #{card_dom_id} .ap-ext-subtitle {{
                font-size:1.02em;
                font-weight:700;
                color:{details_subtitle_color};
                margin-top:12px;
                margin-bottom:4px;
            }}
            #{card_dom_id} .ap-ext-subtitle-l2 {{
                font-size:1.02em;
                font-weight:700;
                margin-top:12px;
                margin-bottom:4px;
            }}
            #{card_dom_id} .ap-ext-subtitle-l3 {{
                font-size:.92em;
                font-weight:700;
                letter-spacing:.01em;
                color:{details_text_color};
                margin-top:9px;
                margin-bottom:5px;
                margin-left:14px;
                padding:4px 9px;
                border-left:3px solid {tagline_color};
                background:rgba(255,255,255,.18);
                border-radius:6px;
            }}
            #{card_dom_id} .ap-ext-subtitle-l3 + .ap-ext-content {{
                margin-left:14px;
                padding-left:8px;
            }}
            #{card_dom_id} .ap-ext-content p {{
                margin:0 0 6px 0;
                font-size:.95em;
                line-height:1.5;
                color:{details_text_color};
            }}
            #{card_dom_id} .ap-ext-zestawy-wrap {{
                display:grid;
                grid-template-columns:minmax(0,1fr) minmax(462px, 40%);
                gap:20px;
                align-items:start;
                margin-top:2px;
            }}
            #{card_dom_id} .ap-ext-zestawy-left {{
                min-width:0;
            }}
            #{card_dom_id} .ap-ext-zestawy-card-wrap {{
                display:flex;
                justify-content:flex-end;
                position:relative;
            }}
            #{card_dom_id} .ap-ext-zestawy-card-link {{
                display:block;
                line-height:0;
                cursor:zoom-in;
            }}
            #{card_dom_id} .ap-ext-zestawy-card {{
                width:100%;
                max-width:572px;
                height:auto;
                display:block;
                border-radius:12px;
                border:1px solid {details_border_color};
                box-shadow:0 8px 18px rgba(0,0,0,.16);
            }}
            #{card_dom_id} .ap-ext-zestawy-card-link:hover .ap-ext-zestawy-card {{
                transform:translateY(-1px);
                box-shadow:0 14px 28px rgba(0,0,0,.24);
                transition:transform .15s ease, box-shadow .15s ease;
            }}
            #{card_dom_id} .ap-ext-card-modal {{
                position:fixed;
                inset:0;
                display:none;
                z-index:99999;
            }}
            #{card_dom_id} .ap-ext-card-modal:target {{
                display:block;
            }}
            #{card_dom_id} .ap-ext-card-modal-backdrop {{
                position:absolute;
                inset:0;
                background:rgba(8,12,20,.74);
                backdrop-filter: blur(2px);
                cursor:zoom-out;
            }}
            #{card_dom_id} .ap-ext-card-modal-content {{
                position:relative;
                z-index:1;
                max-width:min(94vw, 1260px);
                max-height:90vh;
                margin:4.5vh auto 0;
                padding:10px;
                border-radius:12px;
                background:rgba(255,255,255,.08);
                border:1px solid rgba(255,255,255,.24);
                box-shadow:0 22px 60px rgba(0,0,0,.45);
            }}
            #{card_dom_id} .ap-ext-card-modal-img {{
                width:100%;
                height:auto;
                max-height:84vh;
                object-fit:contain;
                display:block;
                border-radius:10px;
                background:#fff;
            }}
            #{card_dom_id} .ap-ext-card-modal-close {{
                position:absolute;
                top:-16px;
                right:-14px;
                width:34px;
                height:34px;
                border-radius:999px;
                display:flex;
                align-items:center;
                justify-content:center;
                text-decoration:none;
                font-size:26px;
                line-height:1;
                color:#fff;
                background:rgba(11,18,32,.88);
                border:1px solid rgba(255,255,255,.28);
                box-shadow:0 4px 12px rgba(0,0,0,.32);
            }}
            @media (max-width: 1280px) {{
                #{card_dom_id} .ap-ext-zestawy-wrap {{
                    grid-template-columns:minmax(0,1fr) minmax(330px, 37%);
                    gap:14px;
                }}
                #{card_dom_id} .ap-ext-zestawy-card {{
                    max-width:396px;
                }}
            }}
            #{card_dom_id} .ap-ext-content .ap-ext-step-head {{
                margin-top:13px;
                margin-bottom:6px;
                font-weight:700;
            }}
            #{card_dom_id} .ap-ext-content .ap-ext-tech-head {{
                margin-top:20px;
                margin-bottom:9px;
            }}
            #{card_dom_id} .ap-ext-content p:first-child.ap-ext-step-head {{
                margin-top:2px;
            }}
            #{card_dom_id} .ap-ext-topic {{
                margin:8px 0 5px 0;
                font-weight:600;
                color:{details_text_color};
            }}
            #{card_dom_id} .ap-ext-topic-strong {{
                margin-top:10px;
                font-weight:760;
                font-size:1.01em;
                color:{details_subtitle_color};
            }}
            #{card_dom_id} .ap-ext-list {{
                margin:1px 0 10px 0;
                padding-left:20px;
            }}
            #{card_dom_id} .ap-ext-list-italic li {{
                font-style:italic;
            }}
            #{card_dom_id} .ap-ext-ol {{
                margin:1px 0 10px 0;
                padding-left:24px;
            }}
            #{card_dom_id} .ap-ext-list-nested {{
                margin-top:4px;
            }}
            #{card_dom_id} .ap-ext-list li,
            #{card_dom_id} .ap-ext-ol li {{
                margin-bottom:5px;
                font-size:.94em;
                line-height:1.42;
                color:{details_text_color};
            }}
            #{card_dom_id} .ap-ext-empty {{
                color:{details_text_color};
                font-size:.95em;
                padding:8px 2px 12px;
            }}
            @media (max-width: 900px) {{
                #{card_dom_id}.ap-card-wrap {{
                    padding:1.35em 1.0em 1.05em 2.15em;
                }}
                #{card_dom_id} .ap-expand-cue {{
                    left:-12px;
                    top:50%;
                    transform:translateY(-50%);
                }}
                #{card_dom_id} .ap-card-head {{
                    flex-direction:column;
                    gap:8px;
                }}
                #{card_dom_id} .ap-card-name {{
                    margin-top:0;
                    font-size:1.82em;
                }}
                #{card_dom_id} .ap-metric-grid {{
                    grid-template-columns:1fr;
                }}
                #{card_dom_id} .ap-ext-zestawy-wrap {{
                    grid-template-columns:1fr;
                    gap:10px;
                }}
                #{card_dom_id} .ap-ext-zestawy-card-wrap {{
                    justify-content:flex-start;
                }}
                #{card_dom_id} .ap-ext-zestawy-card {{
                    max-width:240px;
                }}
                #{card_dom_id} .ap-ext-card-modal-content {{
                    max-width:95vw;
                    margin-top:6vh;
                    padding:8px;
                }}
            }}
        </style>
        <div id="{card_dom_id}" class="ap-card-wrap">
            <div class="ap-card-head">
                <div class="ap-card-icon-wrap" style="flex-shrink:0;">
                    {arche_icon_img_html(archetype_data.get('name', '?'), height_px=146, gender_code=gender_code)}
                </div>
                <div>
                    <div class="ap-card-name">{html.escape(str(archetype_data.get('name', '?')))}</div>
                    <div class="ap-card-tagline">{html.escape(str(core_triplet))}</div>
                </div>
            </div>
            <div class="ap-metric-title">Metryka archetypu</div>
            {metric_rows_html}
            <div class="ap-details-row">
                <div class="ap-expand-cue" aria-hidden="true"><span class="ap-expand-cue-arrow">➜</span></div>
                <details class="ap-details">
                    <summary>
                        <span class="ap-summary-left">
                            <span class="ap-summary-badge">🔎 Rozszerzona analiza</span>
                            <span class="ap-summary-text">
                                <span class="ap-summary-open">Pokaż rozbudowany opis</span>
                                <span class="ap-summary-close">Zwiń rozbudowany opis</span>
                            </span>
                        </span>
                        <span class="ap-summary-arrow">▾</span>
                    </summary>
                    <div class="ap-expanded-wrap">
                        {expanded_html}
                    </div>
                </details>
            </div>
        </div>
    """)
    # Usuń wcięcia na początku linii: inaczej markdown Streamlit traktuje HTML jak code block.
    card_html = "\n".join(line.lstrip() for line in card_html_raw.splitlines()).strip()
    st.markdown(card_html, unsafe_allow_html=True)

# ============ RESZTA PANELU: nagłówki, kolumny, eksporty, wykres, tabele respondentów ============

def show_report(sb, study: dict, wide: bool = True, public_view: bool = False) -> None:
    # stan przełącznika eksportu (bezpieczeństwo przy pierwszym renderze)
    st.session_state.setdefault("prep_docs", False)
    # CSS musi być doładowywany na każdym rerunie (zmiana osoby w selectboxie).
    inject_global_css(GLOBAL_CSS)
    # Szerokość raportu (również dla widoku publicznego).
    st.markdown(
        f"<style>.block-container{{max-width:{'100vw' if wide else '1160px'} !important;}}</style>",
        unsafe_allow_html=True,
    )
    is_mobile = _is_probably_mobile_client()
    public_dark_mode = bool(public_view)
    mobile_table_bg = "transparent" if public_dark_mode else "#ffffff"
    mobile_table_text = "#dce8f8" if public_dark_mode else "#0f172a"
    mobile_table_border = "rgba(148,163,184,.34)" if public_dark_mode else "#e2e8f0"
    mobile_section_margin_top = 30 if is_mobile else 6
    mobile_profile_actions_margin_top = 32 if is_mobile else 8
    radar_plot_size = 430 if is_mobile else 560
    radar_tick_size = 10 if is_mobile else 15
    radar_hover_size = 12 if is_mobile else 14
    radar_margins = dict(l=16, r=16, t=2, b=0) if is_mobile else dict(l=20, r=20, t=20, b=0)
    radar_domain = dict(x=[0.08, 0.92], y=[0.10, 0.99]) if is_mobile else dict(x=[0.10, 0.90], y=[0.06, 0.95])
    wheel_img_width = 360 if is_mobile else 620
    axes_img_width = 360 if is_mobile else 620
    segment_profile_width = 360 if is_mobile else 640
    # TU ZMIENISZ ROZMIAR sekcji "Profil siły archetypów ... (skala: 0-100)" dla desktopu.
    segment_profile_display_width_desktop = 800
    # TU ZMIENISZ ROZMIAR tej sekcji dla mniejszych desktopów (np. <=1920x1200).
    segment_profile_display_width_mid_desktop = 680
    # Mobile-only: poprawa responsywności wykresów/obrazów i czytelności tabeli
    mobile_layout_css = """
        <style>
        @media (max-width: 900px){
          .block-container [data-testid="stHorizontalBlock"]{
            flex-direction:column !important;
            align-items:stretch !important;
            gap:.65rem !important;
          }
          .block-container [data-testid="column"]{
            width:100% !important;
            min-width:0 !important;
            flex:1 1 100% !important;
          }
          html, body, .stApp, [data-testid="stAppViewContainer"], [data-testid="stMain"], .main{
            overflow-x:hidden !important;
            max-width:100% !important;
          }
          .block-container [data-testid="stImage"]{
            width:100% !important;
            overflow:visible !important;
          }
          .block-container [data-testid="stImage"] > div{
            width:100% !important;
            max-width:100% !important;
          }
          .block-container [data-testid="stImage"] img{
            display:block !important;
            width:auto !important;
            max-width:100% !important;
            height:auto !important;
            margin-left:auto !important;
            margin-right:auto !important;
            object-fit:contain !important;
          }
          .block-container .js-plotly-plot,
          .block-container .plotly,
          .block-container .main-svg{
            width:100% !important;
            max-width:100% !important;
          }
          .block-container .plotly-graph-div{
            margin-left:auto !important;
            margin-right:auto !important;
          }
          .block-container .ap-card-wrap{
            width:100% !important;
            max-width:100% !important;
            margin-left:0 !important;
            margin-right:0 !important;
          }
          .ap-table-wrap{
            width:100% !important;
            max-width:100% !important;
            display:block !important;
            overflow-x:auto !important;
            overflow-y:hidden !important;
            -webkit-overflow-scrolling:touch !important;
          }
          .ap-table{
            min-width:690px !important;
            width:max-content !important;
            table-layout:auto !important;
            background:__AP_MOBILE_TABLE_BG__ !important;
          }
          .ap-table,
          .ap-table thead th,
          .ap-table tbody td{
            color:__AP_MOBILE_TABLE_TEXT__ !important;
            background:__AP_MOBILE_TABLE_BG__ !important;
            border-bottom:1px solid __AP_MOBILE_TABLE_BORDER__ !important;
          }
          .ap-table th, .ap-table td{
            font-size:11.6px !important;
            padding:7px 5px !important;
            white-space:nowrap !important;
            word-break:normal !important;
          }
        }
        </style>
    """
    mobile_layout_css = (
        mobile_layout_css
        .replace("__AP_MOBILE_TABLE_BG__", mobile_table_bg)
        .replace("__AP_MOBILE_TABLE_TEXT__", mobile_table_text)
        .replace("__AP_MOBILE_TABLE_BORDER__", mobile_table_border)
    )
    st.markdown(
        mobile_layout_css,
        unsafe_allow_html=True,
    )
    if public_dark_mode:
        st.markdown(
            """
            <style>
            :root{
              --ap-heading-color:#e8f1ff;
              --text-color:#d7e3f5;
            }
            [data-testid="stMarkdownContainer"],
            [data-testid="stMarkdownContainer"] p,
            [data-testid="stMarkdownContainer"] li{
              color:var(--text-color,#d7e3f5) !important;
            }
            .ap-public-heading-title{
              color:var(--ap-heading-color,#e8f1ff) !important;
            }
            .ap-public-heading-count{
              color:#b8d5ff !important;
            }
            .ap-public-heading-count-label{
              color:#d2e1f4 !important;
            }
            </style>
            """,
            unsafe_allow_html=True,
        )
    st.markdown(
        f"""
        <style>
        .ap-strength-wheel-img {{
          width:min(100%, {segment_profile_display_width_desktop}px) !important;
          height:auto !important;
          display:block !important;
          margin-left:0 !important;
          margin-right:0 !important;
        }}
        /* Desktop do około 1920-2048 px szerokości: wyraźnie zmniejsz wykres. */
        @media (min-width:1200px) and (max-width:2050px) {{
          .ap-strength-wheel-img {{
            width:min(100%, {segment_profile_display_width_mid_desktop}px) !important;
          }}
        }}
        @media (max-width:900px) {{
          .ap-strength-wheel-img {{
            width:100% !important;
          }}
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )

    # --- NOWE: płeć + mapowanie nazw do żeńskich ---
    gender_raw = (study.get("gender") or study.get("sex") or study.get("plec") or "").strip().lower()
    IS_FEMALE = gender_raw in {"k", "kobieta", "female", "f"}
    report_gender_code = "K" if IS_FEMALE else "M"

    FEM_NAME_MAP = {
        "Władca": "Władczyni",
        "Bohater": "Bohaterka",
        "Mędrzec": "Mędrczyni",
        "Opiekun": "Opiekunka",
        "Kochanek": "Kochanka",
        "Błazen": "Komiczka",
        "Twórca": "Twórczyni",
        "Odkrywca": "Odkrywczyni",
        "Czarodziej": "Czarodziejka",
        "Towarzysz": "Towarzyszka",
        "Niewinny": "Niewinna",
        "Buntownik": "Buntowniczka",
    }

    def disp_name(name: str) -> str:
        return FEM_NAME_MAP.get(name, name) if IS_FEMALE else name

    # personalizacja z przekazanego rekordu
    personNom = f"{study.get('first_name_nom') or study.get('first_name','')} {study.get('last_name_nom') or study.get('last_name','')}".strip()
    personGen = f"{study.get('first_name_gen') or (study.get('first_name_nom') or study.get('first_name',''))} {study.get('last_name_gen') or (study.get('last_name_nom') or study.get('last_name',''))}".strip()

    # ⬇️ NOWE – wszystkie pozostałe przypadki
    def _join(a, b):
        return f"{(a or '').strip()} {(b or '').strip()}".strip()

    personDat  = _join(study.get("first_name_dat"),  study.get("last_name_dat"))
    personAcc  = _join(study.get("first_name_acc"),  study.get("last_name_acc"))
    personInst = _join(study.get("first_name_ins"),  study.get("last_name_ins"))
    personLoc  = _join(study.get("first_name_loc"),  study.get("last_name_loc"))
    personVoc  = _join(study.get("first_name_voc"),  study.get("last_name_voc"))

    study_id = study["id"]
    demo_page_key = f"personal_demo_page_{study_id}"
    st.session_state.setdefault(demo_page_key, False)
    show_demo_subpage = (not public_view) and bool(st.session_state.get(demo_page_key))
    data = load(study_id)

    # ❗️ZBIERAMY WSZYSTKIE PRZYPADKI DO SŁOWNIKA DLA WORDA
    person = {
        "NOM": personNom,
        "GEN": personGen,
        "DAT": personDat,
        "ACC": personAcc,
        "INS": personInst,
        "LOC": personLoc,
        "VOC": personVoc,
        "CITY_NOM": (study.get("city_nom") or "").strip(),
    }

    # ... tu dalej Twój kod (generowanie wykresów, budowanie contextu itd.)

    num_ankiet = len(data) if not data.empty else 0

    if not public_view and not show_demo_subpage:
        header_col1, header_col2 = st.columns([0.77, 0.23])
        with header_col1:
            st.markdown(
                f"""
                <div style="font-size:2.3em; font-weight:bold; background:#1a93e3; color:#fff; 
                    padding:14px 32px 10px 24px; border-radius:2px; width:fit-content; display:inline-block;">
                    Archetypy {personGen} – panel administratora
                </div>
                """,
                unsafe_allow_html=True
            )
        with header_col2:
            st.markdown(f"""
            <div style="display:flex;align-items:center;justify-content:flex-end;height:100%;"><div style="font-size:1.23em;text-align:right;background:#f3f3fa;padding:12px 29px 8px 29px; border-radius:17px; border:2px solid #d1d9ed;color:#195299;font-weight:600;box-shadow:0 2px 10px 0 #b5c9e399;">
                <span style="font-size:1.8em;font-weight:bold;">{num_ankiet}</span><br/>uczestników badania
            </div></div>
            """, unsafe_allow_html=True)

        st.markdown("""
        <hr style="height:1.3px;background:#eaeaec; margin-top:1.8em; margin-bottom:3.8em; border:none;" />
        """, unsafe_allow_html=True)

    # --- Analiza respondentów i agregacja ---

    if "answers" in data.columns and not data.empty:
        metry_questions = _build_personal_metry_questions(study)

        results = []

        for idx, row in data.iterrows():

            if not isinstance(row.get("answers", None), list):
                continue

            arcsums = archetype_scores(row["answers"])
            arcper = {k: archetype_percent(v) for k, v in arcsums.items()}

            # --- kolory heurystyczne dla tej odpowiedzi ---
            col_scores = color_scores_from_answers(row["answers"])
            col_perc = color_percents_from_scores(col_scores)

            # <<< UNIKALNE NAZWY TYLKO DLA TEGO RESPONDENTA >>>
            main_i, aux_i, supp_i = pick_top_3_archetypes(arcsums, ARCHE_NAMES_ORDER)

            main = get_archetype_payload_for_gender(main_i, report_gender_code)
            second = (
                get_archetype_payload_for_gender(aux_i, report_gender_code)
                if aux_i != main_i
                else {}
            )
            supplement = (
                get_archetype_payload_for_gender(supp_i, report_gender_code)
                if supp_i not in [main_i, aux_i]
                else {}
            )

            # wersje do wyświetlania – żeńskie/męskie
            main_disp = dict(main);
            main_disp["name"] = disp_name(main.get("name", main_i or ""))
            second_disp = dict(second)
            if second:
                second_disp["name"] = disp_name(second.get("name", aux_i or ""))

            supplement_disp = dict(supplement)
            if supplement:
                supplement_disp["name"] = disp_name(supplement.get("name", supp_i or ""))

            czas_ankiety = ""
            if pd.notna(row.get("created_at", None)):
                try:
                    czas_ankiety = row["created_at"].astimezone(pytz.timezone('Europe/Warsaw')).strftime('%Y-%m-%d %H:%M:%S')
                except Exception:
                    czas_ankiety = row["created_at"].strftime('%Y-%m-%d %H:%M:%S')

            metry_payload = _extract_personal_metry_payload(row.get("scores"))
            metry_row: dict[str, str] = {}
            for mq in metry_questions:
                db_col = str(mq.get("db_column") or "").strip()
                qid = str(mq.get("id") or db_col).strip()
                raw_code = str(metry_payload.get(db_col) or metry_payload.get(qid) or "").strip()
                metry_row[f"METRY_{db_col}"] = raw_code

            results.append({
                "Czas ankiety": czas_ankiety,
                **metry_row,
                **arcsums,
                **{f"{k}_%": v for k, v in arcper.items()},
                "Główny archetyp": main_i,
                "Cechy kluczowe": archetype_features.get(main_i, ""),
                "Opis": main.get("description", ""),
                "Storyline": main.get("storyline", ""),
                "Rekomendacje": "\n".join(main.get("recommendations", [])),
                "Archetyp wspierający": aux_i if aux_i != main_i else "",
                "Cechy wspierający": archetype_features.get(aux_i, "") if aux_i != main_i else "",
                "Opis wspierający": second.get("description", "") if aux_i != main_i else "",
                "Storyline wspierający": second.get("storyline", "") if aux_i != main_i else "",
                "Rekomendacje wspierający": "\n".join(second.get("recommendations", [])) if aux_i != main_i else "",
                "Archetyp poboczny": supp_i if supp_i not in [main_i, aux_i] else "",
                "Cechy poboczny": archetype_features.get(supp_i, "") if supp_i not in [main_i, aux_i] else "",
                "Opis poboczny": supplement.get("description", "") if supp_i not in [main_i, aux_i] else "",
                "Storyline poboczny": supplement.get("storyline", "") if supp_i not in [main_i, aux_i] else "",
                "Rekomendacje poboczny": "\n".join(supplement.get("recommendations", [])) if supp_i not in [main_i, aux_i] else "",
                **{f"Kolor_{k}": v for k, v in col_scores.items()},
                **{f"Kolor_{k}_%": v for k, v in col_perc.items()},
            })

        results_df = pd.DataFrame(results)

        if not results_df.empty and "Czas ankiety" in results_df.columns:

            results_df = results_df.sort_values("Czas ankiety", ascending=True)
            archetype_names = KOLO_NAMES_ORDER

            if (not public_view) and bool(st.session_state.get(demo_page_key)):
                _render_personal_demography_subpage(
                    study_id=str(study_id),
                    person_gen=personGen,
                    results_df=results_df,
                    metry_questions=metry_questions,
                    archetype_names=list(archetype_names),
                    disp_name_fn=disp_name,
                    is_mobile=is_mobile,
                    gender_code=report_gender_code,
                )
                return

            if public_view:
                participants_label = "uczestnik badania" if int(num_ankiet or 0) == 1 else "uczestników badania"
                public_heading_css = """
                    <style>
                      .ap-public-heading-row{
                        display:flex;
                        align-items:flex-end;
                        justify-content:space-between;
                        gap:18px;
                        margin-bottom:22px;
                        flex-wrap:wrap;
                      }
                      .ap-public-heading-title{
                        font-size:2.1em;
                        font-weight:600;
                        line-height:1.15;
                        color:var(--ap-heading-color,#1f2937);
                      }
                      .ap-public-heading-count{
                        display:flex;
                        align-items:baseline;
                        gap:8px;
                        white-space:nowrap;
                        color:__AP_PUBLIC_COUNT_COLOR__;
                        font-weight:700;
                      }
                      .ap-public-heading-count-value{
                        font-size:1.7em;
                        line-height:1;
                        font-weight:800;
                      }
                      .ap-public-heading-count-label{
                        font-size:0.98em;
                        color:__AP_PUBLIC_COUNT_LABEL_COLOR__;
                      }
                      @media (max-width: 900px){
                        .ap-public-heading-row{
                          align-items:flex-start;
                        }
                        .ap-public-heading-title{
                          font-size:1.62em;
                        }
                        .ap-public-heading-count{
                          width:100%;
                        }
                      }
                    </style>
                """
                public_heading_css = (
                    public_heading_css
                    .replace("__AP_PUBLIC_COUNT_COLOR__", "#b8d5ff" if public_dark_mode else "#1f4f8d")
                    .replace("__AP_PUBLIC_COUNT_LABEL_COLOR__", "#d2e1f4" if public_dark_mode else "#3f5873")
                )
                st.markdown(
                    public_heading_css,
                    unsafe_allow_html=True,
                )
                st.markdown(
                    f"""
                    <div class="ap-public-heading-row">
                      <div class="ap-public-heading-title">Informacje na temat archetypów {personGen}</div>
                      <div class="ap-public-heading-count">
                        <span class="ap-public-heading-count-value">{int(num_ankiet or 0)}</span>
                        <span class="ap-public-heading-count-label">{participants_label}</span>
                      </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f'<div style="font-size:2.1em;font-weight:600;margin-bottom:22px;">Informacje na temat archetypów {personGen}</div>',
                    unsafe_allow_html=True,
                )

            # --- ⬇️ RANKING I WYKRES NA BAZIE ŚREDNIEJ (0–20), NIE LICZEBNOŚCI! ---
            # Kolejność zgodna z "Koło potrzeb"
            # (od godz. 12, zgodnie z ruchem wskazówek zegara).

            # 1) średnia suma punktów dla każdego archetypu (0–20)
            mean_archetype_scores = {
                k: float(results_df[k].mean()) if k in results_df.columns else 0.0
                for k in archetype_names
            }

            # <<< NAZWY „AVG” – tylko dla ŚREDNICH >>>
            main_avg, aux_avg, supp_avg = pick_top_3_archetypes(mean_archetype_scores,
                                                                archetype_names)

            # ——— PRÓG 70% dla archetypu pobocznego (liczony na średnich % z całego df)
            means_pct_for_threshold = mean_pct_by_archetype_from_df(data)  # {archetyp: %}
            SHOW_SUPP = should_show_supplement(supp_avg, means_pct_for_threshold, threshold=70.0)
            if not SHOW_SUPP:
                # wyłącz poboczny globalnie (radar, koła, karty)
                supp_avg = None


            # Progi widoczności pobocznego: używamy % (0..100) ze średnich
            means_pct_for_threshold = mean_pct_by_archetype_from_df(data)  # {archetyp: %}
            SHOW_SUPP = should_show_supplement(supp_avg, means_pct_for_threshold, threshold=70.0)
            if not SHOW_SUPP:
                # „Wyłącz” poboczny globalnie – dalszy kod dostanie None i nic nie namaluje
                supp_avg = None

            # >>> KARTY I OPISY: przygotuj dane archetypów na podstawie ŚREDNICH <<<
            main_data = get_archetype_payload_for_gender(main_avg, report_gender_code)
            second_data = (
                get_archetype_payload_for_gender(aux_avg, report_gender_code)
                if aux_avg and aux_avg != main_avg
                else {}
            )
            supp_data = (
                get_archetype_payload_for_gender(supp_avg, report_gender_code)
                if supp_avg and supp_avg not in [main_avg, aux_avg]
                else {}
            )

            def _with_core_triplet(payload: dict, arche_name: str | None) -> dict:
                out = dict(payload or {})
                if arche_name and not out.get("core_triplet"):
                    out["core_triplet"] = CORE_TRIPLET_MAP.get(arche_name, "")
                return out

            # wersje z żeńskimi nazwami, jeśli trzeba
            main_disp = _with_core_triplet(main_data, main_avg)
            main_disp["name"] = disp_name(main_avg or "")
            second_disp = _with_core_triplet(second_data, aux_avg)
            if second_data:
                second_disp["name"] = disp_name(aux_avg or "")
            supp_disp = _with_core_triplet(supp_data, supp_avg)
            if supp_data:
                supp_disp["name"] = disp_name(supp_avg or "")

            left_col, col3 = st.columns([0.62, 0.38], gap="large")  #ręczna regulacja w "Informacje na temat archetypów"
            with left_col:
                col1, col2 = st.columns([0.39, 0.61], gap="small")  #ręczna regulacja tabala a radar
            means_pct = mean_pct_by_archetype_from_df(data)

            def _make_desc_result(arche_name: str | None) -> dict[str, object] | None:
                if not arche_name:
                    return None
                return {
                    "label": disp_name(arche_name),
                    "score": float(means_pct.get(arche_name, 0.0) or 0.0),
                }

            primary_name_for_desc = main_avg or ARCHE_NAMES_ORDER[0]
            supporting_name_for_desc = aux_avg or next(
                (name for name in ARCHE_NAMES_ORDER if name != primary_name_for_desc),
                primary_name_for_desc,
            )

            all_archetypes_for_desc = [
                {
                    "label": disp_name(name),
                    "score": float(means_pct.get(name, 0.0) or 0.0),
                }
                for name in ARCHE_NAMES_ORDER
            ]

            primary_for_desc = _make_desc_result(primary_name_for_desc) or {"label": "", "score": 0.0}
            supporting_for_desc = _make_desc_result(supporting_name_for_desc) or {"label": "", "score": 0.0}
            tertiary_for_desc = _make_desc_result(supp_avg)

            try:
                generated_descriptions = generate_archetype_descriptions(
                    {
                        "allArchetypes": all_archetypes_for_desc,
                        "primary": primary_for_desc,
                        "supporting": supporting_for_desc,
                        "tertiary": tertiary_for_desc,
                        "personGenitive": personGen,
                    }
                )
            except Exception:
                generated_descriptions = {
                    "valuesWheelDescription": "",
                    "needsWheelDescription": "",
                    "actionProfileDescription": "",
                }

            def _render_auto_description(text: str) -> None:
                safe_text = str(text or "").strip()
                if not safe_text or safe_text in {"0", "0.0"}:
                    return
                st.markdown(
                    (
                        "<div style='margin-top:10px;margin-bottom:16px;"
                        "font-size:0.92em;line-height:1.58;color:var(--text-color,#334155);'>"
                        f"{html.escape(safe_text)}"
                        "</div>"
                    ),
                    unsafe_allow_html=True,
                )

            def _render_radar_role_legend() -> None:
                st.markdown(
                    """
                    <style>
                      .ap-radar-role-legend{
                        display:flex;
                        justify-content:center;
                        align-items:center;
                        gap:24px;
                        flex-wrap:wrap;
                        /* TU REGULUJESZ ODSTĘP LEGENDY OD WYKRESU "Profil archetypów":
                           ujemny margin-top przybliża legendę do radaru. */
                        margin-top:-12px;
                        margin-bottom:-2px;
                        color:var(--text-color,#334155);
                        font-size:0.90em;
                        font-weight:500;
                      }
                      .ap-radar-role-legend .ap-item{
                        display:inline-flex;
                        align-items:center;
                        gap:8px;
                      }
                      .ap-radar-role-legend .ap-dot{
                        width:21px;
                        height:21px;
                        border-radius:50%;
                        display:inline-block;
                        border:2px solid var(--text-color,#0f172a);
                        box-sizing:border-box;
                      }
                      .ap-radar-role-legend .ap-dot-main{ background:#ff0000; }
                      .ap-radar-role-legend .ap-dot-aux{ background:#FFD22F; }
                      .ap-radar-role-legend .ap-dot-supp{ background:#40b900; }
                      @media (max-width: 900px){
                        .ap-radar-role-legend{
                          gap:12px;
                          margin-top:-16px;
                          margin-bottom:-4px;
                          font-size:0.86em;
                        }
                        .ap-radar-role-legend .ap-dot{
                          width:17px;
                          height:17px;
                        }
                      }
                    </style>
                    <div class="ap-radar-role-legend">
                      <span class="ap-item"><span class="ap-dot ap-dot-main"></span><span>Archetyp główny</span></span>
                      <span class="ap-item"><span class="ap-dot ap-dot-aux"></span><span>Archetyp wspierający</span></span>
                      <span class="ap-item"><span class="ap-dot ap-dot-supp"></span><span>Archetyp poboczny</span></span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            def _render_image_90pct(image_obj, dark_image_obj=None, force_dark: bool = False) -> None:
                _pad_l, _img_col, _pad_r = st.columns([0.05, 0.90, 0.05], gap="small")
                with _img_col:
                    if force_dark and dark_image_obj is not None:
                        st.image(dark_image_obj, use_column_width=True)
                        return
                    if dark_image_obj is not None:
                        light_uri = _img_data_uri_from_pil(image_obj)
                        dark_uri = _img_data_uri_from_pil(dark_image_obj)
                        st.markdown(
                            _theme_image_dual_html(
                                light_uri,
                                dark_uri,
                                style="width:100%;height:auto;display:block;",
                            ),
                            unsafe_allow_html=True,
                        )
                    else:
                        st.image(image_obj, use_column_width=True)

            def _render_theme_path_image(
                light_path: str | Path,
                dark_path: str | Path | None,
                width_px: int | None = None,
                img_css_class: str | None = None,
                force_dark: bool = False,
            ) -> None:
                light_path_obj = Path(light_path)
                dark_path_obj = Path(dark_path) if dark_path else None
                if force_dark and dark_path_obj and dark_path_obj.exists():
                    if width_px:
                        st.image(str(dark_path_obj), width=int(width_px))
                    else:
                        st.image(str(dark_path_obj), use_column_width=True)
                    return
                if dark_path_obj and dark_path_obj.exists() and light_path_obj.exists():
                    light_uri = _img_data_uri_from_path(light_path_obj)
                    dark_uri = _img_data_uri_from_path(dark_path_obj)
                    width_style = (
                        f"width:min(100%, {int(width_px)}px);height:auto;display:block;"
                        if width_px
                        else "width:100%;height:auto;display:block;"
                    )
                    st.markdown(
                        _theme_image_dual_html(
                            light_uri,
                            dark_uri,
                            style=width_style,
                            extra_class=img_css_class,
                        ),
                        unsafe_allow_html=True,
                    )
                    return
                if img_css_class and light_path_obj.exists():
                    light_uri = _img_data_uri_from_path(light_path_obj)
                    width_style = (
                        f"width:min(100%, {int(width_px)}px);height:auto;display:block;"
                        if width_px
                        else "width:100%;height:auto;display:block;"
                    )
                    st.markdown(
                        f"<img src='{light_uri}' class='{img_css_class}' style='{width_style}'/>",
                        unsafe_allow_html=True,
                    )
                    return
                if width_px:
                    st.image(str(light_path_obj), width=int(width_px))
                else:
                    st.image(str(light_path_obj), use_column_width=True)

            # --- LICZEBNOŚCI TYLKO DO TABELI (NIE DO RANKINGU/WYKRESU) ---
            counts_main = (
                results_df['Główny archetyp'].map(normalize)
                .value_counts().reindex(archetype_names, fill_value=0)
            )

            counts_aux = (
                results_df['Archetyp wspierający'].map(normalize)
                .value_counts().reindex(archetype_names, fill_value=0)
            )

            counts_supp = (
                results_df['Archetyp poboczny']
                .map(normalize)
                .value_counts()
                .reindex(archetype_names, fill_value=0)
            )

            # wykres skumulowany (PNG) dla Worda/PDF
            stacked_png_path = make_stacked_bar_png_for_word(
                archetype_names=ARCHE_NAMES_ORDER,  # lub Twoja lista 'archetype_names'
                counts_main=counts_main,
                counts_aux=counts_aux,
                counts_supp=counts_supp,
                out_path="archetypes_stacked.png",
            )

            with col1:
                st.markdown(
                    ap_section_heading(
                        "Liczebność i natężenie archetypów",
                        center=False,
                        margin_bottom_px=8,
                        margin_top_px=mobile_section_margin_top,
                    ),
                    unsafe_allow_html=True
                )

                # 2) kolejność wierszy – tie-break:
                #    % natężenia (1 miejsce po przecinku) ↓,
                #    liczba wskazań: główny ↓, wspierający ↓, poboczny ↓,
                #    na końcu alfabetycznie ↑.
                def _summary_rank_key(arche_name: str) -> tuple[float, int, int, int, str]:
                    pct_val = round(float(means_pct.get(arche_name, 0.0) or 0.0), 1)
                    main_cnt = int(counts_main.get(normalize(arche_name), 0) or 0)
                    aux_cnt = int(counts_aux.get(normalize(arche_name), 0) or 0)
                    supp_cnt = int(counts_supp.get(normalize(arche_name), 0) or 0)
                    alpha_key = _slug_pl(disp_name(arche_name) or arche_name)
                    return (-pct_val, -main_cnt, -aux_cnt, -supp_cnt, alpha_key)

                ordered_names = sorted(archetype_names, key=_summary_rank_key)

                # 3) budowa tabeli z nową kolumną „% natężenie archetypu”
                # 1) nagłówek grupy + link „i” (bez JS – otwiera modal CSS)
                NAT_GRP = (
                    "<a href='#ap-intensity-modal' class='ap-int-info' title='Co oznaczają progi?'>"
                    "<img src='https://cdn3.iconfinder.com/data/icons/thunderstorm-5/80/5-32-512.png' "
                    "alt='info' style='width:16px;height:16px'/>"
                    "</a>"
                )

                # 2) wartości % i opisy słowne (z ikonką – kolorowe KWADRATY, nie kółka)
                _pct_vals = [float(means_pct.get(n, 0.0)) for n in ordered_names]
                _pct_strs = [f"{v:.1f}%" for v in _pct_vals]
                _labels = [interpret_archetype_intensity(v)["short"] for v in _pct_vals]
                _labels_ico = [intensity_icon_html(s) for s in _labels]

                # 3) MultiIndex → powstaje nagłówek dwupoziomowy z 2 „podkolumnami”
                _cols = pd.MultiIndex.from_tuples([
                    ("", "Archetyp"),
                    ("", "Główny"),
                    ("", "Wspierający"),
                    ("", "Poboczny"),
                    (NAT_GRP, "%"),
                    (NAT_GRP, "opis"),
                ])

                archetype_table = pd.DataFrame({
                    ("", "Archetyp"): [
                        (
                            f"<span class='ap-arch-emoji'>{html.escape(get_emoji(n))}</span>"
                            f"<span class='ap-arch-name'>{html.escape(disp_name(n))}</span>"
                        )
                        for n in ordered_names
                    ],
                    ("", "Główny"): [zero_to_dash(counts_main.get(normalize(n), 0)) for n in ordered_names],
                    ("", "Wspierający"): [zero_to_dash(counts_aux.get(normalize(n), 0)) for n in ordered_names],
                    ("", "Poboczny"): [zero_to_dash(counts_supp.get(normalize(n), 0)) for n in ordered_names],
                    (NAT_GRP, "%"): _pct_strs,
                    (NAT_GRP, "opis"): _labels_ico,
                }, columns=_cols)

                # 4) kolejność jest już nadana przez ordered_names (z tie-break),
                #    więc nie wykonujemy dodatkowego sortowania DataFrame.
                archetype_table = archetype_table.reset_index(drop=True)

                # 5) HTML + CSS tabeli -> tabela: Liczebność i natężenie archetypów
                # --- ŁATWE DO ZMIANY SZEROKOŚCI (procenty) ---
                COL_W = {"c1": "27%",
                         "c2": "7%",
                         "c3": "7%",
                         "c4": "7%",
                         "c5": "3%",
                         "c6": "52%"}

                # Budujemy body tabeli bez nagłówka (header=False), a nagłówek zrobimy ręcznie (rowspan/colspan).
                _body = (
                    archetype_table.to_html(index=False, header=False, escape=False, border=0)
                    .replace('class="dataframe"', 'class="ap-table"')
                    .replace('border="1"', 'border="0"')
                )

                # Ręcznie złożony nagłówek (scalenia jak na screenie: 4 kolumny z rowspan=2 i grupa 2-kolumnowa)
                thead_html = f"""
                <thead>
                  <tr>
                    <th rowspan="2" style="width:{COL_W['c1']}">Archetyp</th>
                    <th rowspan="2" class="ap-vert-head" style="width:{COL_W['c2']}"><span class="ap-vert-col">Główny</span></th>
                    <th rowspan="2" class="ap-vert-head" style="width:{COL_W['c3']}"><span class="ap-vert-col">Wspierający</span></th>
                    <th rowspan="2" class="ap-vert-head" style="width:{COL_W['c4']}"><span class="ap-vert-col">Poboczny</span></th>
                    <th colspan="2" style="width:calc({COL_W['c5']} + {COL_W['c6']})">
                      % natężenie archetypu&nbsp;{NAT_GRP}
                    </th>
                  </tr>
                  <tr>
                    <th style="width:{COL_W['c5']}">%</th>
                    <th style="width:{COL_W['c6']}">opis</th>
                  </tr>
                </thead>
                """

                # Wstrzykujemy thead przed <tbody>
                html_table = _body.replace("<tbody>", thead_html + "<tbody>", 1)
                table_theme_override_css = ""
                if public_dark_mode:
                    table_theme_override_css = """
                      .ap-table{
                        background:transparent !important;
                        color:#dce8f8 !important;
                      }
                      .ap-table th, .ap-table td{
                        color:#dce8f8 !important;
                        background:transparent !important;
                        border-bottom:1px solid rgba(148,163,184,.34) !important;
                      }
                      .ap-table thead th{
                        color:#e9f2ff !important;
                      }
                      .ap-table .ap-int-ico{
                        border-color:rgba(148,163,184,.44) !important;
                      }
                    """

                st.markdown(f"""
                <style>
                  .ap-table {{
                    table-layout: fixed; width: 100%; border-collapse: collapse;
                    font-family: 'Segoe UI', system-ui, -apple-system, Arial, sans-serif;
                    font-size: 14px; margin-top: 3px;
                  }}
                  .ap-table th, .ap-table td {{
                    padding: 13px 10px; border-bottom: 1px solid #eaeaea;
                    vertical-align: middle; line-height: 1.22; text-align: center;
                  }}
                  .ap-table thead th {{ font-weight: 700; }}
                  .ap-table thead th.ap-vert-head {{ padding:6px 4px; }}
                  .ap-table .ap-vert-col {{
                    display:inline-block;
                    writing-mode:vertical-rl;
                    transform:rotate(180deg);
                    white-space:nowrap;
                    letter-spacing:0;
                    line-height:1;
                    min-height:78px;
                    font-family:'Segoe UI', system-ui, -apple-system, Arial, sans-serif !important;
                    font-size:14px !important;
                    font-weight:700 !important;
                  }}
                  .ap-table .ap-arch-emoji {{
                    display:inline-block;
                    margin-right:6px;
                  }}
                  .ap-table .ap-arch-name {{
                    display:inline-block;
                  }}
                  @media (max-width:1920px), (max-height:1200px) {{
                    .ap-table .ap-arch-emoji {{
                      display:none !important;
                      margin-right:0 !important;
                    }}
                  }}

                  /* Wyrównania tylko dla WIERZY (tbody) w kolumnach 1 i 6 */
                  .ap-table tbody td:nth-child(1),
                  .ap-table tbody td:nth-child(6) {{ text-align: left !important; }}

                  /* Delikatna, szara ikonka „info” */
                  .ap-int-info img {{
                    opacity:.65; filter:grayscale(100%); vertical-align:-3px;
                    width:16px; height:16px;
                  }}
                  .ap-int-info:hover img {{ opacity:1; filter:none; }}

                  /* Kwadracik przy opisie natężenia */
                  .ap-int-ico {{
                    display:inline-block; width:12px; height:12px; border-radius:3px;
                    border:1px solid #d1d5db; margin-right:5px; vertical-align:-2px;
                  }}

                  .ap-table tbody td:nth-child(6) {{
                    white-space: nowrap !important;
                  }}

                  @media (max-width:1920px), (max-height:1200px) {{
                    .ap-table {{
                      font-size: 13px !important;
                    }}
                    .ap-table th, .ap-table td {{
                      padding: 11px 7px !important;
                    }}
                  }}
                  {table_theme_override_css}
                </style>
                """, unsafe_allow_html=True)

                # 👉 Auto-wysokość i brak iframa (Streamlit ≥ 1.38 ma st.html)
                _table_rows = len(ordered_names) if isinstance(ordered_names, list) else 12
                # trochę większy zapas na nagłówek i odstępy
                _table_height = 240 + 56 * _table_rows

                # --- NOWY FRAGMENT DO WSTAWIENIA ---
                _html_block = f"""
                <style>
                  .ap-table-wrap {{
                    width: 100%;
                    overflow-x: hidden;
                  }}
                  .ap-table {{
                    table-layout: fixed; width: 100%; border-collapse: collapse;
                    font-family: 'Segoe UI', system-ui, -apple-system, Arial, sans-serif;
                    font-size: 14px; margin-top: 3px;
                  }}
                  .ap-table th, .ap-table td {{
                    padding: 13px 10px; border-bottom: 1px solid #eaeaea;
                    vertical-align: middle; line-height: 1.22; text-align: center;
                  }}
                  .ap-table thead th {{ font-weight: 700; }}
                  .ap-table thead th.ap-vert-head {{ padding:6px 4px; }}
                  .ap-table .ap-vert-col {{
                    display:inline-block;
                    writing-mode:vertical-rl;
                    transform:rotate(180deg);
                    white-space:nowrap;
                    letter-spacing:0;
                    line-height:1;
                    min-height:78px;
                    font-family:'Segoe UI', system-ui, -apple-system, Arial, sans-serif !important;
                    font-size:14px !important;
                    font-weight:700 !important;
                  }}
                  .ap-table .ap-arch-emoji {{
                    display:inline-block;
                    margin-right:6px;
                  }}
                  .ap-table .ap-arch-name {{
                    display:inline-block;
                  }}
                  @media (max-width:1920px), (max-height:1200px) {{
                    .ap-table .ap-arch-emoji {{
                      display:none !important;
                      margin-right:0 !important;
                    }}
                  }}
                  .ap-table tbody td:nth-child(1),
                  .ap-table tbody td:nth-child(6) {{ text-align: left !important; }}

                  .ap-int-info img {{
                    opacity:.65; filter:grayscale(100%); vertical-align:-3px; width:16px; height:16px;
                  }}
                  .ap-int-info:hover img {{ opacity:1; filter:none; }}

                  .ap-int-ico {{
                    display:inline-block; width:12px; height:12px; border-radius:3px;
                    border:1px solid #d1d5db; margin-right:5px; vertical-align:-2px;
                  }}

                  .ap-table tbody td:nth-child(6) {{
                    white-space: nowrap !important;
                  }}

                  @media (max-width:1920px), (max-height:1200px) {{
                    .ap-table {{
                      font-size: 13px !important;
                    }}
                    .ap-table th, .ap-table td {{
                      padding: 11px 7px !important;
                    }}
                  }}

                  /* 👉 KONFIGURACJA — tu zmieniasz marginesy/paddingi */
                  :root{{
                    --ap-tip-offset: 10px;   /* odległość dymka od ikonki */
                    --ap-tip-pad-v: 10px;    /* padding GÓRA/DÓŁ w dymku */
                    --ap-tip-pad-h: 12px;    /* padding LEWO/PRAWO w dymku */
                  }}

                  /* kontener tooltipa */
                  .ap-tip{{ position:relative; display:inline-block; }}

                  /* sam dymek */
                  .ap-tip-box{{
                    position:absolute;
                    left:50%; transform:translateX(-50%) translateY(4px);
                    top: calc(100% + var(--ap-tip-offset));
                    background:#111827; color:#fff;
                    padding: var(--ap-tip-pad-v) var(--ap-tip-pad-h);
                    border-radius:8px; box-shadow:0 6px 18px rgba(0,0,0,.18);
                    white-space:nowrap; font-size:12px; line-height:1.35;
                    z-index:15; opacity:0; pointer-events:none;
                    transition:opacity .15s ease, transform .15s ease;
                  }}

                  /* strzałka */
                  .ap-tip-box::after{{
                    content:""; position:absolute; top:-6px; left:50%; transform:translateX(-50%);
                    border:6px solid transparent; border-bottom-color:#111827;
                  }}

                  /* pokaż przy najechaniu */
                  .ap-tip:hover .ap-tip-box{{ opacity:1; transform:translateX(-50%) translateY(0); }}

                  /* opcja: dymek NAD ikoną */
                  .ap-tip.ap-tip--above .ap-tip-box{{
                    bottom: calc(100% + var(--ap-tip-offset)); top:auto;
                  }}
                  .ap-tip.ap-tip--above .ap-tip-box::after{{
                    top:auto; bottom:-6px; border-bottom-color:transparent; border-top-color:#111827;
                  }}

                  /* jeśli w tooltipie masz mini-tabelkę */
                    /* 👉 USTAWIENIA — tu ręcznie regulujesz odstępy w wierszach */
                    :root{{
                      --ap-row-pad-top: 18px;       /* górny akapit (padding TOP) w wierszu */
                      --ap-row-pad-bottom: 18px;    /* dolny akapit (padding BOTTOM) w wierszu */
                      --ap-cell-pad-h: 12px;        /* poziomy padding w komórkach tabeli */
                    
                      /* paleta kolorów dla kwadracików (natężenia) — możesz zmieniać */
                      --ap-col-1: #e5e7eb; /* 0–29%  */
                      --ap-col-2: #bfdbfe; /* 30–49% */
                      --ap-col-3: #fde68a; /* 50–59% */
                      --ap-col-4: #fdba74; /* 60–69% */
                      --ap-col-5: #f87171; /* 70–79% */
                      --ap-col-6: #c084fc; /* 80–89% */
                      --ap-col-7: #111827; /* 90–100% */
                    }}
                    
                    /* tabela w modalu z interpretacją */
                    .ap-tip-box .ap-int-table{{
                      border-collapse: collapse; margin-block: 6px; width: 100%;
                    }}
                    .ap-tip-box .ap-int-table th,
                    .ap-tip-box .ap-int-table td{{
                      padding: var(--ap-row-pad-top) var(--ap-cell-pad-h) var(--ap-row-pad-bottom) var(--ap-cell-pad-h);
                      border-bottom: 1px solid #e5e7eb;
                      vertical-align: top; text-align: left;
                    }}
                    .ap-tip-box .ap-int-table thead th{{ font-weight: 700; }}
                    
                    /* szerokości kolumn – opcjonalnie dopasuj pod siebie */
                    .ap-tip-box .ap-int-table th:nth-child(1),
                    .ap-tip-box .ap-int-table td:nth-child(1){{ width: 100px; text-align: center; }}
                    .ap-tip-box .ap-int-table th:nth-child(2),
                    .ap-tip-box .ap-int-table td:nth-child(2){{ width: 450px; }}
                    
                    /* kwadraciki natężenia */
                    .ap-tip-box .ap-int-ico{{
                      display:inline-block; width:12px; height:12px; border-radius:3px;
                      border:1px solid #d1d5db; margin-right:8px; vertical-align:-2px;
                    }}
                    .ap-tip-box .ap-i--1{{ background:var(--ap-col-1); }}
                    .ap-tip-box .ap-i--2{{ background:var(--ap-col-2); }}
                    .ap-tip-box .ap-i--3{{ background:var(--ap-col-3); }}
                    .ap-tip-box .ap-i--4{{ background:var(--ap-col-4); }}
                    .ap-tip-box .ap-i--5{{ background:var(--ap-col-5); }}
                    .ap-tip-box .ap-i--6{{ background:var(--ap-col-6); }}
                    .ap-tip-box .ap-i--7{{ background:var(--ap-col-7); border-color:var(--ap-col-7); }}
                    
                    .ap-tip-box .ap-int-label{{ font-weight: 600; }}

                    /* Tooltip nad kolorowym paskiem heurystyki */
                    .ap-hc-fill[data-tip]{{ position: relative; }}
                    .ap-hc-fill[data-tip]:hover::after{{
                      content: attr(data-tip);
                      position: absolute;
                      left: 14px;                 /* startuje lekko od lewej krawędzi paska */
                      bottom: calc(100% + 8px);   /* nad paskiem */
                      max-width: 520px;
                      background:#111827; color:#fff;
                      padding:8px 10px; border-radius:8px;
                      font:500 13px/1.35 'Segoe UI', system-ui, Arial;
                      box-shadow:0 8px 24px rgba(0,0,0,.2);
                      white-space:normal; z-index:10;
                    }}
                    .ap-hc-fill[data-tip]:hover::before{{
                      content:"";
                      position:absolute; left: 20px; bottom: 100%;
                      border:6px solid transparent;
                      border-top-color:#111827;   /* „trójkącik” */
                      transform: translateY(2px);
                    }}

                    @media (max-width: 900px){{
                      .ap-table {{
                        min-width: 680px !important;
                        width: max-content !important;
                        table-layout: fixed !important;
                      }}
                      .ap-table th, .ap-table td {{
                        padding: 7px 5px !important;
                        font-size: 11.6px !important;
                        white-space: nowrap !important;
                        word-break: normal !important;
                      }}
                      .ap-table .ap-vert-col {{
                        min-height: 66px !important;
                        font-size: 12px !important;
                      }}
                    }}
                    {table_theme_override_css}

                </style>
                """ + f"<div class='ap-table-wrap'>{html_table}</div>" + intensity_help_modal_html()

                # jeśli masz nowe Streamlit: prawdziwy, „wbudowany” HTML bez iframa
                if hasattr(st, "html"):
                    st.html(_html_block)  # auto-dopasowanie wysokości, brak dodatkowego scrolla
                else:
                    # starszy Streamlit – zostaje components, ale bez przewijania (duża wysokość)
                    components.html(_html_block, height=_table_height, scrolling=False)

            with col2:
                radar_base_label_color = "#c9d8ee" if public_dark_mode else "#656565"
                radar_marker_border_color = "#dbe7f8" if public_dark_mode else "black"
                radar_grid_color = "rgba(148,163,184,0.46)" if public_dark_mode else "rgba(148,163,184,0.35)"
                radar_tick_color = "#eef6ff" if public_dark_mode else "#475569"
                radar_radial_tick_color = "#deebfb" if public_dark_mode else "#64748b"
                theta_labels = []
                for n in archetype_names:
                    label = disp_name(n)
                    if n == main_avg:
                        theta_labels.append(f"<b><span style='color:red;'>{label}</span></b>")
                    elif n == aux_avg:
                        theta_labels.append(f"<b><span style='color:#FFD22F;'>{label}</span></b>")
                    elif n == supp_avg:
                        theta_labels.append(f"<b><span style='color:#40b900;'>{label}</span></b>")
                    else:
                        theta_labels.append(f"<span style='color:{radar_base_label_color};'>{label}</span>")

                # markery TOP-3 z mean_archetype_scores
                highlight_r = []
                highlight_marker_color = []
                for name in archetype_names:
                    if name == main_avg:
                        highlight_r.append(mean_archetype_scores.get(name, 0.0))
                        highlight_marker_color.append("red")
                    elif name == aux_avg:
                        highlight_r.append(mean_archetype_scores.get(name, 0.0))
                        highlight_marker_color.append("#FFD22F")
                    elif name == supp_avg:
                        highlight_r.append(mean_archetype_scores.get(name, 0.0))
                        highlight_marker_color.append("#40b900")
                    else:
                        highlight_r.append(None)
                        highlight_marker_color.append("rgba(0,0,0,0)")

                # ŚREDNIE (0–20) w KOLEJNOŚCI archetype_names
                mean_vals_ordered = [mean_archetype_scores.get(n, 0.0) for n in archetype_names]
                theta_display = [disp_name(n) for n in archetype_names]

                fig = go.Figure(
                    data=[
                        go.Scatterpolar(
                            r=mean_vals_ordered + [mean_vals_ordered[0]],
                            theta=theta_display + [theta_display[0]],
                            fill='toself',
                            name='średnia wszystkich',
                            line=dict(color="royalblue", width=3),
                            marker=dict(size=6),
                            # 👇 własny tooltip: bez "r:" i "θ:", zaokrąglenie do 2 miejsc
                            hovertemplate="<b>%{theta}</b><br>średnia: %{r:.2f}<extra></extra>",
                        ),
                        go.Scatterpolar(
                            r=highlight_r,
                            theta=theta_display,
                            mode='markers',
                            marker=dict(size=18, color=highlight_marker_color, opacity=0.90,
                                        line=dict(color=radar_marker_border_color, width=3)),
                            name='Archetyp główny/wspierający/poboczny',
                            showlegend=False,
                            # 👇 spójny tooltip z 2 miejscami po przecinku
                            hovertemplate="<b>%{theta}</b><br>wartość: %{r:.2f}<extra></extra>",
                        )
                    ],
                    layout=go.Layout(
                        polar=dict(
                            radialaxis=dict(
                                visible=True,
                                range=[0, 20],
                                gridcolor=radar_grid_color,
                                linecolor=radar_grid_color,
                                tickfont=dict(color=radar_radial_tick_color),
                                tickcolor=radar_grid_color,
                            ),
                            angularaxis=dict(
                                tickfont=dict(size=19, color=radar_tick_color),
                                tickvals=theta_display,
                                ticktext=theta_labels,
                                rotation=90,
                                direction="clockwise",
                                gridcolor=radar_grid_color,
                                linecolor=radar_grid_color,
                            )
                        ),
                        width=400, height=400,
                        margin=dict(l=20, r=20, t=32, b=32),
                        showlegend=False
                    )
                )

                # PRZEZROCZYSTE TŁO RADARU — bez zmian
                fig.update_layout(
                    paper_bgcolor="rgba(0,0,0,0)",
                    polar=dict(
                        domain=radar_domain,
                        bgcolor="rgba(0,0,0,0)",
                        radialaxis=dict(
                            visible=True,
                            range=[0, 20],
                            gridcolor=radar_grid_color,
                            linecolor=radar_grid_color,
                            tickfont=dict(color=radar_radial_tick_color),
                            tickcolor=radar_grid_color,
                        ),
                        angularaxis=dict(
                            tickfont=dict(size=radar_tick_size, color=radar_tick_color),
                            tickvals=theta_display,
                            ticktext=theta_labels,
                            rotation=90,
                            direction="clockwise",
                            gridcolor=radar_grid_color,
                            linecolor=radar_grid_color,
                        ),
                    ),
                    autosize=True,
                    height=radar_plot_size,
                    margin=radar_margins,
                    showlegend=False,
                )

                # 👇 większa czcionka w dymkach hover
                fig.update_layout(hoverlabel=dict(font=dict(size=radar_hover_size)))

                radar_config = {
                    "displaylogo": False,
                    "displayModeBar": (not is_mobile),
                    "responsive": True,
                }
                radar_heading_mb = 2 if is_mobile else 8

                if is_mobile:
                    st.markdown(
                        ap_section_heading(
                            f"Profil archetypów {personGen}",
                            center=True,
                            margin_bottom_px=radar_heading_mb,
                            margin_top_px=mobile_section_margin_top,
                        ),
                        unsafe_allow_html=True,
                    )
                    st.plotly_chart(
                        fig,
                        use_container_width=True,
                        config=radar_config,
                        key=f"radar-{study_id}",
                    )
                    _render_radar_role_legend()
                else:
                    st.markdown(
                        ap_section_heading(
                            f"Profil archetypów {personGen}",
                            center=True,
                            margin_bottom_px=radar_heading_mb,
                            margin_top_px=mobile_section_margin_top,
                        ),
                        unsafe_allow_html=True,
                    )
                    st.plotly_chart(
                        fig,
                        use_container_width=True,
                        config=radar_config,
                        key=f"radar-{study_id}",
                    )
                    _render_radar_role_legend()

            # --- Heurystyczna analiza koloru + profil podsumowania pod tabelą ---
            color_pcts = calc_color_percentages_from_df(data)
            dom_name, dom_pct = max(color_pcts.items(), key=lambda kv: kv[1])

            # Dane opisowe dominującego koloru do Worda
            dom_meta = COLOR_LONG[dom_name]  # masz już COLOR_LONG w pliku
            dom_color = {
                "name": dom_name,
                "pct": dom_pct,
                "emoji": COLOR_EMOJI[dom_name],
                "title": dom_meta["title"],
                "orient": dom_meta["orient"],
                "body": dom_meta["body"],
                "politics": dom_meta["politics"],
                "hex": dom_meta["hex"],
            }

            segment_profile_png_path = make_segment_profile_wheel_png(
                mean_scores=means_pct,
                out_path=f"segment_profile_{study_id}.png",
                gender_code=report_gender_code,
                dark_mode=False,
            )
            segment_profile_png_path_dark = make_segment_profile_wheel_png(
                mean_scores=means_pct,
                out_path=f"segment_profile_{study_id}_dark.png",
                gender_code=report_gender_code,
                dark_mode=True,
            )

            with left_col:
                st.markdown("<div style='height:22px;'></div>", unsafe_allow_html=True)
                st.markdown(
                    ap_section_heading(
                        "Heurystyczna analiza koloru psychologicznego",
                        center=False,
                        margin_bottom_px=8,
                        margin_top_px=mobile_section_margin_top,
                    ),
                    unsafe_allow_html=True,
                )
                components.html(
                    color_progress_bars_html(
                        color_pcts,
                        order="desc",
                        label_color=("#d4e2f4" if public_dark_mode else "#31333F"),
                        track_color=("#dbe2ec" if public_dark_mode else "#eef2f7"),
                    ),
                    height=280,
                    scrolling=False,
                )
                st.markdown("<style>.cp-row{margin:15px 0 !important}</style>", unsafe_allow_html=True)
                st.markdown(
                    f"<div style='text-align:center; font:680 20px/1.30 \"Roboto\",\"Segoe UI\",\"Arial\",system-ui,sans-serif; color:{'#dbe8f8' if public_dark_mode else '#222'}; margin: -15px 0 60px;'>"
                    f"Dominujący kolor: <span style='color:{COLOR_HEX[dom_name]}'>{dom_name}</span></div>",
                    unsafe_allow_html=True,
                )
                st.markdown(color_explainer_one_html(dom_name, dom_pct, dark_mode=public_dark_mode), unsafe_allow_html=True)

                st.markdown("<div style='height:24px;'></div>", unsafe_allow_html=True)
                st.markdown(
                    ap_section_heading(
                        f"Profil siły archetypów {personGen} (skala: 0-100)",
                        center=False,
                        margin_bottom_px=12,
                        margin_top_px=mobile_section_margin_top,
                    ),
                    unsafe_allow_html=True,
                )
                if is_mobile:
                    _render_theme_path_image(
                        segment_profile_png_path,
                        segment_profile_png_path_dark,
                        width_px=None,
                        img_css_class="ap-strength-wheel-img",
                        force_dark=public_dark_mode,
                    )
                else:
                    _render_theme_path_image(
                        segment_profile_png_path,
                        segment_profile_png_path_dark,
                        width_px=segment_profile_display_width_desktop,
                        img_css_class="ap-strength-wheel-img",
                        force_dark=public_dark_mode,
                    )
                st.markdown(
                    """
                    <div style="display:flex;gap:24px;flex-wrap:wrap;align-items:center;justify-content:flex-start;margin-top:8px;margin-bottom:6px;font-size:1.03em;font-weight:600;color:var(--text-color,#475569);">
                      <span style="display:inline-flex;align-items:center;gap:7px;"><span style="width:11px;height:11px;background:#de4b43;border-radius:2px;display:inline-block;"></span>Zmiana</span>
                      <span style="display:inline-flex;align-items:center;gap:7px;"><span style="width:11px;height:11px;background:#2d5ad5;border-radius:2px;display:inline-block;"></span>Ludzie</span>
                      <span style="display:inline-flex;align-items:center;gap:7px;"><span style="width:11px;height:11px;background:#2f8a45;border-radius:2px;display:inline-block;"></span>Porządek</span>
                      <span style="display:inline-flex;align-items:center;gap:7px;"><span style="width:11px;height:11px;background:#6f53d4;border-radius:2px;display:inline-block;"></span>Niezależność</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )


            with col3:
                if is_mobile:
                    st.markdown(
                        ap_section_heading(
                            "Koło pragnień i wartości",
                            center=True,
                            margin_bottom_px=8,
                            margin_top_px=mobile_section_margin_top,
                            shift_x_px=0,
                        ),
                        unsafe_allow_html=True,
                    )
                    if main_avg is not None:
                        idx_main_wheel = archetype_name_to_img_idx(main_avg)
                        idx_aux_wheel = archetype_name_to_img_idx(aux_avg) if aux_avg != main_avg else None
                        idx_supp_wheel = (
                            archetype_name_to_img_idx(supp_avg) if supp_avg not in [main_avg, aux_avg] else None
                        )
                        try:
                            kola_img = compose_archetype_highlight(
                                idx_main_wheel,
                                idx_aux_wheel,
                                idx_supp_wheel,
                                gender_code=report_gender_code,
                                dark_mode=False,
                            )
                            kola_img_dark = compose_archetype_highlight(
                                idx_main_wheel,
                                idx_aux_wheel,
                                idx_supp_wheel,
                                gender_code=report_gender_code,
                                dark_mode=True,
                            )
                            if not isinstance(kola_img, Image.Image):
                                raise TypeError("compose_archetype_highlight nie zwrócił obrazu PIL")
                        except Exception:
                            kola_img = load_base_arche_img(gender_code=report_gender_code, dark_mode=False)
                            kola_img_dark = load_base_arche_img(gender_code=report_gender_code, dark_mode=True)
                        _render_image_90pct(kola_img, dark_image_obj=kola_img_dark, force_dark=public_dark_mode)
                        st.markdown(
                            "<div style='margin:6px auto 6px auto;width:fit-content;max-width:100%;font-size:0.88em;color:var(--text-color,#64748b);text-align:center;'>"
                            "Podświetlenie: główny – czerwony, wspierający – żółty, poboczny – zielony"
                            "</div>",
                            unsafe_allow_html=True,
                        )
                        _render_auto_description(generated_descriptions["valuesWheelDescription"])
                else:
                    st.markdown(
                        ap_section_heading(
                            "Koło pragnień i wartości",
                            center=True,
                            margin_bottom_px=8,
                            margin_top_px=mobile_section_margin_top,
                            shift_x_px=0,
                        ),
                        unsafe_allow_html=True,
                    )
                    if main_avg is not None:
                        idx_main_wheel = archetype_name_to_img_idx(main_avg)
                        idx_aux_wheel = archetype_name_to_img_idx(aux_avg) if aux_avg != main_avg else None
                        idx_supp_wheel = (
                            archetype_name_to_img_idx(supp_avg) if supp_avg not in [main_avg, aux_avg] else None
                        )
                        try:
                            kola_img = compose_archetype_highlight(
                                idx_main_wheel,
                                idx_aux_wheel,
                                idx_supp_wheel,
                                gender_code=report_gender_code,
                                dark_mode=False,
                            )
                            kola_img_dark = compose_archetype_highlight(
                                idx_main_wheel,
                                idx_aux_wheel,
                                idx_supp_wheel,
                                gender_code=report_gender_code,
                                dark_mode=True,
                            )
                            if not isinstance(kola_img, Image.Image):
                                raise TypeError("compose_archetype_highlight nie zwrócił obrazu PIL")
                        except Exception:
                            kola_img = load_base_arche_img(gender_code=report_gender_code, dark_mode=False)
                            kola_img_dark = load_base_arche_img(gender_code=report_gender_code, dark_mode=True)
                        _render_image_90pct(kola_img, dark_image_obj=kola_img_dark, force_dark=public_dark_mode)
                        st.markdown(
                            "<div style='margin:6px auto 6px auto;width:fit-content;max-width:100%;font-size:0.88em;color:var(--text-color,#64748b);text-align:center;'>"
                            "Podświetlenie: główny – czerwony, wspierający – żółty, poboczny – zielony"
                            "</div>",
                            unsafe_allow_html=True,
                        )
                        _render_auto_description(generated_descriptions["valuesWheelDescription"])

                if is_mobile:
                    st.markdown("<div style='height:24px;'></div>", unsafe_allow_html=True)
                    st.markdown(
                        ap_section_heading(
                            "Koło potrzeb",
                            center=True,
                            margin_bottom_px=8,
                            margin_top_px=mobile_section_margin_top,
                            shift_x_px=0,
                        ),
                        unsafe_allow_html=True,
                    )
                    aux = aux_avg if aux_avg != main_avg else None
                    supp = supp_avg if supp_avg not in [main_avg, aux_avg] else None
                    kolo_axes_img = compose_axes_wheel_highlight(
                        main_avg,
                        aux,
                        supp,
                        gender_code=report_gender_code,
                        dark_mode=False,
                    )
                    kolo_axes_img_dark = compose_axes_wheel_highlight(
                        main_avg,
                        aux,
                        supp,
                        gender_code=report_gender_code,
                        dark_mode=True,
                    )
                    _render_image_90pct(kolo_axes_img, dark_image_obj=kolo_axes_img_dark, force_dark=public_dark_mode)
                    _render_auto_description(generated_descriptions["needsWheelDescription"])
                else:
                    st.markdown("<div style='height:24px;'></div>", unsafe_allow_html=True)
                    st.markdown(
                        ap_section_heading(
                            "Koło potrzeb",
                            center=True,
                            margin_bottom_px=8,
                            margin_top_px=mobile_section_margin_top,
                            shift_x_px=0,
                        ),
                        unsafe_allow_html=True,
                    )
                    aux = aux_avg if aux_avg != main_avg else None
                    supp = supp_avg if supp_avg not in [main_avg, aux_avg] else None
                    kolo_axes_img = compose_axes_wheel_highlight(
                        main_avg,
                        aux,
                        supp,
                        gender_code=report_gender_code,
                        dark_mode=False,
                    )
                    kolo_axes_img_dark = compose_axes_wheel_highlight(
                        main_avg,
                        aux,
                        supp,
                        gender_code=report_gender_code,
                        dark_mode=True,
                    )
                    _render_image_90pct(kolo_axes_img, dark_image_obj=kolo_axes_img_dark, force_dark=public_dark_mode)
                    _render_auto_description(generated_descriptions["needsWheelDescription"])

            top_profile_archetypes: list[str] = [main_avg]
            if aux_avg and aux_avg != main_avg:
                top_profile_archetypes.append(aux_avg)
            if supp_avg and supp_avg not in [main_avg, aux_avg]:
                top_profile_archetypes.append(supp_avg)

            profile_cards_for_top: list[tuple[int, str, str, Path, Path | None]] = []
            def _role_for_rank(rank: int) -> str:
                if rank == 1:
                    return "Archetyp główny"
                if rank == 2:
                    return "Archetyp wspierający"
                return "Archetyp poboczny"

            for rank, arch_name in enumerate(top_profile_archetypes, start=1):
                profile_path = _profile_card_file_for(arch_name, gender_code=report_gender_code)
                if profile_path and Path(profile_path).exists():
                    light_path = Path(profile_path)
                    dark_path = _profile_card_dark_variant(light_path)
                    profile_cards_for_top.append((rank, _role_for_rank(rank), disp_name(arch_name), light_path, dark_path))
            profile_cards_export: list[dict[str, str]] = [
                {
                    "role": role_label,
                    "name": title,
                    "path": str(light_path),
                    "dark_path": str(dark_path) if dark_path else "",
                }
                for _, role_label, title, light_path, dark_path in profile_cards_for_top
            ]

            def _render_top_profile_cards(items: list[tuple[int, str, str, Path, Path | None]]) -> None:
                if not items:
                    return
                st.markdown("<div style='height:28px;'></div>", unsafe_allow_html=True)
                st.markdown(
                    ap_section_heading(
                        f"Profile działania archetypów {personGen}",
                        center=False,
                        margin_bottom_px=12,
                        margin_top_px=mobile_profile_actions_margin_top,
                    ),
                    unsafe_allow_html=True,
                )
                for rank, role_label, title, light_path, dark_path in items:
                    st.markdown(
                        (
                            "<div style='font-size:.9em;font-weight:700;color:var(--text-color,#475569);"
                            "margin-top:14px;margin-bottom:10px;'>"
                            f"{html.escape(role_label)}: {html.escape(title)} - profil działania"
                            "</div>"
                        ),
                        unsafe_allow_html=True,
                    )
                    if public_dark_mode and dark_path and dark_path.exists():
                        st.image(str(dark_path), use_column_width=True)
                    elif dark_path and dark_path.exists():
                        light_uri = _img_data_uri_from_path(light_path)
                        dark_uri = _img_data_uri_from_path(dark_path)
                        st.markdown(
                            _theme_image_dual_html(
                                light_uri,
                                dark_uri,
                                style="width:100%;height:auto;display:block;",
                            ),
                            unsafe_allow_html=True,
                        )
                    else:
                        st.image(str(light_path), use_column_width=True)
                    st.markdown("<div style='height:18px;'></div>", unsafe_allow_html=True)
                _render_auto_description(generated_descriptions["actionProfileDescription"])

            with col3:
                _render_top_profile_cards(profile_cards_for_top)

            st.markdown("""
            <hr style="height:1px; border:none; background:#eee; margin-top:34px; margin-bottom:19px;" />
            """, unsafe_allow_html=True)
            st.markdown("<div id='opisy'></div>", unsafe_allow_html=True)
            st.markdown(f'<div style="font-size:2.1em;font-weight:700;margin-bottom:16px;">Archetyp główny {personGen}</div>', unsafe_allow_html=True)
            render_archetype_card(main_disp, main=True, gender_code=report_gender_code)

            if aux_avg and aux_avg != main_avg:
                st.markdown("<div style='height:35px;'></div>", unsafe_allow_html=True)
                st.markdown("""<hr style="height:1.1px; border:none; background:#ddd; margin-top:6px; margin-bottom:18px;" />""", unsafe_allow_html=True)
                st.markdown(f"<div style='font-size:1.63em;font-weight:700;margin-bottom:15px;'>Archetyp wspierający {personGen}</div>", unsafe_allow_html=True)
                render_archetype_card(second_disp, main=False, gender_code=report_gender_code)

            if supp_avg and supp_avg not in [main_avg, aux_avg]:
                st.markdown("<div style='height:35px;'></div>", unsafe_allow_html=True)
                st.markdown("""<hr style="height:1.1px; border:none; background:#ddd; margin-top:6px; margin-bottom:18px;" />""", unsafe_allow_html=True)
                st.markdown(f"<div style='font-size:1.63em;font-weight:700;margin-bottom:15px;'>Archetyp poboczny {personGen}</div>", unsafe_allow_html=True)
                render_archetype_card(supp_disp, main=False, supplement=True, gender_code=report_gender_code)

            if public_view:
                return

            st.markdown("<div id='raport'></div>", unsafe_allow_html=True)

            st.markdown(f"""
            <div style='height:44px;'></div>
            <hr style="height:1px; border:none; background:#e5e5e5; margin-bottom:26px;" />
            <div style="font-size:1.2em; font-weight:600; margin-bottom:23px;">
                Pobierz raporty archetypu {personGen}
            </div>
            """, unsafe_allow_html=True)

            # GENEROWANIE OBRAZU PANELU DYNAMICZNIE
            idx_main = archetype_name_to_img_idx(main_avg)
            idx_aux = archetype_name_to_img_idx(aux_avg) if aux_avg != main_avg else None
            idx_supplement = archetype_name_to_img_idx(supp_avg) if supp_avg not in [main_avg, aux_avg] else None

            # === Przygotowanie eksportów tylko na żądanie ===
            prep = st.toggle(
                "Przygotuj pliki Word/PDF (kliknij tylko gdy chcesz pobrać)",
                key="prep_docs",
                value=st.session_state.get("prep_docs", False)
            )

            # Ustal nazwy plików z góry (przydadzą się też przy cache)
            DOCX_FILENAME_FULL, PDF_FILENAME_FULL = build_report_filenames(study)
            DOCX_FILENAME_SHORT, PDF_FILENAME_SHORT = build_short_report_filenames(study)
            cache_key = f"exports_{study_id}"

            def _build_exports() -> dict[str, bytes]:
                """Liczy wszystkie obrazy do Worda, buduje raport full + short (DOCX/PDF)."""
                # 1) Zapisz PNG z radaru (do DOCX)
                try:
                    fig.write_image("radar.png", scale=4)
                except Exception:
                    pass

                # 2) Koło osi potrzeb → PNG
                aux = aux_avg if aux_avg != main_avg else None
                supp = supp_avg if supp_avg not in [main_avg, aux_avg] else None
                kolo_axes_img = compose_axes_wheel_highlight(main_avg, aux, supp, gender_code=report_gender_code)
                kolo_axes_img.save("axes_wheel.png")

                # 3) Dominujący pierścień koloru → SVG/PNG
                big_color = max(color_pcts.items(), key=lambda kv: kv[1])[0]
                big_svg = _ring_svg(color_pcts[big_color], COLOR_HEX[big_color], size=600,
                                    stroke=48)
                # Nie zapisujemy pliku pośredniego SVG na dysku.
                cairosvg.svg2png(bytestring=big_svg.encode("utf-8"), write_to="color_ring.png")

                # 4) Pastylki kolorów (PNG)
                progress_png_path = make_color_progress_png_for_word(
                    color_pcts, width_px=1600, pad=32, bar_h=66, bar_gap=30,
                    dot_radius=10, label_gap_px=35, label_font_size=36, pct_font_size=32,
                    pct_margin=14
                )

                # 5) Skumulowany wykres liczebności (PNG)
                stacked_png_path = make_stacked_bar_png_for_word(
                    archetype_names=ARCHE_NAMES_ORDER,
                    counts_main=counts_main, counts_aux=counts_aux, counts_supp=counts_supp,
                    out_path="archetypes_stacked.png",
                )

                # 6) Kapsuły średnich (PNG)
                means_pct = mean_pct_by_archetype_from_df(data)
                capsules_path = make_capsule_columns_png_for_word(
                    means_pct, out_path="arche_capsules.png", top_title=None
                )

                # 7) Panel (czerwony/żółty/zielony) → PNG
                idx_main = archetype_name_to_img_idx(main_avg)
                idx_aux = archetype_name_to_img_idx(aux_avg) if aux_avg != main_avg else None
                idx_supp = archetype_name_to_img_idx(supp_avg) if supp_avg not in [main_avg,
                                                                                   aux_avg] else None
                panel_img = compose_archetype_highlight(
                    idx_main,
                    idx_aux,
                    idx_supp,
                    gender_code=report_gender_code,
                )
                panel_img_path = f"panel_{(main_avg or '').lower()}_{(aux_avg or '').lower()}_{(supp_avg or '').lower()}.png"
                panel_img.save(panel_img_path)

                # 8) FULL DOCX
                docx_full_io = export_word_docxtpl(
                    main_avg,
                    aux_avg,
                    supp_avg,
                    archetype_features,
                    main_disp,
                    second_disp,
                    supp_disp,
                    mean_scores=means_pct,  # ⬅️ przekażemy % do wyliczenia natężenia
                    radar_img_path="radar.png",
                    archetype_table=archetype_table,
                    num_ankiet=num_ankiet,
                    panel_img_path=panel_img_path,
                    person=person,
                    gender_code=report_gender_code,
                    axes_wheel_img_path="axes_wheel.png",
                    dom_color=dom_color,
                    color_progress_img_path=progress_png_path,
                    archetype_stacked_img_path=stacked_png_path,
                    capsule_columns_img_path=capsules_path,
                    segment_profile_img_path=segment_profile_png_path,
                    action_profiles=profile_cards_export,
                    generated_descriptions=generated_descriptions,
                    show_supplement=SHOW_SUPP
                )

                # 9) SHORT DOCX (tylko metryka)
                docx_short_io = export_word_metrics_only(
                    main_avg,
                    aux_avg,
                    supp_avg,
                    main_disp,
                    second_disp,
                    supp_disp,
                    person=person,
                    show_supplement=SHOW_SUPP,
                    segment_profile_img_path=segment_profile_png_path,
                    features=archetype_features,
                    mean_scores=means_pct,
                    radar_img_path="radar.png",
                    archetype_table=archetype_table,
                    num_ankiet=num_ankiet,
                    panel_img_path=panel_img_path,
                    gender_code=report_gender_code,
                    axes_wheel_img_path="axes_wheel.png",
                    dom_color=dom_color,
                    color_progress_img_path=progress_png_path,
                    archetype_stacked_img_path=stacked_png_path,
                    capsule_columns_img_path=capsules_path,
                )

                # 10) PDF – platformowo:
                import sys as _sys
                if _sys.platform.startswith("win"):
                    pdf_full_io = word_to_pdf(docx_full_io)
                    pdf_short_io = word_to_pdf(docx_short_io)
                else:
                    soffice_path = _find_soffice()
                    if not soffice_path:
                        raise RuntimeError("LibreOffice (soffice) nie jest dostępny w systemie.")
                    pdf_full_io = word_to_pdf(docx_full_io, soffice_bin=soffice_path)
                    pdf_short_io = word_to_pdf(docx_short_io, soffice_bin=soffice_path)

                return {
                    "docx_full": docx_full_io.getvalue(),
                    "pdf_full": pdf_full_io.getvalue(),
                    "docx_short": docx_short_io.getvalue(),
                    "pdf_short": pdf_short_io.getvalue(),
                }

            # — logika UI — generuj tylko na żądanie, ale pozwól pobrać poprzednie
            if prep:
                with st.spinner("Przygotowuję raport Word/PDF…"):
                    try:
                        exports = _build_exports()
                        st.session_state[cache_key] = {
                            **exports,
                            "docx_name_full": DOCX_FILENAME_FULL,
                            "pdf_name_full": PDF_FILENAME_FULL,
                            "docx_name_short": DOCX_FILENAME_SHORT,
                            "pdf_name_short": PDF_FILENAME_SHORT,
                        }
                        st.success("Gotowe. Możesz pobrać pliki poniżej.")
                    except Exception as e:
                        st.error(f"Nie udało się przygotować raportu: {e}")

            cache = st.session_state.get(cache_key)

            word_icon = "<svg width='21' height='21' viewBox='0 0 32 32' style='vertical-align:middle;margin-right:7px;margin-bottom:2px;'><rect width='32' height='32' rx='4' fill='#185abd'/><text x='16' y='22' text-anchor='middle' font-family='Segoe UI,Arial' font-size='16' fill='#fff' font-weight='bold'>W</text></svg>"
            pdf_icon = "<svg width='21' height='21' viewBox='0 0 32 32' style='vertical-align:middle;margin-right:7px;margin-bottom:2px;'><rect width='32' height='32' rx='4' fill='#d32f2f'/><text x='16' y='22' text-anchor='middle' font-family='Segoe UI,Arial' font-size='16' fill='#fff' font-weight='bold'>PDF</text></svg>"

            st.markdown(
                f"<div style='margin-bottom:11px;'>{word_icon}<b>Raport pełny (.docx)</b></div>",
                unsafe_allow_html=True)
            st.download_button(
                "Pobierz raport pełny (Word)",
                data=(cache["docx_full"] if cache else b""),
                file_name=(cache["docx_name_full"] if cache else DOCX_FILENAME_FULL),
                disabled=not bool(cache),
                key="word_button_full"
            )

            st.markdown(
                f"<div style='margin-top:21px; margin-bottom:11px;'>{pdf_icon}<b>Raport pełny (.pdf)</b></div>",
                unsafe_allow_html=True)
            st.download_button(
                "Pobierz raport pełny (PDF)",
                data=(cache["pdf_full"] if cache else b""),
                file_name=(cache["pdf_name_full"] if cache else PDF_FILENAME_FULL),
                disabled=not bool(cache),
                key="pdf_button_full"
            )

            st.markdown(
                f"<div style='margin-top:21px; margin-bottom:11px;'>{word_icon}<b>Raport skrócony (.docx)</b></div>",
                unsafe_allow_html=True)
            st.download_button(
                "Pobierz raport skrócony (Word)",
                data=(cache["docx_short"] if cache else b""),
                file_name=(cache["docx_name_short"] if cache else DOCX_FILENAME_SHORT),
                disabled=not bool(cache),
                key="word_button_short"
            )

            st.markdown(
                f"<div style='margin-top:21px; margin-bottom:11px;'>{pdf_icon}<b>Raport skrócony (.pdf)</b></div>",
                unsafe_allow_html=True)
            st.download_button(
                "Pobierz raport skrócony (PDF)",
                data=(cache["pdf_short"] if cache else b""),
                file_name=(cache["pdf_name_short"] if cache else PDF_FILENAME_SHORT),
                disabled=not bool(cache),
                key="pdf_button_short"
            )

            if not cache and not prep:
                st.caption(
                    "Pliki nie są jeszcze przygotowane. Włącz przełącznik powyżej, żeby je wygenerować.")

            st.markdown("""
            <hr style="height:1px; border:none; background:#eee; margin-top:38px; margin-bottom:24px;" />
            """, unsafe_allow_html=True)

            st.markdown("<div id='tabela'></div>", unsafe_allow_html=True)

            st.markdown('<div style="font-size:1.13em;font-weight:600;margin-bottom:13px;">Tabela odpowiedzi respondentów (pełne wyniki)</div>', unsafe_allow_html=True)
            final_df = results_df.copy()
            try:
                col_to_exclude = [
                    "Czas ankiety", "Archetyp", "Główny archetyp", "Cechy kluczowe", "Opis", "Storyline",
                    "Rekomendacje", "Archetyp wspierający", "Cechy wspierający", "Opis wspierający",
                    "Storyline wspierający", "Rekomendacje wspierający"
                ]
                means = final_df.drop(columns=col_to_exclude, errors="ignore").mean(numeric_only=True)
                summary_row = {col: round(means[col], 2) if col in means else "-" for col in final_df.columns}
                summary_row["Czas ankiety"] = "ŚREDNIA"
                final_df = pd.concat([final_df, pd.DataFrame([summary_row])], ignore_index=True)
            except Exception as e:
                pass
            st.dataframe(final_df, hide_index=True)
            st.caption("Eksport poniżej dotyczy tabeli odpowiedzi respondentów (CSV/XLSX), a nie raportów DOCX/PDF.")
            export_col_csv, export_col_xlsx, _export_spacer = st.columns([1, 1, 4])
            first_nom = (study.get("first_name_nom") or study.get("first_name") or "").strip()
            last_nom = (study.get("last_name_nom") or study.get("last_name") or "").strip()
            person_slug = _slug_pl(f"{first_nom} {last_nom}".strip() or "osoba")
            csv_table_name = f"{person_slug}_baza-odpowiedzi.csv"
            xlsx_table_name = f"{person_slug}_baza-odpowiedzi.xlsx"
            with export_col_csv:
                st.download_button(
                    "Pobierz tabelę odpowiedzi (CSV)",
                    final_df.to_csv(index=False),
                    csv_table_name,
                )
            with export_col_xlsx:
                buffer = io.BytesIO()
                final_df.to_excel(buffer, index=False)
                st.download_button(
                    label="Pobierz tabelę odpowiedzi (XLSX)",
                    data=buffer.getvalue(),
                    file_name=xlsx_table_name,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
    else:
        st.info("Brak danych – nie ma żadnych odpowiedzi w tym badaniu.")

